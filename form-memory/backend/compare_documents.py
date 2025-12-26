"""
Compare template and output DOCX files to identify differences in styles, formatting, and structure.
"""

from docx import Document
from pathlib import Path
from collections import defaultdict
import re

def compare_documents(template_path: str, output_path: str):
    """Compare template and output documents in detail."""
    
    template_doc = Document(template_path)
    output_doc = Document(output_path)
    
    print("=" * 80)
    print("DOCUMENT COMPARISON ANALYSIS")
    print("=" * 80)
    print(f"\nTemplate: {template_path}")
    print(f"Output: {output_path}\n")
    
    # 1. Basic Statistics
    print("=" * 80)
    print("1. BASIC STATISTICS")
    print("=" * 80)
    print(f"Template paragraphs: {len(template_doc.paragraphs)}")
    print(f"Output paragraphs: {len(output_doc.paragraphs)}")
    print(f"Template styles: {len(template_doc.styles)}")
    print(f"Output styles: {len(output_doc.styles)}")
    print(f"Template tables: {len(template_doc.tables)}")
    print(f"Output tables: {len(output_doc.tables)}")
    
    # 2. Style Comparison
    print("\n" + "=" * 80)
    print("2. STYLE COMPARISON")
    print("=" * 80)
    
    template_styles = {s.name: s for s in template_doc.styles}
    output_styles = {s.name: s for s in output_doc.styles}
    
    missing_styles = []
    for style_name in template_styles:
        if style_name not in output_styles:
            missing_styles.append(style_name)
    
    if missing_styles:
        print(f"\n[X] MISSING STYLES IN OUTPUT ({len(missing_styles)}):")
        for style in missing_styles:
            print(f"   - {style}")
    else:
        print("\n[OK] All template styles are present in output")
    
    added_styles = []
    for style_name in output_styles:
        if style_name not in template_styles:
            added_styles.append(style_name)
    
    if added_styles:
        print(f"\n[!] ADDED STYLES IN OUTPUT ({len(added_styles)}):")
        for style in added_styles[:10]:  # Show first 10
            print(f"   - {style}")
        if len(added_styles) > 10:
            print(f"   ... and {len(added_styles) - 10} more")
    
    # 3. Style Usage Analysis
    print("\n" + "=" * 80)
    print("3. STYLE USAGE IN TEMPLATE")
    print("=" * 80)
    
    template_style_usage = defaultdict(int)
    for para in template_doc.paragraphs:
        if para.style:
            template_style_usage[para.style.name] += 1
    
    print("\nMost used styles in template:")
    for style_name, count in sorted(template_style_usage.items(), key=lambda x: x[1], reverse=True)[:15]:
        print(f"   {style_name}: {count} paragraphs")
    
    print("\n" + "=" * 80)
    print("4. STYLE USAGE IN OUTPUT")
    print("=" * 80)
    
    output_style_usage = defaultdict(int)
    for para in output_doc.paragraphs:
        if para.style:
            output_style_usage[para.style.name] += 1
    
    print("\nMost used styles in output:")
    for style_name, count in sorted(output_style_usage.items(), key=lambda x: x[1], reverse=True)[:15]:
        print(f"   {style_name}: {count} paragraphs")
    
    # 4. Paragraph Style Differences
    print("\n" + "=" * 80)
    print("5. PARAGRAPH STYLE DIFFERENCES")
    print("=" * 80)
    
    # Analyze key sections
    key_sections = {
        'BAB': [],
        'Subbab': [],
        'Content': []
    }
    
    for i, para in enumerate(template_doc.paragraphs):
        text = para.text.strip()
        if re.match(r'^BAB\s+[IVX\d]+', text, re.I):
            key_sections['BAB'].append((i, text[:50], para.style.name if para.style else 'None'))
        elif re.search(r'\b(Subbab|Anak Subbab)\b', text, re.I) or re.match(r'^\d+\.\d+', text):
            key_sections['Subbab'].append((i, text[:50], para.style.name if para.style else 'None'))
        elif len(text) > 50 and para.style:
            key_sections['Content'].append((i, text[:50], para.style.name))
    
    print("\nTemplate structure:")
    print(f"   BAB headings: {len(key_sections['BAB'])}")
    print(f"   Subbab headings: {len(key_sections['Subbab'])}")
    print(f"   Content paragraphs: {len(key_sections['Content'])}")
    
    # Check output structure
    output_bab = []
    output_subbab = []
    output_content = []
    
    for i, para in enumerate(output_doc.paragraphs):
        text = para.text.strip()
        if re.match(r'^BAB\s+[IVX\d]+', text, re.I):
            output_bab.append((i, text[:50], para.style.name if para.style else 'None'))
        elif re.search(r'\b(Subbab|Anak Subbab)\b', text, re.I) or re.match(r'^\d+\.\d+', text):
            output_subbab.append((i, text[:50], para.style.name if para.style else 'None'))
        elif len(text) > 50 and para.style:
            output_content.append((i, text[:50], para.style.name))
    
    print("\nOutput structure:")
    print(f"   BAB headings: {len(output_bab)}")
    print(f"   Subbab headings: {len(output_subbab)}")
    print(f"   Content paragraphs: {len(output_content)}")
    
    # 5. Formatting Analysis
    print("\n" + "=" * 80)
    print("6. FORMATTING ANALYSIS")
    print("=" * 80)
    
    # Check paragraph formatting
    template_formatting = defaultdict(list)
    output_formatting = defaultdict(list)
    
    for para in template_doc.paragraphs[:100]:  # Sample first 100
        if para.text.strip() and para.style:
            style_name = para.style.name
            pf = para.paragraph_format
            
            template_formatting[style_name].append({
                'line_spacing': pf.line_spacing,
                'first_line_indent': pf.first_line_indent,
                'left_indent': pf.left_indent,
                'alignment': str(pf.alignment),
                'space_before': pf.space_before,
                'space_after': pf.space_after
            })
    
    for para in output_doc.paragraphs[:100]:  # Sample first 100
        if para.text.strip() and para.style:
            style_name = para.style.name
            pf = para.paragraph_format
            
            output_formatting[style_name].append({
                'line_spacing': pf.line_spacing,
                'first_line_indent': pf.first_line_indent,
                'left_indent': pf.left_indent,
                'alignment': str(pf.alignment),
                'space_before': pf.space_before,
                'space_after': pf.space_after
            })
    
    # Compare formatting for common styles
    common_styles = set(template_formatting.keys()) & set(output_formatting.keys())
    
    if common_styles:
        print("\nFormatting differences for common styles:")
        for style_name in sorted(common_styles)[:10]:
            template_fmt = template_formatting[style_name][0] if template_formatting[style_name] else {}
            output_fmt = output_formatting[style_name][0] if output_formatting[style_name] else {}
            
            differences = []
            for key in ['line_spacing', 'first_line_indent', 'left_indent', 'alignment']:
                if template_fmt.get(key) != output_fmt.get(key):
                    differences.append(f"{key}: template={template_fmt.get(key)} vs output={output_fmt.get(key)}")
            
            if differences:
                print(f"\n   [X] {style_name}:")
                for diff in differences:
                    print(f"      - {diff}")
    
    # 6. Font Analysis
    print("\n" + "=" * 80)
    print("7. FONT ANALYSIS")
    print("=" * 80)
    
    template_fonts = defaultdict(int)
    output_fonts = defaultdict(int)
    
    for para in template_doc.paragraphs[:200]:
        for run in para.runs:
            if run.font.name:
                template_fonts[run.font.name] += 1
    
    for para in output_doc.paragraphs[:200]:
        for run in para.runs:
            if run.font.name:
                output_fonts[run.font.name] += 1
    
    print("\nTemplate fonts:")
    for font, count in sorted(template_fonts.items(), key=lambda x: x[1], reverse=True):
        print(f"   {font}: {count} runs")
    
    print("\nOutput fonts:")
    for font, count in sorted(output_fonts.items(), key=lambda x: x[1], reverse=True):
        print(f"   {font}: {count} runs")
    
    # 7. Specific Issues
    print("\n" + "=" * 80)
    print("8. SPECIFIC ISSUES DETECTED")
    print("=" * 80)
    
    issues = []
    
    # Check for missing subsections
    if len(output_subbab) < len(key_sections['Subbab']):
        issues.append(f"Missing subsections: Template has {len(key_sections['Subbab'])} but output has {len(output_subbab)}")
    
    # Check for style mismatches in BAB headings
    template_bab_styles = {text[:20]: style for _, text, style in key_sections['BAB']}
    output_bab_styles = {text[:20]: style for _, text, style in output_bab}
    
    for bab_text, template_style in template_bab_styles.items():
        if bab_text in output_bab_styles:
            output_style = output_bab_styles[bab_text]
            if template_style != output_style:
                issues.append(f"BAB heading style mismatch for '{bab_text}': template uses '{template_style}', output uses '{output_style}'")
    
    # Check for content paragraphs using wrong styles
    template_content_styles = set(style for _, _, style in key_sections['Content'])
    output_content_styles = set(style for _, _, style in output_content)
    
    if template_content_styles and output_content_styles:
        missing_content_styles = template_content_styles - output_content_styles
        if missing_content_styles:
            issues.append(f"Content paragraphs not using expected styles: {missing_content_styles}")
    
    if issues:
        for issue in issues:
            print(f"   [X] {issue}")
    else:
        print("   [OK] No major issues detected")
    
    # 8. Recommendations
    print("\n" + "=" * 80)
    print("9. RECOMMENDATIONS")
    print("=" * 80)
    
    recommendations = []
    
    if missing_styles:
        recommendations.append(f"Copy missing styles from template: {', '.join(missing_styles[:5])}")
    
    if len(output_subbab) < len(key_sections['Subbab']):
        recommendations.append("Ensure all subsection headings are properly inserted")
    
    if template_content_styles and output_content_styles:
        if template_content_styles != output_content_styles:
            recommendations.append(f"Apply correct content paragraph styles. Template uses: {template_content_styles}")
    
    if recommendations:
        for i, rec in enumerate(recommendations, 1):
            print(f"   {i}. {rec}")
    else:
        print("   [OK] No specific recommendations")
    
    print("\n" + "=" * 80)
    print("ANALYSIS COMPLETE")
    print("=" * 80)

if __name__ == "__main__":
    template_path = r"C:\Folio\form-memory\storage\references\Template-skripsi-final-versi2020.docx"
    output_path = r"C:\Folio\form-memory\storage\outputs\Skripsi_Tarik.docx"
    
    compare_documents(template_path, output_path)

