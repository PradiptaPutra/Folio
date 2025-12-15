
from skripsi_enforcer import SkripsiEnforcer
from docx import Document
import os

# Create dummy doc
test_path = "test_frontmatter.docx"
doc = Document()
doc.add_paragraph("BAB I PENDAHULUAN")
doc.add_paragraph("Isi pendahuluan...")
doc.save(test_path)

# Enforce
enforcer = SkripsiEnforcer(test_path)
enforcer.phase_3_2_enforce_isiparagraf() # Ensure styles exist
enforcer.execute_phase_4(
    judul="IMPLEMENTASI SISTEM INFORMASI AKADEMIK BERBASIS WEB PADA UNIVERSITAS CONTOH",
    penulis="BUDI SANTOSO",
    nim="12345678",
    universitas="UNIVERSITAS TEKNOLOGI",
    tahun=2025,
    abstrak_teks="Ini adalah abstrak bahasa Indonesia.",
    abstrak_en_teks="This is the abstract in English."
)
enforcer.save()

# Inspection
doc = Document(test_path)
print("<DOCUMENT filename='improved_formatted.doc'>")
current_page = 1
for i, p in enumerate(doc.paragraphs):
    # Detect manual page breaks for visualization
    if 'w:pageBreakBefore' in p._element.xml:
        print(f"\n[PAGE BREAK]\n")
    
    text = p.text.strip()
    if not text:
        continue
        
    style = p.style.name
    print(f"[{style}] {text}")

print("... [Main Content Continues] ...")
print("</DOCUMENT>")

# Cleanup
try:
    os.remove(test_path)
except:
    pass
