"""
Verify that the fixes are working correctly by analyzing the new output.
"""

from docx import Document
from pathlib import Path
import re
from collections import defaultdict

template_path = r"C:\Folio\form-memory\storage\references\Template-skripsi-final-versi2020.docx"
output_path = r"C:\Folio\form-memory\storage\outputs\Skripsi_Redify.docx"

template_doc = Document(template_path)
output_doc = Document(output_path)

print("=" * 80)
print("VERIFICATION OF FIXES - NEW OUTPUT ANALYSIS")
print("=" * 80)
print(f"\nTemplate: {template_path}")
print(f"Output: {output_path}\n")

# 1. Check Subsection Count
print("=" * 80)
print("1. SUBSECTION COUNT VERIFICATION")
print("=" * 80)

template_subbabs = []
for i, para in enumerate(template_doc.paragraphs):
    text = para.text.strip()
    if re.search(r'\b(Subbab|Anak Subbab)\b', text, re.I) or re.match(r'^\d+\.\d+', text):
        template_subbabs.append((i, text[:60]))

output_subbabs = []
for i, para in enumerate(output_doc.paragraphs):
    text = para.text.strip()
    if re.search(r'\b(Subbab|Anak Subbab)\b', text, re.I) or re.match(r'^\d+\.\d+', text):
        output_subbabs.append((i, text[:60]))

print(f"\nTemplate subsections: {len(template_subbabs)}")
print(f"Output subsections: {len(output_subbabs)}")
print(f"Difference: {len(output_subbabs) - len(template_subbabs)}")

if len(output_subbabs) >= len(template_subbabs):
    print("[OK] Subsection count is correct or improved!")
else:
    print(f"[X] Still missing {len(template_subbabs) - len(output_subbabs)} subsections")

# 2. Check "Isi Paragraf" Style Usage
print("\n" + "=" * 80)
print("2. 'ISI PARAGRAF' STYLE USAGE")
print("=" * 80)

template_isi_count = sum(1 for p in template_doc.paragraphs if p.style and 'isi paragraf' in p.style.name.lower())
output_isi_count = sum(1 for p in output_doc.paragraphs if p.style and 'isi paragraf' in p.style.name.lower())

print(f"\nTemplate 'Isi Paragraf' paragraphs: {template_isi_count}")
print(f"Output 'Isi Paragraf' paragraphs: {output_isi_count}")
print(f"Difference: {output_isi_count - template_isi_count}")

if output_isi_count >= template_isi_count * 0.8:  # At least 80% of template
    print("[OK] 'Isi Paragraf' style usage is good!")
else:
    print(f"[X] Still missing {template_isi_count - output_isi_count} 'Isi Paragraf' paragraphs")

# 3. Check Font Usage
print("\n" + "=" * 80)
print("3. FONT VERIFICATION")
print("=" * 80)

template_fonts = defaultdict(int)
for para in template_doc.paragraphs[:100]:
    for run in para.runs:
        if run.font.name:
            template_fonts[run.font.name] += 1

output_fonts = defaultdict(int)
for para in output_doc.paragraphs[:100]:
    for run in para.runs:
        if run.font.name:
            output_fonts[run.font.name] += 1

print("\nTemplate fonts (first 100 paragraphs):")
for font, count in sorted(template_fonts.items(), key=lambda x: x[1], reverse=True):
    print(f"   {font}: {count} runs")

print("\nOutput fonts (first 100 paragraphs):")
for font, count in sorted(output_fonts.items(), key=lambda x: x[1], reverse=True):
    print(f"   {font}: {count} runs")

# Check if template font is being used
template_main_font = max(template_fonts.items(), key=lambda x: x[1])[0] if template_fonts else None
output_main_font = max(output_fonts.items(), key=lambda x: x[1])[0] if output_fonts else None

if template_main_font and output_main_font:
    if template_main_font == output_main_font:
        print(f"\n[OK] Font matches template: {template_main_font}")
    else:
        print(f"\n[X] Font mismatch: Template uses {template_main_font}, Output uses {output_main_font}")

# 4. Check Style Usage
print("\n" + "=" * 80)
print("4. STYLE USAGE COMPARISON")
print("=" * 80)

template_styles = defaultdict(int)
for para in template_doc.paragraphs:
    if para.style:
        template_styles[para.style.name] += 1

output_styles = defaultdict(int)
for para in output_doc.paragraphs:
    if para.style:
        output_styles[para.style.name] += 1

print("\nKey style differences:")
key_styles = ['Isi Paragraf', 'Normal', 'Heading 2', 'Heading 3', 'JUDUL']
for style in key_styles:
    template_count = template_styles.get(style, 0)
    output_count = output_styles.get(style, 0)
    diff = output_count - template_count
    status = "[OK]" if abs(diff) <= 5 or (style == 'Isi Paragraf' and output_count >= template_count * 0.7) else "[X]"
    print(f"   {style}: Template={template_count}, Output={output_count}, Diff={diff:+.0f} {status}")

# 5. Check Content Paragraph Styles
print("\n" + "=" * 80)
print("5. CONTENT PARAGRAPH STYLE ANALYSIS")
print("=" * 80)

output_content_styles = defaultdict(int)
for para in output_doc.paragraphs:
    text = para.text.strip()
    # Content paragraphs are long text that aren't headings
    if len(text) > 50 and not re.match(r'^BAB\s+[IVX\d]+', text, re.I) and not re.search(r'\b(Subbab|Anak Subbab)\b', text, re.I):
        if para.style:
            output_content_styles[para.style.name] += 1

print("\nOutput content paragraph styles:")
for style, count in sorted(output_content_styles.items(), key=lambda x: x[1], reverse=True)[:10]:
    print(f"   {style}: {count} paragraphs")

isi_in_content = output_content_styles.get('Isi Paragraf', 0)
normal_in_content = output_content_styles.get('Normal', 0)
total_content = sum(output_content_styles.values())

if total_content > 0:
    isi_percentage = (isi_in_content / total_content) * 100
    print(f"\n'Isi Paragraf' usage in content: {isi_in_content}/{total_content} ({isi_percentage:.1f}%)")
    if isi_percentage >= 50:
        print("[OK] Most content paragraphs use 'Isi Paragraf' style!")
    else:
        print(f"[X] Only {isi_percentage:.1f}% of content uses 'Isi Paragraf' (should be higher)")

# 6. Summary
print("\n" + "=" * 80)
print("6. SUMMARY")
print("=" * 80)

issues_found = []
improvements = []

# Subsection check
if len(output_subbabs) >= len(template_subbabs):
    improvements.append(f"Subsection count: {len(output_subbabs)} (target: {len(template_subbabs)})")
else:
    issues_found.append(f"Missing {len(template_subbabs) - len(output_subbabs)} subsections")

# Isi Paragraf check
if output_isi_count >= template_isi_count * 0.8:
    improvements.append(f"'Isi Paragraf' style: {output_isi_count} paragraphs")
else:
    issues_found.append(f"'Isi Paragraf' style: Only {output_isi_count} (target: {template_isi_count})")

# Font check
if template_main_font and output_main_font and template_main_font == output_main_font:
    improvements.append(f"Font: {output_main_font} (matches template)")
else:
    issues_found.append(f"Font mismatch: {output_main_font} vs {template_main_font}")

# Content style check
if total_content > 0 and isi_percentage >= 50:
    improvements.append(f"Content paragraphs: {isi_percentage:.1f}% use 'Isi Paragraf'")
else:
    issues_found.append(f"Content paragraphs: Only {isi_percentage:.1f}% use 'Isi Paragraf'")

if improvements:
    print("\n[IMPROVEMENTS]:")
    for imp in improvements:
        print(f"   [OK] {imp}")

if issues_found:
    print("\n[ISSUES REMAINING]:")
    for issue in issues_found:
        print(f"   [X] {issue}")
else:
    print("\n[OK] All major issues appear to be resolved!")

print("\n" + "=" * 80)
print("VERIFICATION COMPLETE")
print("=" * 80)

