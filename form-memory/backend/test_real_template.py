#!/usr/bin/env python
from pathlib import Path
from engine.analyzer.complete_thesis_builder import create_complete_thesis
import os
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

def test_real_template():
    """Test with the real Indonesian thesis template and content."""

    # File paths
    template_path = "C:\\Folio\\form-memory\\storage\\references\\Template-skripsi-final-versi2020.docx"
    content_path = "C:\\Folio\\form-memory\\storage\\uploads\\thesis_content.txt"
    output_path = "C:\\Folio\\form-memory\\storage\\outputs\\debug_test.docx"

    # Check if files exist
    print(f"Template exists: {Path(template_path).exists()}")
    print(f"Content exists: {Path(content_path).exists()}")

    if not Path(template_path).exists():
        print(f"ERROR: Template not found at {template_path}")
        return

    if not Path(content_path).exists():
        print(f"ERROR: Content not found at {content_path}")
        return

    # Get API key from environment
    api_key = os.getenv('OPENROUTER_API_KEY')
    print(f"API key present: {api_key is not None}")

    # Test with AI enabled
    print("\n=== TESTING WITH AI ENABLED ===")
    result = create_complete_thesis(
        template_path=template_path,
        content_path=content_path,
        output_path=output_path,
        use_ai=True,
        api_key=api_key,
        user_data={
            "title": "Sistem Informasi Akademik Berbasis Web",
            "author": "Budi Santoso",
            "supervisor": "Dr. Ahmad Rahman, M.T.",
            "year": "2024"
        }
    )

    print("Result:", result)
    if result['status'] == 'success':
        output_file = Path(result['output_file'])
        print(f"Output file size: {output_file.stat().st_size} bytes")
        print(f"Output file exists: {output_file.exists()}")

        # Analyze the final document quality
        print("\n=== DOCUMENT QUALITY ANALYSIS ===")
        from docx import Document
        try:
            generated_doc = Document(str(output_file))
            template_doc = Document(template_path)

            print("TEMPLATE vs GENERATED:")
            print(f"  Paragraphs: {len(template_doc.paragraphs)} -> {len(generated_doc.paragraphs)}")
            print(f"  Styles: {len(template_doc.styles)} -> {len(generated_doc.styles)}")
            print(f"  Tables: {len(template_doc.tables)} -> {len(generated_doc.tables)}")
            print(f"  Sections: {len(template_doc.sections)} -> {len(generated_doc.sections)}")

            # Check if content was actually added
            content_paragraphs = sum(1 for p in generated_doc.paragraphs if p.text.strip())
            template_paragraphs = sum(1 for p in template_doc.paragraphs if p.text.strip())
            print(f"  Content paragraphs: {template_paragraphs} -> {content_paragraphs}")

            # Quality assessment
            size_ratio = output_file.stat().st_size / Path(template_path).stat().st_size
            print(".2f")
            print(f"  Structure preserved: {'YES' if len(generated_doc.tables) >= len(template_doc.tables) else 'PARTIAL'}")
            print(f"  Content added: {'YES' if content_paragraphs > template_paragraphs else 'NO'}")

        except Exception as e:
            print(f"Error analyzing document: {e}")

if __name__ == '__main__':
    test_real_template()