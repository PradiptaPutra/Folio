#!/usr/bin/env python
from pathlib import Path
from docx import Document

def test_template_only():
    """Test just loading and saving the template without any processing."""
    print("Testing template-only operations...")

    template_path = "C:\\Folio\\form-memory\\storage\\references\\Template-skripsi-final-versi2020.docx"
    output_path = "C:\\Folio\\form-memory\\storage\\outputs\\template_only_test.docx"

    try:
        print("Loading template...")
        doc = Document(template_path)
        print(f"Template loaded: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

        print("Saving template as-is...")
        doc.save(output_path)
        print(f"Saved to {output_path}")

        # Verify
        saved_doc = Document(output_path)
        print(f"Saved document: {len(saved_doc.paragraphs)} paragraphs, {len(saved_doc.tables)} tables")

        print("SUCCESS: Template-only test completed")

    except Exception as e:
        print(f"FAILED: {e}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    test_template_only()