
import os
import shutil
from pathlib import Path
from docx import Document
from text_normalizer import normalize_txt_to_markdown
from pandoc_runner import markdown_to_docx
from skripsi_formatter import enforce_skripsi_format
from docx_inspector import extract_docx_styles, detect_style_usage

# Setup paths
BASE_DIR = Path(__file__).resolve().parent
TEST_DIR = BASE_DIR / "tests_temp"
TEST_DIR.mkdir(exist_ok=True)

TEMPLATE_PATH = TEST_DIR / "template.docx"
CONTENT_PATH = TEST_DIR / "content.txt"
OUTPUT_DOCX = TEST_DIR / "output.docx"
OUTPUT_DOC = TEST_DIR / "output.doc"

def create_dummy_template():
    """Create a minimal valid DOCX template."""
    doc = Document()
    doc.add_heading("Template Heading 1", level=1)
    doc.add_heading("Template Heading 2", level=2)
    doc.add_paragraph("Template Normal text.")
    doc.save(TEMPLATE_PATH)
    print(f"Created template at {TEMPLATE_PATH}")

def create_dummy_content():
    """Create raw text content with manual numbering to test normalization."""
    content = """
BAB I PENDAHULUAN
1.1 Latar Belakang
Ini adalah latar belakang masalah. Harusnya tidak ada angka 1.1 di sini.

1.2 Rumusan Masalah
Masalahnya apa?

BAB II PEMBAHASAN
2.1 Teori
Teori A.

2.1.1 Sub-Teori
Detail teori.
"""
    CONTENT_PATH.write_text(content, encoding="utf-8")
    print(f"Created content at {CONTENT_PATH}")

def run_verification():
    print("--- Starting Verification ---")
    
    # 1. Setup
    create_dummy_template()
    create_dummy_content()
    
    # 2. Text Normalization
    print("Running Text Normalization...")
    raw_text = CONTENT_PATH.read_text(encoding="utf-8")
    normalized_md = normalize_txt_to_markdown(raw_text)
    
    md_path = TEST_DIR / "content.md"
    md_path.write_text(normalized_md, encoding="utf-8")
    print(f"Normalized Markdown:\n{normalized_md[:200]}...\n")
    
    # 3. Extract Styles (Simulate Pipeline)
    print("Extracting styles...")
    extracted = extract_docx_styles(TEMPLATE_PATH)
    # Mock usage detection
    style_usage = {"chapter_style": "Heading 1", "body_style": "Normal"}
    
    style_config = {
        "margins": extracted["margins"],
        "paragraph": extracted["styles"].get("Normal", {}).get("paragraph", {}),
        "mapping": style_usage
    }
    
    # 4. Generate DOCX
    print("Generating DOCX...")
    try:
        markdown_to_docx(md_path, TEMPLATE_PATH, OUTPUT_DOCX, style_config)
        print(f"Generated DOCX at {OUTPUT_DOCX}")
    except Exception as e:
        print(f"FAILED to check pandoc/generation: {e}")
        return

    # 5. Check Output
    doc = Document(OUTPUT_DOCX)
    print("Inspecting Output...")
    
    # Check Heading 1
    h1_count = 0
    h1_texts = []
    for p in doc.paragraphs:
        if p.style.name == "Heading 1":
            h1_count += 1
            h1_texts.append(p.text)
            
            # Check Page Break (Phase 4)
            pPr = p._element.get_or_add_pPr()
            pb = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pageBreakBefore")
            if pb is not None and pb.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "on":
                print(f"  [OK] Heading 1 '{p.text}' has PageBreakBefore.")
            else:
                print(f"  [WARN] Heading 1 '{p.text}' MISSING PageBreakBefore.")
                
            # Check Spacing (Phase 3.4)
            spacing = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing")
            if spacing is not None:
                before = spacing.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}before")
                after = spacing.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}after")
                if before == "480" and after == "240":
                    print(f"  [OK] Heading 1 '{p.text}' spacing correct (24pt/12pt).")
                else:
                    print(f"  [FAIL] Heading 1 spacing incorrect: {before}/{after}")
    
    if h1_count >= 2:
        print(f"  [OK] Found {h1_count} Heading 1s: {h1_texts}")
    else:
        print(f"  [FAIL] Expected at least 2 Heading 1s, found {h1_count}")

    # Check Numbering XML existence
    if doc.part.numbering_part:
        print("  [OK] Numbering part exists.")
    else:
        print("  [FAIL] Numbering part missing!")

    print("--- Verification Complete ---")

if __name__ == "__main__":
    run_verification()
