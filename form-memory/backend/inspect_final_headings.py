from docx import Document
from docx.oxml.ns import qn

path = r"c:\Folio\form-memory\storage\outputs\final.docx"
print(f"Inspecting: {path}")

doc = Document(path)

print("\n--- TOC Field ---")
for p in doc.paragraphs[:5]:
    if "TOC" in p._element.xml:
        print(f"Found TOC in paragraph: {p.text[:20]}")
        print(p._element.xml[:200]) # Print snippet

print("\n--- Headings Analysis ---")
for i, p in enumerate(doc.paragraphs):
    if "BAB" in p.text or p.style.name.startswith("Heading"):
        print(f"Text: {p.text[:40]}...")
        print(f"Style: '{p.style.name}'")
        
        # Check outline level in XML
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None:
            outline = pPr.find(qn("w:outlineLvl"))
            if outline is not None:
                print(f"XML Outline Level: {outline.get(qn('w:val'))}")
            else:
                print("XML Outline Level: None")
        else:
            print("No pPr")
        print("-" * 10)
        
    if i > 50: break
