"""
Analyze output files to identify content placement errors.
"""

from docx import Document
import re
from pathlib import Path

output1 = r"C:\Folio\form-memory\storage\outputs\Skripsi_Gawx.docx"
output2 = r"C:\Folio\form-memory\storage\outputs\Skripsi_Arsyad.docx"

def analyze_document(doc_path, name):
    """Analyze a document to find content placement issues."""
    if not Path(doc_path).exists():
        print(f"[SKIP] File not found: {doc_path}")
        return
    
    print("=" * 80)
    print(f"ANALYZING: {name}")
    print("=" * 80)
    
    doc = Document(doc_path)
    
    # Front matter indicators
    front_matter_keywords = [
        'DAFTAR ISI', 'DAFTAR TABEL', 'DAFTAR GAMBAR', 'DAFTAR SINGKATAN',
        'ABSTRACT', 'ABSTRAK', 'KATA PENGANTAR', 'PREFACE',
        'HALAMAN JUDUL', 'HALAMAN PENGESAHAN', 'LEMBAR PERSETUJUAN',
        'TABLE OF CONTENTS', 'LIST OF TABLES', 'LIST OF FIGURES'
    ]
    
    # Main content indicators
    main_content_keywords = [
        'BAB I', 'BAB II', 'BAB III', 'BAB IV', 'BAB V', 'BAB VI',
        'CHAPTER 1', 'CHAPTER 2', 'CHAPTER 3',
        'PENDAHULUAN', 'TINJAUAN PUSTAKA', 'METODOLOGI'
    ]
    
    # Track sections
    front_matter_sections = []
    main_content_sections = []
    misplaced_content = []
    
    in_front_matter = True
    front_matter_end = None
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        text_upper = text.upper()
        
        # Check if we're in front matter
        if any(keyword in text_upper for keyword in front_matter_keywords):
            front_matter_sections.append({
                'para_idx': i,
                'text': text[:80],
                'type': 'front_matter'
            })
            in_front_matter = True
            front_matter_end = i
        
        # Check if we've entered main content
        if any(keyword in text_upper for keyword in main_content_keywords):
            if in_front_matter:
                print(f"[STRUCTURE] Front matter ends around paragraph {i}")
                front_matter_end = i
            in_front_matter = False
            main_content_sections.append({
                'para_idx': i,
                'text': text[:80],
                'type': 'main_content'
            })
        
        # Check for content that looks like it should be in main content but is in front matter
        if in_front_matter and front_matter_end and i > front_matter_end:
            # Check if this looks like chapter content
            if (re.match(r'^\d+\.\d+', text) or  # Subsection numbering
                len(text) > 200 or  # Long paragraphs (likely content)
                any(kw in text_upper for kw in ['LATAR BELAKANG', 'RUMUSAN MASALAH', 'TUJUAN', 'MANFAAT'])):
                misplaced_content.append({
                    'para_idx': i,
                    'text': text[:100],
                    'issue': 'Content in front matter area'
                })
    
    print(f"\n[STRUCTURE ANALYSIS]")
    print(f"  Front matter sections: {len(front_matter_sections)}")
    print(f"  Main content sections: {len(main_content_sections)}")
    print(f"  Potential misplaced content: {len(misplaced_content)}")
    
    if front_matter_sections:
        print(f"\n  Front Matter Sections Found:")
        for section in front_matter_sections[:10]:
            print(f"    Para {section['para_idx']}: {section['text']}")
    
    if main_content_sections:
        print(f"\n  Main Content Sections Found:")
        for section in main_content_sections[:10]:
            print(f"    Para {section['para_idx']}: {section['text']}")
    
    if misplaced_content:
        print(f"\n  [ERROR] MISPLACED CONTENT DETECTED:")
        for item in misplaced_content[:10]:
            print(f"    Para {item['para_idx']}: {item['text']}")
            print(f"      Issue: {item['issue']}")
    
    # Check for content in Daftar Isi area
    print(f"\n  [DAFTAR ISI CHECK]")
    daftar_isi_paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if 'DAFTAR ISI' in text.upper() or 'TABLE OF CONTENTS' in text.upper():
            # Check next 50 paragraphs for content
            for j in range(i, min(i + 50, len(doc.paragraphs))):
                next_para = doc.paragraphs[j]
                next_text = next_para.text.strip()
                # Check if this looks like chapter content
                if (re.match(r'^\d+\.\d+', next_text) or  # Subsection
                    (len(next_text) > 150 and not any(kw in next_text.upper() for kw in ['DAFTAR', 'TABLE', 'LIST', 'ABSTRACT']))):
                    daftar_isi_paragraphs.append({
                        'para_idx': j,
                        'text': next_text[:100],
                        'distance_from_daftar_isi': j - i
                    })
    
    if daftar_isi_paragraphs:
        print(f"    [ERROR] Found {len(daftar_isi_paragraphs)} paragraphs with content-like text near Daftar Isi:")
        for item in daftar_isi_paragraphs[:5]:
            print(f"      Para {item['para_idx']} (distance: {item['distance_from_daftar_isi']}): {item['text']}")
    else:
        print(f"    [OK] No content found in Daftar Isi area")

# Analyze both files
analyze_document(output1, "Skripsi_Gawx.docx")
print("\n\n")
analyze_document(output2, "Skripsi_Arsyad.docx")
