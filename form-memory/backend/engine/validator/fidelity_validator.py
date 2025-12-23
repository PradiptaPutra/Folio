"""
Fidelity Validator
XML-level diff detection between template and output.
No AI - deterministic comparison only.
"""

from docx import Document
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
import xml.etree.ElementTree as ET


class FidelityValidator:
    """Validate template fidelity at XML level."""
    
    def __init__(self, template_path: str, output_path: str):
        """Initialize with template and output paths."""
        self.template_path = Path(template_path)
        self.output_path = Path(output_path)
        
        self.template_doc = Document(str(self.template_path))
        self.output_doc = Document(str(self.output_path))
    
    def validate(self) -> Dict[str, Any]:
        """Perform complete validation."""
        diffs = {
            "styles": self._compare_styles(),
            "numbering": self._compare_numbering(),
            "margins": self._compare_margins(),
            "sections": self._compare_sections(),
            "paragraph_count": self._compare_paragraph_count(),
            "content_changes": self._detect_content_changes(),
        }
        
        return {
            "diffs": diffs,
            "fidelity_score": self._calculate_fidelity_score(diffs),
            "is_valid": self._is_valid(diffs),
        }
    
    def _compare_styles(self) -> List[Dict[str, Any]]:
        """Compare styles.xml between template and output."""
        diffs = []
        
        template_styles = {s.name: s for s in self.template_doc.styles}
        output_styles = {s.name: s for s in self.output_doc.styles}
        
        # Check for removed styles
        for style_name in template_styles:
            if style_name not in output_styles:
                diffs.append({
                    "type": "style_removed",
                    "style": style_name,
                    "severity": "warning"
                })
        
        # Check for added styles (less critical)
        for style_name in output_styles:
            if style_name not in template_styles:
                diffs.append({
                    "type": "style_added",
                    "style": style_name,
                    "severity": "info"
                })
        
        return diffs
    
    def _compare_numbering(self) -> List[Dict[str, Any]]:
        """Compare numbering.xml between template and output."""
        diffs = []
        
        template_numbering = self.template_doc._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numbering')
        output_numbering = self.output_doc._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numbering')
        
        # Check if numbering exists in template but not in output
        if template_numbering is not None and output_numbering is None:
            diffs.append({
                "type": "numbering_removed",
                "severity": "warning"
            })
        
        # Count numbering definitions
        if template_numbering is not None and output_numbering is not None:
            template_count = len(template_numbering.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num'))
            output_count = len(output_numbering.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num'))
            
            if template_count != output_count:
                diffs.append({
                    "type": "numbering_definition_changed",
                    "template_count": template_count,
                    "output_count": output_count,
                    "severity": "info"
                })
        
        return diffs
    
    def _compare_margins(self) -> List[Dict[str, Any]]:
        """Compare section margins."""
        diffs = []
        
        template_margins = self._extract_margins(self.template_doc)
        output_margins = self._extract_margins(self.output_doc)
        
        if template_margins != output_margins:
            diffs.append({
                "type": "margins_changed",
                "template": template_margins,
                "output": output_margins,
                "severity": "critical"
            })
        
        return diffs
    
    def _compare_sections(self) -> List[Dict[str, Any]]:
        """Compare section properties."""
        diffs = []
        
        template_sections = len(self.template_doc.sections)
        output_sections = len(self.output_doc.sections)
        
        if template_sections != output_sections:
            diffs.append({
                "type": "section_count_changed",
                "template_count": template_sections,
                "output_count": output_sections,
                "severity": "warning"
            })
        
        return diffs
    
    def _compare_paragraph_count(self) -> List[Dict[str, Any]]:
        """Compare paragraph counts."""
        diffs = []
        
        template_count = len(self.template_doc.paragraphs)
        output_count = len(self.output_doc.paragraphs)
        
        # Allow some variation for injected content
        if output_count < template_count - 5:
            diffs.append({
                "type": "paragraph_count_decreased",
                "template_count": template_count,
                "output_count": output_count,
                "severity": "critical"
            })
        
        return diffs
    
    def _detect_content_changes(self) -> List[Dict[str, Any]]:
        """Detect significant content changes."""
        diffs = []
        
        template_paras = [p.text for p in self.template_doc.paragraphs]
        output_paras = [p.text for p in self.output_doc.paragraphs]
        
        # Check if all template content still exists
        lost_content = []
        for i, para_text in enumerate(template_paras):
            if para_text.strip() and para_text not in output_paras:
                lost_content.append({
                    "index": i,
                    "text": para_text[:100],
                })
        
        if lost_content:
            diffs.append({
                "type": "template_content_lost",
                "count": len(lost_content),
                "samples": lost_content[:3],
                "severity": "critical"
            })
        
        return diffs
    
    @staticmethod
    def _extract_margins(doc: Document) -> Dict[str, float]:
        """Extract margin values from document."""
        if not doc.sections:
            return {}
        
        section = doc.sections[0]
        return {
            "top": section.top_margin.cm if section.top_margin else 0,
            "bottom": section.bottom_margin.cm if section.bottom_margin else 0,
            "left": section.left_margin.cm if section.left_margin else 0,
            "right": section.right_margin.cm if section.right_margin else 0,
        }
    
    @staticmethod
    def _calculate_fidelity_score(diffs: Dict[str, List]) -> float:
        """Calculate overall fidelity score (0.0-1.0)."""
        penalty = 0.0
        
        for diff_list in diffs.values():
            if not isinstance(diff_list, list):
                continue
            
            for diff in diff_list:
                severity = diff.get("severity", "info")
                
                if severity == "critical":
                    penalty += 0.3
                elif severity == "warning":
                    penalty += 0.1
                elif severity == "info":
                    penalty += 0.02
        
        return max(0.0, 1.0 - penalty)
    
    @staticmethod
    def _is_valid(diffs: Dict[str, List]) -> bool:
        """Check if document is valid (no critical diffs)."""
        for diff_list in diffs.values():
            if not isinstance(diff_list, list):
                continue
            
            for diff in diff_list:
                if diff.get("severity") == "critical":
                    return False
        
        return True
