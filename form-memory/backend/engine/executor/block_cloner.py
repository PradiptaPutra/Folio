"""
Block Cloner
Clones entire paragraph blocks preserving all formatting and structure.
No modification of styles, numbering, or section properties.
"""

from docx import Document
from docx.oxml import parse_xml
from copy import deepcopy
from typing import List, Optional


class BlockCloner:
    """Clones paragraph blocks from template."""
    
    def __init__(self, doc: Document):
        self.doc = doc
        self.body = doc._element.body
    
    def clone_paragraph(self, para_index: int) -> object:
        """Clone a paragraph element preserving all formatting."""
        if para_index >= len(self.doc.paragraphs):
            raise IndexError(f"Paragraph index {para_index} out of range")
        
        source_para = self.doc.paragraphs[para_index]
        source_elem = source_para._element
        
        # Deep copy the XML element
        cloned_elem = deepcopy(source_elem)
        
        return cloned_elem
    
    def clone_paragraph_range(self, start_idx: int, end_idx: int) -> List[object]:
        """Clone range of paragraphs [start_idx, end_idx)."""
        cloned = []
        for idx in range(start_idx, end_idx):
            cloned.append(self.clone_paragraph(idx))
        return cloned
    
    def insert_cloned_paragraph(
        self,
        cloned_elem,
        insert_after_para_index: int
    ) -> None:
        """Insert cloned paragraph after specified paragraph."""
        if insert_after_para_index >= len(self.doc.paragraphs):
            raise IndexError(f"Paragraph index {insert_after_para_index} out of range")
        
        anchor_para = self.doc.paragraphs[insert_after_para_index]
        anchor_elem = anchor_para._element
        
        # Insert after anchor
        anchor_elem.addnext(cloned_elem)
    
    def insert_cloned_paragraphs(
        self,
        cloned_elems: List[object],
        insert_after_para_index: int
    ) -> None:
        """Insert multiple cloned paragraphs after specified paragraph."""
        anchor_para = self.doc.paragraphs[insert_after_para_index]
        anchor_elem = anchor_para._element
        
        # Insert in reverse order so they appear in correct order
        for cloned_elem in reversed(cloned_elems):
            anchor_elem.addnext(cloned_elem)
    
    def clone_and_replace_at(
        self,
        source_para_index: int,
        target_para_index: int
    ) -> None:
        """Clone source paragraph and replace target paragraph."""
        if (source_para_index >= len(self.doc.paragraphs) or 
            target_para_index >= len(self.doc.paragraphs)):
            raise IndexError("Paragraph index out of range")
        
        cloned_elem = self.clone_paragraph(source_para_index)
        target_para = self.doc.paragraphs[target_para_index]
        target_elem = target_para._element
        
        # Replace
        target_elem.getparent().replace(target_elem, cloned_elem)
    
    def clone_section(
        self,
        start_para_idx: int,
        end_para_idx: int
    ) -> List[object]:
        """Clone entire section of paragraphs."""
        cloned = []
        for idx in range(start_para_idx, end_para_idx + 1):
            cloned.append(self.clone_paragraph(idx))
        return cloned
    
    def get_paragraph_runs(self, para_index: int) -> List[object]:
        """Get all run elements from a paragraph."""
        if para_index >= len(self.doc.paragraphs):
            raise IndexError(f"Paragraph index {para_index} out of range")
        
        para = self.doc.paragraphs[para_index]
        return para._element.findall('.//w:r', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
