"""
Template Executor
Generates documents by cloning and injecting content into user-uploaded templates.
Pure template reuse - no document creation, no formatting assumptions.
"""

from docx import Document
from pathlib import Path
from typing import Dict, Any, Optional, List

from .anchor_discovery import AnchorDiscovery, AnchorDescriptor
from .block_cloner import BlockCloner
from .text_replacer import TextReplacer


class TemplateExecutor:
    """
    Executor that generates documents by cloning template blocks.
    
    Rules:
    - Single Document() call only (template loading)
    - All formatting must exist in template
    - Content injection via rule-driven matching
    - No style, margin, or numbering modifications
    """
    
    def __init__(self, template_path: str, compiled_rules: Dict[str, Any]):
        """
        Initialize executor.
        
        Args:
            template_path: Path to user-uploaded template DOCX
            compiled_rules: Rule dictionary defining anchor matching and content injection
        """
        self.template_path = Path(template_path)
        self.compiled_rules = compiled_rules
        
        # Load template once - NO other Document() calls
        self.doc = Document(str(self.template_path))
        
        # Initialize helper engines
        self.anchor_discovery = AnchorDiscovery(self.doc)
        self.block_cloner = BlockCloner(self.doc)
        self.text_replacer = TextReplacer()
        
        # Cache discovered anchors
        self.anchors_cache = {}
    
    def generate(self, structural_doc: Dict[str, Any], output_path: str) -> None:
        """
        Generate output document by injecting content into template.
        
        Args:
            structural_doc: Semantic structure with content sections
            output_path: Output file path
        """
        output_path = Path(output_path)
        
        try:
            # Process each section in structural_doc
            for section_key, section_content in structural_doc.items():
                self._inject_section(section_key, section_content)
            
            # Save document preserving all template properties
            self.doc.save(str(output_path))
            
        except Exception as e:
            raise RuntimeError(f"Document generation failed: {str(e)}")
    
    def _inject_section(self, section_key: str, content: Any) -> None:
        """Inject content for a section using compiled rules."""
        if section_key not in self.compiled_rules:
            # Skip unknown sections
            return
        
        rule = self.compiled_rules[section_key]
        anchor = self._find_anchor(rule)
        
        if anchor is None:
            raise ValueError(f"Anchor not found for section: {section_key}")
        
        if isinstance(content, str):
            self._inject_text(anchor, content)
        elif isinstance(content, list):
            self._inject_list(anchor, content)
        elif isinstance(content, dict):
            self._inject_dict(anchor, content)
    
    def _find_anchor(self, rule: Dict[str, Any]) -> Optional[int]:
        """Find paragraph index matching rule criteria."""
        anchor_type = rule.get('anchor_type')
        
        if anchor_type == 'style':
            return self.anchor_discovery.find_paragraph_by_style(
                rule.get('style_name')
            )
        
        elif anchor_type == 'outline_level':
            return self.anchor_discovery.find_paragraph_by_outline_level(
                rule.get('level')
            )
        
        elif anchor_type == 'style_and_level':
            return self.anchor_discovery.find_paragraph_by_style_and_level(
                rule.get('style_name'),
                rule.get('level')
            )
        
        elif anchor_type == 'text_contains':
            return self.anchor_discovery.find_paragraph_containing_text(
                rule.get('text')
            )
        
        elif anchor_type == 'index':
            idx = rule.get('index')
            if idx < len(self.doc.paragraphs):
                return idx
        
        return None
    
    def _inject_text(self, anchor_idx: int, text: str) -> None:
        """Inject plain text at anchor."""
        if anchor_idx >= len(self.doc.paragraphs):
            return
        
        para = self.doc.paragraphs[anchor_idx]
        self.text_replacer.replace_paragraph_text(para, text)
    
    def _inject_list(self, anchor_idx: int, items: List[str]) -> None:
        """Inject list of items cloning paragraph for each."""
        if anchor_idx >= len(self.doc.paragraphs):
            return
        
        # Set first item at anchor
        if items:
            self.text_replacer.replace_paragraph_text(
                self.doc.paragraphs[anchor_idx],
                items[0]
            )
            
            # Clone paragraph for remaining items
            for item in items[1:]:
                cloned = self.block_cloner.clone_paragraph(anchor_idx)
                self.block_cloner.insert_cloned_paragraph(cloned, anchor_idx)
                
                # Update text in cloned paragraph
                # Refresh paragraph reference after insertion
                anchor_idx += 1
                self.text_replacer.replace_paragraph_text(
                    self.doc.paragraphs[anchor_idx],
                    item
                )
    
    def _inject_dict(self, anchor_idx: int, data: Dict[str, str]) -> None:
        """Inject structured data (placeholders in template)."""
        if anchor_idx >= len(self.doc.paragraphs):
            return
        
        para = self.doc.paragraphs[anchor_idx]
        
        for key, value in data.items():
            placeholder = "{" + key + "}"
            self.text_replacer.replace_text_in_runs(para, placeholder, str(value))
    
    def clone_block(
        self,
        source_para_idx: int,
        num_clones: int,
        insert_after_idx: int
    ) -> List[int]:
        """Clone a paragraph block multiple times."""
        cloned_indices = []
        current_insert_idx = insert_after_idx
        
        for i in range(num_clones):
            cloned = self.block_cloner.clone_paragraph(source_para_idx)
            self.block_cloner.insert_cloned_paragraph(cloned, current_insert_idx)
            current_insert_idx += 1
            cloned_indices.append(current_insert_idx)
        
        return cloned_indices
    
    def find_anchors_by_rule(self, rule: Dict[str, Any]) -> List[int]:
        """Find all paragraphs matching a rule."""
        anchor_type = rule.get('anchor_type')
        
        if anchor_type == 'style':
            return self.anchor_discovery.find_paragraphs_by_style(
                rule.get('style_name')
            )
        
        elif anchor_type == 'outline_level':
            return self.anchor_discovery.find_paragraphs_by_outline_level(
                rule.get('level')
            )
        
        elif anchor_type == 'text_contains':
            return self.anchor_discovery.find_paragraphs_containing_text(
                rule.get('text')
            )
        
        return []
    
    def get_template_metadata(self) -> Dict[str, Any]:
        """Get template metadata for rule compilation."""
        return {
            'paragraph_count': len(self.doc.paragraphs),
            'section_count': len(self.doc.sections),
            'style_names': [p.style.name for p in self.doc.paragraphs if p.style],
            'anchors': self.anchor_discovery.discover_all_anchors(),
        }
