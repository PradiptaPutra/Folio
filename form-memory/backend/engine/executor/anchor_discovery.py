"""
Anchor Discovery Engine
Identifies template blocks via style names, outline levels, numbering patterns, section boundaries.
No campus-specific assumptions.
"""

from docx import Document
from docx.oxml.ns import qn
from typing import List, Dict, Optional, Tuple


class AnchorDescriptor:
    """Describes an anchor block location in the template."""
    
    def __init__(
        self,
        para_index: int,
        style_name: Optional[str] = None,
        outline_level: Optional[int] = None,
        numbering_pattern: Optional[str] = None,
        section_index: Optional[int] = None,
        text_contains: Optional[str] = None,
    ):
        self.para_index = para_index
        self.style_name = style_name
        self.outline_level = outline_level
        self.numbering_pattern = numbering_pattern
        self.section_index = section_index
        self.text_contains = text_contains


class AnchorDiscovery:
    """Discovers anchor blocks in template using rule-driven matching."""
    
    def __init__(self, doc: Document):
        self.doc = doc
        self.paragraphs = doc.paragraphs
        self.sections = doc.sections
    
    def find_paragraph_by_style(self, style_name: str) -> Optional[int]:
        """Find first paragraph with exact style name."""
        for idx, para in enumerate(self.paragraphs):
            if para.style.name == style_name:
                return idx
        return None
    
    def find_paragraphs_by_style(self, style_name: str) -> List[int]:
        """Find all paragraphs with exact style name."""
        indices = []
        for idx, para in enumerate(self.paragraphs):
            if para.style.name == style_name:
                indices.append(idx)
        return indices
    
    def find_paragraph_by_outline_level(self, level: int) -> Optional[int]:
        """Find first paragraph with exact outline level."""
        for idx, para in enumerate(self.paragraphs):
            if self._get_outline_level(para) == level:
                return idx
        return None
    
    def find_paragraphs_by_outline_level(self, level: int) -> List[int]:
        """Find all paragraphs with exact outline level."""
        indices = []
        for idx, para in enumerate(self.paragraphs):
            if self._get_outline_level(para) == level:
                indices.append(idx)
        return indices
    
    def find_paragraph_containing_text(self, text: str) -> Optional[int]:
        """Find first paragraph containing text (case-insensitive)."""
        text_lower = text.lower()
        for idx, para in enumerate(self.paragraphs):
            if text_lower in para.text.lower():
                return idx
        return None
    
    def find_paragraphs_containing_text(self, text: str) -> List[int]:
        """Find all paragraphs containing text (case-insensitive)."""
        indices = []
        text_lower = text.lower()
        for idx, para in enumerate(self.paragraphs):
            if text_lower in para.text.lower():
                indices.append(idx)
        return indices
    
    def find_paragraph_by_style_and_level(
        self, 
        style_name: str, 
        outline_level: int
    ) -> Optional[int]:
        """Find paragraph matching both style and outline level."""
        for idx, para in enumerate(self.paragraphs):
            if (para.style.name == style_name and 
                self._get_outline_level(para) == outline_level):
                return idx
        return None
    
    def find_numbered_sequence(
        self, 
        style_name: str, 
        count: int
    ) -> Optional[List[int]]:
        """Find sequence of numbered paragraphs with same style."""
        matches = []
        consecutive = []
        
        for idx, para in enumerate(self.paragraphs):
            if para.style.name == style_name and self._is_numbered(para):
                if not consecutive or consecutive[-1] == idx - 1:
                    consecutive.append(idx)
                    if len(consecutive) == count:
                        return consecutive
                else:
                    consecutive = [idx]
            else:
                consecutive = []
        
        return None
    
    def get_section_for_paragraph(self, para_index: int) -> Optional[int]:
        """Get section index for a paragraph."""
        para = self.paragraphs[para_index]
        para_elem = para._element
        
        for sect_idx, section in enumerate(self.sections):
            sect_elem = section._sectPr
            if sect_elem is None:
                continue
            
            # Check if paragraph is in this section
            body = self.doc._element.body
            sect_parent = sect_elem.getparent()
            
            # Section contains all elements up to its sectPr
            if sect_parent is body:
                sect_idx_found = list(body).index(sect_parent)
                para_idx_found = list(body).index(para_elem)
                if para_idx_found <= sect_idx_found:
                    return sect_idx
        
        return 0
    
    def get_anchor_descriptor(self, para_index: int) -> AnchorDescriptor:
        """Build descriptor for a paragraph."""
        if para_index >= len(self.paragraphs):
            return None
        
        para = self.paragraphs[para_index]
        
        return AnchorDescriptor(
            para_index=para_index,
            style_name=para.style.name,
            outline_level=self._get_outline_level(para),
            numbering_pattern=self._detect_numbering_pattern(para),
            section_index=self.get_section_for_paragraph(para_index),
            text_contains=para.text[:50] if para.text else None,
        )
    
    def discover_all_anchors(self) -> List[AnchorDescriptor]:
        """Return descriptors for all paragraphs with identifiable anchors."""
        anchors = []
        for idx in range(len(self.paragraphs)):
            desc = self.get_anchor_descriptor(idx)
            if (desc.style_name or desc.outline_level is not None or 
                desc.numbering_pattern or desc.text_contains):
                anchors.append(desc)
        return anchors
    
    @staticmethod
    def _get_outline_level(para) -> Optional[int]:
        """Extract outline level from paragraph properties."""
        pPr = para._element.pPr
        if pPr is None:
            return None
        
        outlineLvl = pPr.find(qn('w:outlineLvl'))
        if outlineLvl is not None:
            try:
                return int(outlineLvl.get(qn('w:val')))
            except (ValueError, TypeError):
                return None
        
        return None
    
    @staticmethod
    def _is_numbered(para) -> bool:
        """Check if paragraph has numbering."""
        pPr = para._element.pPr
        if pPr is None:
            return False
        
        numPr = pPr.find(qn('w:numPr'))
        return numPr is not None
    
    @staticmethod
    def _detect_numbering_pattern(para) -> Optional[str]:
        """Detect if paragraph has numbering (simplified)."""
        pPr = para._element.pPr
        if pPr is None:
            return None
        
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            return "numbered"
        
        return None
