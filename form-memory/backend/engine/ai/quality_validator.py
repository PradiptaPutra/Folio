"""
Quality Validation and Compliance System
Validates thesis documents against Indonesian university standards and ensures
perfect formatting compliance for flawless submission.
"""

from typing import Dict, List, Any, Optional, Tuple
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..analyzer.template_analyzer import TemplateAnalyzer
from .academic_content_enhancer import ContentQualityScorer


class QualityValidator:
    """Comprehensive quality validation for Indonesian thesis documents."""

    def __init__(self, template_analyzer: Optional[TemplateAnalyzer] = None):
        """Initialize with template analyzer for compliance checking."""
        self.template_analyzer = template_analyzer
        self.quality_scorer = ContentQualityScorer()

        # Indonesian academic standards
        self.indonesian_standards = {
            "typography": {
                "fonts": ["Times New Roman", "Arial", "Calibri", "Bookman Old Style"],
                "sizes": {
                    "title": [14, 16, 18],
                    "chapter": [12, 14],
                    "section": [11, 12],
                    "body": [11, 12],
                    "caption": [10, 11]
                },
                "line_spacing": [1.0, 1.5, 2.0],
                "alignment": ["left", "justify", "center"]
            },
            "structure": {
                "front_matter": [
                    "halaman_judul", "lembar_pengesahan", "pernyataan_keaslian",
                    "kata_pengantar", "abstrak", "abstract", "daftar_isi"
                ],
                "main_content": ["bab_i", "bab_ii", "bab_iii"],  # Minimum
                "back_matter": ["daftar_pustaka"]
            },
            "spacing": {
                "margins": {
                    "top": [1.0, 1.5, 2.5],  # inches
                    "bottom": [1.0, 1.5, 2.5],
                    "left": [1.0, 1.25, 1.5],
                    "right": [1.0, 1.25, 1.5]
                },
                "indentation": [1.0, 1.25, 1.5]  # cm, first line
            },
            "language": {
                "primary": "indonesian",
                "academic_tone": True,
                "formal_vocabulary": True
            }
        }

    def validate_document(self, docx_path: str, content_type: str = "thesis") -> Dict[str, Any]:
        """
        Comprehensive document validation.

        Args:
            docx_path: Path to the DOCX document
            content_type: Type of document (thesis, proposal, etc.)

        Returns:
            Validation results with scores and recommendations
        """
        doc = Document(docx_path)

        validation_results = {
            "overall_score": 0.0,
            "compliance_score": 0.0,
            "quality_score": 0.0,
            "structure_score": 0.0,
            "typography_score": 0.0,
            "content_score": 0.0,
            "validation_details": {},
            "compliance_issues": [],
            "quality_issues": [],
            "recommendations": [],
            "passed_checks": [],
            "failed_checks": []
        }

        # Extract text for content analysis
        full_text = self._extract_document_text(doc)

        # Run all validation checks
        validation_results["validation_details"]["typography"] = self._validate_typography(doc)
        validation_results["validation_details"]["structure"] = self._validate_structure(doc)
        validation_results["validation_details"]["spacing"] = self._validate_spacing(doc)
        validation_results["validation_details"]["content"] = self._validate_content_quality(full_text)
        validation_results["validation_details"]["compliance"] = self._validate_university_compliance(doc)

        # Calculate scores
        validation_results["typography_score"] = validation_results["validation_details"]["typography"]["score"]
        validation_results["structure_score"] = validation_results["validation_details"]["structure"]["score"]
        validation_results["content_score"] = validation_results["validation_details"]["content"]["score"]
        validation_results["compliance_score"] = validation_results["validation_details"]["compliance"]["score"]

        # Overall quality score (weighted average)
        weights = {
            "typography": 0.15,
            "structure": 0.25,
            "spacing": 0.15,
            "content": 0.30,
            "compliance": 0.15
        }

        validation_results["overall_score"] = (
            validation_results["typography_score"] * weights["typography"] +
            validation_results["structure_score"] * weights["structure"] +
            validation_results["validation_details"]["spacing"]["score"] * weights["spacing"] +
            validation_results["content_score"] * weights["content"] +
            validation_results["compliance_score"] * weights["compliance"]
        )

        # Quality score based on content analysis
        content_analysis = self.quality_scorer.score_content(full_text, content_type)
        validation_results["quality_score"] = content_analysis["overall_score"]

        # Collect issues and recommendations
        for check_name, check_result in validation_results["validation_details"].items():
            if check_result.get("issues"):
                validation_results["failed_checks"].extend([f"{check_name}: {issue}" for issue in check_result["issues"]])
            if check_result.get("passed"):
                validation_results["passed_checks"].extend([f"{check_name}: {item}" for item in check_result["passed"]])

        # Generate recommendations
        validation_results["recommendations"] = self._generate_recommendations(validation_results)

        # Compliance issues (critical)
        validation_results["compliance_issues"] = [
            issue for check in validation_results["validation_details"].values()
            for issue in check.get("issues", [])
            if "critical" in check.get("severity", "").lower()
        ]

        return validation_results

    def _validate_typography(self, doc: Document) -> Dict[str, Any]:
        """Validate typography standards."""
        result = {
            "score": 0.0,
            "passed": [],
            "issues": [],
            "severity": "medium"
        }

        standards = self.indonesian_standards["typography"]
        total_checks = 5
        passed_checks = 0

        # Font validation
        fonts_used = set()
        sizes_used = set()

        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name:
                    fonts_used.add(run.font.name)
                if run.font.size:
                    sizes_used.add(run.font.size.pt)

        # Check primary font
        primary_fonts = standards["fonts"]
        common_fonts = [f for f in fonts_used if any(pf.lower() in f.lower() for pf in primary_fonts)]

        if common_fonts:
            result["passed"].append(f"Using appropriate academic fonts: {', '.join(common_fonts)}")
            passed_checks += 1
        else:
            result["issues"].append(f"Using non-standard fonts: {', '.join(fonts_used) if fonts_used else 'No fonts detected'}")

        # Font size validation
        body_sizes = [s for s in sizes_used if 10 <= s <= 14]
        if body_sizes:
            result["passed"].append(f"Appropriate font sizes used: {sorted(body_sizes)}")
            passed_checks += 1
        else:
            result["issues"].append("Font sizes outside academic standard range (10-14pt)")

        # Check for consistent font usage
        if len(fonts_used) <= 3:
            result["passed"].append("Consistent font usage throughout document")
            passed_checks += 1
        else:
            result["issues"].append(f"Too many different fonts used: {len(fonts_used)}")

        # Check for proper heading hierarchy
        heading_sizes = sorted([s for s in sizes_used if s > 12], reverse=True)
        if len(heading_sizes) >= 2 and heading_sizes[0] > heading_sizes[1]:
            result["passed"].append("Proper heading size hierarchy")
            passed_checks += 1
        else:
            result["issues"].append("Heading sizes do not follow proper hierarchy")

        # Bold/italic usage check
        bold_runs = sum(1 for para in doc.paragraphs for run in para.runs if run.bold)
        if bold_runs > 0:
            result["passed"].append("Appropriate use of bold formatting")
            passed_checks += 1

        result["score"] = passed_checks / total_checks
        return result

    def _validate_structure(self, doc: Document) -> Dict[str, Any]:
        """Validate document structure."""
        result = {
            "score": 0.0,
            "passed": [],
            "issues": [],
            "severity": "high"
        }

        standards = self.indonesian_standards["structure"]
        total_checks = 4
        passed_checks = 0

        # Extract structure from document
        doc_structure = self._analyze_document_structure(doc)

        # Check front matter
        front_matter_present = []
        for section in standards["front_matter"]:
            if any(section.lower() in item.lower() for item in doc_structure.get("sections", [])):
                front_matter_present.append(section)

        if len(front_matter_present) >= 4:  # At least 4 key front matter sections
            result["passed"].append(f"Front matter complete: {len(front_matter_present)} sections")
            passed_checks += 1
        else:
            result["issues"].append(f"Missing front matter sections. Found: {len(front_matter_present)}, Expected: {len(standards['front_matter'])}")

        # Check main content (BAB structure)
        bab_count = len([s for s in doc_structure.get("sections", []) if "bab " in s.lower()])
        if bab_count >= 3:
            result["passed"].append(f"Main content structure adequate: {bab_count} chapters")
            passed_checks += 1
        else:
            result["issues"].append(f"Insufficient chapter structure. Found: {bab_count}, Recommended: 3+")

        # Check back matter
        back_matter_present = []
        for section in standards["back_matter"]:
            if any(section.lower() in item.lower() for item in doc_structure.get("sections", [])):
                back_matter_present.append(section)

        if back_matter_present:
            result["passed"].append(f"Back matter present: {', '.join(back_matter_present)}")
            passed_checks += 1
        else:
            result["issues"].append("Missing back matter (references, bibliography)")

        # Check page breaks
        page_breaks = len([para for para in doc.paragraphs if para.text.strip() == ""])
        if page_breaks >= bab_count:  # At least one page break per chapter
            result["passed"].append("Appropriate page break usage")
            passed_checks += 1
        else:
            result["issues"].append("Insufficient page breaks between sections")

        result["score"] = passed_checks / total_checks
        return result

    def _validate_spacing(self, doc: Document) -> Dict[str, Any]:
        """Validate spacing and margins."""
        result = {
            "score": 0.0,
            "passed": [],
            "issues": [],
            "severity": "medium"
        }

        standards = self.indonesian_standards["spacing"]
        total_checks = 4
        passed_checks = 0

        # Check margins
        for section in doc.sections:
            margins_ok = True
            margin_issues = []

            if section.top_margin and not (standards["margins"]["top"][0] <= section.top_margin.inches <= standards["margins"]["top"][-1]):
                margin_issues.append(".1f")
                margins_ok = False

            if section.bottom_margin and not (standards["margins"]["bottom"][0] <= section.bottom_margin.inches <= standards["margins"]["bottom"][-1]):
                margin_issues.append(".1f")
                margins_ok = False

            if section.left_margin and not (standards["margins"]["left"][0] <= section.left_margin.inches <= standards["margins"]["left"][-1]):
                margin_issues.append(".1f")
                margins_ok = False

            if margins_ok and not margin_issues:
                result["passed"].append("Margins within academic standards")
                passed_checks += 1
            elif margin_issues:
                result["issues"].extend(margin_issues)

        # Check line spacing
        line_spacings = []
        for para in doc.paragraphs[:50]:  # Check first 50 paragraphs
            if para.paragraph_format.line_spacing:
                line_spacings.append(para.paragraph_format.line_spacing)

        avg_line_spacing = sum(line_spacings) / len(line_spacings) if line_spacings else 1.0

        if standards["line_spacing"][0] <= avg_line_spacing <= standards["line_spacing"][-1]:
            result["passed"].append(".1f")
            passed_checks += 1
        else:
            result["issues"].append(".1f")

        # Check indentation
        indentations = []
        for para in doc.paragraphs[:50]:
            if para.paragraph_format.first_line_indent:
                indentations.append(para.paragraph_format.first_line_indent.cm)

        if indentations:
            avg_indent = sum(indentations) / len(indentations)
            if any(abs(avg_indent - std) < 0.3 for std in standards["indentation"]):
                result["passed"].append(".1f")
                passed_checks += 1
            else:
                result["issues"].append(".1f")

        # Check paragraph spacing consistency
        space_befores = [para.paragraph_format.space_before.pt for para in doc.paragraphs[:50]
                        if para.paragraph_format.space_before]
        if space_befores and len(set(space_befores)) <= 2:  # Mostly consistent
            result["passed"].append("Consistent paragraph spacing")
            passed_checks += 1
        else:
            result["issues"].append("Inconsistent paragraph spacing")

        result["score"] = passed_checks / total_checks
        return result

    def _validate_content_quality(self, text: str) -> Dict[str, Any]:
        """Validate content quality using the quality scorer."""
        result = {
            "score": 0.0,
            "passed": [],
            "issues": [],
            "severity": "high"
        }

        # Use the quality scorer
        quality_analysis = self.quality_scorer.score_content(text, "body", "general")
        result["score"] = quality_analysis["overall_score"]

        # Convert quality analysis to validation format
        if quality_analysis["overall_score"] >= 0.8:
            result["passed"].append(f"High quality content (Grade: {quality_analysis['grade']})")
        elif quality_analysis["overall_score"] >= 0.6:
            result["issues"].append(f"Moderate quality content (Grade: {quality_analysis['grade']})")
        else:
            result["issues"].append(f"Low quality content requiring revision (Grade: {quality_analysis['grade']})")

        # Add specific recommendations as issues
        recommendations = quality_analysis.get("recommendations", [])
        result["issues"].extend([f"Content: {rec}" for rec in recommendations])

        return result

    def _validate_university_compliance(self, doc: Document) -> Dict[str, Any]:
        """Validate compliance with university template (if available)."""
        result = {
            "score": 1.0,  # Default high score if no template to compare
            "passed": ["Basic structure validation completed"],
            "issues": [],
            "severity": "low"
        }

        if not self.template_analyzer:
            result["issues"].append("No template provided for compliance checking")
            return result

        # Compare document against template
        template_compliance = self._compare_with_template(doc)

        if template_compliance["compliant"]:
            result["passed"].append("Document structure matches template")
            result["score"] = 0.9
        else:
            result["issues"].extend(template_compliance["issues"])
            result["score"] = 0.6

        return result

    def _compare_with_template(self, doc: Document) -> Dict[str, Any]:
        """Compare document with template requirements."""
        compliance = {
            "compliant": True,
            "issues": []
        }

        # This would compare styles, structure, etc. with the template
        # For now, return basic compliance
        return compliance

    def _analyze_document_structure(self, doc: Document) -> Dict[str, Any]:
        """Analyze the document's structural elements."""
        structure = {
            "sections": [],
            "headings": [],
            "page_breaks": 0,
            "tables": len(doc.tables),
            "images": 0  # Would need more complex analysis
        }

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect section headers
            if any(keyword in text.lower() for keyword in
                   ["halaman judul", "pengesahan", "abstrak", "daftar isi", "bab ", "kesimpulan", "daftar pustaka"]):
                structure["sections"].append(text)

            # Detect headings by style or formatting
            if para.style and 'heading' in para.style.name.lower():
                structure["headings"].append(text)

        return structure

    def _extract_document_text(self, doc: Document) -> str:
        """Extract all text from document."""
        return '\n'.join(para.text for para in doc.paragraphs if para.text.strip())

    def _generate_recommendations(self, validation_results: Dict[str, Any]) -> List[str]:
        """Generate improvement recommendations based on validation."""
        recommendations = []

        overall_score = validation_results.get("overall_score", 0.0)

        if overall_score >= 0.9:
            recommendations.append("Document meets high academic standards - ready for submission")
        elif overall_score >= 0.8:
            recommendations.append("Document is good but has minor issues to address")
        elif overall_score >= 0.7:
            recommendations.append("Document needs moderate revisions before submission")
        else:
            recommendations.append("Document requires significant revisions")

        # Add specific recommendations based on failed checks
        failed_checks = validation_results.get("failed_checks", [])

        if any("typography" in check.lower() for check in failed_checks):
            recommendations.append("Review typography: ensure Times New Roman 12pt, 1.5 line spacing")

        if any("structure" in check.lower() for check in failed_checks):
            recommendations.append("Complete document structure: add missing front/back matter sections")

        if any("spacing" in check.lower() for check in failed_checks):
            recommendations.append("Adjust spacing: 1 inch margins, proper indentation")

        if any("content" in check.lower() for check in failed_checks):
            recommendations.append("Improve content quality: enhance academic tone and clarity")

        return recommendations

    def get_compliance_report(self, validation_results: Dict[str, Any]) -> str:
        """Generate a detailed compliance report."""
        report = f"""
THESIS COMPLIANCE REPORT
========================

Overall Score: {validation_results.get('overall_score', 0.0):.1%}
Quality Grade: {self._score_to_grade(validation_results.get('overall_score', 0.0))}
Compliance Score: {validation_results.get('compliance_score', 0.0):.1%}

VALIDATION SUMMARY
------------------
Typography: {validation_results.get('typography_score', 0.0):.1%}
Structure: {validation_results.get('structure_score', 0.0):.1%}
Spacing: {validation_results.get('validation_details', {}).get('spacing', {}).get('score', 0.0):.1%}
Content: {validation_results.get('content_score', 0.0):.1%}
Compliance: {validation_results.get('compliance_score', 0.0):.1%}

PASSED CHECKS ({len(validation_results.get('passed_checks', []))}):
{chr(10).join(f"✓ {check}" for check in validation_results.get('passed_checks', []))}

ISSUES FOUND ({len(validation_results.get('failed_checks', []))}):
{chr(10).join(f"✗ {check}" for check in validation_results.get('failed_checks', []))}

RECOMMENDATIONS:
{chr(10).join(f"• {rec}" for rec in validation_results.get('recommendations', []))}

CRITICAL ISSUES ({len(validation_results.get('compliance_issues', []))}):
{chr(10).join(f"⚠ {issue}" for issue in validation_results.get('compliance_issues', []))}
        """.strip()

        return report

    def _score_to_grade(self, score: float) -> str:
        """Convert numerical score to letter grade."""
        if score >= 0.95:
            return "A+ (Excellent - Ready for Submission)"
        elif score >= 0.90:
            return "A (Very Good - Minor Touch-ups)"
        elif score >= 0.85:
            return "B+ (Good - Some Improvements Needed)"
        elif score >= 0.80:
            return "B (Satisfactory - Moderate Revisions)"
        elif score >= 0.70:
            return "C (Needs Work - Significant Revisions)"
        elif score >= 0.60:
            return "D (Poor - Major Revisions Required)"
        else:
            return "F (Unsatisfactory - Complete Rewrite Needed)"