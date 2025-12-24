"""
Perfect Template Adapter
Intelligently adapts enhanced content to match university templates exactly,
ensuring pixel-perfect formatting and structure compliance.
"""

from typing import Dict, List, Any, Optional, Tuple
import re
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from ..analyzer.template_analyzer import TemplateAnalyzer
from .academic_content_enhancer import AcademicContentEnhancer


class PerfectTemplateAdapter:
    """Adapts content to match templates with perfect formatting accuracy."""

    def __init__(self, template_analyzer: TemplateAnalyzer,
                 content_enhancer: Optional[AcademicContentEnhancer] = None):
        """Initialize with template analysis and content enhancer."""
        self.template = template_analyzer
        self.enhancer = content_enhancer
        self.analysis = template_analyzer.analysis

        # Template adaptation rules
        self.adaptation_rules = self._generate_adaptation_rules()

    def adapt_content_to_template(self, raw_content: str, user_data: Optional[Dict[str, Any]] = None) -> Document:
        """
        Adapt content to match template perfectly.

        Args:
            raw_content: Raw thesis content
            user_data: User-specific data (title, author, etc.)

        Returns:
            Perfectly formatted DOCX document matching the template
        """
        user_data = user_data or {}

        # Step 1: Enhance content if enhancer available
        if self.enhancer:
            enhanced_result = self.enhancer.enhance_content(raw_content, "body", "general")
            processed_content = enhanced_result["enhanced_content"]
        else:
            processed_content = raw_content

        # Step 2: Create new document with template properties
        adapted_doc = self._create_template_base_document()

        # Step 3: Apply perfect formatting adaptation
        self._apply_perfect_formatting(adapted_doc, processed_content, user_data)

        return adapted_doc

    def _generate_adaptation_rules(self) -> Dict[str, Any]:
        """Generate rules for perfect template adaptation."""
        rules = {
            "font_mapping": self._create_font_mapping_rules(),
            "style_mapping": self._create_style_mapping_rules(),
            "spacing_rules": self._create_spacing_rules(),
            "structure_rules": self._create_structure_rules(),
            "formatting_rules": self._create_formatting_rules()
        }
        return rules

    def _create_font_mapping_rules(self) -> Dict[str, Any]:
        """Create font mapping rules for perfect typography matching."""
        template_fonts = self.analysis.get('formatting_rules', {}).get('common_font', 'Times New Roman')

        # Indonesian academic standard fonts
        font_rules = {
            "primary_font": template_fonts,
            "fallback_fonts": ["Times New Roman", "Arial", "Calibri"],
            "size_mapping": {
                "title": 16,  # Title page
                "chapter": 14,  # Chapter headings
                "section": 12,  # Section headings
                "body": 11,  # Main text
                "caption": 10  # Captions
            },
            "font_variants": {
                "bold": True,  # For headings
                "italic": False,  # Generally avoided in academic text
                "underline": False  # Generally avoided
            }
        }

        return font_rules

    def _create_style_mapping_rules(self) -> Dict[str, Any]:
        """Create style mapping rules for perfect style application."""
        template_styles = self.analysis.get('styles', {})

        style_rules = {
            "heading_styles": {},
            "paragraph_styles": {},
            "list_styles": {},
            "custom_styles": {}
        }

        # Map heading styles
        for style_name, style_info in template_styles.items():
            if 'heading' in style_name.lower():
                level = style_info.get('outline_level', 1)
                style_rules["heading_styles"][f"heading_{level}"] = {
                    "template_style": style_name,
                    "font_size": style_info.get('font', {}).get('size', 12),
                    "bold": style_info.get('font', {}).get('bold', True),
                    "alignment": style_info.get('paragraph', {}).get('alignment', 'left')
                }

        # Map paragraph styles
        for style_name, style_info in template_styles.items():
            if style_info.get('type') == 'paragraph':
                style_rules["paragraph_styles"]["body"] = {
                    "template_style": style_name,
                    "font_size": style_info.get('font', {}).get('size', 11),
                    "line_spacing": style_info.get('paragraph', {}).get('line_spacing', 1.5),
                    "indentation": style_info.get('paragraph', {}).get('first_line_indent', 1.0)
                }

        return style_rules

    def _create_spacing_rules(self) -> Dict[str, Any]:
        """Create spacing rules for perfect layout matching."""
        template_margins = self.analysis.get('margins', {})
        template_spacing = self.analysis.get('formatting_rules', {})

        spacing_rules = {
            "margins": {
                "top": template_margins.get('top', 1.0),
                "bottom": template_margins.get('bottom', 1.0),
                "left": template_margins.get('left', 1.0),
                "right": template_margins.get('right', 1.0)
            },
            "paragraph_spacing": {
                "line_spacing": template_spacing.get('line_spacing_pattern', 1.5),
                "space_before": 0,  # Academic standard
                "space_after": 0,   # Academic standard
                "first_line_indent": template_spacing.get('indentation_pattern', {}).get('common_first_line', 1.0)
            },
            "page_spacing": {
                "header_margin": 0.5,
                "footer_margin": 0.5,
                "gutter_margin": template_margins.get('gutter_margin', 0)
            }
        }

        return spacing_rules

    def _create_structure_rules(self) -> Dict[str, Any]:
        """Create structure rules for perfect document organization."""
        structure_rules = {
            "front_matter_order": [
                "title_page",
                "approval_page",
                "originality_statement",
                "dedication",
                "motto",
                "preface",
                "abstract_id",
                "abstract_en",
                "glossary",
                "table_of_contents",
                "list_of_tables",
                "list_of_figures"
            ],
            "main_content_structure": {
                "chapter_pattern": r"^BAB\s+[IVX]+\s*[:-]?\s*(.+)",
                "section_pattern": r"^([0-9]+\.[0-9]+)\s+(.+)",
                "subsection_pattern": r"^([0-9]+\.[0-9]+\.[0-9]+)\s+(.+)"
            },
            "back_matter_order": [
                "references",
                "appendices",
                "index"
            ],
            "page_break_rules": {
                "before_chapters": True,
                "after_front_matter": True,
                "before_back_matter": True
            }
        }

        return structure_rules

    def _create_formatting_rules(self) -> Dict[str, Any]:
        """Create detailed formatting rules."""
        formatting_rules = {
            "alignment_rules": {
                "title_page": "center",
                "headings": "left",
                "body_text": "justify",
                "captions": "center"
            },
            "numbering_rules": {
                "chapters": "roman_upper",  # BAB I, II, III
                "sections": "arabic",       # 1.1, 1.2, 2.1
                "subsections": "arabic",    # 1.1.1, 1.1.2
                "lists": "arabic",          # 1., 2., 3.
                "figures": "arabic",        # Gambar 1.1, Gambar 1.2
                "tables": "arabic"          # Tabel 1.1, Tabel 1.2
            },
            "language_rules": {
                "primary_language": "id",   # Indonesian
                "secondary_language": "en", # English abstracts
                "formal_tone": True,
                "academic_vocabulary": True
            },
            "citation_rules": {
                "style": "APA",  # Common in Indonesian academia
                "in_text_format": "(Author, Year)",
                "reference_format": "Author. (Year). Title. Publisher."
            }
        }

        return formatting_rules

    def _create_template_base_document(self) -> Document:
        """Create a new document with template properties."""
        doc = Document()

        # Apply template margins
        spacing_rules = self.adaptation_rules["spacing_rules"]
        margins = spacing_rules["margins"]

        for section in doc.sections:
            section.top_margin = Inches(margins["top"])
            section.bottom_margin = Inches(margins["bottom"])
            section.left_margin = Inches(margins["left"])
            section.right_margin = Inches(margins["right"])

        # Copy template styles
        self._copy_template_styles(doc)

        return doc

    def _copy_template_styles(self, target_doc: Document) -> None:
        """Copy styles from template to target document."""
        template_doc = Document(str(self.template.template_path))

        for style in template_doc.styles:
            if style.name not in target_doc.styles and hasattr(style, 'type'):
                try:
                    # Create corresponding style in target document
                    if style.type == WD_STYLE_TYPE.PARAGRAPH:
                        new_style = target_doc.styles.add_style(style.name, WD_STYLE_TYPE.PARAGRAPH)

                        # Copy font properties
                        if hasattr(style, 'font') and style.font:
                            if style.font.name:
                                new_style.font.name = style.font.name
                            if style.font.size:
                                new_style.font.size = style.font.size
                            if style.font.bold is not None:
                                new_style.font.bold = style.font.bold

                        # Copy paragraph properties
                        if hasattr(style, 'paragraph_format') and style.paragraph_format:
                            if style.paragraph_format.alignment:
                                new_style.paragraph_format.alignment = style.paragraph_format.alignment
                            if style.paragraph_format.left_indent:
                                new_style.paragraph_format.left_indent = style.paragraph_format.left_indent
                            if style.paragraph_format.first_line_indent:
                                new_style.paragraph_format.first_line_indent = style.paragraph_format.first_line_indent
                            if style.paragraph_format.line_spacing:
                                new_style.paragraph_format.line_spacing = style.paragraph_format.line_spacing

                except Exception:
                    # Skip problematic styles
                    continue

    def _apply_perfect_formatting(self, doc: Document, content: str, user_data: Dict[str, Any]) -> None:
        """Apply perfect formatting to match template exactly."""
        # Step 1: Add front matter
        self._add_front_matter_perfect(doc, user_data)

        # Step 2: Process and add main content
        self._add_main_content_perfect(doc, content, user_data)

        # Step 3: Add back matter
        self._add_back_matter_perfect(doc, user_data)

    def _add_front_matter_perfect(self, doc: Document, user_data: Dict[str, Any]) -> None:
        """Add front matter with perfect template matching."""
        structure_rules = self.adaptation_rules["structure_rules"]

        for section_type in structure_rules["front_matter_order"]:
            if section_type == "title_page":
                self._add_title_page_perfect(doc, user_data)
            elif section_type == "approval_page":
                self._add_approval_page_perfect(doc, user_data)
            elif section_type == "preface":
                self._add_preface_perfect(doc, user_data)
            elif section_type == "abstract_id":
                self._add_abstract_perfect(doc, user_data, "id")
            elif section_type == "abstract_en":
                self._add_abstract_perfect(doc, user_data, "en")
            elif section_type == "table_of_contents":
                self._add_toc_placeholder(doc)

            # Add page breaks between sections
            if section_type != structure_rules["front_matter_order"][-1]:
                doc.add_page_break()

    def _add_title_page_perfect(self, doc: Document, user_data: Dict[str, Any]) -> None:
        """Add title page with perfect formatting."""
        font_rules = self.adaptation_rules["font_mapping"]
        formatting_rules = self.adaptation_rules["formatting_rules"]

        # Title
        title = user_data.get("title", "JUDUL SKRIPSI")
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.size = Pt(font_rules["size_mapping"]["title"])
        title_run.font.bold = True
        title_run.font.name = font_rules["primary_font"]

        # Add spacing
        for _ in range(4):
            doc.add_paragraph()

        # Author info
        author = user_data.get("author", "Nama Penulis")
        nim = user_data.get("nim", "NIM")
        author_text = f"Oleh:\n{author}\nNIM: {nim}"
        author_para = doc.add_paragraph(author_text)
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in author_para.runs:
            run.font.size = Pt(font_rules["size_mapping"]["body"])
            run.font.name = font_rules["primary_font"]

        # Add spacing
        for _ in range(6):
            doc.add_paragraph()

        # Institution
        institution = user_data.get("institution", "Universitas")
        date = user_data.get("date", "2025")
        institution_para = doc.add_paragraph(f"{institution}\n{date}")
        institution_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in institution_para.runs:
            run.font.size = Pt(font_rules["size_mapping"]["body"])
            run.font.name = font_rules["primary_font"]

    def _add_approval_page_perfect(self, doc: Document, user_data: Dict[str, Any]) -> None:
        """Add approval page with perfect formatting."""
        font_rules = self.adaptation_rules["font_mapping"]

        # Title
        title_para = doc.add_paragraph("HALAMAN PENGESAHAN")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.bold = True
        title_run.font.size = Pt(font_rules["size_mapping"]["chapter"])
        title_run.font.name = font_rules["primary_font"]

        doc.add_paragraph()

        # Content
        title = user_data.get("title", "Judul Skripsi")
        author = user_data.get("author", "Nama Penulis")
        nim = user_data.get("nim", "NIM")
        advisor = user_data.get("advisor", "Nama Dosen Pembimbing")

        approval_text = f"""Judul: {title}
Penulis: {author}
NIM: {nim}

Dosen Pembimbing:

_________________________
{advisor}"""

        approval_para = doc.add_paragraph(approval_text)
        approval_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in approval_para.runs:
            run.font.size = Pt(font_rules["size_mapping"]["body"])
            run.font.name = font_rules["primary_font"]

    def _add_preface_perfect(self, doc: Document, user_data: Dict[str, Any]) -> None:
        """Add preface with perfect formatting."""
        font_rules = self.adaptation_rules["font_mapping"]
        spacing_rules = self.adaptation_rules["spacing_rules"]

        # Title
        title_para = doc.add_paragraph("KATA PENGANTAR")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.bold = True
        title_run.font.size = Pt(font_rules["size_mapping"]["chapter"])
        title_run.font.name = font_rules["primary_font"]

        doc.add_paragraph()

        # Preface content
        preface_text = user_data.get("preface")
        if not preface_text and self.enhancer:
            # Generate AI-enhanced preface
            preface_prompt = f"""
            Write a proper Indonesian academic preface for a thesis.
            Title: {user_data.get('title', 'Thesis Title')}
            Author: {user_data.get('author', 'Author Name')}
            Advisor: {user_data.get('advisor', 'Advisor Name')}
            Institution: {user_data.get('institution', 'University')}
            Research focus: {user_data.get('thesis_focus', 'Research topic')}
            """
            enhanced_preface = self.enhancer.enhance_content(preface_prompt, "preface", "general")
            preface_text = enhanced_preface["enhanced_content"]

        if not preface_text:
            preface_text = f"""Puji syukur kami panjatkan kepada Tuhan Yang Maha Esa atas segala rahmat dan hidayahnya sehingga penulis dapat menyelesaikan skripsi ini dengan baik.

Skripsi dengan judul "{user_data.get('title', 'Judul Skripsi')}" ini dibuat sebagai salah satu persyaratan untuk memperoleh gelar sarjana pada Program Studi [Program Studi] di {user_data.get('institution', 'Universitas')}.

Penulis mengucapkan terima kasih kepada semua pihak yang telah membantu dalam penyelesaian skripsi ini, khususnya kepada:

1. {user_data.get('advisor', 'Dosen Pembimbing')} selaku dosen pembimbing
2. Orang tua dan keluarga yang memberikan dukungan moral dan material
3. Teman-teman yang telah memberikan inspirasi dan semangat

Penulis menyadari bahwa skripsi ini masih memiliki banyak kekurangan. Oleh karena itu, penulis dengan hati yang tulus menerima segala kritik dan saran untuk penyempurnaan skripsi ini.

Semoga skripsi ini dapat memberikan manfaat bagi pengembangan ilmu pengetahuan.

Jakarta, {user_data.get('date', '2025')}

_________________________
{user_data.get('author', 'Nama Penulis')}"""

        preface_para = doc.add_paragraph(preface_text)
        preface_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Apply perfect paragraph formatting
        para_format = preface_para.paragraph_format
        para_format.left_indent = Inches(spacing_rules["paragraph_spacing"]["first_line_indent"])
        para_format.line_spacing = spacing_rules["paragraph_spacing"]["line_spacing"]
        para_format.space_before = Pt(0)
        para_format.space_after = Pt(0)

        for run in preface_para.runs:
            run.font.size = Pt(font_rules["size_mapping"]["body"])
            run.font.name = font_rules["primary_font"]

    def _add_abstract_perfect(self, doc: Document, user_data: Dict[str, Any], language: str) -> None:
        """Add abstract with perfect formatting."""
        font_rules = self.adaptation_rules["font_mapping"]
        spacing_rules = self.adaptation_rules["spacing_rules"]

        # Title
        title_text = "SARI" if language == "id" else "ABSTRACT"
        title_para = doc.add_paragraph(title_text)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.bold = True
        title_run.font.size = Pt(font_rules["size_mapping"]["chapter"])
        title_run.font.name = font_rules["primary_font"]

        doc.add_paragraph()

        # Abstract content
        abstract_key = "abstract_id" if language == "id" else "abstract_en"
        abstract_text = user_data.get(abstract_key)

        if not abstract_text and self.enhancer:
            # Generate AI-enhanced abstract
            abstract_prompt = f"""
            Write a proper {'Indonesian' if language == 'id' else 'English'} academic abstract for a thesis.
            Title: {user_data.get('title', 'Thesis Title')}
            Objectives: {user_data.get('objectives', 'Research objectives')}
            Methods: {user_data.get('methods', 'Research methods')}
            Results: {user_data.get('results', 'Research results')}
            Focus: {user_data.get('thesis_focus', 'Research focus')}
            """
            enhanced_abstract = self.enhancer.enhance_content(abstract_prompt, "abstract", "general")
            abstract_text = enhanced_abstract["enhanced_content"]

        if not abstract_text:
            if language == "id":
                abstract_text = "[Tuliskan abstrak dalam bahasa Indonesia di sini. Maksimal 250 kata yang merangkum latar belakang, metodologi, hasil, dan kesimpulan penelitian.]"
            else:
                abstract_text = "[Write the abstract in English here. Maximum 250 words summarizing the background, methodology, results, and conclusions of the research.]"

        abstract_para = doc.add_paragraph(abstract_text)
        abstract_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Apply perfect paragraph formatting
        para_format = abstract_para.paragraph_format
        para_format.left_indent = Inches(spacing_rules["paragraph_spacing"]["first_line_indent"])
        para_format.line_spacing = spacing_rules["paragraph_spacing"]["line_spacing"]

        for run in abstract_para.runs:
            run.font.size = Pt(font_rules["size_mapping"]["body"])
            run.font.name = font_rules["primary_font"]

        # Keywords
        doc.add_paragraph()
        keywords = user_data.get("keywords", [])
        if isinstance(keywords, list):
            keywords_text = ", ".join(keywords)
        else:
            keywords_text = str(keywords)

        if not keywords_text or keywords_text == "[]":
            keywords_text = "[tuliskan kata kunci dipisahkan dengan koma]" if language == "id" else "[write keywords separated by commas]"

        keywords_para = doc.add_paragraph(f"{'Kata Kunci' if language == 'id' else 'Keywords'}: {keywords_text}")
        for run in keywords_para.runs:
            run.font.size = Pt(font_rules["size_mapping"]["body"])
            run.font.name = font_rules["primary_font"]

    def _add_toc_placeholder(self, doc: Document) -> None:
        """Add table of contents placeholder."""
        font_rules = self.adaptation_rules["font_mapping"]

        # Title
        title_para = doc.add_paragraph("DAFTAR ISI")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.bold = True
        title_run.font.size = Pt(font_rules["size_mapping"]["chapter"])
        title_run.font.name = font_rules["primary_font"]

        doc.add_paragraph()

        # Placeholder text
        placeholder_text = "[Daftar isi akan diperbarui secara otomatis. Tekan Ctrl+A kemudian F9 di Microsoft Word untuk memperbarui.]"
        placeholder_para = doc.add_paragraph(placeholder_text)
        placeholder_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in placeholder_para.runs:
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.name = font_rules["primary_font"]

    def _add_main_content_perfect(self, doc: Document, content: str, user_data: Dict[str, Any]) -> None:
        """Add main content with perfect formatting."""
        structure_rules = self.adaptation_rules["structure_rules"]
        font_rules = self.adaptation_rules["font_mapping"]
        spacing_rules = self.adaptation_rules["spacing_rules"]

        # Parse content into chapters and sections
        chapters = self._parse_content_structure(content)

        for i, chapter in enumerate(chapters):
            # Page break before each chapter (except first)
            if i > 0 and structure_rules["page_break_rules"]["before_chapters"]:
                doc.add_page_break()

            # Chapter title
            if "title" in chapter:
                chapter_title = chapter["title"]
                chapter_para = doc.add_paragraph(chapter_title)
                chapter_para.style = "Heading 1"
                chapter_run = chapter_para.runs[0]
                chapter_run.font.size = Pt(font_rules["size_mapping"]["chapter"])
                chapter_run.font.bold = True
                chapter_run.font.name = font_rules["primary_font"]

            # Chapter content
            if "content" in chapter:
                for paragraph_text in chapter["content"]:
                    if paragraph_text.strip():
                        para = doc.add_paragraph(paragraph_text)

                        # Apply perfect paragraph formatting
                        para_format = para.paragraph_format
                        para_format.left_indent = Inches(spacing_rules["paragraph_spacing"]["first_line_indent"])
                        para_format.line_spacing = spacing_rules["paragraph_spacing"]["line_spacing"]
                        para_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                        for run in para.runs:
                            run.font.size = Pt(font_rules["size_mapping"]["body"])
                            run.font.name = font_rules["primary_font"]

    def _add_back_matter_perfect(self, doc: Document, user_data: Dict[str, Any]) -> None:
        """Add back matter with perfect formatting."""
        structure_rules = self.adaptation_rules["structure_rules"]
        font_rules = self.adaptation_rules["font_mapping"]

        # References/Bibliography
        references = user_data.get("references", [])
        if references:
            doc.add_page_break()

            title_para = doc.add_paragraph("DAFTAR PUSTAKA")
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.runs[0]
            title_run.font.bold = True
            title_run.font.size = Pt(font_rules["size_mapping"]["chapter"])
            title_run.font.name = font_rules["primary_font"]

            doc.add_paragraph()

            for ref in references:
                if isinstance(ref, dict):
                    ref_text = f"{ref.get('author', '')}. ({ref.get('year', '')}). {ref.get('title', '')}. {ref.get('publisher', '')}."
                else:
                    ref_text = str(ref)

                ref_para = doc.add_paragraph(ref_text)
                for run in ref_para.runs:
                    run.font.size = Pt(font_rules["size_mapping"]["body"])
                    run.font.name = font_rules["primary_font"]

        # Appendices
        appendices = user_data.get("appendices", [])
        if appendices:
            doc.add_page_break()

            title_para = doc.add_paragraph("LAMPIRAN")
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.runs[0]
            title_run.font.bold = True
            title_run.font.size = Pt(font_rules["size_mapping"]["chapter"])
            title_run.font.name = font_rules["primary_font"]

            doc.add_paragraph()

            for appendix in appendices:
                if isinstance(appendix, dict):
                    app_title = appendix.get("title", "")
                    app_content = appendix.get("content", "")

                    if app_title:
                        app_title_para = doc.add_paragraph(app_title)
                        app_title_para.style = "Heading 2"
                        title_run = app_title_para.runs[0]
                        title_run.font.size = Pt(font_rules["size_mapping"]["section"])
                        title_run.font.bold = True

                    if app_content:
                        app_para = doc.add_paragraph(app_content)
                        for run in app_para.runs:
                            run.font.size = Pt(font_rules["size_mapping"]["body"])
                            run.font.name = font_rules["primary_font"]

    def _parse_content_structure(self, content: str) -> List[Dict[str, Any]]:
        """Parse content into structured chapters and sections."""
        structure_rules = self.adaptation_rules["structure_rules"]
        lines = content.split('\n')
        chapters = []
        current_chapter = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check for chapter pattern
            chapter_match = re.match(structure_rules["main_content_structure"]["chapter_pattern"], line, re.IGNORECASE)
            if chapter_match:
                # Save previous chapter
                if current_chapter:
                    chapters.append(current_chapter)

                # Start new chapter
                chapter_title = chapter_match.group(0).strip()
                current_chapter = {
                    "title": chapter_title,
                    "content": []
                }
                continue

            # Check for section patterns
            section_match = re.match(structure_rules["main_content_structure"]["section_pattern"], line)
            subsection_match = re.match(structure_rules["main_content_structure"]["subsection_pattern"], line)

            if section_match or subsection_match:
                # This is a section heading
                if current_chapter:
                    current_chapter["content"].append(line)
            else:
                # Regular content
                if current_chapter:
                    current_chapter["content"].append(line)

        # Don't forget the last chapter
        if current_chapter:
            chapters.append(current_chapter)

        return chapters if chapters else [{"title": "BAB I PENDAHULUAN", "content": [content]}]