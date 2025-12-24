"""
Enhanced Thesis Builder with AI-Powered Templating
Uses docxtpl for advanced template population with AI content enhancement.
"""

from typing import Dict, List, Any, Optional
from pathlib import Path
import re
from datetime import datetime

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Inches
from docxtpl import DocxTemplate  # type: ignore

from .template_analyzer import TemplateAnalyzer
from .content_extractor import ContentExtractor
from .ai_enhanced_extractor import AIEnhancedContentExtractor
from .content_mapper import ContentMapper
from .front_matter_generator import FrontMatterGenerator


class EnhancedThesisBuilder:
    """Enhanced thesis builder with AI-powered templating using docxtpl."""

    def __init__(self, template_path: str, content_path: str, output_path: str, use_ai: bool = True):
        """Initialize the enhanced thesis builder.

        Args:
            template_path: Path to DOCX template (can contain placeholders)
            content_path: Path to content file
            output_path: Path for output DOCX
            use_ai: Whether to use AI semantic analysis
        """
        self.template_path = Path(template_path)
        self.content_path = Path(content_path)
        self.output_path = Path(output_path)
        self.use_ai = use_ai

        # Initialize analyzers
        self.analyzer = TemplateAnalyzer(str(self.template_path))
        self.extractor = ContentExtractor(str(self.content_path))
        self.ai_extractor = AIEnhancedContentExtractor(str(self.content_path), use_ai=use_ai) if use_ai else None

        print(f"[EnhancedThesisBuilder] Initialized with template: {template_path}")
        print(f"[EnhancedThesisBuilder] AI enabled: {use_ai}")

    def build(self, user_data: Optional[Dict[str, Any]] = None) -> Path:
        """Build the enhanced thesis document.

        Args:
            user_data: User metadata and content

        Returns:
            Path to created DOCX file
        """
        user_data = user_data or {}
        print(f"[EnhancedThesisBuilder] Building document for user: {user_data.get('penulis', 'Unknown')}")

        try:
            # Check if template has placeholders
            if self._has_placeholders():
                print("[EnhancedThesisBuilder] Template has placeholders, using docxtpl")
                return self._build_with_docxtpl(user_data)
            else:
                print("[EnhancedThesisBuilder] Template has no placeholders, using fallback method")
                return self._build_fallback(user_data)

        except Exception as e:
            print(f"[EnhancedThesisBuilder] Error during build: {e}")
            # Fallback to basic method
            return self._build_fallback(user_data)

    def _has_placeholders(self) -> bool:
        """Check if template contains docxtpl placeholders."""
        try:
            # Try to load with docxtpl
            doc = DocxTemplate(str(self.template_path))
            # Check for common placeholder patterns
            content = str(doc.docx)
            placeholder_patterns = [
                r'\{\{.*?\}\}',  # {{variable}}
                r'\{\%.*?\%\}',  # {% for %} etc.
            ]

            for pattern in placeholder_patterns:
                if re.search(pattern, content):
                    return True

            return False

        except Exception:
            return False

    def _build_with_docxtpl(self, user_data: Dict[str, Any]) -> Path:
        """Build document using docxtpl templating."""
        print("[EnhancedThesisBuilder] Building with docxtpl")

        # Load template
        doc = DocxTemplate(str(self.template_path))  # type: ignore

        # Prepare context with AI-enhanced content
        context = self._prepare_ai_context(user_data)

        # Render template
        doc.render(context)

        # Save document
        doc.save(str(self.output_path))

        print(f"[EnhancedThesisBuilder] Document saved to: {self.output_path}")
        return self.output_path

    def _prepare_ai_context(self, user_data: Optional[Dict[str, Any]]) -> Dict[str, Any]:
        """Prepare context dictionary with AI-enhanced content."""
        user_data = user_data or {}
        context = {}

        # Basic user data
        context.update({
            'title': user_data.get('judul', 'JUDUL SKRIPSI'),
            'author': user_data.get('penulis', 'Nama Penulis'),
            'nim': user_data.get('nim', 'NIM'),
            'advisor': user_data.get('dosen_pembimbing', 'Nama Dosen Pembimbing'),
            'institution': user_data.get('universitas', 'Universitas'),
            'year': user_data.get('tahun', str(datetime.now().year)),
            'date': user_data.get('tahun', str(datetime.now().year)),
        })

        # Abstracts
        context.update({
            'abstract_id': user_data.get('abstrak_teks', ''),
            'abstract_en': user_data.get('abstrak_en_teks', ''),
        })

        # Keywords
        keywords = user_data.get('kata_kunci', '')
        if isinstance(keywords, str):
            context['keywords'] = keywords.split(',')
        else:
            context['keywords'] = keywords or []

        # Get content sections
        sections = self.extractor.get_sections()
        ai_sections = self.ai_extractor.get_sections() if self.ai_extractor else sections

        # Process chapters
        chapters = []
        for section in ai_sections:
            if isinstance(section, dict) and section.get('type') == 'chapter':
                chapter = {
                    'title': section.get('title', ''),
                    'content': section.get('content', []),
                    'level': section.get('level', 1),
                }
                chapters.append(chapter)

        context['chapters'] = chapters
        context['sections'] = ai_sections

        # Generate front matter if AI available
        if self.use_ai and self.ai_extractor:
            try:
                # Generate enhanced front matter
                front_matter_gen = FrontMatterGenerator(user_data)
                context['front_matter'] = self._generate_front_matter_context(front_matter_gen, user_data)
            except Exception as e:
                print(f"[EnhancedThesisBuilder] Front matter generation failed: {e}")

        # Add metadata
        context['generated_date'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        context['ai_enhanced'] = self.use_ai

        print(f"[EnhancedThesisBuilder] Context prepared with {len(chapters)} chapters")
        return context

    def _generate_front_matter_context(self, generator: FrontMatterGenerator, user_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate front matter context for templating."""
        front_matter = {}

        # Basic front matter sections
        front_matter.update({
            'title_page': {
                'title': user_data.get('judul', ''),
                'author': user_data.get('penulis', ''),
                'nim': user_data.get('nim', ''),
                'advisor': user_data.get('dosen_pembimbing', ''),
                'institution': user_data.get('universitas', ''),
                'year': user_data.get('tahun', ''),
            },
            'abstract_indonesian': user_data.get('abstrak_teks', ''),
            'abstract_english': user_data.get('abstrak_en_teks', ''),
            'keywords': user_data.get('kata_kunci', '').split(',') if user_data.get('kata_kunci') else [],
        })

        return front_matter

    def _build_fallback(self, user_data: Optional[Dict[str, Any]]) -> Path:
        """Fallback method when docxtpl cannot be used."""
        print("[EnhancedThesisBuilder] Using fallback build method")

        # Create new document
        doc = Document()

        # Basic content addition (similar to original builder)
        self._add_basic_content(doc, user_data)

        # Save
        doc.save(str(self.output_path))

        print(f"[EnhancedThesisBuilder] Fallback document saved to: {self.output_path}")
        return self.output_path

    def _add_basic_content(self, doc: Document, user_data: Optional[Dict[str, Any]]) -> None:
        """Add basic content when templating is not available."""
        user_data = user_data or {}

        # Add title
        title = doc.add_heading(user_data.get('judul', 'JUDUL SKRIPSI'), 0)
        title.alignment = 1  # Center

        # Add author info
        author_para = doc.add_paragraph(f"Oleh: {user_data.get('penulis', 'Nama Penulis')}")
        author_para.alignment = 1

        # Add chapters
        sections = self.extractor.get_sections()
        for section in sections:
            if isinstance(section, dict) and section.get('type') == 'chapter':
                # Add chapter title
                doc.add_heading(section.get('title', ''), level=1)

                # Add content
                content = section.get('content', [])
                if isinstance(content, list):
                    for line in content:
                        if isinstance(line, str) and line.strip():
                            doc.add_paragraph(line)
                elif isinstance(content, str):
                    doc.add_paragraph(content)

                # Add page break between chapters
                doc.add_page_break()


def create_enhanced_thesis(
    template_path: str,
    content_path: str,
    output_path: str,
    user_data: Optional[Dict[str, Any]] = None,
    use_ai: bool = True,
    include_frontmatter: bool = True
) -> Dict[str, Any]:
    """Convenience function to create an enhanced thesis in one call.

    Args:
        template_path: Path to DOCX template
        content_path: Path to content file
        output_path: Path for output DOCX
        user_data: User data dictionary
        use_ai: Whether to use AI semantic analysis
        include_frontmatter: Whether to include front matter sections

    Returns:
        Dictionary with result information
    """
    try:
        # Verify files exist
        template_p = Path(template_path)
        content_p = Path(content_path)

        if not template_p.exists():
            return {
                "status": "error",
                "message": f"Template file not found: {template_path}",
            }

        if not content_p.exists():
            return {
                "status": "error",
                "message": f"Content file not found: {content_path}",
            }

        # Create enhanced builder
        builder = EnhancedThesisBuilder(template_path, content_path, output_path, use_ai=use_ai)

        # Build the thesis
        output = builder.build(user_data or {})

        return {
            "status": "success",
            "output_file": str(output),
            "message": "Enhanced thesis document created successfully",
            "method": "docxtpl" if builder._has_placeholders() else "fallback",
            "ai_enhanced": use_ai,
            "file_size": output.stat().st_size if output.exists() else 0,
        }

    except Exception as e:
        import traceback
        return {
            "status": "error",
            "message": f"Failed to create enhanced thesis: {str(e)}",
            "error_details": traceback.format_exc(),
        }