"""
Complete Thesis Document Builder
Orchestrates the creation of a complete, properly structured thesis document
with all front matter, main content, and back matter.
"""

from typing import Dict, List, Any, Optional
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
import re
from .template_analyzer import TemplateAnalyzer
from .content_extractor import ContentExtractor
from .ai_enhanced_extractor import AIEnhancedContentExtractor
from .content_mapper import ContentMapper
from .document_merger import DocumentMerger
from .front_matter_generator import FrontMatterGenerator, BackMatterGenerator
from ..parser.normalized_extractor import extract_normalized_structure
from ..ai.thesis_rewriter import ThesisRewriter

# Advanced Template Intelligence System (New)
try:
    from .advanced_template_analyzer import AdvancedTemplateAnalyzer
    from .content_zone_mapper import ContentZoneMapper
    from .style_inheritance_engine import StyleInheritanceEngine
    from .adaptive_insertion_engine import AdaptiveInsertionEngine, InsertionContext
    from .dynamic_content_generator import DynamicContentGenerator, GeneratedContent
    ADVANCED_SYSTEM_AVAILABLE = True
except ImportError:
    ADVANCED_SYSTEM_AVAILABLE = False
    print("[WARNING] Advanced template intelligence system not available, using legacy system")


class CompleteThesisBuilder:
    """Builds complete thesis documents from templates and content."""

    # Indonesian university template configurations
    UNIVERSITY_CONFIGS = {
        'indonesian_standard': {
            'chapter_prefix': 'BAB',
            'chapter_style': 'roman_upper',  # BAB I, II, III
            'chapter_titles': {
                1: 'PENDAHULUAN',
                2: 'TINJAUAN PUSTAKA',
                3: 'METODOLOGI PENELITIAN',
                4: 'ANALISIS DAN PERANCANGAN',
                5: 'IMPLEMENTASI DAN PENGUJIAN',
                6: 'PENUTUP'
            },
            'abstract_keywords': ['ABSTRAK', 'SARI', 'ABSTRAK INDONESIA'],
            'english_abstract_keywords': ['ABSTRACT', 'ABSTRAK BAHASA INGGRIS'],
            'references_keyword': 'DAFTAR PUSTAKA',
            'author_placeholder': ['N a m a', 'Nama Mahasiswa', 'Nama Lengkap'],
            'nim_placeholder': ['NIM', '94523999', 'Nomor Induk Mahasiswa'],
            'language': 'id',
            'formatting': {
                'paragraph_indent': 1.0,  # cm
                'line_spacing': 1.5,
                'font': 'Times New Roman',
                'font_size': 11,
                'alignment': 'justify'
            }
        },

        # Alternative Indonesian configurations for different universities
        'indonesian_ui': {  # Universitas Indonesia style
            'chapter_prefix': 'BAB',
            'chapter_style': 'roman_upper',
            'chapter_titles': {
                1: 'PENDAHULUAN',
                2: 'TINJAUAN PUSTAKA',
                3: 'METODOLOGI',
                4: 'HASIL DAN PEMBAHASAN',
                5: 'KESIMPULAN'
            },
            'abstract_keywords': ['ABSTRAK'],
            'english_abstract_keywords': ['ABSTRACT'],
            'references_keyword': 'DAFTAR PUSTAKA',
            'author_placeholder': ['Nama', 'Nama Mahasiswa'],
            'nim_placeholder': ['NIM', 'NPM'],
            'language': 'id',
            'formatting': {
                'paragraph_indent': 1.27,  # 1.27 cm (UI standard)
                'line_spacing': 1.5,
                'font': 'Times New Roman',
                'font_size': 12,
                'alignment': 'justify'
            }
        },

        'indonesian_ugm': {  # Universitas Gadjah Mada style
            'chapter_prefix': 'BAB',
            'chapter_style': 'roman_upper',
            'chapter_titles': {
                1: 'PENDAHULUAN',
                2: 'TINJAUAN PUSTAKA',
                3: 'METODE PENELITIAN',
                4: 'HASIL PENELITIAN DAN PEMBAHASAN',
                5: 'PENUTUP'
            },
            'abstract_keywords': ['ABSTRAK'],
            'english_abstract_keywords': ['ABSTRACT'],
            'references_keyword': 'DAFTAR PUSTAKA',
            'author_placeholder': ['Nama Mahasiswa', 'Nama Lengkap'],
            'nim_placeholder': ['NIM', 'Nomor Pokok Mahasiswa'],
            'language': 'id',
            'formatting': {
                'paragraph_indent': 1.0,
                'line_spacing': 1.5,
                'font': 'Times New Roman',
                'font_size': 11,
                'alignment': 'justify'
            }
        },

        'indonesian_itb': {  # Institut Teknologi Bandung style
            'chapter_prefix': 'BAB',
            'chapter_style': 'roman_upper',
            'chapter_titles': {
                1: 'PENDAHULUAN',
                2: 'TINJAUAN PUSTAKA',
                3: 'METODOLOGI',
                4: 'ANALISIS DAN PERANCANGAN',
                5: 'IMPLEMENTASI DAN PENGUJIAN',
                6: 'KESIMPULAN DAN SARAN'
            },
            'abstract_keywords': ['ABSTRAK'],
            'english_abstract_keywords': ['ABSTRACT'],
            'references_keyword': 'DAFTAR PUSTAKA',
            'author_placeholder': ['Nama', 'Nama Mahasiswa'],
            'nim_placeholder': ['NIM', 'NRP'],
            'language': 'id',
            'formatting': {
                'paragraph_indent': 1.0,
                'line_spacing': 1.5,
                'font': 'Times New Roman',
                'font_size': 12,
                'alignment': 'justify'
            }
        },

        'indonesian_arabic': {  # Some universities use Arabic numbers
            'chapter_prefix': 'BAB',
            'chapter_style': 'arabic',  # BAB 1, 2, 3
            'chapter_titles': {
                1: 'PENDAHULUAN',
                2: 'TINJAUAN PUSTAKA',
                3: 'METODOLOGI PENELITIAN',
                4: 'ANALISIS DAN PERANCANGAN',
                5: 'IMPLEMENTASI DAN PENGUJIAN',
                6: 'PENUTUP'
            },
            'abstract_keywords': ['ABSTRAK', 'SARI'],
            'english_abstract_keywords': ['ABSTRACT'],
            'references_keyword': 'DAFTAR PUSTAKA',
            'author_placeholder': ['Nama Mahasiswa', 'Nama Lengkap'],
            'nim_placeholder': ['NIM', 'Nomor Induk'],
            'language': 'id',
            'formatting': {
                'paragraph_indent': 1.0,
                'line_spacing': 1.5,
                'font': 'Times New Roman',
                'font_size': 11,
                'alignment': 'justify'
            }
        },

        'indonesian_ipb': {  # Institut Pertanian Bogor style
            'chapter_prefix': 'BAB',
            'chapter_style': 'roman_upper',
            'chapter_titles': {
                1: 'PENDAHULUAN',
                2: 'TINJAUAN PUSTAKA',
                3: 'BAHAN DAN METODE',
                4: 'HASIL DAN PEMBAHASAN',
                5: 'KESIMPULAN DAN SARAN'
            },
            'abstract_keywords': ['ABSTRAK'],
            'english_abstract_keywords': ['ABSTRACT'],
            'references_keyword': 'DAFTAR PUSTAKA',
            'author_placeholder': ['Nama Mahasiswa', 'Nama Lengkap'],
            'nim_placeholder': ['NIM', 'NRP'],
            'language': 'id',
            'formatting': {
                'paragraph_indent': 1.25,  # 1.25 cm
                'line_spacing': 1.5,
                'font': 'Times New Roman',
                'font_size': 12,
                'alignment': 'justify'
            }
        }
    }

    def __init__(self, template_path: str, content_path: str, output_path: str, use_ai: bool = True, api_key: Optional[str] = None, include_frontmatter: bool = True, university_config: str = 'indonesian_standard'):
        """Initialize with paths and options.

        Args:
            template_path: Path to DOCX template file
            content_path: Path to content file (DOCX/TXT)
            output_path: Path for output DOCX file
            use_ai: Whether to use AI enhancement
            api_key: OpenRouter API key for AI features
            include_frontmatter: Whether to include front matter
            university_config: University template configuration ('indonesian_standard', 'english_standard', 'international')
        """
        self.template_path = Path(template_path)
        self.content_path = Path(content_path)
        self.output_path = Path(output_path)
        self.use_ai = use_ai
        self.include_frontmatter = include_frontmatter
        self.api_key = api_key

        # Load university configuration
        self.config = self.UNIVERSITY_CONFIGS.get(university_config, self.UNIVERSITY_CONFIGS['indonesian_standard'])
        print(f"[INIT] Using university configuration: {university_config}")

        # Initialize components
        self.analyzer = TemplateAnalyzer(str(self.template_path))

        # Use AI-enhanced extractor when available
        self.ai_extractor = AIEnhancedContentExtractor(str(self.content_path), use_ai=use_ai, api_key=api_key)
        self.extractor = ContentExtractor(str(self.content_path))
        self.ai_data = {}  # Will be populated during build
        self.mapper = ContentMapper(self.analyzer, self.extractor, self.ai_data)

        # Initialize AI-powered components for perfect formatting (optional)
        self.perfect_adapter = None
        self.quality_validator = None

        # Initialize AI components dynamically in build method
        self.thesis_rewriter = ThesisRewriter(api_key=api_key) if api_key else None
        self.ai_data = {}  # Will be populated during build

        # Initialize semantic validation
        self.semantic_validator = None

        # AI extraction will be run during build process, not in init

        # Initialize AI-powered components for perfect formatting (optional)
        self.perfect_adapter = None
        self.quality_validator = None

        # Initialize AI components dynamically in build method
        self.thesis_rewriter = ThesisRewriter(api_key=api_key) if api_key else None
        self.ai_data = {}  # Will be populated during build

        # Semantic validation
        self.semantic_validation = self.ai_extractor.get_semantic_validation() if use_ai else None

    def _copy_styles_from_template(self, target_doc: Document, template_doc: Document) -> None:
        """Copy styles from template document to target document."""
        try:
            # Copy paragraph styles
            for style in template_doc.styles:
                if style.name not in target_doc.styles:
                    try:
                        # Create a copy of the style in the target document
                        new_style = target_doc.styles.add_style(style.name, style.type)
                        if hasattr(style, 'font') and style.font:
                            new_style.font.name = style.font.name
                            new_style.font.size = style.font.size
                            new_style.font.bold = style.font.bold
                            new_style.font.italic = style.font.italic
                    except Exception:
                        # Skip styles that can't be copied
                        pass
        except Exception as e:
            print(f"[WARNING] Failed to copy styles from template: {e}")
    
    def build(self, user_data: Optional[Dict[str, Any]] = None) -> Path:
        """Build the complete thesis document with perfect formatting.

        Args:
            user_data: Dictionary with user data for personalization

        Returns:
            Path to created DOCX file
        """
        user_data = user_data or {}
        print(f"[DEBUG] Starting build with use_ai={self.use_ai}, api_key={'present' if self.api_key else 'missing'}")

        # CRITICAL FIX: Check if we have substantial analyzed_data from first AI call
        analyzed_data = None
        if hasattr(self.ai_extractor, 'analyzed_data') and self.ai_extractor.analyzed_data:
            analyzed_data = self.ai_extractor.analyzed_data
            user_data['analyzed_data'] = analyzed_data

            # Check content quality
            total_chars = sum(
                len(str(content)) for chapter in [analyzed_data.get(f'chapter{i}', {}) for i in range(1, 7)]
                for content in chapter.values() if content
            )

            if total_chars > 1000:  # Substantial content threshold
                print(f"[INFO] Found substantial analyzed content ({total_chars} chars), using direct content insertion")
                return self._build_with_analyzed_data_direct(user_data, analyzed_data)

        # Try advanced template intelligence system
        if ADVANCED_SYSTEM_AVAILABLE and self._should_use_advanced_system():
            try:
                return self._build_with_advanced_system(user_data)
            except Exception as e:
                print(f"[WARNING] Advanced system failed ({e}), falling back to legacy system")

        # Legacy system with critical fixes
        return self._build_with_legacy_system(user_data)

    def _should_use_advanced_system(self) -> bool:
        """Determine if we should use the advanced template intelligence system"""
        return bool(ADVANCED_SYSTEM_AVAILABLE and
                   self.use_ai and
                   hasattr(self, 'api_key') and self.api_key)

    def _build_with_advanced_system(self, user_data: Dict[str, Any]) -> Path:
        """Build using the advanced template intelligence system with dynamic content generation"""
        print("[INFO] 🚀 Using Advanced Template Intelligence System v2.0")

        # Step 1: Analyze template with advanced analyzer
        template_analyzer = AdvancedTemplateAnalyzer(str(self.template_path))
        template_structure = template_analyzer.analyze_template_comprehensive()

        # Step 2: Generate dynamic AI content based on template analysis
        content_generator = DynamicContentGenerator()

        # Get user text (from content_path or user_data)
        user_text = self._get_user_text_for_ai(user_data)
        if not user_text:
            print("[WARNING] No user text found, using minimal content generation")
            user_text = "Generate a basic academic thesis structure."

        # Generate content with template awareness
        generated_content = content_generator.generate_content(
            user_text=user_text,
            template_structure=template_structure,
            user_metadata=user_data,
            api_key=self.api_key or ""
        )

        analyzed_data = generated_content.content
        print(f"[INFO] Generated content: {len(analyzed_data)} chapters with quality score {generated_content.quality_metrics.get('overall_score', 0):.1f}")

        # Step 3: Map content to zones
        zone_mapper = ContentZoneMapper(template_structure)
        insertion_plan = zone_mapper.map_ai_content_to_zones(analyzed_data)

        # Step 4: Load and prepare document
        doc = Document(str(self.template_path))

        # Step 5: Apply metadata and front matter
        self._apply_metadata_and_front_matter(doc, user_data, analyzed_data)

        # Step 6: Execute intelligent insertion
        style_engine = StyleInheritanceEngine()
        insertion_context = InsertionContext(
            document=doc,
            template_structure=template_structure,
            insertion_plan=insertion_plan,
            style_engine=style_engine
        )

        insertion_engine = AdaptiveInsertionEngine()
        result = insertion_engine.execute_insertion_plan(insertion_context)

        if not result.success:
            error_msg = f"Content insertion failed: {', '.join(result.errors)}"
            raise ValueError(error_msg)

        # Step 7: Save and return
        output_path = self._get_output_path(user_data)
        doc.save(str(output_path))

        print(f"[SUCCESS] Advanced system v2.0 completed. Output: {output_path}")
        print(f"[METRICS] Content quality: {generated_content.quality_metrics.get('overall_score', 0):.1f}")
        return output_path

    def _get_user_text_for_ai(self, user_data: Dict[str, Any]) -> str:
        """Get user text for AI content generation"""
        # Try to get from user_data first (direct input)
        if 'content' in user_data and user_data['content']:
            return user_data['content']

        # Try to read from content file
        if hasattr(self, 'content_path') and self.content_path and self.content_path.exists():
            try:
                with open(self.content_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except Exception as e:
                print(f"[WARNING] Could not read content file: {e}")

        # Fallback to minimal content
        return "Generate a comprehensive academic thesis on information systems development."

    def _get_ai_content(self, user_data: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        """Get AI-generated content for processing"""
        if hasattr(self.ai_extractor, 'analyzed_data') and self.ai_extractor.analyzed_data:
            analyzed_data = self.ai_extractor.analyzed_data
            user_data['analyzed_data'] = analyzed_data

            # Validate content quality
            total_chars = 0
            for chapter_key in ['chapter1', 'chapter2', 'chapter3', 'chapter4', 'chapter5', 'chapter6']:
                if chapter_key in analyzed_data:
                    chapter = analyzed_data[chapter_key]
                    chapter_chars = sum(len(str(content)) for content in chapter.values())
                    total_chars += chapter_chars

            if total_chars > 5000:  # Substantial content threshold
                print(f"[INFO] Found substantial AI content ({total_chars} chars)")
                return analyzed_data
            else:
                print(f"[WARNING] AI content too short ({total_chars} chars)")
                return None

        return None

    def _apply_metadata_and_front_matter(self, doc, user_data: Dict[str, Any], analyzed_data: Dict[str, Any]):
        """Apply metadata and front matter to document"""
        # Reuse existing methods
        self._clean_template_instructions(doc, user_data, self.config)
        self._insert_abstract(doc, analyzed_data, self.config)

    def _build_with_analyzed_data_direct(self, user_data: Dict[str, Any], analyzed_data: Dict[str, Any]) -> Path:
        """Build document using comprehensive analyzed_data from first AI call."""
        import re
        print("\n" + "="*60)
        print("BUILDING DOCUMENT WITH ANALYZED CONTENT")
        print("="*60)
        
        # Define subsection titles for heading replacement
        subsection_titles = {
            'latar_belakang': 'Latar Belakang',
            'rumusan_masalah': 'Rumusan Masalah',
            'tujuan_penelitian': 'Tujuan Penelitian',
            'manfaat_penelitian': 'Manfaat Penelitian',
            'batasan_masalah': 'Batasan Masalah dan Ruang Lingkup',
            'landasan_teori': 'Landasan Teori',
            'penelitian_terkait': 'Penelitian Terkait',
            'kerangka_pemikiran': 'Kerangka Pemikiran',
            'desain_penelitian': 'Desain Penelitian',
            'metode_pengumpulan_data': 'Metode Pengumpulan Data',
            'metode_analisis': 'Metode Analisis Data',
            'tools': 'Alat dan Bahan',
            'analisis_kebutuhan': 'Analisis Kebutuhan Sistem',
            'perancangan_sistem': 'Perancangan Sistem',
            'perancangan_interface': 'Perancangan Antarmuka',
            'implementasi': 'Implementasi Sistem',
            'hasil_pengujian': 'Hasil Pengujian',
            'pembahasan': 'Pembahasan Hasil',
            'evaluasi': 'Evaluasi Sistem',
            'kesimpulan': 'Kesimpulan',
            'saran': 'Saran'
        }

        # Verify analyzed_data has content
        print(f"[VERIFY] Analyzed data keys: {list(analyzed_data.keys())}")

        # Check each chapter content
        for i in range(1, 7):
            chapter_key = f'chapter{i}'
            if chapter_key in analyzed_data:
                chapter = analyzed_data[chapter_key]
                print(f"\n[VERIFY] Chapter {i}:")
                for subsection, content in chapter.items():
                    content_length = len(content) if content else 0
                    preview = content[:100] if content else "[EMPTY]"
                    # Handle Unicode characters safely
                    try:
                        print(f"  - {subsection}: {content_length} chars | Preview: {preview}...")
                    except UnicodeEncodeError:
                        safe_preview = preview.encode('ascii', 'replace').decode('ascii')
                        print(f"  - {subsection}: {content_length} chars | Preview: {safe_preview}...")

                    if content_length < 200:
                        print(f"  [WARNING] {subsection} content is too short ({content_length} chars)")
            else:
                print(f"\n[VERIFY] Chapter {i}: NOT FOUND")

        # Load template
        try:
            doc = Document(str(self.template_path))
            print(f"\n[INFO] Template loaded: {len(doc.paragraphs)} paragraphs")
        except Exception as e:
            print(f"[ERROR] Failed to load template: {e}")
            return self.output_path

        # PHASE 1: Structural Landmark Analysis (BEFORE cleaning)
        print("[INFO] Phase 1: Analyzing template structural landmarks...")
        landmark_chapters = {} # {num: paragraph_obj}
        landmark_subsections = [] # list of (paragraph_obj, chapter_num, original_text)
        
        current_chapter = 0
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            text_upper = text.upper()
            
            # 1. Detect Chapters
            is_ch = False
            ch_num = 0
            
            # Pattern: BAB I, BAB 1, CHAPTER 1 (handle Shift+Enter \v or \n)
            text_norm = text_upper.replace('\v', ' ').replace('\n', ' ')
            ch_match = re.match(r'^(BAB|CHAPTER)\s+([IVX\d]+)', text_norm)
            if ch_match:
                num_str = ch_match.group(2)
                try:
                    if num_str.isdigit(): ch_num = int(num_str)
                    else: ch_num = self._roman_to_int(num_str)
                    is_ch = True
                except: pass
            
            # Alternative: Style-based for auto-numbered headers
            elif (para.style and 'Heading 1' in str(para.style.name)) or (text_upper.startswith('BAB') and len(text) < 100):
                if any(kw in text_upper for kw in ['PENDAHULUAN', 'LATAR BELAKANG']): ch_num = 1
                elif any(kw in text_upper for kw in ['PUSTAKA', 'TEORI']): ch_num = 2
                elif any(kw in text_upper for kw in ['METODOLOGI', 'METODE']): ch_num = 3
                elif any(kw in text_upper for kw in ['ANALISIS', 'PERANCANGAN']): ch_num = 4
                elif any(kw in text_upper for kw in ['IMPLEMENTASI', 'HASIL', 'PEMBAHASAN']): ch_num = 5
                elif any(kw in text_upper for kw in ['PENUTUP', 'KESIMPULAN']): ch_num = 6
                
                if ch_num > 0 and ch_num not in landmark_chapters:
                    is_ch = True

            if is_ch:
                landmark_chapters[ch_num] = para
                current_chapter = ch_num
                print(f"[DEBUG] Landmark found: Chapter {ch_num} at paragraph {i}")
                continue

            # 2. Detect Subsections (Subbab)
            # UII template often uses "Subbab" or "Anak Subbab" as placeholders
            text_clean = text.replace('[', '').replace(']', '').strip()
            if text_clean == 'Subbab' or text_clean == 'Anak Subbab' or (text_clean.startswith('Subbab') and len(text_clean) < 20):
                landmark_subsections.append({
                    'para': para,
                    'chapter': current_chapter,
                    'is_anak': 'Anak' in text_clean
                })
                print(f"[DEBUG] Landmark found: Subsection anchor at paragraph {i} (Chapter {current_chapter})")

        print(f"[INFO] Structure Analysis Complete: Found {len(landmark_chapters)} chapters and {len(landmark_subsections)} subsection anchors")

        # PHASE 2: Metadata and Abstract (Standard logic)
        metadata = analyzed_data.get('metadata', {})
        title = metadata.get('title', user_data.get('title', 'Untitled Thesis'))
        author = metadata.get('author', user_data.get('author', 'Unknown'))
        nim = metadata.get('nim', user_data.get('nim', '00000000'))

        # Metadata replacement is now handled inside _clean_template_instructions
        # with more robust and selective logic.
        
        # Insert abstract
        self._insert_abstract(doc, analyzed_data, self.config)

        # PHASE 3: Clean template instructions (now protects landmarks)
        print("[INFO] Phase 3: Cleaning template instructions...")
        self._clean_template_instructions(doc, user_data, self.config)

        # PHASE 4: Targeted Content Insertion
        print("\n[INFO] Phase 4: Targeted Content Insertion...")
        total_insertions = 0
        
        # RE-SCAN Landmarks after cleaning to ensure object validity
        landmark_chapters = {}
        landmark_subsections = []
        current_chapter = 0
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            text_upper = para.text.upper().replace('\v', ' ').replace('\n', ' ')
            
            # Detect Chapters
            ch_match = re.match(r'^(BAB|CHAPTER)\s+([IVX\d]+)', text_upper)
            if ch_match:
                num_str = ch_match.group(2)
                try:
                    ch_num = int(num_str) if num_str.isdigit() else self._roman_to_int(num_str)
                    landmark_chapters[ch_num] = para
                    current_chapter = ch_num
                except: pass
            elif (para.style and 'Heading 1' in str(para.style.name)) or (text_upper.startswith('BAB') and len(text) < 100):
                 if any(kw in text_upper for kw in ['PENDAHULUAN', 'LATAR BELAKANG']): ch_num = 1
                 elif any(kw in text_upper for kw in ['PUSTAKA', 'TEORI']): ch_num = 2
                 elif any(kw in text_upper for kw in ['METODOLOGI', 'METODE']): ch_num = 3
                 elif any(kw in text_upper for kw in ['ANALISIS', 'PERANCANGAN']): ch_num = 4
                 elif any(kw in text_upper for kw in ['IMPLEMENTASI', 'HASIL', 'PEMBAHASAN']): ch_num = 5
                 elif any(kw in text_upper for kw in ['PENUTUP', 'KESIMPULAN']): ch_num = 6
                 if ch_num > 0:
                     landmark_chapters[ch_num] = para
                     current_chapter = ch_num

            # Detect Subsections
            text_clean = text.replace('[', '').replace(']', '').strip()
            if text_clean == 'Subbab' or text_clean == 'Anak Subbab' or (text_clean.startswith('Subbab') and len(text_clean) < 20):
                landmark_subsections.append({
                    'para': para,
                    'chapter': current_chapter,
                    'is_anak': 'Anak' in text_clean
                })

        # Map AI content keys to logical order per chapter
        chapter_key_order = {
            1: ['latar_belakang', 'rumusan_masalah', 'tujuan_penelitian', 'manfaat_penelitian', 'batasan_masalah'],
            2: ['landasan_teori', 'penelitian_terkait', 'kerangka_pemikiran'],
            3: ['desain_penelitian', 'metode_pengumpulan_data', 'metode_analisis', 'tools'],
            4: ['analisis_kebutuhan', 'perancangan_sistem', 'perancangan_interface'],
            5: ['implementasi', 'hasil_pengujian', 'pembahasan', 'evaluasi'],
            6: ['kesimpulan', 'saran']
        }

        # Insert chapter content
        chapters_data = {
            1: analyzed_data.get('chapter1', {}),
            2: analyzed_data.get('chapter2', {}),
            3: analyzed_data.get('chapter3', {}),
            4: analyzed_data.get('chapter4', {}),
            5: analyzed_data.get('chapter5', {}),
            6: analyzed_data.get('chapter6', {})
        }

        for chapter_num in range(1, 7):
            chapter_content = chapters_data.get(chapter_num, {})
            if not chapter_content:
                print(f"[DEBUG] No content keys for Chapter {chapter_num}")
                continue
            
            # Find subsections for THIS chapter
            current_sub_anchors = [s for s in landmark_subsections if s['chapter'] == chapter_num and not s['is_anak']]
            print(f"[DEBUG] Chapter {chapter_num} has {len(current_sub_anchors)} sub-anchors")
            
            # Find the chapter heading to update its text (remove "TULISKAN JUDUL...")
            last_inserted_para = None
            if chapter_num in landmark_chapters:
                ch_para = landmark_chapters[chapter_num]
                last_inserted_para = ch_para
                ch_text = ch_para.text
                if 'TULISKAN' in ch_text.upper() or 'JUDUL BAB' in ch_text.upper():
                    # Preserve "BAB I" part but replace the rest
                    main_title = subsection_titles.get(chapter_key_order[chapter_num][0], "PENDAHULUAN").upper()
                    if chapter_num == 2: main_title = "LANDASAN TEORI"
                    elif chapter_num == 3: main_title = "METODOLOGI PENELITIAN"
                    elif chapter_num == 4: main_title = "ANALISIS DAN PERANCANGAN"
                    elif chapter_num == 5: main_title = "HASIL DAN PEMBAHASAN"
                    elif chapter_num == 6: main_title = "KESIMPULAN DAN SARAN"
                    
                    prefix = f"BAB {self._to_roman(chapter_num)}"
                    ch_para.clear()
                    run = ch_para.add_run(prefix)
                    run.bold = True
                    ch_para.add_run("\n") # Shift+Enter
                    run2 = ch_para.add_run(main_title)
                    run2.bold = True

            # Insert subsections
            for i, key in enumerate(chapter_key_order.get(chapter_num, [])):
                content = chapter_content.get(key)
                if not content or len(content.strip()) < 50:
                    print(f"[DEBUG] Skipping {key}: content too short or missing")
                    continue
                
                target_para = None
                
                # Case A: We have a "Subbab" anchor for this position
                if i < len(current_sub_anchors):
                    anchor = current_sub_anchors[i]
                    anchor_para = anchor['para']
                    
                    # 1. Update Anchor Text
                    title_text = subsection_titles.get(key, key.replace('_', ' ').title())
                    prefix = f"{chapter_num}.{i+1} " # Force clean numbering
                    anchor_para.text = f"{prefix}{title_text}"
                    anchor_para.bold = True
                    
                    # 2. Find target paragraph (the one after the anchor)
                    try:
                        idx = doc.paragraphs.index(anchor_para)
                        next_para = doc.paragraphs[idx+1] if idx+1 < len(doc.paragraphs) else None
                        
                        # If next para is another anchor or BAB heading, we MUST insert a new one
                        if next_para and (next_para.text.strip() == "" or len(next_para.text) < 100) and \
                           not ('Subbab' in next_para.text or 'BAB' in next_para.text.upper()):
                            target_para = next_para
                        else:
                            target_para = self._insert_paragraph_after(anchor_para, content)
                            last_inserted_para = target_para
                            total_insertions += 1
                            print(f"[SUCCESS] Created new paragraph for {key} in Chapter {chapter_num}")
                            continue
                    except:
                        pass
                
                # Case B: No anchor found, or anchor logic failed - insert after last para
                if not target_para:
                    if last_inserted_para:
                        target_para = self._insert_paragraph_after(last_inserted_para, content)
                        print(f"[SUCCESS] Inserted {key} after previous paragraph in Chapter {chapter_num}")
                    else:
                        print(f"[WARNING] Nowhere to put {key} in Chapter {chapter_num}")

                if target_para:
                    target_para.clear()
                    target_para.add_run(content)
                    self._apply_paragraph_formatting(target_para)
                    last_inserted_para = target_para
                    total_insertions += 1
                    print(f"[SUCCESS] Filled {key} in Chapter {chapter_num}")

        print(f"\n[FINAL] Total content insertions: {total_insertions}")

        if total_insertions == 0:
            print("[ERROR] NO CONTENT WAS INSERTED!")
            print("[ERROR] Check that analyzed_data has content and template has content paragraphs")
            return self.output_path

        # APPLY CRITICAL FORMATTING FIXES
        print("[INFO] Applying critical formatting fixes...")
        for para in doc.paragraphs:
            if para.text.strip():
                self._apply_paragraph_formatting(para)

        self._apply_list_formatting(doc)
        self._apply_heading_formatting(doc)

        # Final save
        doc.save(str(self.output_path))
        print(f"[INFO] Document saved to: {self.output_path}")

        final_size = self.output_path.stat().st_size if self.output_path.exists() else 0
        print(f"[INFO] Final document size: {final_size} bytes")

        return self.output_path

    def _apply_paragraph_formatting(self, para):
        """Apply standard academic formatting to a paragraph."""
        # CRITICAL: Do NOT format if it's likely a title, heading, or special section
        style_name = str(para.style.name).lower() if para.style else ""
        if any(s in style_name for s in ['heading', 'title', 'subtitle', 'caption', 'toc', 'daftar']):
            return

        # Check if text looks like a heading (all caps, short)
        if para.text.isupper() and len(para.text) < 100:
            return

        # 1. INDENTATION: 1cm first line
        from docx.shared import Inches, Pt

        # 1. PARAGRAPH INDENTATION: First line indent 1 cm
        try:
            para.paragraph_format.first_line_indent = Inches(1.0)  # 1 cm indent
        except:
            pass

        # 2. ALIGNMENT: Justified (Ctrl + J equivalent)
        try:
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        except:
            pass

        # 3. LINE SPACING: 1.5 lines
        try:
            para.paragraph_format.line_spacing = 1.5
        except:
            pass

        # 4. FONT: Times New Roman, 11pt
        for run in para.runs:
            try:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            except:
                pass

        # 5. STYLE: Apply 'isi paragraf' if available
        try:
            if hasattr(para, 'style') and para.style:
                # Keep existing style if it's already 'isi paragraf'
                if 'isi paragraf' in str(para.style).lower():
                    pass  # Keep it
                else:
                    # Try to set to isi paragraf style
                    para.style = 'isi paragraf'
        except:
            pass

    def _apply_list_formatting(self, doc):
        """Apply proper list formatting (CRITICAL FIXES)."""
        from docx.shared import Inches

        # Find all paragraphs that might be lists
        for para in doc.paragraphs:
            text = para.text.strip()

            # Check for numbered lists (1., 2., etc.) or bulleted (•, -, etc.)
            if (text.startswith(('1.', '2.', '3.', '4.', '5.', 'a.', 'b.', 'c.', 'd.', '•', '-', '–'))) or \
               any(text.startswith(f'{i}.') for i in range(1, 21)):

                try:
                    # Level 1 lists (a, b, c) - align with left margin
                    if text.startswith(('a.', 'b.', 'c.', 'd.')):
                        para.paragraph_format.left_indent = Inches(0.0)  # No indentation

                    # Level 2 lists (1), 2), etc.) - indent 0.8 cm
                    elif text.startswith(('1)', '2)', '3)', '4)', '5)')) or \
                         any(text.startswith(f'{i})') for i in range(1, 21)):
                        para.paragraph_format.left_indent = Inches(0.8)  # 0.8 cm indentation

                    # Regular numbered lists - standard indentation
                    elif any(text.startswith(f'{i}.') for i in range(1, 21)):
                        para.paragraph_format.left_indent = Inches(0.3)  # Light indentation

                except:
                    pass

    def _clean_template_instructions(self, doc, user_data, config):
        """Remove template instructional text and replace placeholders with dynamic user metadata detection."""
        replacements = 0

        # Extract comprehensive user metadata with fallbacks
        title = user_data.get('title', user_data.get('judul', ''))
        author = user_data.get('author', user_data.get('nama', user_data.get('penulis', '')))
        nim = user_data.get('nim', user_data.get('nomor_induk', ''))
        university = user_data.get('university', user_data.get('universitas', ''))
        faculty = user_data.get('faculty', user_data.get('fakultas', ''))
        program = user_data.get('program', user_data.get('program_studi', user_data.get('jurusan', '')))
        department = user_data.get('department', user_data.get('departemen', ''))
        supervisor1 = user_data.get('supervisor1', user_data.get('pembimbing1', user_data.get('dosen_pembimbing', '')))
        supervisor2 = user_data.get('supervisor2', user_data.get('pembimbing2', ''))
        examiner1 = user_data.get('examiner1', user_data.get('penguji1', ''))
        examiner2 = user_data.get('examiner2', user_data.get('penguji2', ''))
        city = user_data.get('city', user_data.get('kota', ''))
        year = user_data.get('year', user_data.get('tahun', ''))
        degree = user_data.get('degree', user_data.get('gelar', ''))

        print(f"[INFO] Processing user metadata:")
        print(f"  title: '{title[:50]}...'")
        print(f"  author: '{author}'")
        print(f"  nim: '{nim}'")
        print(f"  university: '{university}'")
        print(f"  faculty: '{faculty}'")
        print(f"  program: '{program}'")
        print(f"  department: '{department}'")
        print(f"  supervisor1: '{supervisor1}'")
        print(f"  supervisor2: '{supervisor2}'")
        print(f"  examiner1: '{examiner1}'")
        print(f"  examiner2: '{examiner2}'")
        print(f"  city: '{city}'")
        print(f"  year: '{year}'")
        print(f"  degree: '{degree}'")

        # Phase 1: Dynamic metadata detection and replacement
        replacements += self._apply_dynamic_metadata_replacement(doc, {
            'title': title,
            'author': author,
            'nim': nim,
            'university': university,
            'faculty': faculty,
            'program': program,
            'department': department,
            'supervisor1': supervisor1,
            'supervisor2': supervisor2,
            'examiner1': examiner1,
            'examiner2': examiner2,
            'city': city,
            'year': year,
            'degree': degree
        })

        # Phase 2: Remove instructional text (keep existing logic)
        replacements += self._remove_instructional_text(doc, config)

        print(f"[INFO] Cleaned {replacements} template instructions and placeholders")
        return replacements

    def _extract_template_metadata(self, doc) -> Dict[str, str]:
        """Extract university, faculty, program, and other metadata from template."""
        template_metadata = {
            'university': '',
            'faculty': '',
            'program': '',
            'department': '',
            'supervisor1': '',
            'supervisor2': '',
            'examiner1': '',
            'examiner2': '',
            'city': '',
            'year': '',
            'degree': ''
        }

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            text_upper = text.upper()

            # Extract university information
            if not template_metadata['university']:
                university_patterns = [
                    'UNIVERSITAS ISLAM INDONESIA', 'UII',
                    'UNIVERSITAS GADJAH MADA', 'UGM',
                    'UNIVERSITAS INDONESIA', 'UI',
                    'INSTITUT TEKNOLOGI BANDUNG', 'ITB',
                    'INSTITUT PERTANIAN BOGOR', 'IPB'
                ]
                for pattern in university_patterns:
                    if pattern in text_upper:
                        template_metadata['university'] = text.strip()
                        break

            # Extract faculty information
            if not template_metadata['faculty']:
                faculty_patterns = [
                    'FAKULTAS TEKNIK', 'FAKULTAS MIPA', 'FAKULTAS EKONOMI',
                    'FAKULTAS KEDOKTERAN', 'FAKULTAS HUKUM', 'FAKULTAS SASTRA'
                ]
                for pattern in faculty_patterns:
                    if pattern in text_upper:
                        template_metadata['faculty'] = text.strip()
                        break

            # Extract program information
            if not template_metadata['program']:
                program_patterns = [
                    'PROGRAM STUDI INFORMATIKA', 'TEKNIK INFORMATIKA',
                    'PROGRAM STUDI TEKNIK', 'SISTEM INFORMASI',
                    'PROGRAM STUDI EKONOMI', 'PROGRAM STUDI MANAJEMEN'
                ]
                for pattern in program_patterns:
                    if pattern in text_upper:
                        template_metadata['program'] = text.strip()
                        break

            # Extract department information
            if not template_metadata['department']:
                dept_patterns = ['JURUSAN', 'DEPARTEMEN', 'DEPARTMENT']
                for pattern in dept_patterns:
                    if pattern in text_upper and len(text.split()) <= 5:
                        template_metadata['department'] = text.strip()
                        break

            # Extract city information
            if not template_metadata['city']:
                city_patterns = ['YOGYAKARTA', 'JAKARTA', 'BANDUNG', 'SURABAYA', 'SEMARANG', 'MALANG']
                for pattern in city_patterns:
                    if pattern in text_upper and ',' in text:
                        # Extract city from "City, Date" format
                        city_part = text.split(',')[0].strip()
                        if city_part.upper() in city_patterns:
                            template_metadata['city'] = city_part
                            break

            # Extract year information
            if not template_metadata['year']:
                year_match = re.search(r'\b(20\d{2})\b', text)
                if year_match:
                    template_metadata['year'] = year_match.group(1)

            # Extract supervisor names (look for patterns)
            if not template_metadata['supervisor1']:
                # Look for supervisor patterns
                supervisor_patterns = [
                    r'\(\s*([^,]+),\s*S\.T[^)]*\)',  # (Name, S.T.)
                    r'Pembimbing\s*:\s*([^,\n]+)',   # Pembimbing: Name
                ]
                for pattern in supervisor_patterns:
                    match = re.search(pattern, text, re.IGNORECASE)
                    if match:
                        template_metadata['supervisor1'] = match.group(1).strip()
                        break

            # Extract degree information
            if not template_metadata['degree']:
                degree_patterns = [
                    'SARJANA KOMPUTER', 'SARJANA TEKNIK', 'SARJANA SAINS',
                    'MAGISTER', 'DOKTOR', 'S.KOM', 'S.T.', 'M.KOM', 'M.T.'
                ]
                for pattern in degree_patterns:
                    if pattern in text_upper:
                        # Look for degree context
                        if 'gelar' in text.lower() or 'sarjana' in text.lower():
                            template_metadata['degree'] = text.strip()
                            break

        return template_metadata

    def _apply_dynamic_metadata_replacement(self, doc, user_metadata: Dict[str, str]) -> int:
        """Dynamically detect and replace user metadata using pattern recognition."""
        replacements = 0

        # First, extract template metadata to fill in gaps
        template_metadata = self._extract_template_metadata(doc)

        # Merge user metadata with template-extracted metadata
        complete_metadata = {}
        complete_metadata.update(template_metadata)  # Template defaults
        complete_metadata.update(user_metadata)      # User overrides

        print(f"[INFO] Complete metadata: university='{complete_metadata.get('university', '')}', "
              f"faculty='{complete_metadata.get('faculty', '')}', "
              f"program='{complete_metadata.get('program', '')}', "
              f"supervisor1='{complete_metadata.get('supervisor1', '')}'")

        # Process paragraphs for metadata replacement
        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue

            text = para.text.strip()
            original_text = text

            # TITLE DETECTION AND REPLACEMENT (EXTREMELY SELECTIVE)
            if user_metadata.get('title'):
                # ONLY replace very specific title placeholders that are clearly instructional

                # Pattern 1: Explicit instructional title text (very specific)
                instructional_phrases = [
                    'tuliskan judul',
                    'tulis judul',
                    'judul skripsi',
                    'judul tesis',
                    'bagian ini adalah bagian judul',
                    'tuliskan judul dengan pola'
                ]

                # EXCLUDE chapter headers that might contain "judul"
                is_chapter_header = any(term in text.upper() for term in ['BAB I', 'BAB II', 'BAB III', 'BAB IV', 'BAB V', 'BAB VI', 'CHAPTER'])

                if any(phrase in text.lower() for phrase in instructional_phrases) and not is_chapter_header:
                    # Preserve prefix if exists
                    if ':' in text:
                        prefix = text.split(':')[0] + ': '
                        para.text = prefix + user_metadata['title']
                    else:
                        para.text = user_metadata['title']
                    replacements += 1
                    print(f"[METADATA] Replaced title (instructional pattern): '{original_text[:50]}...' -> '{user_metadata['title'][:50]}...'")
                    continue

                # Pattern 2: ONLY replace very specific short title placeholders (exclude chapter headers and other document elements)
                words = text.lower().split()
                if (len(words) <= 4 and
                    any(word in ['judul', 'title'] for word in words) and
                    not is_chapter_header and
                    not any(term in text.upper() for term in ['HALAMAN', 'LEMBAR', 'KATA', 'DAFTAR', 'TABEL', 'GAMBAR', 'FAKULTAS', 'UNIVERSITAS', 'PROGRAM', 'PENGESAHAN', 'PENGANTAR'])):
                    para.text = user_metadata['title']
                    replacements += 1
                    print(f"[METADATA] Replaced title (short title placeholder): '{original_text}' -> '{user_metadata['title'][:50]}...'")
                    continue

            # AUTHOR/NAME DETECTION AND REPLACEMENT
            if user_metadata.get('author'):
                # Pattern 1: Specific name placeholders that are clearly meant to be replaced
                name_placeholders = [
                    'nama mahasiswa', 'nama lengkap mahasiswa', 'student name',
                    'author name', 'john doe', 'jane smith'
                ]

                if any(placeholder.lower() in text.lower() for placeholder in name_placeholders):
                    if ':' in text:
                        prefix = text.split(':')[0] + ': '
                        para.text = prefix + user_metadata['author']
                    else:
                        para.text = user_metadata['author']
                    replacements += 1
                    print(f"[METADATA] Replaced author (placeholder pattern): '{original_text}' -> '{user_metadata['author']}'")
                    continue

                # Pattern 2: Parenthesized placeholders like "(Nama Mahasiswa)"
                if text.strip().startswith('(') and text.strip().endswith(')') and 'mahasiswa' in text.lower():
                    para.text = f"({user_metadata['author']})"
                    replacements += 1
                    print(f"[METADATA] Replaced author (parentheses pattern): '{original_text}' -> '({user_metadata['author']})'")
                    continue

            # NIM/Student ID DETECTION AND REPLACEMENT
            if user_metadata.get('nim'):
                # Pattern 1: 8-10 digit numbers (common NIM length)
                nim_pattern = re.compile(r'\b\d{8,10}\b')
                if nim_pattern.search(text):
                    # Replace the number, keep surrounding text
                    new_text = nim_pattern.sub(user_metadata['nim'], text)
                    if new_text != text:
                        para.text = new_text
                        replacements += 1
                        print(f"[METADATA] Replaced NIM (digit pattern): '{original_text}' -> '{new_text}'")
                        continue

                # Pattern 2: NIM-related keywords
                nim_keywords = ['nim', 'nomor induk', 'student id', 'no. mahasiswa', 'nrp', 'npm']
                if any(keyword.lower() in text.lower() for keyword in nim_keywords):
                    # Look for numbers in this paragraph or nearby
                    for check_i in range(max(0, i-1), min(len(doc.paragraphs), i+2)):
                        check_text = doc.paragraphs[check_i].text
                        if nim_pattern.search(check_text):
                            new_text = nim_pattern.sub(user_metadata['nim'], check_text)
                            doc.paragraphs[check_i].text = new_text
                            replacements += 1
                            print(f"[METADATA] Replaced NIM (keyword pattern): '{check_text}' -> '{new_text}'")
                            break
                    continue

            # UNIVERSITY/INSTITUTION DETECTION AND REPLACEMENT
            if complete_metadata.get('university') and complete_metadata['university'] != template_metadata.get('university'):
                university_keywords = ['universitas', 'university', 'institut', 'institute', 'sekolah tinggi']
                if any(keyword.lower() in text.lower() for keyword in university_keywords):
                    target = complete_metadata['university']
                    if ':' in text:
                        para.text = text.split(':')[0] + ': ' + target
                    else:
                        para.text = target
                    replacements += 1
                    continue

            # FACULTY DETECTION AND REPLACEMENT
            if complete_metadata.get('faculty') and complete_metadata['faculty'] != template_metadata.get('faculty'):
                faculty_keywords = ['fakultas', 'faculty', 'jurusan', 'department']
                if any(keyword.lower() in text.lower() for keyword in faculty_keywords):
                    target = complete_metadata['faculty']
                    if ':' in text:
                        para.text = text.split(':')[0] + ': ' + target
                    else:
                        para.text = target
                    replacements += 1
                    continue

            # PROGRAM/STUDY PROGRAM DETECTION AND REPLACEMENT
            if complete_metadata.get('program') and complete_metadata['program'] != template_metadata.get('program'):
                program_keywords = ['program studi', 'prodi', 'study program', 'program pendidikan']
                if any(keyword.lower() in text.lower() for keyword in program_keywords):
                    target = complete_metadata['program']
                    if ':' in text:
                        para.text = text.split(':')[0] + ': ' + target
                    else:
                        para.text = target
                    replacements += 1
                    continue

            # SUPERVISOR/EXAMINER NAME DETECTION (More Robust)
            supervisor_val = complete_metadata.get('supervisor1')
            if supervisor_val and any(k in text.upper() for k in ['PEMBIMBING', 'SUPERVISOR', 'DOSEN']):
                sig_pattern = re.compile(r'(\(\s*[^)]+\s*\)|:\s*[^,\n]+)')
                match = sig_pattern.search(text)
                if match and any(p in text.upper() for p in ['( NAMA', ': NAMA', 'DR.', 'PEMBIMBING']):
                    if '(' in match.group(1):
                        para.text = text.replace(match.group(1), f"( {supervisor_val} )")
                    else:
                        para.text = text.replace(match.group(1), f": {supervisor_val}")
                    replacements += 1
                    continue


            # EXAMINER NAME DETECTION AND REPLACEMENT
            examiner_val = complete_metadata.get('examiner1')
            if examiner_val and any(k in text.upper() for k in ['PENGUJI', 'ANGGOTA', 'EXAMINER']):
                 sig_pattern = re.compile(r'(\(\s*[^)]+\s*\)|:\s*[^,\n]+)')
                 match = sig_pattern.search(text)
                 if match and any(p in text.upper() for p in ['( NAMA', ': NAMA', 'DR.', 'ANGGOTA', 'PENGUJI']):
                    if '(' in match.group(1):
                        para.text = text.replace(match.group(1), f"( {examiner_val} )")
                    else:
                        para.text = text.replace(match.group(1), f": {examiner_val}")
                    replacements += 1
                    continue


            # CITY DETECTION AND REPLACEMENT
            if complete_metadata.get('city'):
                # Look for city in date lines like "City, Date"
                if ',' in text and any(month in text.lower() for month in ['januari', 'februari', 'maret', 'april', 'mei', 'juni', 'juli', 'agustus', 'september', 'oktober', 'november', 'desember', 'january', 'february', 'march']):
                    city_part = text.split(',')[0].strip()
                    if city_part and city_part != complete_metadata['city']:
                        new_text = text.replace(city_part, complete_metadata['city'])
                        para.text = new_text
                        replacements += 1
                        print(f"[METADATA] Replaced city: '{original_text}' -> '{new_text}'")
                        continue

            # YEAR DETECTION AND REPLACEMENT
            if complete_metadata.get('year'):
                year_pattern = re.compile(r'\b(20\d{2})\b')
                if year_pattern.search(text):
                    current_year = year_pattern.search(text).group(1)
                    if current_year != complete_metadata['year']:
                        new_text = year_pattern.sub(complete_metadata['year'], text)
                        para.text = new_text
                        replacements += 1
                        print(f"[METADATA] Replaced year: '{original_text}' -> '{new_text}'")
                        continue

            # YEAR DETECTION AND REPLACEMENT
            if user_metadata.get('year'):
                # Look for 4-digit years (current year ± 5 years)
                current_year = 2024
                year_pattern = re.compile(r'\b(2019|2020|2021|2022|2023|2024|2025|2026|2027|2028|2029)\b')

                if year_pattern.search(text):
                    new_text = year_pattern.sub(user_metadata['year'], text)
                    if new_text != text:
                        para.text = new_text
                        replacements += 1
                        print(f"[METADATA] Replaced year: '{original_text}' -> '{new_text}'")
                        continue

        # Process TABLES for metadata replacement (CRITICAL FIX)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if not text: continue

                        # Use the same prefix-aware logic as paragraphs
                        if complete_metadata.get('title') and any(phrase in text.lower() for phrase in ['tuliskan judul', 'tulis judul']):
                             para.text = (text.split(':')[0] + ': ' if ':' in text else '') + complete_metadata['title']
                             replacements += 1
                        
                        if complete_metadata.get('author') and any(p in text.lower() for p in ['nama mahasiswa', 'nama penulis', 'n a m a']):
                             para.text = (text.split(':')[0] + ': ' if ':' in text else '') + complete_metadata['author']
                             replacements += 1

                        if complete_metadata.get('nim') and any(p in text.lower() for p in ['nim', 'nomor induk', 'n i m']):
                             nim_pattern = re.compile(r'\b(94523999|\d{8,10})\b')
                             if nim_pattern.search(text):
                                 para.text = nim_pattern.sub(complete_metadata['nim'], text)
                             else:
                                 para.text = (text.split(':')[0] + ': ' if ':' in text else '') + complete_metadata['nim']
                             replacements += 1

                        if complete_metadata.get('program') and any(p in text.lower() for p in ['program studi', 'prodi']):
                             para.text = (text.split(':')[0] + ': ' if ':' in text else '') + complete_metadata['program']
                             replacements += 1
                        
                        if complete_metadata.get('faculty') and 'fakultas' in text.lower():
                             para.text = (text.split(':')[0] + ': ' if ':' in text else '') + complete_metadata['faculty']
                             replacements += 1

        return replacements

    def _remove_instructional_text(self, doc, config) -> int:
        """Remove template instructional text (existing logic)."""
        replacements = 0
        paragraphs_to_remove = []

        # Student data for replacement (keep as fallback)
        author = config.get('author_placeholder', ['Nama Mahasiswa'])[0]
        nim = config.get('nim_placeholder', ['94523999'])[0]

        for para in doc.paragraphs:
            text = para.text
            text_strip = text.strip()
            text_upper = text_strip.upper()

            # PROTECT: Do not delete structural elements or headings
            if (text_upper.startswith(('BAB ', 'CHAPTER ', 'ANNEX ', 'DAFTAR ', 'KATA PENGANTAR', 'SARI', 'ABSTRAK')) or 
                'SUBBAB' in text_upper or 'ANAK SUBBAB' in text_upper or
                (para.style and 'Heading' in str(para.style.name))):
                continue

            # Remove instructional text from title/approval pages (language-specific)
            if config.get('language') == 'id':
                instructional_phrases = [
                    "The title and each student's name may differ",
                    "Adjust the layout to maintain neat formatting",
                    "Replace the date and names in this template",
                    "Tuliskan judul skripsi di sini",
                    "TULISKAN JUDUL BAB DI BARIS INI",
                    # Indonesian academic template instructions
                    "Format paragraf dengan style",
                    "Format paragraf dengan",
                    # "Subbab",  # REMOVED - needed for content insertion
                    # "Anak Subbab",  # REMOVED - needed for content insertion
                    "Cucu Subbab",
                    "Tuliskan isi bab",
                    "Isi bab di sini",
                    "Ketik isi bab",
                    "Tulis konten bab",
                    "BAB PENGANTAR",
                    "BAB PENUTUP",
                    "TINJAUAN PUSTAKA",
                    "METODOLOGI PENELITIAN",
                    "HASIL DAN PEMBAHASAN",
                    "KESIMPULAN DAN SARAN",
                    # Additional template formatting instructions
                    "Penyebutan dengan nomor",
                    "Penomoran level",
                    "Jika di dalam penyebutan",
                    "Notasi algoritmik",
                    "kode program atau pseudocode",
                    "Gunakan reference manager",
                    "Glosarium memuat",
                    "Compile proses",
                    "Debug langkah",
                    "Untuk membuat persamaan",
                    "Lampiran tidak perlu",
                    "Untuk mengacu ke gambar",
                    "Untuk mengacu ke tabel",
                    "Silakan copy paste",
                    "Cara copy paste persamaan",
                    "Contoh kode program yang dianggap",
                    "Contoh persamaan",
                    "Contoh tabel yang dibuat",
                    "Gunakan cara yang sama",
                    "Diharapkan dengan adanya Gambar",
                    "Sumber: Gunakan menu References",
                    "Untuk mengacu ke tabel",
                    # CRITICAL: Remove title subtitle instructions
                    "DENGAN POLA PIRAMIDA TERBALIK",
                    "LEBIH PANJANG DARI BARIS BAWAH)",
                    "PENGARAN KEBALIK",
                    "(BARIS ATAS",
                    "BARIS BAWAH)",
                    "HALAMAN JUDUL",
                    "HALAMAN PENGESAHAN",
                    "PERNYATAAN KEASLIAN",
                    "Pernyataan ini harus ditandatangani",
                    "TUGAS AKHIR",
                    "Bagian ini bebas untuk diisikan",
                    "Idealnya halaman persembahan",
                    "Idealnya halaman moto",
                    "SARI",
                    "Kata kunci:",
                    "DAFTAR ISI",
                    "DAFTAR TABEL",
                    "DAFTAR GAMBAR",
                    "GLOSARIUM",
                    "LAMPIRAN",
                    "pola piramida terbalik",
                ]
            else:
                instructional_phrases = [
                    "Replace this text with your thesis title",
                    "Replace with your name",
                    "Replace with your student ID"
                ]

            for phrase in instructional_phrases:
                if phrase.lower() in text.lower():
                    # CRITICAL: If it's a short instructional paragraph, remove it
                    if len(text) < 300: # Slightly larger window for instructions
                        # Store for deletion later to avoid iterator issues
                        paragraphs_to_remove.append(para)
                        replacements += 1
                        break

        # Remove the marked paragraphs in reverse order
        for para in reversed(paragraphs_to_remove):
            try:
                p = para._element
                if p is not None and p.getparent() is not None:
                    p.getparent().remove(p)
                para._p = para._element = None
            except:
                try:
                    para.text = "" # Fallback to clearing
                except:
                    pass

        # Process TABLES for instructional text removal
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()

                    # Remove instructional text from table cells
                    for phrase in instructional_phrases:
                        if phrase.lower() in cell_text.lower():
                            cell.text = ''
                            replacements += 1
                            break

        return replacements

    def _apply_heading_formatting(self, doc):
        """Apply proper heading formatting and numbering."""
        from docx.shared import Pt

        for para in doc.paragraphs:
            # Level 3 headings (subsection like 1.1, 2.1, etc.)
            if para.style and 'heading 3' in str(para.style).lower():
                # Ensure bold formatting
                for run in para.runs:
                    try:
                        run.font.bold = True
                        run.font.size = Pt(12)  # Slightly larger than body text
                    except:
                        pass

    def _fix_chapter_title_formatting(self, doc):
        """Fix chapter title formatting using Shift+Enter instead of separate paragraphs."""
        # This is complex to implement programmatically
        # For now, we'll ensure proper spacing and formatting

        chapter_pattern = None
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().upper()

            if text.startswith('BAB ') and any(c.isdigit() or c in 'IVX' for c in text.split()[1][:3] if len(text.split()) > 1):
                chapter_pattern = i

                # Look for the next paragraph which should be the chapter title
                if i + 1 < len(doc.paragraphs):
                    title_para = doc.paragraphs[i + 1]
                    title_text = title_para.text.strip()

                    # If it's a chapter title, ensure proper formatting
                    if title_text and len(title_text) > 3 and not title_text.startswith('BAB'):
                        # Combine with chapter header if they're separate paragraphs
                        # This simulates Shift+Enter behavior
                        try:
                            # Remove extra spacing between chapter and title
                            para.paragraph_format.space_after = Pt(0)
                            title_para.paragraph_format.space_before = Pt(0)
                        except:
                            pass

    def _to_roman(self, num):
        """Convert number to Roman numeral."""
        val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
        syms = ['M', 'CM', 'D', 'CD', 'C', 'XC', 'L', 'XL', 'X', 'IX', 'V', 'IV', 'I']
        roman_num = ''
        i = 0
        while num > 0:
            for _ in range(num // val[i]):
                roman_num += syms[i]
                num -= val[i]
            i += 1
        return roman_num

    def _populate_template_with_ai_content(self, doc, ai_sections, user_data):
        """Populate template placeholders with AI-analyzed content while preserving formatting."""
        replacements_made = 0

        print(f"[DEBUG] Starting template population with {len(ai_sections)} AI sections")

        # DEBUG: Log what AI sections contain
        for i, section in enumerate(ai_sections[:3]):  # First 3 sections
            content = section.get('content', [])
            print(f"[DEBUG] AI Section {i}: title='{section.get('title', 'no title')}', content_lines={len(content)}")
            if content:
                print(f"[DEBUG] Content preview: {content[0][:100]}..." if content[0] else "Empty content")

        # Replace user data placeholders first (title, author, etc.)
        replacements_made += self._replace_user_data_placeholders(doc, user_data)

        # Replace chapter content placeholders with AI sections
        replacements_made += self._replace_chapter_placeholders(doc, ai_sections)

        # Clean up any remaining generic placeholders
        replacements_made += self._clean_generic_placeholders(doc)

        print(f"[DEBUG] Made {replacements_made} intelligent template replacements")
        return replacements_made

    def _replace_user_data_placeholders(self, doc, user_data):
        """Replace user data placeholders like title, author, etc."""
        replacements = 0

        # Title placeholders
        title_placeholders = [
            r'TULIS JUDUL', r'JUDUL SKRIPSI', r'JUDUL TUGAS AKHIR',
            r'BAGIAN JUDUL', r'HALAMAN JUDUL'
        ]

        title = user_data.get('title', '')
        if title:
            for para in doc.paragraphs:
                text = para.text.upper()
                for placeholder in title_placeholders:
                    if placeholder.upper() in text:
                        para.text = title
                        replacements += 1
                        print(f"[DEBUG] Replaced title placeholder with: '{title}'")
                        break

        # Author placeholders
        author_placeholders = [r'NAMA PENULIS', r'NAMA MAHASISWA', r'AUTHOR']
        author = user_data.get('author', '')
        if author:
            for para in doc.paragraphs:
                text = para.text.upper()
                for placeholder in author_placeholders:
                    if placeholder in text:
                        para.text = author
                        replacements += 1
                        print(f"[DEBUG] Replaced author placeholder with: '{author}'")
                        break

        # NIM placeholders
        nim_placeholders = [r'NIM', r'NOMOR INDUK MAHASISWA']
        nim = user_data.get('nim', '')
        if nim:
            for para in doc.paragraphs:
                text = para.text.upper()
                for placeholder in nim_placeholders:
                    if placeholder in text:
                        para.text = nim
                        replacements += 1
                        print(f"[DEBUG] Replaced NIM placeholder with: '{nim}'")
                        break

        return replacements

    def _replace_chapter_placeholders(self, doc, ai_sections):
        """Replace chapter content placeholders with AI-analyzed sections."""
        replacements = 0

        # Map AI sections to chapter numbers - more flexible mapping
        chapter_mapping = {}
        for i, section in enumerate(ai_sections):
            title = section.get('title', '').upper()
            content = section.get('content', [])

            # Map based on position and keywords
            if i == 0 or 'BAB I' in title or 'PENDAHULUAN' in title or 'INTRODUCTION' in title:
                chapter_mapping['I'] = section
            elif i == 1 or 'BAB II' in title or 'TINJAUAN PUSTAKA' in title or 'LITERATURE' in title:
                chapter_mapping['II'] = section
            elif i == 2 or 'BAB III' in title or 'METODOLOGI' in title or 'METHODOLOGY' in title:
                chapter_mapping['III'] = section
            elif i == 3 or 'BAB IV' in title or 'HASIL' in title or 'ANALYSIS' in title or 'RESULTS' in title:
                chapter_mapping['IV'] = section
            elif i == 4 or 'BAB V' in title or 'PENUTUP' in title or 'CONCLUSION' in title:
                chapter_mapping['V'] = section
            elif i == 5 or 'BAB VI' in title:
                chapter_mapping['VI'] = section

        print(f"[DEBUG] Chapter mapping: {list(chapter_mapping.keys())}")

        # Replace content in ALL paragraphs that look like placeholders
        # This is the key fix - insert content instead of clearing
        content_inserted = 0
        placeholders_cleared = 0

        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            upper_text = text.upper()

            # Skip if paragraph already has substantial content (likely already processed)
            if len(text) > 200:  # Increased threshold
                continue

            # Check for placeholder patterns that should be replaced with content
            is_placeholder = (
                'SUBBAB' in upper_text or
                'ANAK SUBBAB' in upper_text or
                'CUCU SUBBAB' in upper_text or
                'TULISKAN' in upper_text or
                'ISI BAB' in upper_text or
                'CONTENT' in upper_text or
                'KONTEN' in upper_text or
                (len(text) < 50 and len(text) > 0)  # Short placeholder text
            )

            if is_placeholder:
                print(f"[DEBUG] Processing paragraph {para_idx}: '{text[:50]}...' (placeholder: {is_placeholder})")

                # Find which chapter this belongs to
                chapter_num = self._find_parent_chapter(doc, para)
                print(f"[DEBUG] Chapter context: {chapter_num}")

                if chapter_num and chapter_num in chapter_mapping:
                    ai_section = chapter_mapping[chapter_num]
                    content_lines = ai_section.get('content', [])
                    print(f"[DEBUG] AI section has {len(content_lines)} content lines")

                    # CRITICAL FIX: Insert actual content instead of clearing
                    if content_lines and len(content_lines) > 0:
                        # Join all content lines into coherent paragraphs
                        # But split into multiple paragraphs if very long
                        full_content = ' '.join([line.strip() for line in content_lines if line.strip()])

                        if len(full_content) > 1000:  # Split long content
                            # Split into paragraphs at sentence boundaries
                            sentences = full_content.replace('?', '.').replace('!', '.').split('.')
                            paragraphs = []
                            current_para = ""
                            for sentence in sentences:
                                sentence = sentence.strip()
                                if not sentence:
                                    continue
                                if len(current_para + sentence) > 500:
                                    if current_para:
                                        paragraphs.append(current_para + '.')
                                    current_para = sentence
                                else:
                                    current_para += ('. ' if current_para else '') + sentence

                            if current_para:
                                paragraphs.append(current_para + '.')

                            # Insert first paragraph into current paragraph
                            para.text = paragraphs[0]
                            content_inserted += 1

                            # Add additional paragraphs
                            current_para_obj = para
                            for additional_para in paragraphs[1:]:
                                if additional_para.strip():
                                    new_para = doc.add_paragraph(additional_para)
                                    # Copy style from original
                                    new_para.style = para.style
                                    if hasattr(para, 'paragraph_format'):
                                        new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
                                        new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
                                        new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
                                    current_para_obj = new_para
                                    content_inserted += 1

                            print(f"[DEBUG] INSERTED MULTI-PARAGRAPH content for chapter {chapter_num}: {len(full_content)} chars total")
                        elif full_content:
                            para.text = full_content
                            content_inserted += 1
                            print(f"[DEBUG] INSERTED content for chapter {chapter_num}: {len(full_content)} chars")
                            print(f"[DEBUG] Content preview: {full_content[:100]}...")
                        else:
                            para.text = ""  # Clear if no content
                            placeholders_cleared += 1
                    else:
                        para.text = ""  # Clear if no content available
                        placeholders_cleared += 1
                        print(f"[DEBUG] No content available for chapter {chapter_num}")
                else:
                    # If we can't map to a chapter, clear the placeholder
                    para.text = ""
                    placeholders_cleared += 1
                    print(f"[DEBUG] Cleared unmapped placeholder: '{text[:50]}...'")
            else:
                print(f"[DEBUG] Skipping paragraph {para_idx}: not a placeholder")

        print(f"[DEBUG] Content insertion summary: {content_inserted} paragraphs filled, {placeholders_cleared} placeholders cleared")
        replacements += content_inserted

        print(f"[DEBUG] Chapter content replacement completed: {replacements} insertions")
        return replacements

    def _find_parent_chapter(self, doc, target_para):
        """Find the chapter number for a given paragraph by looking backwards."""
        para_index = None
        for i, para in enumerate(doc.paragraphs):
            if para == target_para:
                para_index = i
                break

        if para_index is None:
            return None

        # Look backwards for chapter header
        for i in range(para_index - 1, max(-1, para_index - 10), -1):
            para = doc.paragraphs[i]
            text = para.text.strip().upper()
            if text.startswith('BAB '):
                parts = text.split()
                if len(parts) >= 2:
                    chapter_marker = parts[1].strip('.:')
                    if chapter_marker in ['I', 'II', 'III', 'IV', 'V', 'VI']:
                        return chapter_marker

        return None

    def _clean_generic_placeholders(self, doc):
        """Clean up any remaining generic placeholders."""
        replacements = 0

        placeholders_to_clear = [
            r'TULISKAN', r'ISI BAB', r'SUBBAB', r'ANAK SUBBAB', r'CUCU SUBBAB',
            r'KONTEN', r'CONTENT', r'PLACEHOLDER'
        ]

        for para in doc.paragraphs:
            text = para.text.upper()
            for placeholder in placeholders_to_clear:
                if placeholder in text and len(text) < 100:  # Only clear short placeholder text
                    para.text = ""
                    replacements += 1
                    print(f"[DEBUG] Cleared generic placeholder: '{text[:50]}...'")
                    break

        return replacements

    def _populate_template_structured(self, doc, ai_sections, user_data):
        """Fallback method: Populate template in a structured way if intelligent replacement fails."""
        print("[DEBUG] Using structured population approach")

        # Find main content area (after front matter)
        main_content_start = self._find_main_content_start(doc)

        # Insert AI content in structured sections
        insert_position = main_content_start
        for idx, section in enumerate(ai_sections, 1):
            if idx > 1:
                # Add page break between chapters
                if insert_position < len(doc.paragraphs):
                    doc.paragraphs[insert_position].insert_page_break_before()

            # Add chapter title
            title = section.get('title', f'Chapter {idx}')
            if insert_position < len(doc.paragraphs):
                doc.paragraphs[insert_position].text = title
                insert_position += 1
            else:
                new_para = doc.add_paragraph(title)
                new_para.style = "Heading 1"
                insert_position = len(doc.paragraphs)

            # Add chapter content
            content = section.get('content', [])
            for line in content:
                if isinstance(line, str) and line.strip():
                    if insert_position < len(doc.paragraphs):
                        doc.paragraphs[insert_position].text = line
                        insert_position += 1
                    else:
                        new_para = doc.add_paragraph(line)
                        new_para.style = "Normal"
                        insert_position = len(doc.paragraphs)

    def _find_main_content_start(self, doc):
        """Find where main content should start in the template."""
        # Look for the first chapter or main content area
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().upper()
            if 'BAB I' in text or 'BAB 1' in text or 'CHAPTER 1' in text:
                return i
            # Look for table of contents end
            if 'DAFTAR ISI' in text and i < len(doc.paragraphs) - 1:
                # Skip a few paragraphs after TOC
                return min(i + 5, len(doc.paragraphs))

        # Default to middle of document
        return len(doc.paragraphs) // 2

    def _add_ai_main_content(self, doc: Document, user_data: Dict[str, Any], ai_content: str, structure_mapping: Dict[str, Any]) -> None:
        """Add main content using AI-enhanced formatting."""
        print(f"[DEBUG] Processing AI content with structure mapping: {list(structure_mapping.keys())}")

        # If we have structured AI content, use it
        if ai_content and structure_mapping:
            # Parse AI-enhanced content into sections
            sections = self._parse_ai_content_sections(ai_content)
            print(f"[DEBUG] Parsed {len(sections)} sections from AI content")

            for idx, section in enumerate(sections, 1):
                # Add page break before chapter (except first)
                if idx > 1:
                    doc.add_page_break()

                # Add chapter heading
                title = section.get("title", f"Chapter {idx}")
                heading_para = doc.add_paragraph(title)
                heading_para.style = "Heading 1"
                heading_run = heading_para.runs[0]
                heading_run.font.bold = True
                heading_run.font.size = Pt(14)

                # Add section content with proper formatting
                content = section.get("content", [])
                if isinstance(content, list):
                    for line in content:
                        if isinstance(line, str) and line.strip():
                            para = doc.add_paragraph(line)
                            para.style = "Normal"
                            para.paragraph_format.left_indent = Inches(0.0)
                            para.paragraph_format.first_line_indent = Inches(1.0)
                            para.paragraph_format.line_spacing = 1.5
                            for run in para.runs:
                                run.font.size = Pt(11)
                                run.font.name = "Times New Roman"
        else:
            # Fallback to standard content processing
            print("[DEBUG] No AI structure mapping, falling back to standard content")
            try:
                normalized = extract_normalized_structure(str(self.content_path))
                self._add_main_content(doc, user_data, normalized)
            except Exception:
                self._add_main_content(doc, user_data, None)

    def _parse_ai_content_sections(self, ai_content: str) -> List[Dict[str, Any]]:
        """Parse AI-enhanced content into structured sections."""
        sections = []
        lines = ai_content.split('\n')
        current_section = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Detect section headers (enhanced patterns)
            if self._is_ai_section_header(line):
                if current_section:
                    sections.append(current_section)

                current_section = {
                    'title': line,
                    'content': [],
                    'type': 'chapter'
                }
            elif current_section:
                current_section['content'].append(line)

        if current_section:
            sections.append(current_section)

        return sections

    def _is_ai_section_header(self, line: str) -> bool:
        """Check if line is a section header in AI-enhanced content."""
        # Enhanced patterns for AI-generated content
        patterns = [
            r'^BAB\s+[IVX]+\b',  # BAB I, BAB II, etc.
            r'^BAB\s+\d+\b',     # BAB 1, BAB 2, etc.
            r'^Chapter\s+\d+',   # Chapter 1, etc.
            r'^\d+\.\s+',        # 1. , 2. , etc.
            r'^[A-Z][^.!?]*:$', # SECTION NAME:
        ]

        return any(re.match(pattern, line, re.IGNORECASE) for pattern in patterns)

    def _apply_intelligent_content_replacement(self, doc: Document, ai_content: str, structure_mapping: Dict[str, Any], user_data: Dict[str, Any]) -> Document:
        """Apply intelligent content replacement to template document."""
        print(f"[DEBUG] Applying intelligent content replacement to {len(doc.paragraphs)} paragraphs")

        # Parse AI content into structured sections
        ai_sections = self._parse_ai_content_sections(ai_content)
        print(f"[DEBUG] AI content parsed into {len(ai_sections)} sections")

        # Strategy: Replace template placeholders with real content
        replacements_made = 0

        # Look for common Indonesian thesis chapter patterns in the template
        chapter_patterns = [
            r'^BAB\s+I\b.*',  # BAB I - Pendahuluan
            r'^BAB\s+II\b.*', # BAB II - Tinjauan Pustaka
            r'^BAB\s+III\b.*', # BAB III - Metodologi
            r'^BAB\s+IV\b.*', # BAB IV - Hasil dan Pembahasan
            r'^BAB\s+V\b.*',  # BAB V - Penutup
            r'^BAB\s+VI\b.*', # BAB VI (sometimes included)
        ]

        # Simple and safe chapter title replacement
        print(f"[DEBUG] Starting chapter replacement with {len(ai_sections)} AI sections")
        for ai_section in ai_sections[:3]:  # Debug first few
            print(f"[DEBUG] AI section: {ai_section.get('title', 'no title')}")

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            # Check if this paragraph matches a chapter pattern
            for pattern in chapter_patterns:
                if re.match(pattern, text.upper()):
                    chapter_num = text.split()[1] if len(text.split()) > 1 else ""
                    print(f"[DEBUG] Found chapter {chapter_num} at paragraph {i}")

                    # Find corresponding AI content - try simpler matching
                    ai_content_section = None
                    for ai_section in ai_sections:
                        section_title = ai_section.get('title', '').upper()
                        # More flexible matching
                        if chapter_num in section_title or f'BAB {chapter_num}' in section_title:
                            ai_content_section = ai_section
                            break

                    if ai_content_section:
                        # Simply replace the chapter title
                        new_title = ai_content_section.get('title', text.split('\n')[0])
                        print(f"[DEBUG] Setting paragraph {i} text to: '{new_title}'")
                        try:
                            para.text = new_title
                            replacements_made += 1
                            print(f"[DEBUG] Replaced chapter {chapter_num} title successfully")
                        except Exception as e:
                            print(f"[ERROR] Failed to set paragraph text: {e}")
                    else:
                        print(f"[DEBUG] No matching AI content for chapter {chapter_num}")
                    break

        print(f"[DEBUG] Chapter replacement completed. Document state: {type(doc)}")
        print(f"[DEBUG] Document has {len(doc.paragraphs) if doc else 0} paragraphs")

        # Skip all placeholder cleaning to avoid document corruption
        print(f"[DEBUG] Skipping all placeholder cleaning")

    def _apply_simple_chapter_titles(self, doc: Document) -> None:
        """Apply standard chapter titles by directly replacing template placeholders."""
        # Standard chapter titles
        chapter_titles = {
            'I': 'BAB I: PENDAHULUAN',
            'II': 'BAB II: TINJAUAN PUSTAKA',
            'III': 'BAB III: METODOLOGI PENELITIAN',
            'IV': 'BAB IV: HASIL DAN PEMBAHASAN',
            'V': 'BAB V: PENUTUP',
            'VI': 'BAB VI: KESIMPULAN DAN SARAN'
        }

        replacements_made = 0

        # Find and replace each chapter by looking for the specific template pattern
        for para in doc.paragraphs:
            text = para.text.strip()

            # Look for the exact template pattern: "BAB X\nTULISKAN JUDUL BAB DI BARIS INI"
            if '\n' in text and 'TULISKAN JUDUL BAB' in text.upper():
                lines = text.split('\n')
                first_line = lines[0].strip()

                # Extract chapter number
                if first_line.startswith('BAB '):
                    parts = first_line.split()
                    if len(parts) >= 2:
                        chapter_marker = parts[1].strip('.:')
                        if chapter_marker in chapter_titles:
                            # Replace with proper title
                            new_title = chapter_titles[chapter_marker]
                            para.text = new_title
                            replacements_made += 1
                            print(f"[DEBUG] Replaced template chapter {chapter_marker} with: '{new_title}'")

        print(f"[DEBUG] Made {replacements_made} template chapter replacements")

        # Also replace any remaining single-line chapter headers
        for para in doc.paragraphs:
            text = para.text.strip()
            if text.startswith('BAB ') and 'TULISKAN' in text.upper():
                parts = text.split()
                if len(parts) >= 2:
                    chapter_marker = parts[1].strip('.:')
                    if chapter_marker in chapter_titles:
                        new_title = chapter_titles[chapter_marker]
                        para.text = new_title
                        replacements_made += 1
                        print(f"[DEBUG] Replaced single-line chapter {chapter_marker} with: '{new_title}'")

        print(f"[DEBUG] Total chapter replacements: {replacements_made}")

    def _apply_safe_chapter_titles(self, doc: Document, ai_content: str, structure_mapping: Dict[str, Any]) -> None:
        """Safely replace chapter titles without risking document corruption."""
        # Parse AI content sections
        ai_sections = self._parse_ai_content_sections(ai_content)

        print(f"[DEBUG] AI content parsed into {len(ai_sections)} sections")
        for i, section in enumerate(ai_sections[:5]):  # Show first 5 for debugging
            print(f"[DEBUG] AI section {i+1}: '{section.get('title', 'no title')}'")

        # Standard chapter titles to use when AI content doesn't match
        standard_titles = {
            'I': 'BAB I: PENDAHULUAN',
            'II': 'BAB II: TINJAUAN PUSTAKA',
            'III': 'BAB III: METODOLOGI PENELITIAN',
            'IV': 'BAB IV: HASIL DAN PEMBAHASAN',
            'V': 'BAB V: PENUTUP'
        }

        # Simple chapter title mapping
        chapter_patterns = [
            r'^BAB\s+I\b.*',
            r'^BAB\s+II\b.*',
            r'^BAB\s+III\b.*',
            r'^BAB\s+IV\b.*',
            r'^BAB\s+V\b.*',
        ]

        replacements_made = 0

        # Apply title replacements safely
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            # Check if this paragraph matches a chapter pattern
            for pattern_idx, pattern in enumerate(chapter_patterns):
                if re.match(pattern, text.upper()):
                    chapter_num = ['I', 'II', 'III', 'IV', 'V'][pattern_idx]
                    print(f"[DEBUG] Found chapter {chapter_num} at paragraph {i}")

                    # Try to find corresponding AI content first
                    ai_content_section = None
                    for ai_section in ai_sections:
                        section_title = ai_section.get('title', '').upper()
                        # More flexible matching
                        if chapter_num in section_title or f'BAB {chapter_num}' in section_title:
                            ai_content_section = ai_section
                            break

                    if ai_content_section:
                        # Use AI-generated title
                        new_title = ai_content_section.get('title', text.split('\n')[0])
                        print(f"[DEBUG] Using AI title: '{new_title}'")
                    else:
                        # Use standard title
                        new_title = standard_titles.get(chapter_num, text.split('\n')[0])
                        print(f"[DEBUG] Using standard title: '{new_title}'")

                    try:
                        para.text = new_title
                        replacements_made += 1
                        print(f"[DEBUG] Replaced chapter {chapter_num} title successfully")
                    except Exception as e:
                        print(f"[ERROR] Failed to set paragraph text: {e}")
                    break  # Break out of pattern matching loop

        print(f"[DEBUG] Total chapter replacements made: {replacements_made}")

    def _clean_title_placeholders_only(self, doc: Document, user_data: Dict[str, Any]) -> None:
        """Safely clean only title placeholders without corrupting the document."""
        if not user_data.get('title'):
            return

        # Only replace the main title placeholders
        title_placeholders = [
            r'.*BAGIAN INI ADALAH BAGIAN JUDUL.*TULIS JUDUL.*',
            r'.*HALAMAN JUDUL.*'
        ]

        for para in doc.paragraphs:
            text = para.text.strip()
            for pattern in title_placeholders:
                if re.match(pattern, text.upper(), re.IGNORECASE | re.DOTALL):
                    para.text = user_data['title']
                    break

    def _clean_remaining_placeholders(self, doc: Document, user_data: Dict[str, Any], ai_sections: List[Dict[str, Any]], replacements_made: int) -> None:
        """Clean up any remaining placeholders in the document."""

        # Title placeholders
        title_placeholders = [
            r'.*TULIS.*JUDUL.*',
            r'.*BAGIAN.*JUDUL.*',
            r'.*HALAMAN.*JUDUL.*',
        ]

        for para in doc.paragraphs:
            text = para.text.strip()
            for pattern in title_placeholders:
                if re.match(pattern, text.upper(), re.IGNORECASE):
                    if user_data.get('title'):
                        para.text = user_data['title']
                        replacements_made += 1
                        print(f"[DEBUG] Replaced title placeholder with: '{user_data['title']}'")
                    break

        # Chapter content placeholders
        chapter_placeholders = [
            r'.*TULISKAN.*BAB.*',
            r'.*ISI.*BAB.*',
            r'.*SUBBAB.*',
            r'.*ANAK SUBBAB.*',
            r'.*CUCU SUBBAB.*',
        ]

        for para in doc.paragraphs:
            text = para.text.strip()
            for pattern in chapter_placeholders:
                if re.match(pattern, text.upper(), re.IGNORECASE):
                    # Clear placeholder content
                    para.text = ""
                    replacements_made += 1
                    print(f"[DEBUG] Cleared placeholder content: '{text[:50]}...'")
                    break

        # Front matter placeholders
        front_matter_placeholders = [
            r'.*TUGAS AKHIR.*JUDUL.*',
            r'.*BAGIAN.*BEBAS.*',
            r'.*GAMBAR.*JUDUL.*',
        ]

        for para in doc.paragraphs:
            text = para.text.strip()
            for pattern in front_matter_placeholders:
                if re.match(pattern, text.upper(), re.IGNORECASE):
                    # Replace with appropriate content or clear
                    if 'JUDUL' in text.upper() and user_data.get('title'):
                        para.text = user_data['title']
                    else:
                        para.text = ""  # Clear generic placeholders
                    replacements_made += 1
                    print(f"[DEBUG] Cleaned front matter placeholder")
                    break

        print(f"[DEBUG] Made {replacements_made} content replacements")

        # If still no meaningful replacements, add content in main sections
        if replacements_made < 3:
            print("[DEBUG] Insufficient replacements, adding content to main body")
            self._add_content_to_main_body(doc, ai_sections, user_data)

        return doc

    def _add_content_to_main_body(self, doc: Document, ai_sections: List[Dict[str, Any]], user_data: Dict[str, Any]) -> None:
        """Add content to the main body of the document when intelligent replacement fails."""
        print("[DEBUG] Adding content to main body sections")

        # Find the main content area (after front matter, before back matter)
        main_content_start = None

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().upper()
            # Look for the start of main content (BAB I or similar)
            if 'BAB I' in text or 'BAB 1' in text or 'CHAPTER 1' in text:
                main_content_start = i
                break

        if main_content_start is None:
            # If no chapter marker found, add after front matter
            main_content_start = 50  # Rough estimate for front matter length

        # Insert content sections
        insert_position = main_content_start
        for section in ai_sections:
            # Insert chapter title
            title = section.get('title', 'Chapter')
            doc.paragraphs[insert_position].text = title

            # Insert content
            content_lines = section.get('content', [])
            for line in content_lines:
                if isinstance(line, str) and line.strip():
                    new_para = doc.add_paragraph(line)
                    new_para.style = "Normal"

            insert_position += 1

    def _replace_chapter_content(self, doc: Document, chapter_start_idx: int, ai_section: Dict[str, Any]) -> None:
        """Replace the content area of a chapter with clean AI content."""
        content_lines = ai_section.get('content', [])
        if not content_lines:
            return

        # Find the end of this chapter (next chapter or major section break)
        chapter_end_idx = self._find_chapter_end(doc, chapter_start_idx)

        print(f"[DEBUG] Chapter content area: paragraphs {chapter_start_idx + 1} to {chapter_end_idx}")

        # Clear existing content in this range (keep the chapter title)
        insert_idx = chapter_start_idx + 1
        content_idx = 0

        # Replace existing paragraphs with new content
        while insert_idx < chapter_end_idx and content_idx < len(content_lines):
            line = content_lines[content_idx].strip()
            if line and len(line) > 10:  # Only substantial content
                para = doc.paragraphs[insert_idx]
                para.text = line

                # Apply consistent formatting
                para.style = "Normal"
                para.paragraph_format.left_indent = Inches(0.0)
                para.paragraph_format.first_line_indent = Inches(1.0)
                para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    run.font.size = Pt(11)
                    run.font.name = "Times New Roman"

                content_idx += 1
                insert_idx += 1
            else:
                content_idx += 1

        # If we have more content, add new paragraphs
        while content_idx < len(content_lines):
            line = content_lines[content_idx].strip()
            if line and len(line) > 10:
                new_para = doc.add_paragraph(line)
                new_para.style = "Normal"
                new_para.paragraph_format.left_indent = Inches(0.0)
                new_para.paragraph_format.first_line_indent = Inches(1.0)
                new_para.paragraph_format.line_spacing = 1.5
                for run in new_para.runs:
                    run.font.size = Pt(11)
                    run.font.name = "Times New Roman"
            content_idx += 1

    def _find_chapter_end(self, doc: Document, chapter_start_idx: int) -> int:
        """Find the end of a chapter section."""
        # Look for the next chapter or major section break
        for i in range(chapter_start_idx + 1, len(doc.paragraphs)):
            text = doc.paragraphs[i].text.strip().upper()

            # Next chapter starts
            if text.startswith('BAB ') and any(c.isdigit() or c in 'IVX' for c in text.split()[1][:3] if len(text.split()) > 1):
                return i

            # Major section break (table of contents, references, etc.)
            if any(keyword in text for keyword in ['DAFTAR PUSTAKA', 'LAMPIRAN', 'TABEL', 'GAMBAR']):
                return i

        return len(doc.paragraphs)  # End of document

    def _replace_template_section(self, doc: Document, section_name: str, ai_section: Dict[str, Any]) -> int:
        """Replace a template section with AI content."""
        replacements = 0

        # Find paragraphs containing the section name
        for para in doc.paragraphs:
            if section_name.lower() in para.text.lower():
                print(f"[DEBUG] Found section '{section_name}' in paragraph")
                # Replace the content
                para.text = ai_section.get('title', section_name)

                # Add the section content after this paragraph
                content = ai_section.get('content', [])
                current_para = para
                for line in content:
                    if isinstance(line, str) and line.strip():
                        new_para = doc.add_paragraph(line)
                        new_para.style = "Normal"
                        # Copy formatting from nearby paragraphs
                        if hasattr(current_para, 'paragraph_format'):
                            new_para.paragraph_format.left_indent = current_para.paragraph_format.left_indent
                            new_para.paragraph_format.first_line_indent = current_para.paragraph_format.first_line_indent
                            new_para.paragraph_format.line_spacing = current_para.paragraph_format.line_spacing
                        for run in new_para.runs:
                            run.font.size = Pt(11)
                            run.font.name = "Times New Roman"
                        current_para = new_para
                replacements += 1
                break  # Only replace first occurrence

        return replacements

    def _append_to_template_section(self, doc: Document, section_name: str, ai_section: Dict[str, Any]) -> int:
        """Append AI content to an existing template section."""
        replacements = 0

        # Find paragraphs containing the section name
        for para in doc.paragraphs:
            if section_name.lower() in para.text.lower():
                print(f"[DEBUG] Found section '{section_name}' to append to")
                # Add content after this paragraph
                content = ai_section.get('content', [])
                current_para = para
                for line in content:
                    if isinstance(line, str) and line.strip():
                        new_para = doc.add_paragraph(line)
                        new_para.style = "Normal"
                        # Copy formatting
                        if hasattr(current_para, 'paragraph_format'):
                            new_para.paragraph_format.left_indent = current_para.paragraph_format.left_indent
                            new_para.paragraph_format.first_line_indent = current_para.paragraph_format.first_line_indent
                            new_para.paragraph_format.line_spacing = current_para.paragraph_format.line_spacing
                        for run in new_para.runs:
                            run.font.size = Pt(11)
                            run.font.name = "Times New Roman"
                        current_para = new_para
                replacements += 1
                break  # Only append to first occurrence

        return replacements

    def _add_front_matter(self, doc: Document, user_data: Dict[str, Any]) -> None:
        """Add all front matter sections."""
        generator = FrontMatterGenerator(user_data)
        
        # Title Page
        generator.create_title_page(doc)
        
        # Approval Pages
        generator.create_approval_page(doc, approval_type="supervisor")
        generator.create_approval_page(doc, approval_type="examiner")
        
        # Originality Statement
        generator.create_originality_statement(doc)
        
        # Optional: Dedication
        if user_data and isinstance(user_data, dict) and user_data.get("dedication"):
            generator.create_dedication_page(doc, user_data["dedication"])
        
        # Optional: Motto
        if user_data and isinstance(user_data, dict) and user_data.get("motto"):
            generator.create_motto_page(doc, user_data["motto"])
        
        # Preface
        generator.create_preface(doc)
        
        # Abstracts
        generator.create_abstract(doc, language="id")
        if user_data and isinstance(user_data, dict) and user_data.get("abstract_en"):
            generator.create_abstract(doc, language="en")
        
        # Table of Contents
        generator.create_table_of_contents(doc)
        
        # Glossary (if provided)
        if user_data and isinstance(user_data, dict) and user_data.get("glossary"):
            generator.create_glossary(doc, user_data["glossary"])
    
    def _add_main_content(self, doc: Document, user_data: Dict[str, Any], normalized: Dict[str, Any] | None = None) -> None:
        """Add main content (chapters/sections from extractor or normalized)."""
        if normalized and isinstance(normalized, dict) and normalized.get("chapters"):
            sections = []
            for ch in normalized.get("chapters", []):
                sec = {
                    "title": ch.get("title") or (f"BAB {ch.get('number', '')}" if ch.get('number') else "Chapter"),
                    "content": list(ch.get("content", [])),
                    "lists": list(ch.get("lists", [])),
                }
                for s in ch.get("sections", []):
                    sec["content"].extend(s.get("content", []))
                    for lst in s.get("lists", []):
                        sec.setdefault("lists", []).append(lst)
                sections.append(sec)
        else:
            sections = self.extractor.get_sections()
        
        for idx, section in enumerate(sections, 1):
            # Skip non-dict sections
            if not isinstance(section, dict):
                continue
            
            # Add page break before chapter (except first)
            if idx > 1:
                doc.add_page_break()
            
            # Add chapter heading
            title = section.get("title", f"Chapter {idx}")
            heading_para = doc.add_paragraph(title)
            heading_para.style = "Heading 1"
            heading_run = heading_para.runs[0]
            heading_run.font.bold = True
            heading_run.font.size = Pt(14)
            
            # Add section content
            content = section.get("content", [])
            if isinstance(content, list):
                for line in content:
                    if isinstance(line, str) and line.strip():
                        para = doc.add_paragraph(line)
                        para.style = "Normal"
                        # Apply paragraph formatting
                        para.paragraph_format.left_indent = Inches(0.0)
                        para.paragraph_format.first_line_indent = Inches(1.0)
                        para.paragraph_format.line_spacing = 1.5
                        for run in para.runs:
                            run.font.size = Pt(11)
                            run.font.name = "Times New Roman"
            else:
                # Content is a single string
                para = doc.add_paragraph(str(content))
                para.style = "Normal"
                para.paragraph_format.left_indent = Inches(0.0)
                para.paragraph_format.first_line_indent = Inches(1.0)
                para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    run.font.size = Pt(11)
                    run.font.name = "Times New Roman"

            # Add lists if present
            for lst in section.get("lists", []):
                style_name = "List Number" if lst.get("ordered") else "List Bullet"
                for item in lst.get("items", []):
                    para = doc.add_paragraph(item)
                    para.style = style_name
    
    def _add_back_matter(self, doc: Document, user_data: Dict[str, Any]) -> None:
        """Add all back matter sections."""
        generator = BackMatterGenerator(user_data)
        
        # List of Tables
        generator.create_list_of_tables(doc)
        
        # List of Figures
        generator.create_list_of_figures(doc)
        
        # References/Bibliography
        references = user_data.get("references", []) if user_data and isinstance(user_data, dict) else []
        try:
            generator.create_references(doc, references)
        except Exception as e:
            print(f"[WARNING] Failed to create references: {e}")
            import traceback
            traceback.print_exc()
        
        # Appendices
        appendices = user_data.get("appendices", []) if user_data and isinstance(user_data, dict) else []
        if appendices:
            try:
                # Ensure appendices is a list
                if not isinstance(appendices, list):
                    appendices = [appendices]
                generator.create_appendices(doc, appendices)
            except Exception as e:
                print(f"[WARNING] Failed to create appendices: {e}")
                import traceback
                traceback.print_exc()
    
    def get_analysis_report(self) -> Dict[str, Any]:
        """Get comprehensive analysis report of template, content, and semantic structure."""
        report = {
            "template_analysis": self.analyzer.get_summary(),
            "content_summary": self.extractor.get_summary(),
            "mapping_summary": self.mapper.get_summary(),
            "sections_found": len(self.extractor.get_sections()),
            "styles_detected": len(self.analyzer.analysis.get("styles", {})),
        }
        
        # Add AI semantic analysis if available
        if self.use_ai:
            sections = self.ai_extractor.get_sections()
            semantic_types = []
            for s in sections:
                if isinstance(s, dict):
                    semantic_types.append(s.get("semantic_type", "unknown"))
            
            report["ai_semantic_analysis"] = {
                "extraction_summary": self.ai_extractor.get_summary(),
                "semantic_validation": self.semantic_validation,
                "semantic_types_found": list(set(semantic_types)),
            }
        
        return report

    def _build_standard(self, user_data: Dict[str, Any]) -> Path:
        """Build thesis using standard processing (fallback method)."""
        print(f"[DEBUG] Starting standard build with user_data keys: {list(user_data.keys())}")
        print(f"[DEBUG] Template path: {self.template_path}, exists: {self.template_path.exists()}")
        print(f"[DEBUG] Content path: {self.content_path}, exists: {self.content_path.exists()}")

        # Load template document and extract styles, then create new document
        try:
            template_doc = Document(str(self.template_path))
            print(f"[DEBUG] Template loaded successfully, {len(template_doc.paragraphs)} paragraphs, {len(template_doc.styles)} styles")

            # Create new document (don't keep template content)
            doc = Document()
            print(f"[DEBUG] New document created with {len(doc.styles)} default styles")

            # Copy styles from template to new document
            self._copy_styles_from_template(doc, template_doc)
            print(f"[DEBUG] After style copying: {len(doc.styles)} styles in new document")

        except Exception as e:
            print(f"[WARNING] Failed to load template, creating new document: {e}")
            doc = Document()

        # Build sections in order
        if self.include_frontmatter:
            print("[DEBUG] Adding front matter...")
            self._add_front_matter(doc, user_data)
            print(f"[DEBUG] Front matter added, document now has {len(doc.paragraphs)} paragraphs")

        # Prefer normalized extractor where available
        try:
            print("[DEBUG] Attempting normalized extraction...")
            normalized = extract_normalized_structure(str(self.content_path))
            print(f"[DEBUG] Normalized extraction successful: {normalized is not None}")
            if normalized:
                print(f"[DEBUG] Normalized structure keys: {list(normalized.keys())}")
                if 'chapters' in normalized:
                    print(f"[DEBUG] Found {len(normalized['chapters'])} chapters")
        except Exception as e:
            print(f"[DEBUG] Normalized extraction failed: {e}")
            normalized = None

        print("[DEBUG] Adding main content...")
        self._add_main_content(doc, user_data, normalized)
        print(f"[DEBUG] Main content added, document now has {len(doc.paragraphs)} paragraphs")

        print("[DEBUG] Adding back matter...")
        self._add_back_matter(doc, user_data)
        print(f"[DEBUG] Back matter added, document now has {len(doc.paragraphs)} paragraphs")

        # Save
        print(f"[DEBUG] Saving document to: {self.output_path}")
        doc.save(str(self.output_path))

        final_size = self.output_path.stat().st_size if self.output_path.exists() else 0
        print(f"[DEBUG] Document saved successfully, size: {final_size} bytes")

        return self.output_path

    def _to_roman(self, num):
        """Convert number to Roman numeral."""
        val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
        syms = ['M', 'CM', 'D', 'CD', 'C', 'XC', 'L', 'XL', 'X', 'IX', 'V', 'IV', 'I']
        roman_num = ''
        i = 0
        while num > 0:
            for _ in range(num // val[i]):
                roman_num += syms[i]
                num -= val[i]
            i += 1
        return roman_num

    def _insert_abstract(self, doc, analyzed_data, config):
        """Insert abstract and keywords with university-specific configuration."""
        abstract_data = analyzed_data.get('abstract', {})
        if not abstract_data:
            print("[WARNING] No abstract data found in analyzed_data")
            return 0

        insertions = 0
        print(f"[INFO] Starting abstract insertion with keys: {list(abstract_data.keys())}")

        # Use university-specific abstract keywords
        indonesian_patterns = config.get('abstract_keywords', ['ABSTRAK', 'SARI'])
        english_patterns = config.get('english_abstract_keywords', ['ABSTRACT'])

        # Insert Indonesian Abstract
        indonesian_inserted = self._insert_single_abstract(
            doc, abstract_data, 'indonesian', indonesian_patterns, 'Kata kunci'
        )
        insertions += indonesian_inserted

        # Insert English Abstract
        english_inserted = self._insert_single_abstract(
            doc, abstract_data, 'english', english_patterns, 'Keywords'
        )
        insertions += english_inserted

        print(f"[INFO] Abstract insertion complete: {insertions} sections inserted")
        return insertions

    def _insert_single_abstract(self, doc, abstract_data, lang_key, section_patterns, keyword_label):
        """Insert a single abstract (Indonesian or English) with keywords."""
        insertions = 0
        abstract_text = abstract_data.get(lang_key, '')
        # Fix: Use correct keywords key
        keywords_key = 'keywords_id' if lang_key == 'indonesian' else 'keywords_en'
        keywords = abstract_data.get(keywords_key, [])

        if not abstract_text:
            print(f"[WARNING] No {lang_key} abstract content found")
            return 0

        print(f"[INFO] Looking for {lang_key} abstract section with patterns: {section_patterns}")

        # Find the abstract section
        for para_index, para in enumerate(doc.paragraphs):
            para_text_upper = para.text.upper()
            found_section = False

            # Check if this paragraph matches any of the section patterns
            for pattern in section_patterns:
                if pattern in para_text_upper:
                    found_section = True
                    print(f"[INFO] Found {lang_key} abstract section: '{para.text.strip()}'")
                    break

            if found_section:
                content_inserted = False
                keywords_inserted = False

                # Look for content insertion point (next few paragraphs)
                for i in range(para_index + 1, min(para_index + 15, len(doc.paragraphs))):
                    content_para = doc.paragraphs[i]

                    # Skip if this is another section header
                    if any(header in content_para.text.upper() for header in
                           ['BAB ', 'CHAPTER', 'DAFTAR', 'LAMPIRAN', 'ABSTRAK', 'ABSTRACT']):
                        continue

                    # Check if this is a suitable content paragraph
                    if (content_para.style and 'isi paragraf' in str(content_para.style).lower()) or \
                       'Format paragraf' in content_para.text or \
                       len(content_para.text.strip()) < 200:  # Likely placeholder or short content

                        # Insert abstract content
                        old_content = content_para.text[:50] if content_para.text else "[empty]"
                        content_para.clear()
                        content_para.add_run(abstract_text)

                        # Preserve or set style
                        try:
                            if hasattr(content_para, 'style') and content_para.style:
                                # Keep existing style
                                pass
                            else:
                                content_para.style = 'isi paragraf'
                        except:
                            pass

                        insertions += 1
                        content_inserted = True
                        print(f"[SUCCESS] Inserted {lang_key} abstract ({len(abstract_text)} chars)")
                        print(f"[SUCCESS] Replaced: '{old_content}...' with abstract content")

                        # Look for keywords insertion point (next paragraph)
                        if keywords and i + 1 < len(doc.paragraphs):
                            keyword_para = doc.paragraphs[i + 1]
                            keyword_text_lower = keyword_para.text.lower()

                            if ('kata kunci' in keyword_text_lower or
                                'keywords' in keyword_text_lower or
                                'key words' in keyword_text_lower or
                                len(keyword_para.text.strip()) < 50):  # Likely placeholder

                                keywords_text = f"{keyword_label}: {', '.join(keywords)}"
                                keyword_para.clear()
                                keyword_para.add_run(keywords_text)

                                insertions += 1
                                keywords_inserted = True
                                print(f"[SUCCESS] Inserted {lang_key} keywords: {keywords_text}")

                        break

                if not content_inserted:
                    print(f"[WARNING] Could not find content insertion point for {lang_key} abstract")
                if keywords and not keywords_inserted:
                    print(f"[WARNING] Could not find keywords insertion point for {lang_key}")

                break  # Stop looking for more sections

        return insertions

    def _insert_chapter_with_headings(self, doc, chapter_num, chapter_content):
        """Insert chapter content with proper subsection headings."""
        insertions = 0

        # Subsection title mapping
        subsection_titles = {
            'latar_belakang': 'Latar Belakang',
            'rumusan_masalah': 'Rumusan Masalah',
            'tujuan_penelitian': 'Tujuan Penelitian',
            'manfaat_penelitian': 'Manfaat Penelitian',
            'batasan_masalah': 'Batasan Masalah dan Ruang Lingkup',
            'landasan_teori': 'Landasan Teori',
            'penelitian_terkait': 'Penelitian Terkait',
            'kerangka_pemikiran': 'Kerangka Pemikiran',
            'desain_penelitian': 'Desain Penelitian',
            'metode_pengumpulan_data': 'Metode Pengumpulan Data',
            'metode_analisis': 'Metode Analisis Data',
            'tools': 'Alat dan Bahan',
            'analisis_kebutuhan': 'Analisis Kebutuhan Sistem',
            'perancangan_sistem': 'Perancangan Sistem',
            'perancangan_interface': 'Perancangan Antarmuka',
            'implementasi': 'Implementasi Sistem',
            'hasil_pengujian': 'Hasil Pengujian',
            'pembahasan': 'Pembahasan Hasil',
            'evaluasi': 'Evaluasi Sistem',
            'kesimpulan': 'Kesimpulan',
            'saran': 'Saran'
        }

        # Find chapter heading
        chapter_heading_index = None
        for i, para in enumerate(doc.paragraphs):
            if f"BAB {self._to_roman(chapter_num)}" in para.text.upper():
                chapter_heading_index = i
                print(f"[INFO] Found Chapter {chapter_num} heading at paragraph {i}")
                break

        if chapter_heading_index is None:
            print(f"[WARNING] Chapter {chapter_num} heading not found")
            return 0

        # Insert each subsection with heading
        subsection_counter = 1
        current_index = chapter_heading_index + 1

        for subsection_key, subsection_content in chapter_content.items():
            if not subsection_content or len(subsection_content.strip()) < 50:
                continue

            # Create subsection heading (e.g., "1.1 Latar Belakang")
            heading_text = f"{chapter_num}.{subsection_counter} {subsection_titles.get(subsection_key, subsection_key.title())}"

            # Find or create heading paragraph
            heading_para = self._find_or_create_heading(doc, current_index, heading_text)

            if heading_para:
                # Find content paragraph after heading
                heading_index = doc.paragraphs.index(heading_para)
                content_para = self._find_content_paragraph_after(doc, heading_index)

                if content_para:
                    # Insert content
                    old_text = content_para.text[:50] if content_para.text else "[empty]"
                    content_para.clear()
                    content_para.add_run(subsection_content)

                    # Apply proper academic formatting (CRITICAL FIXES)
                    self._apply_paragraph_formatting(content_para)

                    print(f"[SUCCESS] {heading_text}: {len(subsection_content)} chars")
                    insertions += 1
                    current_index = doc.paragraphs.index(content_para) + 1

            subsection_counter += 1

        return insertions

    def _find_or_create_heading(self, doc, start_index, heading_text):
        """Find existing heading or create new one."""
        # Search for existing similar heading
        for i in range(start_index, min(start_index + 15, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            if (para.style.name == 'Heading 3' or
                'Subbab' in para.text or
                heading_text.split()[0] in para.text):  # e.g., "1.1" in existing text

                # Update with proper heading
                para.clear()
                para.add_run(heading_text)
                para.style = 'Heading 3'
                return para

        # If not found, try to find a placeholder paragraph
        for i in range(start_index, min(start_index + 15, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            if ('TULISKAN' in para.text.upper() or
                len(para.text.strip()) < 50):  # Likely placeholder

                para.clear()
                para.add_run(heading_text)
                para.style = 'Heading 3'
                return para

        print(f"[WARNING] Could not find or create heading for: {heading_text}")
        return None

    def _find_content_paragraph_after(self, doc, heading_index):
        """Find content paragraph after heading."""
        for i in range(heading_index + 1, min(heading_index + 10, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            if ('isi paragraf' in str(para.style).lower() or
                'Format paragraf' in para.text or
                len(para.text.strip()) < 100):  # Likely content area

                return para

        print(f"[WARNING] No content paragraph found after heading at index {heading_index}")
        return None

    def _insert_references(self, doc, analyzed_data):
        """Insert references into DAFTAR PUSTAKA section."""
        references = analyzed_data.get('references', [])
        if not references:
            return 0

        insertions = 0

        # Find DAFTAR PUSTAKA section
        for para in doc.paragraphs:
            if 'DAFTAR PUSTAKA' in para.text.upper():
                para_index = doc.paragraphs.index(para)
                print(f"[INFO] Found DAFTAR PUSTAKA at paragraph {para_index}")

                # Clear existing placeholder content
                cleared_count = 0
                for i in range(para_index + 1, min(para_index + 20, len(doc.paragraphs))):
                    ref_para = doc.paragraphs[i]
                    if ('Gunakan reference manager' in ref_para.text or
                        'DAFTAR PUSTAKA' in ref_para.text.upper() or
                        len(ref_para.text.strip()) < 50):

                        ref_para.clear()
                        cleared_count += 1
                        if cleared_count >= 5:  # Clear a few placeholder entries
                            break

                # Insert references
                insert_index = para_index + 1
                for ref in references[:15]:  # Limit to 15 references
                    if insert_index < len(doc.paragraphs):
                        target_para = doc.paragraphs[insert_index]
                        target_para.clear()
                        target_para.add_run(ref)
                        target_para.style = 'Normal'
                        insertions += 1
                        insert_index += 1
                    else:
                        # Add new paragraph
                        new_para = doc.add_paragraph(ref)
                        new_para.style = 'Normal'
                        insertions += 1

                print(f"[SUCCESS] Inserted {insertions} references")
                break

        return insertions

    def _replace_user_metadata_in_doc(self, doc, title, author, nim):
        """Helper to replace core metadata everywhere."""
        metadata_replacements = 0
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            # Title
            if any(phrase in text.upper() for phrase in ['TULIS JUDUL', 'JUDUL SKRIPSI', 'BAGIAN JUDUL']):
                is_title_page = False
                for i in range(max(0, para_idx - 5), min(para_idx + 5, len(doc.paragraphs))):
                    if 'HALAMAN JUDUL' in doc.paragraphs[i].text.upper():
                        is_title_page = True; break
                if is_title_page:
                    self._replace_text_preserve_formatting(para, para.text, title)
                    metadata_replacements += 1
            # Author
            if any(phrase in text.upper() for phrase in ['NAMA MAHASISWA', 'NAMA PENULIS', 'N A M A']):
                self._replace_text_preserve_formatting(para, para.text, author)
                metadata_replacements += 1
            # NIM
            if text == 'NIM' or (len(text) == 8 and text.isdigit()) or '94523999' in text or 'N I M' in text.upper():
                self._replace_text_preserve_formatting(para, para.text, nim)
                metadata_replacements += 1
        
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text_upper = para.text.upper()
                        if any(p in text_upper for p in ['NAMA MAHASISWA', 'NAMA PENULIS', 'N A M A']):
                            val = f"{para.text.split(':')[0]}: {author}" if ':' in para.text else author
                            self._replace_text_preserve_formatting(para, para.text, val)
                            metadata_replacements += 1
                        if any(p in text_upper for p in ['NIM', '94523999', 'N I M']):
                            val = f"{para.text.split(':')[0]}: {nim}" if ':' in para.text else nim
                            self._replace_text_preserve_formatting(para, para.text, val)
                            metadata_replacements += 1
        return metadata_replacements

    def _insert_paragraph_after(self, paragraph, text=None, style=None):
        """Insert a new paragraph after the given paragraph."""
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph
        
        new_p = OxmlElement('w:p')
        paragraph._element.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if text:
            new_para.add_run(text)
        if style:
            try:
                new_para.style = style
            except:
                pass
        return new_para

    def _roman_to_int(self, s):
        """Convert a Roman numeral string to an integer."""
        roman = {'I':1,'V':5,'X':10,'L':50,'C':100,'D':500,'M':1000}
        res = 0
        s = s.upper()
        for i in range(len(s)):
            if i > 0 and roman[s[i]] > roman[s[i-1]]:
                res += roman[s[i]] - 2 * roman[s[i-1]]
            else:
                res += roman[s[i]]
        return res

    def _replace_text_preserve_formatting(self, para, old_text, new_text):
        """Replace text while preserving run-level formatting (bold, italic, etc.)."""
        if not para.runs:
            para.add_run(new_text)
            return

        # Strategy: Replace text in the first run, clear other runs
        # This preserves the overall paragraph style and the first run's appearance
        para.runs[0].text = new_text
        for i in range(1, len(para.runs)):
            # Don't delete images or special elements in other runs if possible
            # But usually metadata fields are pure text
            if para.runs[i].text:
                para.runs[i].text = ""

def create_complete_thesis(
    template_path: str,
    content_path: str,
    output_path: str,
    user_data: Optional[Dict[str, Any]] = None,
    use_ai: bool = True,
    include_frontmatter: bool = True,
    api_key: Optional[str] = None,
    university_config: str = 'indonesian_standard'
) -> Dict[str, Any]:
    """Convenience function to create a complete thesis in one call.

    Args:
        template_path: Path to DOCX template
        content_path: Path to content file
        output_path: Path for output DOCX
        user_data: User data dictionary
        use_ai: Whether to use AI enhancement
        include_frontmatter: Whether to include front matter
        api_key: OpenRouter API key
        university_config: University template configuration
        use_ai: Whether to use AI semantic analysis
        include_frontmatter: Whether to include front matter sections
    
    Returns:
        Dictionary with:
            - output_file: Path to created document
            - status: "success" or "error"
            - message: Status message
            - report: Analysis report
    """
    try:
        # Verify files exist before creating builder
        template_p = Path(template_path)
        content_p = Path(content_path)
        
        if not template_p.exists():
            return {
                "status": "error",
                "message": f"Template file not found: {template_path}",
                "error_details": f"File does not exist: {template_path}",
            }
        
        if not content_p.exists():
            return {
                "status": "error",
                "message": f"Content file not found: {content_path}",
                "error_details": f"File does not exist: {content_path}",
            }
        
        builder = CompleteThesisBuilder(template_path, content_path, output_path, use_ai=use_ai, include_frontmatter=include_frontmatter, api_key=api_key, university_config=university_config)
        
        # Get analysis before building
        report = builder.get_analysis_report()
        
        # Build the thesis
        output = builder.build(user_data or {})
        
        return {
            "status": "success",
            "output_file": str(output),
            "message": "Complete thesis document created successfully",
            "report": report,
            "file_size": output.stat().st_size if output.exists() else 0,
        }
    except Exception as e:
        import traceback
        return {
            "status": "error",
            "message": f"Failed to create thesis: {str(e)}",
            "error_details": traceback.format_exc(),
        }
