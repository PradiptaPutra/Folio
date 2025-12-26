"""
Adaptive Insertion Engine
Intelligent content insertion with multiple strategies and fallbacks for robust template handling
"""

from typing import Dict, List, Any, Optional, Tuple, Callable
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from .advanced_template_analyzer import TemplateStructure, ZoneType
from .content_zone_mapper import InsertionPlan, ContentItem, ContentType
from .style_inheritance_engine import StyleInheritanceEngine, StyleRules


class InsertionStrategy(Enum):
    """Available insertion strategies"""
    DIRECT_ZONE_REPLACEMENT = "direct_zone_replacement"  # Replace placeholders directly
    SECTION_AWARE_INSERTION = "section_aware_insertion"  # Insert within section boundaries
    SEQUENTIAL_ZONE_FILLING = "sequential_zone_filling"  # Fill zones in order
    HYBRID_ADAPTIVE = "hybrid_adaptive"                # Combine multiple approaches


@dataclass
class InsertionResult:
    """Result of an insertion operation"""
    success: bool
    strategy_used: InsertionStrategy
    zones_processed: int
    content_inserted: int
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)


@dataclass
class InsertionContext:
    """Context for insertion operations"""
    document: Document
    template_structure: TemplateStructure
    insertion_plan: InsertionPlan
    style_engine: StyleInheritanceEngine


class AdaptiveInsertionEngine:
    """
    Adaptive engine that selects and executes the best insertion strategy
    based on template analysis and content characteristics
    """

    def __init__(self):
        self.strategies = {
            InsertionStrategy.DIRECT_ZONE_REPLACEMENT: self._execute_direct_replacement,
            InsertionStrategy.SECTION_AWARE_INSERTION: self._execute_section_aware_insertion,
            InsertionStrategy.SEQUENTIAL_ZONE_FILLING: self._execute_sequential_filling,
            InsertionStrategy.HYBRID_ADAPTIVE: self._execute_hybrid_adaptive,
        }

        self.style_engine = StyleInheritanceEngine()

    def execute_insertion_plan(self, context: InsertionContext) -> InsertionResult:
        """
        Execute content insertion using the most appropriate strategy

        Args:
            context: InsertionContext with document, template, plan, and style engine

        Returns:
            InsertionResult with success status and details
        """
        print("[INFO] Starting adaptive content insertion...")

        # Analyze template and content to select best strategy
        best_strategy = self._select_optimal_strategy(context)

        print(f"[INFO] Selected insertion strategy: {best_strategy.value}")

        # Execute the chosen strategy
        strategy_function = self.strategies[best_strategy]
        result = strategy_function(context)

        # Validate and finalize
        self._validate_insertion_result(context, result)
        self._finalize_document(context.document)

        print(f"[SUCCESS] Insertion completed. Strategy: {best_strategy.value}, "
              f"Zones: {result.zones_processed}, Content: {result.content_inserted}")

        return result

    def _select_optimal_strategy(self, context: InsertionContext) -> InsertionStrategy:
        """
        Select the optimal insertion strategy based on template and content analysis
        """
        template = context.template_structure
        plan = context.insertion_plan

        # Analyze template characteristics
        has_clear_placeholders = len(template.academic_patterns.get('placeholders', [])) > 5
        has_structured_zones = len([z for z in template.zones.values()
                                   if z.zone_type in [ZoneType.HEADER, ZoneType.CONTENT]]) > 10
        confidence_score = template.confidence_score

        # Analyze content characteristics
        total_mappings = len(plan.mappings)
        direct_mappings = len([m for m in plan.mappings.values()
                              if self._is_direct_mappable(context, m)])

        # Decision logic
        if confidence_score > 0.8 and has_clear_placeholders and direct_mappings > total_mappings * 0.7:
            return InsertionStrategy.DIRECT_ZONE_REPLACEMENT
        elif confidence_score > 0.6 and has_structured_zones:
            return InsertionStrategy.SECTION_AWARE_INSERTION
        elif confidence_score > 0.4:
            return InsertionStrategy.SEQUENTIAL_ZONE_FILLING
        else:
            return InsertionStrategy.HYBRID_ADAPTIVE

    def _execute_direct_replacement(self, context: InsertionContext) -> InsertionResult:
        """
        Strategy 1: Direct zone replacement (ideal for templates with clear placeholders)
        """
        result = InsertionResult(
            success=True,
            strategy_used=InsertionStrategy.DIRECT_ZONE_REPLACEMENT,
            zones_processed=0,
            content_inserted=0
        )

        for zone_id, content_id in context.insertion_plan.mappings.items():
            try:
                zone = context.template_structure.zones.get(zone_id)
                if not zone:
                    result.errors.append(f"Zone {zone_id} not found in template")
                    continue

                # Find the corresponding content
                content_item = self._find_content_item(context, content_id)
                if not content_item:
                    result.warnings.append(f"Content {content_id} not found")
                    continue

                # Apply styles
                style_rules = context.insertion_plan.style_mappings.get(content_id, {})

                # Execute direct replacement
                success = self._replace_zone_content(
                    context.document, zone, content_item, style_rules
                )

                if success:
                    result.zones_processed += 1
                    result.content_inserted += 1
                else:
                    result.errors.append(f"Failed to replace content in zone {zone_id}")

            except Exception as e:
                result.errors.append(f"Error processing zone {zone_id}: {str(e)}")

        return result

    def _execute_section_aware_insertion(self, context: InsertionContext) -> InsertionResult:
        """
        Strategy 2: Section-aware insertion (respects document hierarchy)
        """
        result = InsertionResult(
            success=True,
            strategy_used=InsertionStrategy.SECTION_AWARE_INSERTION,
            zones_processed=0,
            content_inserted=0
        )

        # Group mappings by hierarchy level
        hierarchy_groups = self._group_mappings_by_hierarchy(context)

        for level, mappings in hierarchy_groups.items():
            print(f"[INFO] Processing hierarchy level {level} with {len(mappings)} mappings")

            for zone_id, content_id in mappings:
                try:
                    zone = context.template_structure.zones.get(zone_id)
                    content_item = self._find_content_item(context, content_id)

                    if zone and content_item:
                        # Apply section-aware insertion
                        success = self._insert_with_hierarchy_awareness(
                            context.document, zone, content_item, context.template_structure
                        )

                        if success:
                            result.zones_processed += 1
                            result.content_inserted += 1
                        else:
                            result.warnings.append(f"Section-aware insertion failed for {zone_id}")

                except Exception as e:
                    result.errors.append(f"Hierarchy insertion error for {zone_id}: {str(e)}")

        return result

    def _execute_sequential_filling(self, context: InsertionContext) -> InsertionResult:
        """
        Strategy 3: Sequential zone filling (fills available zones in order)
        """
        result = InsertionResult(
            success=True,
            strategy_used=InsertionStrategy.SEQUENTIAL_ZONE_FILLING,
            zones_processed=0,
            content_inserted=0
        )

        # Get available zones sorted by position
        available_zones = []
        for zone in context.template_structure.zones.values():
            if zone.zone_type in [ZoneType.CONTENT, ZoneType.PLACEHOLDER]:
                available_zones.append(zone)

        available_zones.sort(key=lambda z: z.start_paragraph)

        # Get content items sorted by hierarchy
        content_items = []
        for content_id in context.insertion_plan.mappings.values():
            content_item = self._find_content_item(context, content_id)
            if content_item:
                content_items.append(content_item)

        content_items.sort(key=lambda c: (c.hierarchy_level, c.metadata.get('chapter_num', 0)))

        # Map sequentially
        for zone, content_item in zip(available_zones, content_items):
            try:
                style_rules = context.insertion_plan.style_mappings.get(content_item.content_id, {})

                success = self._replace_zone_content(
                    context.document, zone, content_item, style_rules
                )

                if success:
                    result.zones_processed += 1
                    result.content_inserted += 1
                else:
                    result.warnings.append(f"Sequential insertion failed for zone {zone.zone_id}")

            except Exception as e:
                result.errors.append(f"Sequential insertion error: {str(e)}")

        return result

    def _execute_hybrid_adaptive(self, context: InsertionContext) -> InsertionResult:
        """
        Strategy 4: Hybrid adaptive (combines multiple approaches)
        """
        result = InsertionResult(
            success=True,
            strategy_used=InsertionStrategy.HYBRID_ADAPTIVE,
            zones_processed=0,
            content_inserted=0
        )

        # Phase 1: Try direct replacement for high-confidence mappings
        direct_result = self._execute_direct_replacement(context)
        result.zones_processed += direct_result.zones_processed
        result.content_inserted += direct_result.content_inserted
        result.errors.extend(direct_result.errors)
        result.warnings.extend(direct_result.warnings)

        # Phase 2: Use section-aware for remaining content
        if len(context.insertion_plan.mappings) > result.content_inserted:
            section_result = self._execute_section_aware_insertion(context)
            # Avoid double-counting
            additional_zones = section_result.zones_processed
            additional_content = section_result.content_inserted
            result.errors.extend(section_result.errors)
            result.warnings.extend(section_result.warnings)

            # Only count truly additional insertions
            result.zones_processed += additional_zones
            result.content_inserted += additional_content

        # Phase 3: Fallback to sequential for any remaining
        if len(context.insertion_plan.mappings) > result.content_inserted:
            sequential_result = self._execute_sequential_filling(context)
            result.zones_processed += sequential_result.zones_processed
            result.content_inserted += sequential_result.content_inserted
            result.errors.extend(sequential_result.errors)
            result.warnings.extend(sequential_result.warnings)

        return result

    def _replace_zone_content(self, doc: Document, zone, content_item: ContentItem,
                             style_rules: Dict[str, Any]) -> bool:
        """
        Replace content in a specific zone with proper styling
        """
        try:
            # Find the target paragraph
            if zone.start_paragraph >= len(doc.paragraphs):
                return False

            target_para = doc.paragraphs[zone.start_paragraph]

            # Clear existing content
            target_para.clear()

            # Add new content
            run = target_para.add_run(content_item.content)

            # Apply styling
            self._apply_styling_to_run(run, style_rules, content_item.content_type)

            # Apply paragraph-level formatting
            self._apply_paragraph_styling(target_para, style_rules)

            return True

        except Exception as e:
            print(f"[ERROR] Content replacement failed: {e}")
            return False

    def _insert_with_hierarchy_awareness(self, doc: Document, zone, content_item: ContentItem,
                                       template_structure: TemplateStructure) -> bool:
        """
        Insert content while respecting document hierarchy
        """
        try:
            # Find insertion point considering hierarchy
            insert_position = self._find_hierarchy_aware_position(zone, template_structure)

            if insert_position >= len(doc.paragraphs):
                # Append to document
                new_para = doc.add_paragraph(content_item.content)
                target_para = new_para
            else:
                target_para = doc.paragraphs[insert_position]
                target_para.clear()
                run = target_para.add_run(content_item.content)

            # Apply hierarchy-appropriate styling
            hierarchy_styles = self._get_hierarchy_styles(content_item.hierarchy_level)
            self._apply_styling_to_paragraph(target_para, hierarchy_styles)

            return True

        except Exception as e:
            print(f"[ERROR] Hierarchy-aware insertion failed: {e}")
            return False

    def _apply_styling_to_run(self, run, style_rules: Dict[str, Any], content_type: ContentType):
        """Apply styling to a document run"""
        try:
            # Font properties
            if 'font_family' in style_rules:
                run.font.name = style_rules['font_family']
            if 'font_size' in style_rules:
                run.font.size = Pt(style_rules['font_size'])
            if style_rules.get('bold', False):
                run.font.bold = True
            if style_rules.get('italic', False):
                run.font.italic = True

            # Special handling for headers
            if content_type == ContentType.CHAPTER_TITLE:
                run.font.bold = True
                run.font.size = Pt(14)
            elif content_type == ContentType.SUBSECTION_TITLE:
                run.font.bold = True
                run.font.size = Pt(12)

        except Exception as e:
            print(f"[WARNING] Run styling failed: {e}")

    def _apply_paragraph_styling(self, para, style_rules: Dict[str, Any]):
        """Apply paragraph-level styling"""
        try:
            # Alignment
            alignment = style_rules.get('alignment', 'justify')
            if alignment == 'justify':
                para.paragraph_format.alignment = 3  # WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            elif alignment == 'center':
                para.paragraph_format.alignment = 1  # WD_PARAGRAPH_ALIGNMENT.CENTER
            elif alignment == 'right':
                para.paragraph_format.alignment = 2  # WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:  # left
                para.paragraph_format.alignment = 0  # WD_PARAGRAPH_ALIGNMENT.LEFT

            # Indentation
            if 'first_line_indent' in style_rules:
                para.paragraph_format.first_line_indent = Inches(style_rules['first_line_indent'])
            if 'left_indent' in style_rules:
                para.paragraph_format.left_indent = Inches(style_rules['left_indent'])

            # Spacing
            if 'line_spacing' in style_rules:
                para.paragraph_format.line_spacing = style_rules['line_spacing']
            if 'space_before' in style_rules:
                para.paragraph_format.space_before = Pt(style_rules['space_before'])
            if 'space_after' in style_rules:
                para.paragraph_format.space_after = Pt(style_rules['space_after'])

        except Exception as e:
            print(f"[WARNING] Paragraph styling failed: {e}")

    def _apply_styling_to_paragraph(self, para, style_rules: Dict[str, Any]):
        """Apply comprehensive styling to paragraph"""
        # Apply to the first run if it exists
        if para.runs:
            self._apply_styling_to_run(para.runs[0], style_rules, ContentType.PARAGRAPH)

        # Apply paragraph formatting
        self._apply_paragraph_styling(para, style_rules)

    def _find_content_item(self, context: InsertionContext, content_id: str) -> Optional[ContentItem]:
        """Find a content item by ID"""
        # This would need to be implemented based on how content items are stored
        # For now, return None
        return None

    def _is_direct_mappable(self, context: InsertionContext, content_id: str) -> bool:
        """Check if content can be directly mapped to a zone"""
        # Implementation would check if the content and zone are clearly compatible
        return True

    def _group_mappings_by_hierarchy(self, context: InsertionContext) -> Dict[int, List[Tuple[str, str]]]:
        """Group mappings by hierarchy level"""
        groups = {}

        for zone_id, content_id in context.insertion_plan.mappings.items():
            zone = context.template_structure.zones.get(zone_id)
            if zone:
                level = zone.hierarchy_level
                if level not in groups:
                    groups[level] = []
                groups[level].append((zone_id, content_id))

        return groups

    def _find_hierarchy_aware_position(self, zone, template_structure: TemplateStructure) -> int:
        """Find insertion position considering document hierarchy"""
        # Start with zone's natural position
        position = zone.start_paragraph

        # Adjust based on hierarchy relationships
        if zone.parent_zone:
            parent = template_structure.zones.get(zone.parent_zone)
            if parent:
                position = max(position, parent.end_paragraph + 1)

        return position

    def _get_hierarchy_styles(self, hierarchy_level: int) -> Dict[str, Any]:
        """Get appropriate styles for hierarchy level"""
        base_styles = {
            'font_family': 'Times New Roman',
            'font_size': 11,
            'alignment': 'justify',
            'first_line_indent': 1.0,
            'line_spacing': 1.5
        }

        if hierarchy_level == 0:  # Chapter level
            base_styles.update({
                'font_size': 14,
                'bold': True,
                'alignment': 'center'
            })
        elif hierarchy_level == 1:  # Major section
            base_styles.update({
                'font_size': 12,
                'bold': True,
                'alignment': 'left'
            })

        return base_styles

    def _validate_insertion_result(self, context: InsertionContext, result: InsertionResult):
        """Validate the insertion result and add warnings/errors"""
        # Check for unmapped content
        mapped_content = set(context.insertion_plan.mappings.values())
        total_expected = len(context.insertion_plan.mappings)

        if result.content_inserted < total_expected:
            result.warnings.append(
                f"Only {result.content_inserted}/{total_expected} content items were inserted"
            )

        # Check document integrity
        try:
            # Basic validation that document is still readable
            para_count = len(context.document.paragraphs)
            if para_count < 10:
                result.errors.append("Document appears corrupted after insertion")
        except Exception as e:
            result.errors.append(f"Document validation failed: {e}")

    def _finalize_document(self, doc: Document):
        """Final document cleanup and optimization"""
        try:
            # Update table of contents if present
            # This would require additional libraries for TOC updating

            # Clean up any remaining placeholders
            for para in doc.paragraphs:
                if para.text.strip() in ['[empty]', 'TULISKAN ISI', 'Format paragraf dengan style']:
                    para.clear()

            print("[INFO] Document finalization completed")

        except Exception as e:
            print(f"[WARNING] Document finalization failed: {e}")