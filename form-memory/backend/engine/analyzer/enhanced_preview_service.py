"""
Enhanced Document Preview Service
Provides high-quality DOCX-to-HTML conversion with styling preservation.
"""

from typing import Dict, Any, Optional
from pathlib import Path
import mammoth
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX


class EnhancedPreviewService:
    """Service for creating high-quality document previews."""

    def __init__(self):
        self.base_css = """
        body {
            font-family: 'Times New Roman', serif;
            margin: 40px;
            line-height: 1.6;
            color: #000000;
            background: #ffffff;
        }

        .document-container {
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }

        h1, h2, h3, h4, h5, h6 {
            color: #2c3e50;
            margin-top: 24px;
            margin-bottom: 12px;
            font-weight: bold;
        }

        h1 { font-size: 24pt; border-bottom: 2px solid #3498db; padding-bottom: 8px; }
        h2 { font-size: 18pt; border-bottom: 1px solid #bdc3c7; padding-bottom: 4px; }
        h3 { font-size: 14pt; }
        h4 { font-size: 12pt; }

        p {
            margin: 8px 0;
            text-align: justify;
            text-indent: 1cm;
        }

        .center { text-align: center; text-indent: 0; }
        .right { text-align: right; text-indent: 0; }
        .left { text-align: left; text-indent: 0; }

        ul, ol {
            margin: 12px 0;
            padding-left: 30px;
        }

        li { margin: 4px 0; }

        table {
            border-collapse: collapse;
            width: 100%;
            margin: 16px 0;
            border: 1px solid #bdc3c7;
        }

        th, td {
            border: 1px solid #bdc3c7;
            padding: 8px 12px;
            text-align: left;
        }

        th {
            background-color: #ecf0f1;
            font-weight: bold;
        }

        blockquote {
            margin: 16px 0;
            padding: 12px 20px;
            border-left: 4px solid #3498db;
            background-color: #ecf0f1;
            font-style: italic;
        }

        .highlight { background-color: #ffff00; }
        .bold { font-weight: bold; }
        .italic { font-style: italic; }
        .underline { text-decoration: underline; }

        .page-break {
            page-break-before: always;
            border-top: 1px dashed #bdc3c7;
            margin: 40px 0;
            padding-top: 20px;
        }

        @media print {
            body { margin: 20px; }
            .document-container { box-shadow: none; }
        }
        """

    def generate_preview(self, docx_path: str) -> Dict[str, Any]:
        """Generate enhanced HTML preview from DOCX file."""
        try:
            docx_file = Path(docx_path)
            if not docx_file.exists():
                return {
                    "status": "error",
                    "message": f"Document not found: {docx_path}",
                    "html_content": ""
                }

            # Extract document styles and structure
            doc = Document(str(docx_file))

            # Generate HTML with custom conversion
            html_content = self._convert_to_html(doc)

            # Extract document metadata
            metadata = self._extract_metadata(doc)

            return {
                "status": "success",
                "html_content": html_content,
                "metadata": metadata,
                "css_styles": self.base_css
            }

        except Exception as e:
            return {
                "status": "error",
                "message": f"Preview generation failed: {str(e)}",
                "html_content": f"<div style='color: red; padding: 20px;'>Error generating preview: {str(e)}</div>"
            }

    def _convert_to_html(self, doc: Document) -> str:
        """Convert DOCX document to styled HTML."""
        html_parts = []

        # Add document container
        html_parts.append('<div class="document-container">')

        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            # Determine paragraph class based on style and alignment
            para_class = self._get_paragraph_class(para)

            # Convert paragraph content
            para_html = self._convert_paragraph_content(para, para_class)
            html_parts.append(para_html)

        # Add tables
        for table in doc.tables:
            table_html = self._convert_table(table)
            html_parts.append(table_html)

        html_parts.append('</div>')

        return '\n'.join(html_parts)

    def _get_paragraph_class(self, para) -> str:
        """Determine CSS class for paragraph based on style and formatting."""
        classes = []

        # Check alignment
        if para.alignment == 1:  # Center
            classes.append("center")
        elif para.alignment == 2:  # Right
            classes.append("right")
        elif para.alignment == 0:  # Left
            classes.append("left")

        # Check style name for headings
        if para.style:
            style_name = para.style.name.lower()
            if 'heading 1' in style_name or 'title' in style_name:
                return "h1"
            elif 'heading 2' in style_name:
                return "h2"
            elif 'heading 3' in style_name:
                return "h3"
            elif 'heading 4' in style_name:
                return "h4"
            elif 'heading 5' in style_name:
                return "h5"
            elif 'heading 6' in style_name:
                return "h6"

        return " ".join(classes) if classes else "normal"

    def _convert_paragraph_content(self, para, para_class: str) -> str:
        """Convert paragraph content to HTML with formatting."""
        content_parts = []

        # Handle different paragraph types
        if para_class.startswith('h'):
            tag = para_class
            content_parts.append(f'<{tag}>{para.text}</{tag}>')
        else:
            # Regular paragraph
            content_parts.append(f'<p class="{para_class}">')

            # Process runs for formatting
            for run in para.runs:
                run_text = run.text
                if not run_text:
                    continue

                # Apply formatting
                formatted_text = run_text
                if run.bold:
                    formatted_text = f'<strong>{formatted_text}</strong>'
                if run.italic:
                    formatted_text = f'<em>{formatted_text}</em>'
                if run.underline:
                    formatted_text = f'<u>{formatted_text}</u>'

                # Handle highlighting
                if hasattr(run.font, 'highlight_color') and run.font.highlight_color:
                    formatted_text = f'<span class="highlight">{formatted_text}</span>'

                content_parts.append(formatted_text)

            content_parts.append('</p>')

        return ''.join(content_parts)

    def _convert_table(self, table) -> str:
        """Convert DOCX table to HTML table."""
        html_parts = ['<table>']

        # Add header row if it exists
        if table.rows:
            html_parts.append('<thead>')
            # Simple approach: assume first row is header
            header_row = table.rows[0]
            html_parts.append('<tr>')
            for cell in header_row.cells:
                html_parts.append(f'<th>{cell.text}</th>')
            html_parts.append('</tr>')
            html_parts.append('</thead>')

            # Add data rows
            if len(table.rows) > 1:
                html_parts.append('<tbody>')
                for row in table.rows[1:]:
                    html_parts.append('<tr>')
                    for cell in row.cells:
                        html_parts.append(f'<td>{cell.text}</td>')
                    html_parts.append('</tr>')
                html_parts.append('</tbody>')

        html_parts.append('</table>')
        return '\n'.join(html_parts)

    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        """Extract document metadata."""
        metadata = {
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
            "word_count": 0,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }

        # Count words
        word_count = 0
        for para in doc.paragraphs:
            words = para.text.split()
            word_count += len(words)

        metadata["word_count"] = word_count

        return metadata


# Global instance
preview_service = EnhancedPreviewService()


def generate_enhanced_preview(docx_path: str) -> Dict[str, Any]:
    """Convenience function to generate enhanced document preview."""
    return preview_service.generate_preview(docx_path)