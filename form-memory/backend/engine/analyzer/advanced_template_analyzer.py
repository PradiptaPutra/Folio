"""
Advanced Indonesian University Template Analyzer
Multi-layer analysis for comprehensive DOCX template understanding:
- Structure detection with academic pattern recognition
- Content zone classification (headers/content/placeholders)
- Style inheritance and formatting rules
- Indonesian academic standards compliance
"""

from typing import Dict, List, Any, Optional, Set, Tuple, Union
from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.style import WD_STYLE_TYPE
from dataclasses import dataclass, field
from enum import Enum


class ZoneType(Enum):
    """Content zone classification types"""
    HEADER = "header"           # Chapter titles, section headers
    CONTENT = "content"         # Where content should be inserted
    PLACEHOLDER = "placeholder" # Template instructions to remove
    STRUCTURAL = "structural"   # Tables, figures, references
    FRONT_MATTER = "front_matter"  # Title page, abstract, etc.
    BACK_MATTER = "back_matter"    # References, appendices


@dataclass
class ContentZone:
    """Represents a content zone in the template"""
    zone_id: str
    zone_type: ZoneType
    start_paragraph: int
    end_paragraph: int
    content: str
    style_info: Dict[str, Any] = field(default_factory=dict)
    hierarchy_level: int = 0
    parent_zone: Optional[str] = None
    child_zones: List[str] = field(default_factory=list)


@dataclass
class TemplateStructure:
    """Complete template structure analysis"""
    zones: Dict[str, ContentZone] = field(default_factory=dict)
    hierarchy: Dict[str, List[str]] = field(default_factory=dict)
    style_rules: Dict[str, Any] = field(default_factory=dict)
    academic_patterns: Dict[str, Any] = field(default_factory=dict)
    confidence_score: float = 0.0


class AdvancedTemplateAnalyzer:
    """
    Advanced analyzer for Indonesian university DOCX templates.
    Performs multi-pass analysis to understand template structure and content zones.
    """

    # Comprehensive Indonesian academic patterns for recognition
    INDONESIAN_PATTERNS = {
        'chapters': [
            # Standard patterns
            r'^BAB\s+[IVXLCDM]+\s*[:-]?\s*(.+)$',      # BAB I - PENDAHULUAN
            r'^BAB\s+\d+\s*[:-]?\s*(.+)$',             # BAB 1 - PENDAHULUAN
            r'^Chapter\s+\d+\s*[:-]?\s*(.+)$',         # Chapter 1 - Introduction
            r'^Bab\s+[ivxlcdm]+\s*[:-]?\s*(.+)$',      # lowercase variants

            # University-specific variations
            r'^BAB\s+[IVXLCDM]+\s+(.+)$',              # BAB I PENDAHULUAN (no dash)
            r'^BAB\s+\d+\s+(.+)$',                     # BAB 1 PENDAHULUAN (no dash)
            r'^([IVXLCDM]+)\.\s+(.+)$',                # I. PENDAHULUAN
            r'^(\d+)\.\s+(.+)$',                       # 1. PENDAHULUAN

            # Alternative chapter titles (different universities)
            r'.*PENDAHULUAN.*',                        # Contains PENDAHULUAN
            r'.*LATAR\s+BELAKANG.*',                   # Contains LATAR BELAKANG
            r'.*TINJAUAN\s+PUSTAKA.*',                 # Contains TINJAUAN PUSTAKA
            r'.*KAJIAN\s+PUSTAKA.*',                   # Contains KAJIAN PUSTAKA
            r'.*METODOLOGI.*',                         # Contains METODOLOGI
            r'.*METODE.*',                             # Contains METODE
            r'.*ANALISIS.*',                           # Contains ANALISIS
            r'.*PEMBAHASAN.*',                         # Contains PEMBAHASAN
            r'.*KESIMPULAN.*',                         # Contains KESIMPULAN
            r'.*PENUTUP.*',                            # Contains PENUTUP
        ],
        'subsections': [
            # Standard numbering
            r'^(\d+)\.(\d+)\.?\s+(.+)$',               # 1.1 Latar Belakang
            r'^(\d+)\.(\d+)\.(\d+)\.?\s+(.+)$',        # 1.1.1 Pendahuluan
            r'^(\d+)\.(\d+)\.(\d+)\.(\d+)\.?\s+(.+)$', # 1.1.1.1 Sub-sub

            # Letter numbering
            r'^([A-Z])\.\s+(.+)$',                     # A. Alternatif
            r'^([a-z])\.\s+(.+)$',                     # a. alternatif
            r'^([A-Z])\)\s+(.+)$',                     # A) Alternatif
            r'^([a-z])\)\s+(.+)$',                     # a) alternatif

            # Mixed numbering
            r'^(\d+)\.([A-Z])\.?\s+(.+)$',             # 1.A Alternatif
            r'^(\d+)\.([a-z])\.?\s+(.+)$',             # 1.a alternatif

            # No numbering (keyword-based)
            r'.*LATAR\s+BELAKANG.*',                   # Contains LATAR BELAKANG
            r'.*RUMUSAN\s+MASALAH.*',                  # Contains RUMUSAN MASALAH
            r'.*TUJUAN.*',                             # Contains TUJUAN
            r'.*MANFAAT.*',                            # Contains MANFAAT
            r'.*BATASAN.*',                            # Contains BATASAN
            r'.*LANDASAN\s+TEORI.*',                   # Contains LANDASAN TEORI
            r'.*PENELITIAN\s+TERKAIT.*',               # Contains PENELITIAN TERKAIT
            r'.*KERANGKA\s+PIKIR.*',                   # Contains KERANGKA PIKIR
            r'.*METODE.*',                             # Contains METODE
            r'.*TEKNIK.*',                             # Contains TEKNIK
            r'.*ANALISIS.*',                           # Contains ANALISIS
            r'.*HASIL.*',                              # Contains HASIL
            r'.*PEMBAHASAN.*',                         # Contains PEMBAHASAN
            r'.*KESIMPULAN.*',                         # Contains KESIMPULAN
            r'.*SARAN.*',                              # Contains SARAN
        ],
        'content_placeholders': [
            # Explicit instructions
            r'Format paragraf dengan style',
            r'TULISKAN\s+(.+)',
            r'Tulis\s+(.+)',
            r'Isi\s+(.+)',
            r'Isikan\s+(.+)',
            r'Deskripsikan\s+(.+)',
            r'Jelaskan\s+(.+)',
            r'Uraikan\s+(.+)',
            r'Bagian\s+ini\s+berisi',
            r'Di\s+bagian\s+ini\s+dijelaskan',

            # Placeholder markers
            r'\[empty\]',
            r'\[isi\]',
            r'\[content\]',
            r'\[teks\]',
            r'\[paragraf\]',

            # Sample content
            r'Lorem\s+ipsum',
            r'Contoh\s+teks',
            r'Sample\s+text',
            r'Dummy\s+text',
            r'Testing',

            # Length indicators
            r'\d+\s*-\s*\d+\s*kata',                    # "100-150 kata"
            r'\d+\s*halaman',                           # "2 halaman"
            r'minimal\s+\d+',                           # "minimal 500"
            r'maksimal\s+\d+',                          # "maksimal 1000"

            # Instructional language
            r'Pada\s+bab\s+ini',
            r'Dalam\s+bab\s+ini',
            r'Bab\s+ini\s+menjelaskan',
            r'Bab\s+ini\s+membahas',
            r'Bab\s+ini\s+berisi',
        ],
        'subsections': [
            r'^(\d+)\.(\d+)\.?\s+(.+)$',               # 1.1 Latar Belakang
            r'^(\d+)\.(\d+)\.(\d+)\.?\s+(.+)$',        # 1.1.1 Pendahuluan
            r'^([A-Z])\.\s+(.+)$',                     # A. Alternatif
            r'^([a-z])\)\s+(.+)$',                     # a) Alternatif
        ],
        'content_placeholders': [
            r'Format paragraf dengan style',
            r'TULISKAN\s+(.+)',
            r'\[empty\]',
            r'Deskripsikan\s+(.+)',
            r'Isikan\s+(.+)',
            r'Jelaskan\s+(.+)',
            r'(?:Tulis|Buat)\s+(.+)',
        ],
        'front_matter': [
            r'^HALAMAN\s+',
            r'^LEMBAR\s+',
            r'^KATA\s+PENGANTAR',
            r'^ABSTRAK',
            r'^ABSTRACT',
            r'^DAFTAR\s+ISI',
            r'^DAFTAR\s+TABEL',
            r'^DAFTAR\s+GAMBAR',
        ],
        'back_matter': [
            r'^DAFTAR\s+PUSTAKA',
            r'^REFERENCES',
            r'^LAMPIRAN',
            r'^APPENDICES',
        ]
    }

    def __init__(self, template_path: Union[str, Path]):
        self.template_path = Path(template_path)
        self.doc: Optional[Document] = None
        self.structure = TemplateStructure()

        # Load and validate template
        self._load_template()
        self._validate_template()

    def _load_template(self) -> None:
        """Load and validate the DOCX template"""
        try:
            self.doc = Document(str(self.template_path))
            if self.doc:
                print(f"[INFO] Loaded template: {len(self.doc.paragraphs)} paragraphs, "
                      f"{len(self.doc.tables)} tables, {len(self.doc.styles)} styles")
        except Exception as e:
            raise ValueError(f"Failed to load template {self.template_path}: {e}")

    def _validate_template(self) -> None:
        """Validate that this is a valid academic template"""
        if not self.doc:
            raise ValueError("No document loaded")

        # Check for academic indicators
        academic_indicators = 0
        text_content = ' '.join([p.text.lower() for p in self.doc.paragraphs[:50]])  # type: ignore

        if any(word in text_content for word in ['bab', 'chapter', 'universitas', 'fakultas']):
            academic_indicators += 1
        if any(word in text_content for word in ['skripsi', 'tesis', 'disertasi']):
            academic_indicators += 1
        if len([p for p in self.doc.paragraphs if len(p.text.strip()) > 100]) > 5:  # type: ignore
            academic_indicators += 1

        if academic_indicators < 2:
            print(f"[WARNING] Template may not be academic: {academic_indicators} indicators found")

    def analyze_template_comprehensive(self) -> TemplateStructure:
        """
        Perform comprehensive multi-pass template analysis.
        Returns complete template structure with zones, hierarchy, and style rules.
        """
        print("[INFO] Starting comprehensive template analysis...")

        # Pass 1: Basic structure detection
        self._analyze_basic_structure()

        # Pass 2: Content zone classification
        self._classify_content_zones()

        # Pass 3: Hierarchy and relationships
        self._build_hierarchy()

        # Pass 4: Style rule extraction
        self._extract_style_rules()

        # Pass 5: Academic pattern recognition
        self._recognize_academic_patterns()

        # Calculate confidence score
        self._calculate_confidence_score()

        print(f"[SUCCESS] Template analysis complete. Found {len(self.structure.zones)} zones "
              f"with {self.structure.confidence_score:.1%} confidence")

        return self.structure

    def _analyze_basic_structure(self) -> None:
        """Pass 1: Analyze basic document structure"""
        print("[INFO] Pass 1: Analyzing basic document structure...")

        if not self.doc:
            return

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # Detect outline level (heading hierarchy)
            outline_level = self._get_outline_level(para)

            # Create basic zone
            zone_id = f"para_{i}"
            zone = ContentZone(
                zone_id=zone_id,
                zone_type=ZoneType.CONTENT,  # Default, will be refined later
                start_paragraph=i,
                end_paragraph=i,
                content=text,
                hierarchy_level=outline_level or 0
            )

            self.structure.zones[zone_id] = zone

    def _classify_content_zones(self) -> None:
        """Pass 2: Classify content zones using pattern recognition"""
        print("[INFO] Pass 2: Classifying content zones...")

        if not self.doc:
            return

        for zone_id, zone in self.structure.zones.items():
            zone.zone_type = self._classify_zone_type(zone.content)

            # Extract additional metadata for headers
            if zone.zone_type == ZoneType.HEADER:
                zone.style_info.update(self._extract_header_info(zone.content))

    def _build_hierarchy(self) -> None:
        """Pass 3: Build content hierarchy and relationships"""
        print("[INFO] Pass 3: Building content hierarchy...")

        current_parent = None
        hierarchy_stack = []

        for zone_id in sorted(self.structure.zones.keys(),
                             key=lambda x: self.structure.zones[x].start_paragraph):

            zone = self.structure.zones[zone_id]

            # Manage hierarchy stack based on level
            while hierarchy_stack and hierarchy_stack[-1][1] >= zone.hierarchy_level:
                hierarchy_stack.pop()

            # Set parent
            if hierarchy_stack:
                parent_id, _ = hierarchy_stack[-1]
                zone.parent_zone = parent_id
                self.structure.zones[parent_id].child_zones.append(zone_id)

            # Add to stack if it's a header
            if zone.zone_type == ZoneType.HEADER:
                hierarchy_stack.append((zone_id, zone.hierarchy_level))
                current_parent = zone_id

            # Build hierarchy map
            if zone.parent_zone:
                if zone.parent_zone not in self.structure.hierarchy:
                    self.structure.hierarchy[zone.parent_zone] = []
                self.structure.hierarchy[zone.parent_zone].append(zone_id)

    def _extract_style_rules(self) -> None:
        """Pass 4: Extract comprehensive style rules"""
        print("[INFO] Pass 4: Extracting style rules...")

        if not self.doc:
            return

        style_rules = {
            'fonts': {},
            'spacing': {},
            'alignment': {},
            'indentation': {},
            'numbering': {},
            'tables': {},
            'headers': {}
        }

        # Analyze paragraph styles
        for para in self.doc.paragraphs:
            if para.style and para.text.strip():
                self._analyze_paragraph_style(para, style_rules)

        # Analyze table styles
        for table in self.doc.tables:
            self._analyze_table_style(table, style_rules)

        self.structure.style_rules = style_rules

    def _recognize_academic_patterns(self) -> None:
        """Pass 5: Recognize Indonesian academic patterns with enhanced chapter mapping"""
        print("[INFO] Pass 5: Recognizing academic patterns...")

        patterns_found = {
            'chapters': [],
            'subsections': [],
            'front_matter': [],
            'back_matter': [],
            'placeholders': [],
            'chapter_mapping': {}  # Map detected chapters to standard research structure
        }

        # Standard research structure mapping
        STANDARD_CHAPTERS = {
            1: ['pendahuluan', 'latar belakang', 'rumusan masalah', 'tujuan', 'manfaat', 'batasan'],
            2: ['tinjauan pustaka', 'kajian pustaka', 'landasan teori', 'penelitian terkait', 'kerangka pikir'],
            3: ['metodologi', 'metode', 'desain penelitian', 'metode pengumpulan data', 'metode analisis', 'tools'],
            4: ['analisis', 'pembahasan', 'hasil', 'perancangan sistem', 'perancangan interface'],
            5: ['implementasi', 'pengujian', 'hasil pengujian', 'pembahasan'],
            6: ['kesimpulan', 'penutup', 'saran', 'rekomendasi']
        }

        for zone_id, zone in self.structure.zones.items():
            text = zone.content.lower()

            # Enhanced chapter pattern detection
            for pattern in self.INDONESIAN_PATTERNS['chapters']:
                match = re.search(pattern, zone.content, re.IGNORECASE)
                if match:
                    chapter_info = {
                        'zone_id': zone_id,
                        'text': zone.content,
                        'pattern': pattern,
                        'chapter_num': self._extract_chapter_number(zone.content),
                        'title': self._extract_chapter_title(zone.content),
                        'standard_mapping': self._map_to_standard_chapter(zone.content, STANDARD_CHAPTERS)
                    }
                    patterns_found['chapters'].append(chapter_info)
                    break

            # Enhanced subsection pattern detection
            for pattern in self.INDONESIAN_PATTERNS['subsections']:
                match = re.search(pattern, zone.content, re.IGNORECASE)
                if match:
                    subsection_info = {
                        'zone_id': zone_id,
                        'text': zone.content,
                        'pattern': pattern,
                        'hierarchy': self._determine_subsection_hierarchy(zone.content),
                        'parent_chapter': self._find_parent_chapter(zone, patterns_found['chapters']),
                        'semantic_type': self._classify_subsection_type(zone.content)
                    }
                    patterns_found['subsections'].append(subsection_info)
                    break

            # Enhanced content placeholder detection
            for pattern in self.INDONESIAN_PATTERNS['content_placeholders']:
                if re.search(pattern, zone.content, re.IGNORECASE):
                    placeholder_info = {
                        'zone_id': zone_id,
                        'text': zone.content,
                        'pattern': pattern,
                        'placeholder_type': self._classify_placeholder_type(zone.content),
                        'estimated_length': self._estimate_content_length(zone.content),
                        'content_type': self._infer_content_type(zone.content)
                    }
                    patterns_found['placeholders'].append(placeholder_info)
                    break

            # Front/back matter detection
            for section_type, patterns in [('front_matter', self.INDONESIAN_PATTERNS['front_matter']),
                                         ('back_matter', self.INDONESIAN_PATTERNS['back_matter'])]:
                for pattern in patterns:
                    if re.search(pattern, zone.content, re.IGNORECASE):
                        matter_info = {
                            'zone_id': zone_id,
                            'text': zone.content,
                            'pattern': pattern,
                            'section_type': section_type
                        }
                        patterns_found[section_type].append(matter_info)
                        break

        self.structure.academic_patterns = patterns_found

    def _extract_chapter_number(self, text: str) -> Optional[int]:
        """Extract chapter number from various formats"""
        # Roman numerals
        roman_map = {'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5, 'VI': 6, 'VII': 7}
        for roman, num in roman_map.items():
            if f'BAB {roman}' in text.upper():
                return num

        # Arabic numerals
        import re
        match = re.search(r'BAB\s+(\d+)', text.upper())
        if match:
            return int(match.group(1))

        # Chapter keyword
        match = re.search(r'Chapter\s+(\d+)', text, re.IGNORECASE)
        if match:
            return int(match.group(1))

        return None

    def _extract_chapter_title(self, text: str) -> str:
        """Extract chapter title from various formats"""
        # Remove chapter numbering prefix
        text = re.sub(r'^BAB\s+[IVXLCDM\d]+\s*[:-]?\s*', '', text, flags=re.IGNORECASE)
        text = re.sub(r'^Chapter\s+\d+\s*[:-]?\s*', '', text, flags=re.IGNORECASE)

        return text.strip()

    def _map_to_standard_chapter(self, text: str, standard_chapters: Dict[int, List[str]]) -> Optional[int]:
        """Map detected chapter to standard research structure"""
        text_lower = text.lower()

        for chapter_num, keywords in standard_chapters.items():
            if any(keyword in text_lower for keyword in keywords):
                return chapter_num

        return None

    def _determine_subsection_hierarchy(self, text: str) -> int:
        """Determine subsection hierarchy level"""
        if re.match(r'^\d+\.\d+\.\d+\.\d+', text):  # 1.1.1.1
            return 4
        elif re.match(r'^\d+\.\d+\.\d+', text):  # 1.1.1
            return 3
        elif re.match(r'^\d+\.\d+', text):  # 1.1
            return 2
        elif re.match(r'^[A-Za-z][\.\)]', text):  # A. or a)
            return 2
        else:
            return 1

    def _find_parent_chapter(self, zone, detected_chapters: List[Dict]) -> Optional[str]:
        """Find parent chapter for a subsection"""
        # Find the closest chapter before this zone
        zone_start = zone.start_paragraph
        closest_chapter = None
        min_distance = float('inf')

        for chapter in detected_chapters:
            chapter_zone = self.structure.zones.get(chapter['zone_id'])
            if chapter_zone and chapter_zone.start_paragraph < zone_start:
                distance = zone_start - chapter_zone.start_paragraph
                if distance < min_distance:
                    min_distance = distance
                    closest_chapter = chapter

        return closest_chapter['zone_id'] if closest_chapter else None

    def _classify_subsection_type(self, text: str) -> str:
        """Classify subsection content type"""
        text_lower = text.lower()

        type_mappings = {
            'latar_belakang': ['latar', 'background'],
            'rumusan_masalah': ['rumusan', 'masalah', 'problem'],
            'tujuan': ['tujuan', 'purpose', 'objective'],
            'manfaat': ['manfaat', 'benefit', 'advantage'],
            'batasan': ['batasan', 'limitation', 'scope'],
            'landasan_teori': ['landasan', 'teori', 'theory'],
            'penelitian_terkait': ['terkait', 'related', 'literature'],
            'kerangka_pikir': ['kerangka', 'pikir', 'framework'],
            'metode': ['metode', 'method', 'technique'],
            'analisis': ['analisis', 'analysis'],
            'hasil': ['hasil', 'result', 'finding'],
            'kesimpulan': ['kesimpulan', 'conclusion'],
            'saran': ['saran', 'recommendation', 'suggestion']
        }

        for content_type, keywords in type_mappings.items():
            if any(keyword in text_lower for keyword in keywords):
                return content_type

        return 'general'

    def _classify_placeholder_type(self, text: str) -> str:
        """Classify placeholder content type"""
        if 'tabel' in text.lower() or 'table' in text.lower():
            return 'table'
        elif 'gambar' in text.lower() or 'figure' in text.lower():
            return 'figure'
        elif 'persamaan' in text.lower() or 'equation' in text.lower():
            return 'equation'
        elif 'kode' in text.lower() or 'code' in text.lower():
            return 'code'
        else:
            return 'text'

    def _estimate_content_length(self, text: str) -> int:
        """Estimate expected content length from placeholder"""
        # Look for length indicators
        import re
        length_match = re.search(r'(\d+)\s*-\s*(\d+)\s*kata', text, re.IGNORECASE)
        if length_match:
            min_words = int(length_match.group(1))
            max_words = int(length_match.group(2))
            return (min_words + max_words) // 2

        # Default estimates based on content type
        if 'pendahuluan' in text.lower() or 'latar' in text.lower():
            return 300  # Background sections are longer
        elif 'kesimpulan' in text.lower():
            return 200  # Conclusions are shorter
        else:
            return 150  # Standard paragraph length

    def _infer_content_type(self, text: str) -> str:
        """Infer content type from placeholder text"""
        text_lower = text.lower()

        if 'jelaskan' in text_lower or 'uraikan' in text_lower:
            return 'explanatory'
        elif 'deskripsikan' in text_lower:
            return 'descriptive'
        elif 'analisis' in text_lower or 'bahas' in text_lower:
            return 'analytical'
        elif 'hasil' in text_lower or 'temuan' in text_lower:
            return 'results'
        elif 'kesimpulan' in text_lower:
            return 'conclusion'
        else:
            return 'narrative'

    def _calculate_confidence_score(self) -> None:
        """Calculate overall analysis confidence"""
        score = 0.0
        max_score = 100.0

        # Chapters found (30% weight)
        chapters_found = len(self.structure.academic_patterns['chapters'])
        score += min(chapters_found * 10, 30)  # Up to 30 points

        # Subsections found (20% weight)
        subsections_found = len(self.structure.academic_patterns['subsections'])
        score += min(subsections_found * 2, 20)  # Up to 20 points

        # Placeholders identified (15% weight)
        placeholders_found = len(self.structure.academic_patterns['placeholders'])
        score += min(placeholders_found, 15)  # Up to 15 points

        # Style rules extracted (15% weight)
        style_rules_count = len(self.structure.style_rules)
        score += min(style_rules_count * 3, 15)  # Up to 15 points

        # Zones classified (20% weight)
        content_zones = len([z for z in self.structure.zones.values()
                           if z.zone_type in [ZoneType.CONTENT, ZoneType.HEADER]])
        score += min(content_zones * 0.5, 20)  # Up to 20 points

        self.structure.confidence_score = score / max_score

    def _get_outline_level(self, para) -> Optional[int]:
        """Get outline level from paragraph style"""
        if not para.style:
            return None

        try:
            outline_level_attr = para.style.element.xpath('.//w:outlineLvl/@w:val')
            if outline_level_attr:
                return int(outline_level_attr[0])
        except Exception:
            pass

        # Fallback: infer from style name
        style_name = para.style.name.lower()
        if 'heading' in style_name:
            try:
                # Extract number from "Heading X"
                match = re.search(r'heading\s+(\d+)', style_name, re.IGNORECASE)
                if match:
                    return int(match.group(1)) - 1  # Heading 1 = level 0
            except:
                pass

        return None

    def _classify_zone_type(self, text: str) -> ZoneType:
        """Classify zone type based on content analysis"""
        text_upper = text.upper()
        text_lower = text.lower()

        # Check for headers
        if self._is_header_text(text):
            return ZoneType.HEADER

        # Check for front/back matter
        for patterns in [self.INDONESIAN_PATTERNS['front_matter'],
                        self.INDONESIAN_PATTERNS['back_matter']]:
            for pattern in patterns:
                if re.search(pattern, text, re.IGNORECASE):
                    return ZoneType.FRONT_MATTER if patterns == self.INDONESIAN_PATTERNS['front_matter'] else ZoneType.BACK_MATTER

        # Check for placeholders
        for pattern in self.INDONESIAN_PATTERNS['content_placeholders']:
            if re.search(pattern, text, re.IGNORECASE):
                return ZoneType.PLACEHOLDER

        # Check for structural elements (tables, figures mentioned in text)
        if any(word in text_lower for word in ['tabel', 'table', 'gambar', 'figure', 'diagram']):
            return ZoneType.STRUCTURAL

        # Default to content
        return ZoneType.CONTENT

    def _is_header_text(self, text: str) -> bool:
        """Check if text appears to be a header"""
        # Check for chapter patterns
        for pattern in self.INDONESIAN_PATTERNS['chapters']:
            if re.match(pattern, text, re.IGNORECASE):
                return True

        # Check for subsection patterns
        for pattern in self.INDONESIAN_PATTERNS['subsections']:
            if re.match(pattern, text, re.IGNORECASE):
                return True

        # Check for outline level in style (handled elsewhere)
        # Check for short, capitalized text
        if len(text) < 100 and text == text.upper() and len(text.split()) < 10:
            return True

        return False

    def _extract_header_info(self, text: str) -> Dict[str, Any]:
        """Extract header-specific information"""
        info = {'level': 0, 'numbering': None, 'title': text}

        # Try to extract numbering
        for pattern in self.INDONESIAN_PATTERNS['chapters'] + self.INDONESIAN_PATTERNS['subsections']:
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                if len(match.groups()) >= 2:
                    info['numbering'] = f"{match.group(1)}.{match.group(2)}"
                    info['title'] = match.group(-1)  # Last group is usually the title
                break

        return info

    def _analyze_paragraph_style(self, para, style_rules: Dict) -> None:
        """Analyze paragraph style for rule extraction"""
        if not para.style:
            return

        style_name = para.style.name

        # Font analysis
        if para.style.font:
            font_info = {
                'name': para.style.font.name or 'Calibri',
                'size': para.style.font.size.pt if para.style.font.size else 12,
                'bold': para.style.font.bold or False,
                'italic': para.style.font.italic or False
            }
            style_rules['fonts'][style_name] = font_info

        # Paragraph formatting
        if para.style.paragraph_format:
            para_info = {
                'alignment': str(para.style.paragraph_format.alignment or 'LEFT'),
                'line_spacing': para.style.paragraph_format.line_spacing or 1.0,
                'space_before': para.style.paragraph_format.space_before.pt if para.style.paragraph_format.space_before else 0,
                'space_after': para.style.paragraph_format.space_after.pt if para.style.paragraph_format.space_after else 0,
                'first_line_indent': para.style.paragraph_format.first_line_indent.pt if para.style.paragraph_format.first_line_indent else 0
            }
            style_rules['spacing'][style_name] = para_info

    def _analyze_table_style(self, table, style_rules: Dict) -> None:
        """Analyze table style for rule extraction"""
        table_info = {
            'rows': len(table.rows),
            'cols': len(table.columns),
            'style': getattr(table.style, 'name', 'Normal Table') if table.style else 'Normal Table'
        }

        # Analyze cell formatting from first few cells
        cell_styles = []
        for row in table.rows[:3]:  # Check first 3 rows
            for cell in row.cells[:3]:  # Check first 3 cells
                if cell.paragraphs and cell.paragraphs[0].style:
                    cell_styles.append(cell.paragraphs[0].style.name)

        if cell_styles:
            table_info['cell_styles'] = list(set(cell_styles))

        style_rules['tables'][f'table_{len(style_rules["tables"])}'] = table_info

    def get_content_zones_for_insertion(self) -> Dict[str, ContentZone]:
        """Get zones suitable for content insertion"""
        return {
            zone_id: zone for zone_id, zone in self.structure.zones.items()
            if zone.zone_type in [ZoneType.CONTENT, ZoneType.PLACEHOLDER]
        }

    def get_header_zones(self) -> Dict[str, ContentZone]:
        """Get header zones for reference"""
        return {
            zone_id: zone for zone_id, zone in self.structure.zones.items()
            if zone.zone_type == ZoneType.HEADER
        }

    def get_style_rules_for_zone(self, zone: ContentZone) -> Dict[str, Any]:
        """Get applicable style rules for a zone"""
        zone_style = zone.style_info.get('style_name', 'Normal')
        return self.structure.style_rules.get(zone_style, {})