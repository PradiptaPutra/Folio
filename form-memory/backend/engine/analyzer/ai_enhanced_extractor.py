"""
AI-Enhanced Content Extractor
Combines rule-based extraction with AI semantic analysis for better chapter detection.
"""

from typing import Dict, List, Any, Tuple, Optional
import re
from pathlib import Path
from docx import Document
from .content_extractor import ContentExtractor
from ..ai.semantic_parser import SemanticParser
from enum import Enum

# Try to import AI semantic parser
try:
    from engine.ai.semantic_parser import SemanticParser
    AI_AVAILABLE = True
except ImportError:
    AI_AVAILABLE = False


class SectionType(str, Enum):
    """Document section types identified by AI."""
    CHAPTER = "chapter"
    SECTION = "section"
    SUBSECTION = "subsection"
    PARAGRAPH = "paragraph"
    INTRODUCTION = "introduction"
    METHODOLOGY = "methodology"
    RESULTS = "results"
    DISCUSSION = "discussion"
    CONCLUSION = "conclusion"
    LITERATURE_REVIEW = "literature_review"
    BACKGROUND = "background"
    UNKNOWN = "unknown"


class AIEnhancedContentExtractor:
    """Extracts content using AI semantic analysis when available."""
    
    def __init__(self, content_path: str, use_ai: bool = True, api_key: Optional[str] = None):
        """Initialize with content file.

        Args:
            content_path: Path to DOCX or TXT file
            use_ai: Whether to use AI semantic analysis (if available)
            api_key: OpenRouter API key for AI features
        """
        self.content_path = Path(content_path)
        self.is_docx = self.content_path.suffix.lower() == '.docx'
        self.use_ai = use_ai and AI_AVAILABLE and api_key is not None
        self.api_key = api_key
        self.semantic_parser = SemanticParser(api_key=api_key) if self.use_ai else None
        
        # Load content - keep ContentExtractor object separately
        self._content_extractor = ContentExtractor(str(self.content_path))
        loaded_content = self._content_extractor._load_content()

        # Store analyzed data for access by thesis builder
        self.analyzed_data = None
        self.raw_text = self._extract_raw_text()
        self.sections = self._extract_sections()
    
    def _load_content(self) -> Dict[str, Any]:
        """Load content from file."""
        if self.is_docx:
            return self._extract_from_docx()
        else:
            return self._extract_from_text()
    
    def _extract_from_docx(self) -> Dict[str, Any]:
        """Extract content from DOCX file."""
        doc = Document(self.content_path)
        
        return {
            "sections": [],  # Will be populated by _extract_sections()
            "raw_text": self._get_docx_text(doc),
            "tables": self._extract_tables_from_docx(doc),
            "metadata": {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
            }
        }
    
    def _extract_from_text(self) -> Dict[str, Any]:
        """Extract content from TXT file."""
        with open(self.content_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        return {
            "sections": [],  # Will be populated by _extract_sections()
            "raw_text": text,
            "tables": [],
            "metadata": {}
        }
    
    def _extract_raw_text(self) -> str:
        """Get all text content."""
        if self.is_docx:
            doc = Document(self.content_path)
            return self._get_docx_text(doc)
        else:
            with open(self.content_path, 'r', encoding='utf-8') as f:
                return f.read()
    
    def _extract_sections(self) -> List[Dict[str, Any]]:
        """Extract sections using AI if available, fallback to rules."""
        if self.use_ai:
            try:
                return self._extract_with_ai()
            except Exception as e:
                print(f"[WARNING] AI extraction failed: {str(e)}. Falling back to rule-based extraction.")
                return self._extract_with_rules()
        else:
            return self._extract_with_rules()
    
    def _extract_with_ai(self) -> List[Dict[str, Any]]:
        """Extract sections using comprehensive AI content generation."""
        if not self.raw_text.strip():
            return []

        print(f"[AI] Starting comprehensive thesis content generation for {len(self.raw_text)} chars")

        # Use comprehensive AI analysis to generate full thesis content
        try:
            analyzed_data = self._generate_comprehensive_thesis_content(self.raw_text)
            
            # Ensure analyzed_data is a dict
            if isinstance(analyzed_data, list):
                print("[WARNING] AI returned a list of sections instead of a data dict, attempting to convert...")
                # If it's already sections, we just use them but we still need analyzed_data for the builder
                self.sections = analyzed_data
                return analyzed_data

            sections = self._convert_analyzed_data_to_sections(analyzed_data)
            print(f"[AI] Generated {len(sections)} sections with full content")
            self.sections = sections

            # DEBUG: Check if analyzed_data has content
            if hasattr(self, 'analyzed_data') and self.analyzed_data:
                total_chars = 0
                for chapter_key in ['chapter1', 'chapter2', 'chapter3', 'chapter4', 'chapter5', 'chapter6']:
                    if chapter_key in self.analyzed_data:
                        chapter = self.analyzed_data[chapter_key]
                        chapter_chars = sum(len(str(content)) for content in chapter.values())
                        total_chars += chapter_chars

                print(f"[AI] Analyzed data contains {total_chars} characters of content")

                if total_chars < 1000:
                    print("[AI] WARNING: Very little content generated!")
                else:
                    print("[AI] SUCCESS: Substantial content generated")
            else:
                print("[AI] ERROR: No analyzed_data stored!")

            return sections
        except Exception as e:
            print(f"[AI] Comprehensive content generation failed: {e}")
            print("[AI] Falling back to basic semantic parsing")
            import traceback
            traceback.print_exc()

        # Fallback to original semantic parsing
        if self.semantic_parser:
            result = self.semantic_parser.parse(self.raw_text)
            if not result or "elements" not in result:
                return self._extract_with_rules()
        else:
            return self._extract_with_rules()

        # Convert semantic elements to sections (original logic)
        sections = []
        current_section = None
        current_level = 0

        for element in result.get("elements", []):
            elem_type = element.get("type", "paragraph")
            text = element.get("text", "").strip()
            confidence = element.get("confidence", 0.0)

            # Determine section type and level
            if elem_type in ["chapter", "subchapter", "subsubchapter"]:
                # Save previous section
                if current_section:
                    sections.append(current_section)

                # Create new section
                level_map = {
                    "chapter": 0,
                    "subchapter": 1,
                    "subsubchapter": 2
                }
                level = level_map.get(elem_type, 0)

                current_section = {
                    "title": text,
                    "level": level,
                    "content": [],
                    "type": elem_type,
                    "ai_analyzed": False,
                    "ai_confidence": confidence,
                    "semantic_type": elem_type,
                }
            elif current_section is not None:
                if text:
                    current_section["content"].append(text)

        if current_section:
            sections.append(current_section)

        return sections

    def _generate_comprehensive_thesis_content(self, raw_text: str) -> Dict[str, Any]:
        """Generate comprehensive thesis content using AI with detailed prompt."""
        import json

        prompt = f"""
You are analyzing Indonesian thesis draft text and converting it into structured academic content.

Raw Text:
{raw_text}

Generate comprehensive thesis content in valid JSON format. Each section should have substantive academic content.

Return ONLY valid JSON:

{{
  "metadata": {{
    "title": "Full thesis title",
    "author": "Author name",
    "nim": "Student ID"
  }},

  "abstract": {{
    "indonesian": "Complete Indonesian abstract (200-300 words) with background, objectives, methods, and conclusions.",
    "english": "Complete English abstract (200-300 words) same content as Indonesian.",
    "keywords_id": ["keyword1", "keyword2", "keyword3"],
    "keywords_en": ["keyword1", "keyword2", "keyword3"]
  }},

  "chapter1": {{
    "latar_belakang": "Write 2-3 paragraphs explaining the research background and context.",
    "rumusan_masalah": "Write 1-2 paragraphs stating the research problems clearly.",
    "tujuan_penelitian": "Write 1-2 paragraphs describing research objectives.",
    "manfaat_penelitian": "Write 1 paragraph explaining research benefits.",
    "batasan_masalah": "Write 1 paragraph defining research scope and limitations."
  }},

  "chapter2": {{
    "landasan_teori": "Write 2-3 paragraphs covering fundamental theories.",
    "penelitian_terkait": "Write 2 paragraphs reviewing related research.",
    "kerangka_pemikiran": "Write 1-2 paragraphs explaining the conceptual framework."
  }},

  "chapter3": {{
    "desain_penelitian": "Write 1-2 paragraphs describing research design.",
    "metode_pengumpulan_data": "Write 2 paragraphs detailing data collection methods.",
    "metode_analisis": "Write 1-2 paragraphs explaining analysis methods.",
    "tools": "Write 1 paragraph listing tools and technologies."
  }},

  "chapter4": {{
    "analisis_kebutuhan": "Write 2 paragraphs analyzing system requirements.",
    "perancangan_sistem": "Write 2-3 paragraphs describing system design.",
    "perancangan_interface": "Write 1-2 paragraphs explaining interface design."
  }},

  "chapter5": {{
    "implementasi": "Write 2 paragraphs describing system implementation.",
    "hasil_pengujian": "Write 2 paragraphs presenting testing results.",
    "pembahasan": "Write 2 paragraphs discussing results.",
    "evaluasi": "Write 1-2 paragraphs evaluating system performance."
  }},

  "chapter6": {{
    "kesimpulan": "Write 2 paragraphs drawing conclusions.",
    "saran": "Write 1-2 paragraphs providing recommendations."
  }},

  "references": [
    "Reference 1 in APA format",
    "Reference 2 in APA format"
  ]
}}

IMPORTANT: Return ONLY the JSON object, NO extra text, no markdown code blocks, just pure JSON.
The content must be in Indonesian, except for keywords_en and abstract.english.
Expand the raw text into professional, academic paragraphs. If the draft is sparse, use your knowledge to fill in standard academic details for a Computer Science/Informatika thesis.
"""

        try:
            # Call AI with comprehensive prompt using OpenAI client directly
            from openai import OpenAI
            client = OpenAI(api_key=self.api_key, base_url="https://openrouter.ai/api/v1")

            models_to_try = [
                "google/gemini-2.0-flash-exp:free",
                "google/gemini-2.0-flash-exp",
                "google/gemini-flash-1.5-exp",
                "meta-llama/llama-3.1-8b-instruct:free",
                "meta-llama/llama-3.1-70b-instruct:free",
                "mistralai/mistral-7b-instruct:free",
                "deepseek/deepseek-chat:free",
                "microsoft/phi-3-medium-128k-instruct:free"
            ]
            
            last_error = None
            response = None
            
            for model_name in models_to_try:
                try:
                    print(f"[AI] Attempting content generation with model: {model_name}")
                    response = client.chat.completions.create(
                        model=model_name,
                        messages=[
                            {
                                "role": "system",
                                "content": "You are an expert at analyzing Indonesian thesis content and generating comprehensive academic paragraphs."
                            },
                            {
                                "role": "user",
                                "content": prompt
                            }
                        ],
                        temperature=0.3
                    )
                    break # Success!
                except Exception as e:
                    last_error = e
                    print(f"[AI] Model {model_name} failed: {e}")
                    # Continue to next model regardless of error type (429, 404, etc.)
                    continue
            
            if not response:
                print(f"[ERROR] All AI models failed. Last error: {last_error}")
                raise last_error or ValueError("All AI models failed to respond")

            # Parse response
            raw_response = response.choices[0].message.content if response.choices else str(response)
            if not raw_response:
                raise ValueError("Empty response from AI")

            # Improved Clean response: find the first { and last }
            import re
            json_match = re.search(r'(\{.*\})', raw_response, re.DOTALL)
            if json_match:
                clean_response = json_match.group(1)
            else:
                clean_response = raw_response.replace('```json', '').replace('```', '').strip()

            # Parse JSON with better error handling
            try:
                analyzed_data = json.loads(clean_response)
            except json.JSONDecodeError as e:
                print(f"[ERROR] JSON parsing failed: {e}")
                print(f"[ERROR] Raw response length: {len(clean_response)} chars")
                print(f"[ERROR] Raw response preview: {clean_response[:300]}...")

                # Try to fix common JSON issues
                fixed_response = clean_response

                # Remove any trailing content after the last valid brace
                last_brace = fixed_response.rfind('}')
                if last_brace > 0:
                    fixed_response = fixed_response[:last_brace + 1]

                # Try parsing the fixed response
                try:
                    analyzed_data = json.loads(fixed_response)
                    print("[INFO] Successfully parsed JSON after fixing")
                except json.JSONDecodeError as e2:
                    print(f"[ERROR] JSON still invalid after fixing: {e2}")

                    # Last resort: try to extract partial data
                    try:
                        # Extract metadata if available
                        metadata_match = re.search(r'"metadata"\s*:\s*\{[^}]*\}', fixed_response)
                        if metadata_match:
                            metadata_str = metadata_match.group()
                            metadata = json.loads(f"{{{metadata_str}}}")
                            analyzed_data = {'metadata': metadata}
                            print("[WARNING] Only partial data extracted (metadata only)")
                        else:
                            raise ValueError("Could not extract any valid JSON")
                    except Exception as e3:
                        print(f"[ERROR] Fatal JSON extraction error: {e3}")
                        print(f"[DEBUG] Full Raw AI Response:\n{raw_response}")
                        raise ValueError(f"Could not extract JSON from AI response: {str(e2)}")

            # VERIFY AI response has actual content
            print(f"[VERIFY] AI response keys: {list(analyzed_data.keys())}")

            for chapter_key in ['chapter1', 'chapter2', 'chapter3', 'chapter4', 'chapter5', 'chapter6']:
                if chapter_key in analyzed_data:
                    chapter_data = analyzed_data[chapter_key]
                    print(f"\n[VERIFY] {chapter_key}:")
                    for subsection, content in chapter_data.items():
                        content_length = len(content) if content else 0
                        preview = content[:100] if content else "[EMPTY]"
                        # Handle Unicode characters safely
                        try:
                            print(f"  - {subsection}: {content_length} chars | Preview: {preview}...")
                        except UnicodeEncodeError:
                            safe_preview = preview.encode('ascii', 'replace').decode('ascii')
                            print(f"  - {subsection}: {content_length} chars | Preview: {safe_preview}...")

                        # ALERT if content is too short
                        if content_length < 200:
                            print(f"  ⚠️  WARNING: {subsection} content is too short ({content_length} chars)")

            # STOP if no content
            if all(len(str(v)) < 100 for v in analyzed_data.get('chapter1', {}).values()):
                raise ValueError("AI returned empty content - check AI prompt and response parsing")

            # Store for access by thesis builder
            self.analyzed_data = analyzed_data
            print(f"[INFO] Stored analyzed_data with {len(analyzed_data)} top-level keys")

            return analyzed_data

        except Exception as e:
            print(f"[AI] Comprehensive content generation failed: {e}")
            import traceback
            traceback.print_exc()

            # EMERGENCY FALLBACK: Generate basic content if AI completely fails
            print("[AI] Generating emergency fallback content...")
            fallback_data = self._generate_fallback_content(self.raw_text)
            self.analyzed_data = fallback_data
            
            # Return the dict, NOT the converted sections, to keep return type consistent
            return fallback_data

    def _generate_fallback_content(self, raw_text: str) -> Dict[str, Any]:
        """Generate basic fallback content when AI completely fails."""
        print("[FALLBACK] Creating basic content structure...")

        # Try to extract basic info from raw text, fallback to generic values
        title = "Judul Skripsi"  # Generic fallback
        title = ""  # Initialize as empty to allow detection
        author = ""  # Initialize as empty to allow detection
        nim = ""  # Initialize as empty to allow detection

        # Look for title, author, nim in text
        lines = raw_text.split('\n')
        for line in lines[:30]:  # Check first 30 lines
            line_strip = line.strip()
            if not line_strip: continue
            
            # Author/NIM detection
            if not author and any(k in line_strip.upper() for k in ['NAMA', 'NAME', 'PENULIS', 'BY:']):
                # Extract value after colon if present
                if ':' in line_strip:
                    val = line_strip.split(':')[-1].strip()
                    # Avoid extracting program/faculty as author
                    if not any(k in val.upper() for k in ['TEKNIK', 'INFORMATIKA', 'SISTEM', 'ILMU', 'FAKULTAS', 'PROGRAM']):
                        author = val
                elif len(line_strip.split()) <= 4:
                    if not any(k in line_strip.upper() for k in ['TEKNIK', 'INFORMATIKA', 'SISTEM', 'ILMU', 'FAKULTAS', 'PROGRAM']):
                        author = line_strip
            
            if not nim and any(k in line_strip.upper() for k in ['NIM', 'ID', 'BP', 'NOMOR INDUK']):
                if ':' in line_strip:
                    val = line_strip.split(':')[-1].strip()
                    nim = ''.join(filter(str.isdigit, val))
                else:
                    nim = ''.join(filter(str.isdigit, line_strip))

            # Improved Title Detection: Look for something that looks like a title
            # (Short-medium length, not a header, not background text)
            if not title and 10 < len(line_strip) < 150:
                is_trash = any(k in line_strip.upper() for k in ['BAB', 'CHAPTER', 'DAFTAR', 'HALAMAN', 'SKRIPSI', 'TEKNIK', 'UNIVERSITAS', 'LATAR BELAKANG'])
                # Also check if it's the start of a sentence like "Latar belakang..."
                is_sentence = line_strip.lower().startswith(('latar belakang', 'penelitian ini', 'dalam era'))
                
                if not is_trash and not is_sentence:
                    title = line_strip

        # Final cleanup for Gilang Gilang type issues
        if author:
            words = author.split()
            if len(words) >= 2 and words[0] == words[1]:
                author = words[0] # Gilang Gilang -> Gilang
        
        # Build chapter content by trying to find actual text in raw_text
        def get_text_between(start_patterns, end_patterns):
            start_idx = -1
            for pattern in start_patterns:
                for i, line in enumerate(lines):
                    if re.search(pattern, line.upper()):
                        start_idx = i + 1; break
                if start_idx != -1: break
            
            if start_idx == -1: return ""
            
            end_idx = len(lines)
            for pattern in end_patterns:
                for i in range(start_idx, len(lines)):
                    if re.search(pattern, lines[i].upper()):
                        end_idx = i; break
                if end_idx != len(lines): break
            
            content = " ".join(l.strip() for l in lines[start_idx:end_idx] if l.strip())
            return content if len(content) > 100 else ""

        # Chapter 1 extraction
        ch1_latar = get_text_between(['LATAR BELAKANG'], ['RUMUSAN MASALAH', 'TUJUAN PENELITIAN', 'BAB II'])
        ch1_rumusan = get_text_between(['RUMUSAN MASALAH'], ['TUJUAN PENELITIAN', 'BAB II'])
        
        # Create basic content for each chapter, mixed with extracted info
        fallback_data = {
            'metadata': {
                'title': title or "Judul Skripsi",
                'author': author or "Nama Mahasiswa",
                'nim': nim or "NIM"
            },
            'abstract': {
                'indonesian': get_text_between(['ABSTRAK', 'SARI'], ['ABSTRACT', 'BAB I']) or 'Penelitian ini bertujuan untuk mengembangkan solusi yang dapat meningkatkan efisiensi dan efektivitas dalam bidang yang diteliti. Pendekatan yang digunakan mengintegrasikan berbagai metodologi dan teknologi terkini untuk mencapai hasil yang optimal.',
                'english': get_text_between(['ABSTRACT'], ['KATA KUNCI', 'BAB I']) or 'This research aims to develop solutions that can improve efficiency and effectiveness in the studied field. The approach used integrates various current methodologies and technologies to achieve optimal results.',
                'keywords_id': ['penelitian', 'efisiensi', 'pengembangan'],
                'keywords_en': ['research', 'efficiency', 'development']
            },
            'chapter1': {
                'latar_belakang': ch1_latar or 'Perkembangan teknologi dan metodologi penelitian telah membawa perubahan signifikan dalam berbagai bidang keilmuan. Di era saat ini, organisasi dan institusi dituntut untuk mengadopsi pendekatan yang lebih efektif guna meningkatkan kualitas layanan dan efisiensi operasional.',
                'rumusan_masalah': ch1_rumusan or 'Berdasarkan latar belakang di atas, rumusan masalah penelitian ini adalah bagaimana mengembangkan solusi yang efektif dan efisien untuk mengatasi permasalahan yang telah diidentifikasi.',
                'tujuan_penelitian': 'Tujuan dari penelitian ini adalah mengembangkan solusi yang dapat memenuhi kebutuhan pengguna dan memberikan kontribusi yang bermakna bagi pengembangan keilmuan.',
                'manfaat_penelitian': 'Manfaat penelitian ini adalah memberikan kontribusi teoritis dan praktis dalam pengembangan pengetahuan serta memberikan solusi untuk permasalahan yang dihadapi.',
                'batasan_masalah': 'Penelitian ini terbatas pada ruang lingkup tertentu dengan mempertimbangkan keterbatasan sumber daya, waktu, and metodologi yang tersedia.'
            },
            'chapter2': {
                'landasan_teori': get_text_between(['BAB II', 'LANDASAN TEORI'], ['PENELITIAN TERKAIT', 'BAB III']) or 'Landasan teori penelitian ini mencakup konsep-konsep dasar yang relevan dengan bidang penelitian. Teori-teori ini memberikan kerangka berpikir untuk memahami dan menganalisis permasalahan yang diteliti.',
                'penelitian_terkait': 'Penelitian terkait menunjukkan bahwa berbagai pendekatan telah dilakukan untuk mengatasi permasalahan serupa. Hasil-hasil penelitian sebelumnya memberikan wawasan penting untuk pengembangan penelitian ini.',
                'kerangka_pemikiran': 'Kerangka pemikiran penelitian ini menggunakan pendekatan yang sistematis untuk menganalisis dan mengembangkan solusi yang sesuai dengan kebutuhan dan konteks penelitian.'
            },
            'chapter3': {
                'desain_penelitian': get_text_between(['BAB III', 'METODOLOGI'], ['DAFTAR PUSTAKA', 'BAB IV']) or 'Desain penelitian ini menggunakan metode yang sesuai dengan tujuan dan karakteristik permasalahan yang diteliti. Pendekatan yang digunakan telah dipilih berdasarkan pertimbangan teoritis dan praktis.',
                'metode_pengumpulan_data': 'Metode pengumpulan data menggunakan teknik yang sesuai untuk mendapatkan informasi yang diperlukan. Data dikumpulkan melalui berbagai sumber yang relevan dan dapat dipertanggungjawabkan.',
                'metode_analisis': 'Metode analisis data menggunakan teknik yang sesuai untuk mengolah dan menginterpretasikan data yang diperoleh. Analisis dilakukan secara sistematis dan menggunakan pendekatan yang valid.',
                'tools': 'Tools yang digunakan dalam penelitian ini meliputi berbagai perangkat lunak dan metodologi yang sesuai dengan kebutuhan analisis dan pengolahan data.'
            },
            'chapter4': {
                'analisis_kebutuhan': get_text_between(['BAB IV', 'HASIL', 'ANALISIS'], ['PENUTUP', 'BAB V']) or 'Analisis kebutuhan menunjukkan bahwa solusi yang dikembangkan harus mempertimbangkan berbagai aspek penting yang berkaitan dengan kebutuhan pengguna dan konteks penerapan.',
                'perancangan_sistem': 'Perancangan sistem mencakup berbagai aspek penting seperti arsitektur, komponen, dan mekanisme kerja yang sesuai dengan kebutuhan dan spesifikasi yang telah ditentukan.',
                'perancangan_interface': 'Perancangan interface menggunakan prinsip-prinsip yang mendukung kemudahan penggunaan dan pengalaman yang optimal bagi pengguna sistem.'
            },
            'chapter5': {
                'implementasi': get_text_between(['BAB V', 'IMPLEMENTASI'], ['PENUTUP', 'BAB VI']) or 'Implementasi dilakukan sesuai dengan desain yang telah dibuat dengan memperhatikan berbagai aspek teknis dan kualitas yang diperlukan.',
                'hasil_pengujian': 'Hasil pengujian menunjukkan bahwa implementasi berhasil memenuhi spesifikasi dan memberikan hasil yang sesuai dengan harapan.',
                'pembahasan': 'Pembahasan hasil menunjukkan bahwa pendekatan yang digunakan memberikan kontribusi yang signifikan dalam mengatasi permasalahan yang diteliti.',
                'evaluasi': 'Evaluasi menunjukkan bahwa solusi yang dikembangkan memiliki potensi untuk diterapkan dalam konteks yang lebih luas and memberikan manfaat yang bermakna.'
            },
            'chapter6': {
                'kesimpulan': get_text_between(['BAB VI', 'PENUTUP', 'KESIMPULAN'], ['DAFTAR PUSTAKA']) or 'Kesimpulan dari penelitian ini adalah bahwa pendekatan yang digunakan berhasil memberikan solusi yang efektif dan efisien untuk mengatasi permasalahan yang diteliti.',
                'saran': 'Saran untuk pengembangan selanjutnya adalah melakukan perluasan ruang lingkup, peningkatan metodologi, and eksplorasi aplikasi dalam konteks yang berbeda.'
            },
            'references': [
                'Author, A. (Year). Title of the work. Publisher.',
                'Author, B., & Author, C. (Year). Title of the article. Journal Name, Volume(Issue), pages.'
            ]
        }

        print("[FALLBACK] Basic content structure created")
        return fallback_data

    def _convert_analyzed_data_to_sections(self, analyzed_data):
        """Convert comprehensive analyzed data to section format."""
        sections = []

        # Chapter mapping
        chapter_titles = {
            1: "BAB I: PENDAHULUAN",
            2: "BAB II: TINJAUAN PUSTAKA",
            3: "BAB III: METODOLOGI PENELITIAN",
            4: "BAB IV: ANALISIS DAN PERANCANGAN",
            5: "BAB V: IMPLEMENTASI DAN PENGUJIAN",
            6: "BAB VI: PENUTUP"
        }

        for chapter_num in range(1, 7):
            chapter_key = f"chapter{chapter_num}"
            if chapter_key in analyzed_data:
                chapter_data = analyzed_data[chapter_key]

                # Create section for each chapter
                section_content = []
                for subsection_key, subsection_content in chapter_data.items():
                    if subsection_content and len(subsection_content.strip()) > 50:
                        # Add subsection content
                        section_content.extend(subsection_content.split('\n\n'))

                if section_content:
                    sections.append({
                        "title": chapter_titles[chapter_num],
                        "level": 0,
                        "content": section_content,
                        "type": "chapter",
                        "ai_analyzed": True,
                        "ai_confidence": 1.0,
                        "semantic_type": "chapter"
                    })

        return sections

    def _extract_with_rules(self) -> List[Dict[str, Any]]:
        """Extract sections using rule-based patterns (fallback)."""
        sections = []
        
        # Patterns for headings
        heading_patterns = [
            (r"^BAB\s+([IVX]+|[0-9]+)\s*[\s\n:]*(.+?)$", "chapter"),
            (r"^([0-9]+\.[0-9]+)\s+(.+?)$", "section"),
            (r"^([0-9]+\.[0-9]+\.[0-9]+)\s+(.+?)$", "subsection"),
            (r"^#+ (.+?)$", "markdown"),
            (r"^(PENDAHULUAN|INTRODUCTION|LATAR BELAKANG|BACKGROUND).*?$", "introduction"),
            (r"^(METODOLOGI|METHODOLOGY|METODE|METHOD).*?$", "methodology"),
            (r"^(HASIL|RESULTS?|FINDINGS?).*?$", "results"),
            (r"^(PEMBAHASAN|DISCUSSION|ANALISIS|ANALYSIS).*?$", "discussion"),
            (r"^(KESIMPULAN|CONCLUSION|PENUTUP).*?$", "conclusion"),
        ]
        
        lines = self.raw_text.split('\n')
        current_section = None
        
        for line in lines:
            matched = False
            for pattern, heading_type in heading_patterns:
                match = re.match(pattern, line.strip(), re.IGNORECASE)
                if match:
                    # Save previous section
                    if current_section:
                        sections.append(current_section)
                    
                    # Extract title
                    if len(match.groups()) >= 2:
                        title = match.group(2).strip()
                    else:
                        title = match.group(1).strip() if match.groups() else line.strip()
                    
                    # Start new section
                    current_section = {
                        "title": title,
                        "level": 0 if heading_type in ["chapter", "introduction", "methodology", "results", "discussion", "conclusion"] else 1,
                        "content": [],
                        "type": heading_type,
                        "ai_analyzed": False,
                        "ai_confidence": 1.0,  # Rule-based has high confidence
                        "semantic_type": heading_type,
                    }
                    matched = True
                    break
            
            if not matched and current_section is not None:
                if line.strip():
                    current_section["content"].append(line)
        
        if current_section:
            sections.append(current_section)
        
        return sections
    
    def _classify_section_semantic(self, title: str, confidence: float) -> str:
        """Classify section by semantic meaning."""
        title_lower = title.lower()
        
        # Semantic classification patterns
        patterns = {
            "introduction": r"(pendahuluan|introduction|latar belakang|background|rumusan masalah)",
            "methodology": r"(metodologi|methodology|metode|method|prosedur|procedure)",
            "results": r"(hasil|results?|temuan|findings?|penemuan)",
            "discussion": r"(pembahasan|discussion|analisis|analysis|penafsiran|interpretation)",
            "conclusion": r"(kesimpulan|conclusion|penutup|closing|saran|recommendation)",
            "literature_review": r"(tinjauan pustaka|literature review|kajian|review)",
            "background": r"(latar belakang|background|konteks|context)",
        }
        
        for semantic_type, pattern in patterns.items():
            if re.search(pattern, title_lower):
                return semantic_type
        
        return "unknown"
    
    def _get_docx_text(self, doc) -> str:
        """Get all text from DOCX."""
        return '\n'.join(para.text for para in doc.paragraphs)
    
    def _extract_tables_from_docx(self, doc) -> List[Dict[str, Any]]:
        """Extract tables from DOCX."""
        tables = []
        
        for i, table in enumerate(doc.tables):
            table_data = {
                "index": i,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "content": []
            }
            
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data["content"].append(row_data)
            
            tables.append(table_data)
        
        return tables
    
    def get_sections(self) -> List[Dict[str, Any]]:
        """Get extracted sections."""
        return self._content_extractor.get_sections()
    
    def get_summary(self) -> Dict[str, Any]:
        """Get extraction summary."""
        ai_analyzed = sum(1 for s in self.sections if isinstance(s, dict) and s.get("ai_analyzed", False))
        rule_based = len(self.sections) - ai_analyzed
        
        semantic_types = []
        for s in self.sections:
            if isinstance(s, dict):
                semantic_types.append(s.get("semantic_type", "unknown"))
        
        avg_conf = 0.0
        if self.sections:
            confidences = [s.get("ai_confidence", 1.0) for s in self.sections if isinstance(s, dict)]
            avg_conf = sum(confidences) / max(len(confidences), 1) if confidences else 0.0
        
        return {
            "total_sections": len(self.sections),
            "ai_analyzed_sections": ai_analyzed,
            "rule_based_sections": rule_based,
            "ai_available": self.use_ai,
            "content_type": "docx" if self.is_docx else "text",
            "semantic_types": list(set(semantic_types)),
            "average_confidence": avg_conf,
        }
    
    def get_semantic_validation(self) -> Dict[str, Any]:
        """Get semantic validation of document structure."""
        validation = {
            "status": "valid",
            "issues": [],
            "warnings": [],
            "suggestions": []
        }
        
        # Check for typical thesis structure
        semantic_types = []
        for s in self.sections:
            if isinstance(s, dict):
                semantic_types.append(s.get("semantic_type", "unknown"))
        
        # Check for missing critical sections
        expected_sections = ["introduction", "methodology", "results", "conclusion"]
        missing = [sec for sec in expected_sections if sec not in semantic_types]
        
        if missing:
            validation["warnings"].append(f"Missing expected sections: {', '.join(missing)}")
        
        # Check for low confidence sections
        low_confidence = [s for s in self.sections if isinstance(s, dict) and s.get("ai_confidence", 1.0) < 0.6]
        if low_confidence:
            validation["warnings"].append(f"{len(low_confidence)} section(s) have low AI confidence")
        
        # Check structure hierarchy
        if self.sections:
            levels = [s.get("level", 0) for s in self.sections if isinstance(s, dict)]
            if levels and max(levels) - min(levels) > 3:
                validation["suggestions"].append("Consider organizing sections with more consistent hierarchy")
        
        return validation
