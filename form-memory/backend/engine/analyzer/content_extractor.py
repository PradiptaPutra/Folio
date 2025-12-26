from typing import Dict, List, Any, Optional
import re
from pathlib import Path
from docx import Document


class ContentExtractor:
    """A lightweight, robust content extractor for DOCX/TXT inputs."""

    def __init__(self, content_path: str, debug_mapping: bool = False):
        self.content_path = Path(content_path)
        self.is_docx = self.content_path.suffix.lower() == ".docx"
        self._debug_mapping = bool(debug_mapping)
        data = self._load_content()
        self._sections = data.get("sections", [])
        self._raw_text = data.get("raw_text", "")

    def _load_content(self) -> Dict[str, Any]:
        if self.is_docx:
            return self._extract_from_docx()
        else:
            return self._extract_from_text()

    def _extract_from_docx(self) -> Dict[str, Any]:
        doc = Document(str(self.content_path))
        # No heavy cleaning; keep parser simple and robust
        cleaned_doc = doc
        sections = self._extract_sections_from_docx(cleaned_doc)
        return {
            "sections": sections,
            "raw_text": self._get_docx_text(cleaned_doc),
            "tables": self._extract_tables_from_docx(cleaned_doc),
            "metadata": {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "formatting_quality": self._assess_docx_quality(doc),
                "needs_cleanup": self._detect_formatting_issues(doc),
            },
        }

    def _extract_from_text(self) -> Dict[str, Any]:
        with open(self.content_path, 'r', encoding='utf-8') as f:
            text = f.read()
        sections = self._extract_sections_from_text(text)
        return {
            "sections": sections,
            "raw_text": text,
            "tables": [],
            "metadata": {},
        }

    def _extract_sections_from_docx(self, doc) -> List[Dict[str, Any]]:
        sections: List[Dict[str, Any]] = []
        current_section: Optional[Dict[str, Any]] = None
        for para in doc.paragraphs:
            outline_level = None
            if para.style:
                try:
                    outline_level_attr = para.style.element.xpath('.//w:outlineLvl/@w:val')
                    if outline_level_attr:
                        outline_level = int(outline_level_attr[0])
                except Exception:
                    outline_level = None

            created = False
            if outline_level is not None:
                if current_section:
                    sections.append(current_section)
                current_section = {
                    "title": para.text,
                    "level": outline_level,
                    "content": [],
                    "style": para.style.name,
                }
                created = True
            else:
                # Fallback 1: heading style name
                if para.style and para.style.name:
                    sl = para.style.name.lower()
                    lvl = None
                    if 'heading 1' in sl:
                        lvl = 1
                    elif 'heading 2' in sl:
                        lvl = 2
                    elif 'heading 3' in sl:
                        lvl = 3
                    if lvl is not None:
                        if current_section:
                            sections.append(current_section)
                        current_section = {
                            "title": para.text,
                            "level": lvl,
                            "content": [],
                            "style": para.style.name,
                        }
                        created = True
                if not created:
                    # Fallback 2: numeric heading like 1.1 Title
                    text = para.text.strip() if para.text else ""
                    if text:
                        m = re.match(r'^(\d+(?:\.\d+)*)\s+(.*)$', text)
                        if m:
                            nums = m.group(1).split('.')
                            lvl = len(nums)
                            if current_section:
                                sections.append(current_section)
                            current_section = {
                                "title": text,
                                "level": lvl,
                                "content": [],
                                "style": para.style.name if para.style else "",
                            }
                            created = True
            if not created:
                if current_section is not None and para.text.strip():
                    current_section["content"].append(para.text)
        if current_section:
            sections.append(current_section)
        return sections

    def _extract_sections_from_text(self, text: str) -> List[Dict[str, Any]]:
        sections: List[Dict[str, Any]] = []
        heading_patterns = [
            (r"^BAB\s+([IVX]+|[0-9]+)[:.]?\s*(.+?)$", "chapter"),
            (r"^([0-9]+\.[0-9]+)\s+(.+?)$", "section"),
            (r"^#+\s+(.+?)$", "markdown"),
        ]
        lines = text.split('\n')
        current_section = None
        for line in lines:
            matched = False
            for pattern, heading_type in heading_patterns:
                match = re.match(pattern, line.strip())
                if match:
                    if current_section:
                        sections.append(current_section)
                    current_section = {
                        "title": line.strip(),
                        "level": 0,
                        "content": [],
                        "type": heading_type,
                    }
                    matched = True
                    break
            if not matched and current_section is not None:
                if line.strip():
                    current_section["content"].append(line)
        if current_section:
            sections.append(current_section)
        return sections

    def _get_docx_text(self, doc) -> str:
        return '\n'.join(para.text for para in doc.paragraphs)

    def _extract_tables_from_docx(self, doc) -> List[Dict[str, Any]]:
        tables = []
        for i, table in enumerate(doc.tables):
            table_data = {
                "index": i,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "content": []
            }
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = '\n'.join(para.text for para in cell.paragraphs)
                    row_data.append(cell_text)
                table_data["content"].append(row_data)
            tables.append(table_data)
        return tables

    def get_sections(self) -> List[Dict[str, Any]]:
        return self._sections if self._sections else []

    def get_raw_text(self) -> str:
        return self._raw_text

    def get_tables(self) -> List[Dict[str, Any]]:
        content_data = self._load_content()
        return content_data.get("tables", [])

    def get_section_by_title(self, title_pattern: str) -> Optional[Dict[str, Any]]:
        for section in self._sections:
            if re.search(title_pattern, section["title"], re.IGNORECASE):
                return section
        return None

    def get_summary(self) -> Dict[str, Any]:
        summary = f"Total Content Length: {len(self.get_raw_text())} characters"
        return {
            "summary": summary.strip(),
            "sections_count": len(self.get_sections()),
            "file_type": "docx" if self.is_docx else "text",
        }

    # Lightweight, non-invasive DOCX cleanup (no destructive editing)
    def _clean_docx_formatting(self, doc) -> Document:
        return doc

    def _copy_styles_to_clean_doc(self, source_doc, target_doc) -> None:
        return None

    def _clean_paragraph_text(self, text: str) -> str:
        if not text:
            return text
        text = re.sub(r'\s+', ' ', text.strip())
        text = re.sub(r'[^\w\s\.,;:!?\-\(\)\[\]{}"\'/]', '', text)
        return text

    def _fix_indonesian_ocr_errors(self, text: str) -> str:
        return text

    def _apply_clean_formatting(self, cleaned_para, original_para) -> None:
        para_format = cleaned_para.paragraph_format
        para_format.line_spacing = 1.5
        para_format.space_before = 0
        para_format.space_after = 0
        if not self._is_heading(original_para):
            try:
                para_format.first_line_indent = 0.5  # approximate; kept for compatibility
            except Exception:
                pass
        para_format.alignment = 3

    def _is_heading(self, para) -> bool:
        if not para.style:
            return False
        style_name = para.style.name.lower()
        return 'heading' in style_name or 'title' in style_name

    def _assess_docx_quality(self, doc) -> str:
        return "unknown"

    def _detect_formatting_issues(self, doc) -> Dict[str, Any]:
        return {"needs_cleanup": False, "issues_found": [], "severity": "low"}
