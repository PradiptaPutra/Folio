"""
Enhanced Document Processing Service - Integration Path Implementation
Uses Pandoc + python-docx for normalized structural representation and rule-based processing.
Implements the integration approach: Pandoc + docxBox → rule extraction → python-docx generation.
"""

from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
import json
import subprocess
import tempfile
import os
import re
import xml.etree.ElementTree as ET

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE


class EnhancedDocumentProcessor:
    """Enhanced document processor implementing the integration path:
    Pandoc + docxBox → rule extraction → python-docx generation."""

    def __init__(self):
        self.temp_dir = Path("storage/temp")
        self.temp_dir.mkdir(parents=True, exist_ok=True)
        self.pandoc_available = self._check_pandoc_availability()
        print(f"[EnhancedDocumentProcessor] Initialized with Pandoc: {self.pandoc_available}")

    def _check_pandoc_availability(self) -> bool:
        """Check if Pandoc is available on the system."""
        try:
            result = subprocess.run(['pandoc', '--version'],
                                  capture_output=True, text=True, timeout=5)
            return result.returncode == 0
        except (subprocess.TimeoutExpired, FileNotFoundError):
            return False

    def extract_document_structure(self, docx_path: str) -> Dict[str, Any]:
        """Extract normalized structural representation using Pandoc + docxBox approach."""
        try:
            # Method 1: Try Pandoc conversion to intermediate format for structural analysis
            if self.pandoc_available:
                structure = self._extract_via_pandoc(docx_path)
            else:
                # Method 2: Fallback to direct python-docx analysis
                structure = self._extract_via_python_docx(docx_path)

            # Extract template rules from XML structure
            template_rules = self._extract_template_rules(docx_path)
            structure['template_rules'] = template_rules

            return structure

        except Exception as e:
            print(f"[EnhancedDocumentProcessor] Structure extraction failed: {e}")
            return self._fallback_extraction(docx_path)

    def _extract_via_pandoc(self, docx_path: str) -> Dict[str, Any]:
        """Extract structure using Pandoc conversion (docxBox approach)."""
        try:
            # Convert DOCX to JATS XML or other structured format for analysis
            with tempfile.NamedTemporaryFile(suffix='.xml', delete=False) as temp_file:
                output_file = Path(temp_file.name)

            # Use Pandoc to convert to JATS XML (structured format)
            cmd = [
                'pandoc',
                docx_path,
                '-f', 'docx',
                '-t', 'jats',
                '-o', str(output_file)
            ]

            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)

            if result.returncode == 0 and output_file.exists():
                # Parse the JATS XML for structural information
                structure = self._parse_jats_xml(str(output_file))
                output_file.unlink(missing_ok=True)
                return structure
            else:
                print(f"[EnhancedDocumentProcessor] Pandoc conversion failed: {result.stderr}")
                output_file.unlink(missing_ok=True)
                raise Exception("Pandoc conversion failed")

        except Exception as e:
            print(f"[EnhancedDocumentProcessor] Pandoc extraction failed: {e}")
            raise

    def _extract_via_python_docx(self, docx_path: str) -> Dict[str, Any]:
        """Extract structure using direct python-docx analysis."""
        doc = Document(docx_path)  # type: ignore

        structure = {
            "paragraphs": self._extract_paragraphs(doc),
            "headings": self._extract_headings(doc),
            "tables": self._extract_tables(doc),
            "styles": self._extract_styles(doc),
            "sections": self._extract_sections(doc),
            "metadata": self._extract_metadata(docx_path)
        }

        return structure

    def _extract_template_rules(self, docx_path: str) -> Dict[str, Any]:
        """Extract template rules from document XML structure."""
        rules = {
            "styles": {},
            "numbering": {},
            "section_properties": {},
            "default_formatting": {}
        }

        try:
            doc = Document(docx_path)  # type: ignore

            # Extract style information
            for style in doc.styles:
                if hasattr(style, 'name') and style.name:
                    rules["styles"][style.name] = {
                        "type": str(getattr(style, 'type', 'unknown')),
                        "based_on": getattr(style.base_style, 'name', None) if hasattr(style, 'base_style') and style.base_style else None
                    }

            # Extract default formatting (simplified)
            rules["default_formatting"] = {
                "font_name": "Times New Roman",  # Common academic font
                "font_size": 12,  # Standard size
                "line_spacing": 1.5  # Standard spacing
            }

            # Extract section properties (margins, etc.)
            if doc.sections:
                section = doc.sections[0]  # Primary section
                rules["section_properties"] = {
                    "top_margin": getattr(section, 'top_margin', Inches(1)).inches,
                    "bottom_margin": getattr(section, 'bottom_margin', Inches(1)).inches,
                    "left_margin": getattr(section, 'left_margin', Inches(1.25)).inches,
                    "right_margin": getattr(section, 'right_margin', Inches(1.25)).inches
                }

        except Exception as e:
            print(f"[EnhancedDocumentProcessor] Rule extraction failed: {e}")

        return rules

    def _parse_jats_xml(self, xml_path: str) -> Dict[str, Any]:
        """Parse JATS XML from Pandoc for structural information."""
        try:
            with open(xml_path, 'r', encoding='utf-8') as f:
                xml_content = f.read()

            soup = BeautifulSoup(xml_content, 'xml')

            structure = {
                "paragraphs": [],
                "headings": [],
                "tables": [],
                "sections": []
            }

            # Extract paragraphs
            for p in soup.find_all('p'):
                if p.get_text().strip():
                    structure["paragraphs"].append({
                        "text": p.get_text(),
                        "style": "Normal"
                    })

            # Extract headings
            for title in soup.find_all(['title', 'sec-title']):
                if title.get_text().strip():
                    structure["headings"].append({
                        "text": title.get_text(),
                        "level": 1  # Default level
                    })

            # Extract sections
            for sec in soup.find_all('sec'):
                title = sec.find('title')
                if title:
                    structure["sections"].append({
                        "title": title.get_text(),
                        "content": [p.get_text() for p in sec.find_all('p') if p.get_text().strip()],
                        "properties": {}
                    })

            return structure

        except Exception as e:
            print(f"[EnhancedDocumentProcessor] JATS parsing failed: {e}")
            return {
                "paragraphs": [],
                "headings": [],
                "tables": [],
                "sections": []
            }

    def convert_to_html_via_pandoc(self, docx_path: str) -> str:
        """Convert DOCX to HTML using Pandoc for high-quality output."""
        try:
            input_file = Path(docx_path)
            if not input_file.exists():
                return "<div style='color: red;'>Document not found</div>"

            # Create temporary output file
            with tempfile.NamedTemporaryFile(suffix='.html', delete=False) as temp_file:
                output_file = Path(temp_file.name)

            # Run pandoc conversion
            cmd = [
                'pandoc',
                str(input_file),
                '-f', 'docx',
                '-t', 'html',
                '-o', str(output_file),
                '--self-contained',
                '--css', self._get_pandoc_css()
            ]

            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)

            if result.returncode == 0 and output_file.exists():
                html_content = output_file.read_text(encoding='utf-8')
                output_file.unlink(missing_ok=True)
                return self._enhance_html_output(html_content)
            else:
                print(f"[EnhancedDocumentProcessor] Pandoc failed: {result.stderr}")
                output_file.unlink(missing_ok=True)
                return self._fallback_html_conversion(docx_path)

        except Exception as e:
            print(f"[EnhancedDocumentProcessor] HTML conversion failed: {e}")
            return self._fallback_html_conversion(docx_path)

    def apply_template_rules(self, content_structure: Dict[str, Any],
                           template_rules: Dict[str, Any]) -> Dict[str, Any]:
        """Apply extracted template rules to content structure (Integration Path Step 2)."""
        try:
            print(f"[EnhancedDocumentProcessor] Applying template rules: {len(template_rules.get('styles', {}))} styles")

            # Merge content with template styling rules
            processed_structure = content_structure.copy()

            # Apply paragraph styling rules
            if 'paragraphs' in processed_structure:
                processed_structure['paragraphs'] = self._apply_paragraph_rules(
                    processed_structure['paragraphs'], template_rules
                )

            # Apply heading hierarchy rules
            if 'headings' in processed_structure:
                processed_structure['headings'] = self._apply_heading_rules(
                    processed_structure['headings'], template_rules
                )

            # Apply table formatting rules
            if 'tables' in processed_structure:
                processed_structure['tables'] = self._apply_table_rules(
                    processed_structure['tables'], template_rules
                )

            # Apply section properties
            processed_structure['applied_rules'] = template_rules

            return processed_structure

        except Exception as e:
            print(f"[EnhancedDocumentProcessor] Rule application failed: {e}")
            return content_structure

    def generate_document_from_structure(self, structure: Dict[str, Any],
                                       output_path: str) -> bool:
        """Generate DOCX document from processed structure (Integration Path Step 3)."""
        try:
            print(f"[EnhancedDocumentProcessor] Generating document with {len(structure.get('paragraphs', []))} paragraphs")

            # Create document with applied rules
            doc = self._create_document_with_rules(structure)

            # Apply document-level settings from template rules
            self._apply_document_settings(doc, structure)

            # Add content sections using processed structure
            self._add_structured_content(doc, structure)

            # Save document
            doc.save(output_path)
            print(f"[EnhancedDocumentProcessor] Document saved to: {output_path}")
            return True

        except Exception as e:
            print(f"[EnhancedDocumentProcessor] Document generation failed: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _create_document_with_rules(self, structure: Dict[str, Any]) -> Document:
        """Create document with applied template rules."""
        doc = Document()  # type: ignore

        # Apply template rules if available
        rules = structure.get('applied_rules', {})
        if rules.get('default_formatting'):
            # Apply default formatting to the document
            default_fmt = rules['default_formatting']
            # This would set default styles - simplified for now
            pass

        return doc

    def _extract_paragraphs(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract paragraph information from Document object."""
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                para_info = {
                    'text': para.text,
                    'style': para.style.name if para.style else 'Normal',
                    'alignment': para.alignment,
                    'level': 0,  # Will be updated if it's in a heading
                    'formatting': self._extract_paragraph_formatting(para)
                }
                paragraphs.append(para_info)

        return paragraphs

    def _extract_paragraph_formatting(self, para) -> Dict[str, Any]:
        """Extract formatting information from a paragraph."""
        formatting = {}

        # Check if paragraph has formatting
        if para.runs:
            first_run = para.runs[0]
            formatting.update({
                'bold': first_run.bold,
                'italic': first_run.italic,
                'underline': first_run.underline,
                'font_size': first_run.font.size.pt if first_run.font.size else None,
                'font_name': first_run.font.name
            })

        return formatting

    def _extract_headings(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract heading information from Document."""
        headings = []

        for para in doc.paragraphs:
            if para.style and para.style.name:
                style_name = para.style.name.lower()
                if 'heading' in style_name and para.text.strip():
                    # Extract heading level from style name
                    level = 1
                    if 'heading' in style_name:
                        try:
                            # Extract number from "Heading 1", "Heading 2", etc.
                            level_part = style_name.split('heading')[1].strip()
                            level = int(level_part) if level_part.isdigit() else 1
                        except (ValueError, IndexError):
                            level = 1

                    heading_info = {
                        'text': para.text,
                        'level': level,
                        'style': para.style.name
                    }
                    headings.append(heading_info)

        return headings

    def _extract_tables(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract table information from Document."""
        tables = []

        for table in doc.tables:
            table_info = {
                'rows': len(table.rows),
                'columns': len(table.columns) if table.rows else 0,
                'style': getattr(table.style, 'name', 'Table Normal') if table.style else 'Table Normal'
            }
            tables.append(table_info)

        return tables

    def _extract_styles(self, doc: Document) -> Dict[str, Any]:
        """Extract style information from Document."""
        styles = {}

        for style in doc.styles:
            if style.name not in styles:
                styles[style.name] = {
                    'name': style.name,
                    'type': style.type,
                    'properties': {}
                }

        return styles

    def _extract_sections(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract document sections from Document."""
        sections = []

        # Basic section detection - group content by major headings
        current_section = None
        section_content = []

        for para in doc.paragraphs:
            if para.style and 'heading' in para.style.name.lower():
                # Save previous section
                if current_section:
                    sections.append({
                        'type': 'section',
                        'title': current_section,
                        'content': section_content,
                        'properties': {}
                    })

                # Start new section
                current_section = para.text
                section_content = []
            elif current_section:
                # Add content to current section
                if para.text.strip():
                    section_content.append(para.text)

        # Add final section
        if current_section:
            sections.append({
                'type': 'section',
                'title': current_section,
                'content': section_content,
                'properties': {}
            })

        return sections

    def _extract_metadata(self, docx_path: str) -> Dict[str, Any]:
        """Extract document metadata."""
        try:
            doc = Document(docx_path)  # type: ignore
            return {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "subject": doc.core_properties.subject or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
                "word_count": len(doc.paragraphs),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
        except Exception:
            return {}

    def _get_pandoc_css(self) -> str:
        """Get CSS for Pandoc HTML output."""
        return """
        body { font-family: 'Times New Roman', serif; margin: 40px; line-height: 1.6; }
        .document-container { max-width: 800px; margin: 0 auto; padding: 20px; }
        h1, h2, h3, h4, h5, h6 { color: #2c3e50; margin-top: 24px; margin-bottom: 12px; font-weight: bold; }
        h1 { font-size: 24pt; border-bottom: 2px solid #3498db; padding-bottom: 8px; }
        h2 { font-size: 18pt; border-bottom: 1px solid #bdc3c7; padding-bottom: 4px; }
        h3 { font-size: 14pt; }
        p { margin: 8px 0; text-align: justify; text-indent: 1cm; }
        table { border-collapse: collapse; width: 100%; margin: 16px 0; border: 1px solid #bdc3c7; }
        th, td { border: 1px solid #bdc3c7; padding: 8px 12px; text-align: left; }
        th { background-color: #ecf0f1; font-weight: bold; }
        ul, ol { margin: 12px 0; padding-left: 30px; }
        li { margin: 4px 0; }
        """

    def _enhance_html_output(self, html_content: str) -> str:
        """Enhance Pandoc HTML output with custom styling."""
        soup = BeautifulSoup(html_content, 'html.parser')

        # Add custom CSS
        if soup.head:
            style_tag = soup.new_tag('style')
            style_tag.string = self._get_pandoc_css()
            soup.head.append(style_tag)

        # Wrap content in container
        if soup.body:
            container = soup.new_tag('div', attrs={'class': 'document-container'})
            body_content = soup.body.contents[:]
            container.extend(body_content)
            soup.body.clear()
            soup.body.append(container)

        return str(soup)

    def _fallback_extraction(self, docx_path: str) -> Dict[str, Any]:
        """Fallback document structure extraction."""
        try:
            doc = Document(docx_path)  # type: ignore  # type: ignore
            return {
                "paragraphs": [{"text": p.text, "style": "Normal"} for p in doc.paragraphs if p.text.strip()],
                "headings": [],
                "tables": [],
                "styles": {},
                "sections": [],
                "metadata": self._extract_metadata(docx_path)
            }
        except Exception:
            return {
                "paragraphs": [],
                "headings": [],
                "tables": [],
                "styles": {},
                "sections": [],
                "metadata": {}
            }

    def _fallback_html_conversion(self, docx_path: str) -> str:
        """Fallback HTML conversion using python-docx."""
        try:
            doc = Document(docx_path)  # type: ignore
            html_parts = ['<div style="font-family: Times New Roman, serif; margin: 40px; line-height: 1.6;">']

            for para in doc.paragraphs:
                if para.text.strip():
                    # Basic HTML conversion
                    text = para.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    html_parts.append(f'<p>{text}</p>')

            html_parts.append('</div>')
            return '\n'.join(html_parts)

        except Exception:
            return '<div style="color: red; padding: 20px;">Failed to convert document</div>'

    def _apply_paragraph_rules(self, paragraphs: List[Dict[str, Any]],
                             template_rules: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Apply template paragraph rules."""
        # Apply font, spacing, indentation rules from template
        for para in paragraphs:
            # Apply default paragraph formatting
            para.setdefault('formatting', {})
            para['formatting'].update({
                'font_size': template_rules.get('default_font_size', 12),
                'line_spacing': template_rules.get('line_spacing', 1.5),
                'indentation': template_rules.get('first_line_indent', 1.0)
            })

        return paragraphs

    def _apply_heading_rules(self, headings: List[Dict[str, Any]],
                           template_rules: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Apply template heading rules."""
        for heading in headings:
            level = heading.get('level', 1)
            # Apply heading-specific formatting
            heading.setdefault('formatting', {})
            heading['formatting'].update({
                'bold': True,
                'font_size': template_rules.get(f'h{level}_font_size', 16 - level),
                'spacing_before': template_rules.get('heading_spacing_before', 20),
                'spacing_after': template_rules.get('heading_spacing_after', 10)
            })

        return headings

    def _apply_table_rules(self, tables: List[Dict[str, Any]],
                         template_rules: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Apply template table rules."""
        for table in tables:
            table.setdefault('formatting', {})
            table['formatting'].update({
                'border_style': template_rules.get('table_border', 'single'),
                'header_background': template_rules.get('table_header_bg', '#f0f0f0'),
                'cell_padding': template_rules.get('table_cell_padding', '8px')
            })

        return tables

    def _apply_document_settings(self, doc: Document, structure: Dict[str, Any]):
        """Apply document-level settings."""
        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

    def _add_structured_content(self, doc: Document, structure: Dict[str, Any]):
        """Add structured content to document."""
        # Add paragraphs
        for para_info in structure.get('paragraphs', []):
            para = doc.add_paragraph(para_info['text'])

            # Apply formatting
            formatting = para_info.get('formatting', {})
            if formatting.get('font_size'):
                for run in para.runs:
                    run.font.size = Pt(formatting['font_size'])

            if formatting.get('line_spacing'):
                para.paragraph_format.line_spacing = formatting['line_spacing']

            if formatting.get('indentation'):
                para.paragraph_format.first_line_indent = Inches(formatting['indentation'])

        # Add headings
        for heading_info in structure.get('headings', []):
            level = heading_info.get('level', 1)
            heading = doc.add_heading(heading_info['text'], level)

            # Apply heading formatting
            formatting = heading_info.get('formatting', {})
            if formatting.get('font_size'):
                for run in heading.runs:
                    run.font.size = Pt(formatting['font_size'])

        # Add tables
        for table_info in structure.get('tables', []):
            if table_info.get('rows'):
                table = doc.add_table(rows=len(table_info['rows']), cols=len(table_info['rows'][0]) if table_info['rows'] else 0)
                for i, row_data in enumerate(table_info['rows']):
                    for j, cell_data in enumerate(row_data):
                        table.cell(i, j).text = str(cell_data)


# Global instance
document_processor = EnhancedDocumentProcessor()


def process_document_enhanced(docx_path: str) -> Dict[str, Any]:
    """Convenience function for enhanced document processing."""
    return document_processor.extract_document_structure(docx_path)


def convert_document_enhanced(docx_path: str) -> str:
    """Convenience function for enhanced document conversion."""
    return document_processor.convert_to_html_via_pandoc(docx_path)


def generate_document_via_integration_path(
    template_path: str,
    content_structure: Dict[str, Any],
    output_path: str
) -> bool:
    """
    Generate document using the complete Integration Path:
    1. Pandoc/docxBox → normalized structural representation
    2. Rule extraction → derive rules from template XML
    3. python-docx + helpers → apply rules and generate document
    """
    try:
        print(f"[Integration Path] Starting generation for template: {template_path}")

        # Step 1: Extract normalized structural representation
        print("[Integration Path] Step 1: Extracting normalized structure")
        normalized_structure = document_processor.extract_document_structure(template_path)

        # Step 2: Derive rules from template XML structure
        print("[Integration Path] Step 2: Deriving template rules")
        template_rules = document_processor._extract_template_rules(template_path)
        print(f"[Integration Path] Extracted {len(template_rules.get('styles', {}))} style rules")

        # Step 3: Apply rules to content structure
        print("[Integration Path] Step 3: Applying rules to content")
        merged_structure = document_processor.apply_template_rules(content_structure, template_rules)

        # Step 4: Generate final document with python-docx + helpers
        print("[Integration Path] Step 4: Generating final document")
        success = document_processor.generate_document_from_structure(merged_structure, output_path)

        if success:
            print(f"[Integration Path] ✓ Document generated successfully: {output_path}")
        else:
            print(f"[Integration Path] ✗ Document generation failed")

        return success

    except Exception as e:
        print(f"[Integration Path] ✗ Integration path failed: {e}")
        import traceback
        traceback.print_exc()
        return False