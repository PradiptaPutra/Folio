"""
Intelligent Template Adapter
AI-powered adaptive system that can work with different thesis templates.
Uses multiple detection strategies and AI assistance to identify content placement zones.
"""

from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass, field
from docx import Document
import re
from pathlib import Path


@dataclass
class TemplatePattern:
    """Detected pattern in template."""
    pattern_type: str  # 'chapter', 'subsection', 'placeholder', 'content_zone'
    location: int  # paragraph index
    text: str
    style: Optional[str] = None
    confidence: float = 0.0
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class TemplateStructure:
    """Adaptive template structure."""
    template_path: str
    chapter_patterns: List[TemplatePattern] = field(default_factory=list)
    subsection_patterns: List[TemplatePattern] = field(default_factory=list)
    placeholder_patterns: List[TemplatePattern] = field(default_factory=list)
    content_zones: List[TemplatePattern] = field(default_factory=list)
    style_mapping: Dict[str, str] = field(default_factory=dict)
    font_info: Dict[str, Any] = field(default_factory=dict)
    template_type: str = "unknown"  # 'uii', 'ui', 'ugm', 'generic'


class IntelligentTemplateAdapter:
    """
    Intelligent adapter that can work with different thesis templates.
    Uses multiple detection strategies and AI assistance.
    """
    
    # Comprehensive chapter detection patterns (multiple universities)
    CHAPTER_PATTERNS = [
        # Standard Indonesian patterns
        (r'^BAB\s+([IVXLCDM]+)\s*[:-]?\s*(.+)$', 'roman', 1.0),
        (r'^BAB\s+(\d+)\s*[:-]?\s*(.+)$', 'numeric', 1.0),
        (r'^BAB\s+([IVXLCDM]+)\s+(.+)$', 'roman_no_dash', 0.9),
        (r'^BAB\s+(\d+)\s+(.+)$', 'numeric_no_dash', 0.9),
        (r'^CHAPTER\s+(\d+)\s*[:-]?\s*(.+)$', 'english', 0.8),
        (r'^([IVXLCDM]+)\.\s+(.+)$', 'roman_dot', 0.7),
        (r'^(\d+)\.\s+(.+)$', 'numeric_dot', 0.7),
        
        # Semantic detection (by keywords)
        (r'.*PENDAHULUAN.*', 'semantic', 0.6),
        (r'.*TINJAUAN\s+PUSTAKA.*', 'semantic', 0.6),
        (r'.*KAJIAN\s+PUSTAKA.*', 'semantic', 0.6),
        (r'.*METODOLOGI.*', 'semantic', 0.6),
        (r'.*METODE\s+PENELITIAN.*', 'semantic', 0.6),
        (r'.*HASIL\s+DAN\s+PEMBAHASAN.*', 'semantic', 0.6),
        (r'.*ANALISIS\s+DAN\s+PERANCANGAN.*', 'semantic', 0.6),
        (r'.*KESIMPULAN.*', 'semantic', 0.6),
        (r'.*PENUTUP.*', 'semantic', 0.6),
    ]
    
    # Comprehensive subsection patterns
    SUBSECTION_PATTERNS = [
        # Numbered subsections
        (r'^(\d+)\.(\d+)\s+(.+)$', 'numbered', 1.0),
        (r'^(\d+)\.(\d+)\.(\d+)\s+(.+)$', 'numbered_3level', 1.0),
        (r'^([A-Z])\.(\d+)\s+(.+)$', 'letter_numbered', 0.9),
        (r'^([a-z])\.(\d+)\s+(.+)$', 'letter_lower', 0.9),
        
        # Explicit subsection markers
        (r'^.*Subbab.*$', 'explicit', 0.8),
        (r'^.*Sub-section.*$', 'explicit', 0.8),
        (r'^.*Subsection.*$', 'explicit', 0.8),
        (r'^.*Anak\s+Subbab.*$', 'explicit_child', 0.8),
        
        # Missing chapter number (needs context)
        (r'^\.(\d+)\s+(.+)$', 'missing_chapter', 0.5),
    ]
    
    # Placeholder patterns
    PLACEHOLDER_PATTERNS = [
        (r'TULISKAN.*', 'instruction', 0.9),
        (r'KETIK.*', 'instruction', 0.9),
        (r'ISI.*', 'instruction', 0.8),
        (r'\[.*\]', 'brackets', 0.9),
        (r'___+', 'underscores', 0.8),
        (r'\.\.\.+', 'dots', 0.7),
        (r'Subbab\s*$', 'empty_subsection', 0.8),
        (r'Anak\s+Subbab\s*$', 'empty_child', 0.8),
    ]
    
    # Content zone indicators
    CONTENT_ZONE_INDICATORS = [
        (r'Format\s+paragraf', 'format_instruction', 0.9),
        (r'isi\s+paragraf', 'content_style', 0.8),
        (r'Paragraf\s+isi', 'content_style', 0.8),
        (r'Body\s+text', 'content_style', 0.7),
    ]
    
    def __init__(self, template_path: str, api_key: Optional[str] = None):
        """Initialize with template path and optional AI API key."""
        self.template_path = Path(template_path)
        self.api_key = api_key
        self.ai_available = api_key is not None
        self.doc = None
        self.structure = None
        
    def analyze_template(self) -> TemplateStructure:
        """
        Perform comprehensive adaptive template analysis.
        
        Returns:
            TemplateStructure with detected patterns and zones
        """
        print("[ADAPTIVE] Starting intelligent template analysis...")
        
        # Load template
        self.doc = Document(str(self.template_path))
        
        # Initialize structure
        self.structure = TemplateStructure(template_path=str(self.template_path))
        
        # Multi-pass analysis
        self._detect_chapters()
        self._detect_subsections()
        self._detect_placeholders()
        self._detect_content_zones()
        self._extract_style_mapping()
        self._extract_font_info()
        self._identify_template_type()
        
        # AI enhancement if available
        if self.ai_available:
            self._ai_enhance_analysis()
        
        print(f"[ADAPTIVE] Analysis complete: {len(self.structure.chapter_patterns)} chapters, "
              f"{len(self.structure.subsection_patterns)} subsections, "
              f"{len(self.structure.placeholder_patterns)} placeholders")
        
        return self.structure
    
    def _detect_chapters(self) -> None:
        """Detect chapter headings using multiple patterns."""
        current_chapter = 0
        
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            text_upper = text.upper().replace('\v', ' ').replace('\n', ' ')
            
            # Try all chapter patterns
            for pattern, pattern_type, confidence in self.CHAPTER_PATTERNS:
                match = re.match(pattern, text_upper, re.IGNORECASE)
                if match:
                    # Extract chapter number
                    chapter_num = self._extract_chapter_number(text, pattern_type, match)
                    
                    if chapter_num:
                        # Extract title from match groups
                        title = text
                        if match.groups():
                            # Get the last group (usually the title)
                            groups = match.groups()
                            if len(groups) > 1:
                                title = groups[-1].strip()
                            elif len(groups) == 1:
                                title = groups[0].strip()
                        
                        pattern_obj = TemplatePattern(
                            pattern_type='chapter',
                            location=i,
                            text=text,
                            style=para.style.name if para.style else None,
                            confidence=confidence,
                            metadata={
                                'chapter_num': chapter_num,
                                'pattern_type': pattern_type,
                                'title': title
                            }
                        )
                        self.structure.chapter_patterns.append(pattern_obj)
                        current_chapter = chapter_num
                        break
        
        # Sort by chapter number
        self.structure.chapter_patterns.sort(key=lambda x: x.metadata.get('chapter_num', 0))
    
    def _extract_chapter_number(self, text: str, pattern_type: str, match: re.Match) -> Optional[int]:
        """Extract chapter number from text."""
        try:
            if pattern_type in ['roman', 'roman_no_dash', 'roman_dot']:
                roman_num = match.group(1)
                return self._roman_to_int(roman_num)
            elif pattern_type in ['numeric', 'numeric_no_dash', 'numeric_dot', 'english']:
                return int(match.group(1))
            elif pattern_type == 'semantic':
                # Map semantic keywords to chapter numbers
                text_upper = text.upper()
                if 'PENDAHULUAN' in text_upper or 'LATAR BELAKANG' in text_upper:
                    return 1
                elif 'PUSTAKA' in text_upper or 'TEORI' in text_upper:
                    return 2
                elif 'METODOLOGI' in text_upper or 'METODE' in text_upper:
                    return 3
                elif 'ANALISIS' in text_upper or 'PERANCANGAN' in text_upper:
                    return 4
                elif 'HASIL' in text_upper or 'PEMBAHASAN' in text_upper or 'IMPLEMENTASI' in text_upper:
                    return 5
                elif 'KESIMPULAN' in text_upper or 'PENUTUP' in text_upper:
                    return 6
        except:
            pass
        return None
    
    def _roman_to_int(self, roman: str) -> int:
        """Convert Roman numeral to integer."""
        roman = roman.upper()
        val_map = {'I': 1, 'V': 5, 'X': 10, 'L': 50, 'C': 100, 'D': 500, 'M': 1000}
        result = 0
        prev = 0
        for char in reversed(roman):
            val = val_map.get(char, 0)
            if val < prev:
                result -= val
            else:
                result += val
            prev = val
        return result
    
    def _detect_subsections(self) -> None:
        """Detect subsection headings."""
        current_chapter = 0
        
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # Update current chapter context
            for chapter in self.structure.chapter_patterns:
                if chapter.location <= i:
                    current_chapter = chapter.metadata.get('chapter_num', 0)
            
            # Try all subsection patterns
            for pattern, pattern_type, confidence in self.SUBSECTION_PATTERNS:
                match = re.match(pattern, text, re.IGNORECASE)
                if match:
                    # Extract subsection info
                    subsection_info = self._extract_subsection_info(text, pattern_type, match, current_chapter)
                    
                    if subsection_info:
                        pattern_obj = TemplatePattern(
                            pattern_type='subsection',
                            location=i,
                            text=text,
                            style=para.style.name if para.style else None,
                            confidence=confidence,
                            metadata=subsection_info
                        )
                        self.structure.subsection_patterns.append(pattern_obj)
                        break
    
    def _extract_subsection_info(self, text: str, pattern_type: str, match: re.Match, current_chapter: int) -> Optional[Dict[str, Any]]:
        """Extract subsection information."""
        try:
            if pattern_type == 'numbered':
                chapter_num = int(match.group(1))
                subsection_num = int(match.group(2))
                title = match.group(3)
                return {
                    'chapter_num': chapter_num,
                    'subsection_num': subsection_num,
                    'title': title,
                    'full_number': f"{chapter_num}.{subsection_num}"
                }
            elif pattern_type == 'numbered_3level':
                chapter_num = int(match.group(1))
                subsection_num = int(match.group(2))
                subsubsection_num = int(match.group(3))
                title = match.group(4)
                return {
                    'chapter_num': chapter_num,
                    'subsection_num': subsection_num,
                    'subsubsection_num': subsubsection_num,
                    'title': title,
                    'full_number': f"{chapter_num}.{subsection_num}.{subsubsection_num}",
                    'is_child': True
                }
            elif pattern_type == 'missing_chapter':
                # Use current chapter context
                subsection_num = int(match.group(1))
                title = match.group(2)
                return {
                    'chapter_num': current_chapter,
                    'subsection_num': subsection_num,
                    'title': title,
                    'full_number': f"{current_chapter}.{subsection_num}",
                    'needs_fix': True
                }
            elif pattern_type in ['explicit', 'explicit_child']:
                return {
                    'chapter_num': current_chapter,
                    'title': text,
                    'is_placeholder': True,
                    'is_child': pattern_type == 'explicit_child'
                }
        except:
            pass
        return None
    
    def _detect_placeholders(self) -> None:
        """Detect placeholder text that needs to be replaced."""
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            for pattern, pattern_type, confidence in self.PLACEHOLDER_PATTERNS:
                if re.search(pattern, text, re.IGNORECASE):
                    pattern_obj = TemplatePattern(
                        pattern_type='placeholder',
                        location=i,
                        text=text,
                        style=para.style.name if para.style else None,
                        confidence=confidence,
                        metadata={'pattern_type': pattern_type}
                    )
                    self.structure.placeholder_patterns.append(pattern_obj)
                    break
    
    def _detect_content_zones(self) -> None:
        """Detect content zones where actual content should be placed."""
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # Check for content zone indicators
            for pattern, zone_type, confidence in self.CONTENT_ZONE_INDICATORS:
                if re.search(pattern, text, re.IGNORECASE):
                    pattern_obj = TemplatePattern(
                        pattern_type='content_zone',
                        location=i,
                        text=text,
                        style=para.style.name if para.style else None,
                        confidence=confidence,
                        metadata={'zone_type': zone_type}
                    )
                    self.structure.content_zones.append(pattern_obj)
                    break
            
            # Also detect by style name
            if para.style:
                style_name = para.style.name.lower()
                if any(keyword in style_name for keyword in ['isi', 'paragraf', 'body', 'content', 'normal']):
                    if len(text) < 100:  # Likely a placeholder or instruction
                        pattern_obj = TemplatePattern(
                            pattern_type='content_zone',
                            location=i,
                            text=text,
                            style=para.style.name,
                            confidence=0.7,
                            metadata={'zone_type': 'style_based'}
                        )
                        self.structure.content_zones.append(pattern_obj)
    
    def _extract_style_mapping(self) -> None:
        """Extract style mapping for content paragraphs."""
        # Find content paragraph style
        for para in self.doc.paragraphs:
            if para.style:
                style_name = para.style.name.lower()
                text = para.text.strip()
                
                # Check if this is a content paragraph style
                if any(keyword in style_name for keyword in ['isi', 'paragraf', 'body', 'content']):
                    if len(text) > 50:  # Has actual content
                        self.structure.style_mapping['content_paragraph'] = para.style.name
                        break
        
        # Find heading styles
        for para in self.doc.paragraphs:
            if para.style and 'heading' in para.style.name.lower():
                text = para.text.strip()
                if re.match(r'^BAB\s+[IVX\d]+', text, re.I):
                    self.structure.style_mapping['chapter_heading'] = para.style.name
                elif re.match(r'^\d+\.\d+', text):
                    self.structure.style_mapping['subsection_heading'] = para.style.name
    
    def _extract_font_info(self) -> None:
        """Extract font information from template."""
        # Check styles first
        for style in self.doc.styles:
            try:
                if hasattr(style, 'font') and style.font and style.font.name:
                    style_name_lower = style.name.lower()
                    if 'normal' in style_name_lower or 'isi' in style_name_lower or 'paragraf' in style_name_lower:
                        self.structure.font_info['default'] = {
                            'name': style.font.name,
                            'size': style.font.size
                        }
                        break
            except:
                continue
        
        # Fallback to paragraph runs
        if 'default' not in self.structure.font_info:
            for para in self.doc.paragraphs[:50]:
                for run in para.runs:
                    if run.font.name:
                        self.structure.font_info['default'] = {
                            'name': run.font.name,
                            'size': run.font.size
                        }
                        return
    
    def _identify_template_type(self) -> None:
        """Identify template type based on patterns."""
        # Check for university-specific patterns
        template_text = ' '.join([p.text for p in self.doc.paragraphs[:100]]).upper()
        
        if 'UNIVERSITAS ISLAM INDONESIA' in template_text or 'UII' in template_text:
            self.structure.template_type = 'uii'
        elif 'UNIVERSITAS INDONESIA' in template_text or 'UI' in template_text:
            self.structure.template_type = 'ui'
        elif 'UNIVERSITAS GADJAH MADA' in template_text or 'UGM' in template_text:
            self.structure.template_type = 'ugm'
        else:
            self.structure.template_type = 'generic'
        
        print(f"[ADAPTIVE] Template type identified: {self.structure.template_type}")
    
    def _ai_enhance_analysis(self) -> None:
        """Use AI to enhance template analysis."""
        if not self.ai_available:
            return
        
        print("[ADAPTIVE] Using AI to enhance template analysis...")
        # TODO: Implement AI enhancement
        # This would use AI to:
        # 1. Identify content placement zones more accurately
        # 2. Suggest optimal insertion points
        # 3. Validate detected patterns
        # 4. Fill in missing information
    
    def get_content_insertion_points(self, chapter_num: int) -> List[Tuple[int, Dict[str, Any]]]:
        """
        Get optimal content insertion points for a chapter.
        
        Args:
            chapter_num: Chapter number
            
        Returns:
            List of (paragraph_index, metadata) tuples for insertion points
        """
        insertion_points = []
        
        # Find chapter
        chapter_pattern = None
        for pattern in self.structure.chapter_patterns:
            if pattern.metadata.get('chapter_num') == chapter_num:
                chapter_pattern = pattern
                break
        
        if not chapter_pattern:
            return insertion_points
        
        chapter_start = chapter_pattern.location
        
        # Find subsections in this chapter
        chapter_subsections = [
            p for p in self.structure.subsection_patterns
            if p.metadata.get('chapter_num') == chapter_num
        ]
        
        # Find placeholders and content zones after chapter
        for i in range(chapter_start + 1, len(self.doc.paragraphs)):
            para = self.doc.paragraphs[i]
            text = para.text.strip()
            
            # Stop if we hit next chapter
            if any(p.location == i and p.pattern_type == 'chapter' 
                   for p in self.structure.chapter_patterns):
                break
            
            # Check if this is a good insertion point
            is_placeholder = any(p.location == i for p in self.structure.placeholder_patterns)
            is_content_zone = any(p.location == i for p in self.structure.content_zones)
            is_subsection = any(p.location == i for p in chapter_subsections)
            
            if is_placeholder or is_content_zone:
                insertion_points.append((i, {
                    'type': 'placeholder' if is_placeholder else 'content_zone',
                    'is_subsection': is_subsection,
                    'text': text[:50]
                }))
        
        return insertion_points

