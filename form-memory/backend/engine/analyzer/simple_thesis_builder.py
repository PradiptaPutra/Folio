"""
Simple Thesis Builder - Clean Document Generation
Builds thesis documents with proper structure without complex template matching.
Focuses on reliability and correctness over template preservation.
"""

import re
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


class SimpleThesisBuilder:
    """Simplified thesis builder that creates clean, well-structured documents."""
    
    # Section detection patterns for Indonesian thesis structure
    CHAPTER_PATTERNS = [
        (r'^BAB\s+([IVX]+|[0-9]+)[\s\-:]*(.*)$', 'chapter'),
        (r'^CHAPTER\s+([0-9]+)[\s\-:]*(.*)$', 'chapter'),
        (r'^(PENDAHULUAN|TINJAUAN PUSTAKA|METODOLOGI|ANALISIS|HASIL|IMPLEMENTASI|PENUTUP|KESIMPULAN)$', 'chapter'),
    ]
    
    SECTION_PATTERNS = [
        (r'^([0-9]+\.[0-9]+)\s+(.+)$', 'section'),  # 1.1 Title
        (r'^([0-9]+\.[0-9]+\.[0-9]+)\s+(.+)$', 'subsection'),  # 1.1.1 Title
    ]
    
    # Content section keywords mapping
    CONTENT_SECTION_MAP = {
        'latar_belakang': ['LATAR BELAKANG', 'BACKGROUND'],
        'rumusan_masalah': ['RUMUSAN MASALAH', 'PROBLEM STATEMENT'],
        'tujuan': ['TUJUAN', 'OBJECTIVES', 'TUJUAN PENELITIAN'],
        'manfaat': ['MANFAAT', 'BENEFITS'],
        'batasan': ['BATASAN', 'LIMITATIONS', 'SCOPE'],
        'tinjauan_pustaka': ['TINJAUAN PUSTAKA', 'LITERATURE REVIEW'],
        'penelitian_terkait': ['PENELITIAN TERKAIT', 'RELATED WORK'],
        'metodologi': ['METODOLOGI', 'METHODOLOGY', 'METODE'],
        'desain': ['DESAIN', 'DESIGN'],
        'implementasi': ['IMPLEMENTASI', 'IMPLEMENTATION'],
        'hasil': ['HASIL', 'RESULTS'],
        'pembahasan': ['PEMBAHASAN', 'DISCUSSION'],
        'kesimpulan': ['KESIMPULAN', 'CONCLUSION'],
        'saran': ['SARAN', 'RECOMMENDATIONS'],
    }
    
    def __init__(self, template_path: str, content_path: str, output_path: str):
        """Initialize the simple thesis builder.
        
        Args:
            template_path: Path to DOCX template (for style reference)
            content_path: Path to content file (TXT or DOCX)
            output_path: Path for output DOCX
        """
        self.template_path = Path(template_path)
        self.content_path = Path(content_path)
        self.output_path = Path(output_path)
        
        self.template_doc = None
        self.content_text = None
        self.content_structure = {}
        
    def build(self, user_data: Optional[Dict[str, Any]] = None) -> Path:
        """Build a complete thesis document.
        
        Args:
            user_data: Metadata for personalization (title, author, nim, etc.)
            
        Returns:
            Path to created document
        """
        user_data = user_data or {}
        
        print("[SIMPLE_BUILDER] Starting clean document build...")
        
        # Step 1: Load template for styles
        try:
            self.template_doc = Document(str(self.template_path))
            print(f"[SIMPLE_BUILDER] Template loaded: {len(self.template_doc.paragraphs)} paragraphs")
        except Exception as e:
            print(f"[SIMPLE_BUILDER] No template available, creating new document: {e}")
            self.template_doc = Document()
        
        # Step 2: Read and parse content
        self._read_and_parse_content()
        
        # Step 3: Create clean document
        doc = Document()
        self._copy_template_styles(doc, self.template_doc)
        
        # Step 4: Build document structure
        self._build_front_matter(doc, user_data)
        self._build_main_content(doc, user_data)
        self._build_back_matter(doc, user_data)
        
        # Step 5: Save document
        doc.save(str(self.output_path))
        print(f"[SIMPLE_BUILDER] Document saved: {self.output_path}")
        
        return self.output_path
    
    def _read_and_parse_content(self) -> None:
        """Read and parse content from file."""
        print(f"[SIMPLE_BUILDER] Reading content from: {self.content_path}")
        
        if self.content_path.suffix.lower() == '.docx':
            # Read from DOCX
            doc = Document(str(self.content_path))
            self.content_text = '\n'.join(p.text for p in doc.paragraphs)
        else:
            # Read from TXT
            self.content_text = self.content_path.read_text(encoding='utf-8')
        
        print(f"[SIMPLE_BUILDER] Content read: {len(self.content_text)} characters")
        
        # Parse content into chapters and sections
        self._parse_content_structure()
    
    def _parse_content_structure(self) -> None:
        """Parse content text into structured format."""
        lines = self.content_text.split('\n')
        
        current_chapter = 0
        current_section = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            line_upper = line.upper()
            
            # Check if this is a chapter header
            is_chapter = False
            for pattern, _ in self.CHAPTER_PATTERNS:
                if re.match(pattern, line_upper):
                    # Save previous content
                    if current_content:
                        key = f"chapter{current_chapter}_{current_section}"
                        self.content_structure[key] = '\n'.join(current_content)
                        current_content = []
                    
                    # Detect chapter number
                    if 'PENDAHULUAN' in line_upper or 'BAB' in line_upper and 'I' in line_upper:
                        current_chapter = 1
                    elif 'PUSTAKA' in line_upper or 'TEORI' in line_upper or 'BAB' in line_upper and 'II' in line_upper:
                        current_chapter = 2
                    elif 'METODOLOGI' in line_upper or 'METODE' in line_upper or 'BAB' in line_upper and 'III' in line_upper:
                        current_chapter = 3
                    elif 'ANALISIS' in line_upper or 'PERANCANGAN' in line_upper or 'BAB' in line_upper and 'IV' in line_upper:
                        current_chapter = 4
                    elif 'IMPLEMENTASI' in line_upper or 'HASIL' in line_upper or 'PEMBAHASAN' in line_upper or 'BAB' in line_upper and 'V' in line_upper:
                        current_chapter = 5
                    elif 'PENUTUP' in line_upper or 'KESIMPULAN' in line_upper or 'BAB' in line_upper and 'VI' in line_upper:
                        current_chapter = 6
                    
                    current_section = 'intro'
                    is_chapter = True
                    break
            
            if is_chapter:
                continue
            
            # Check if this is a section header
            for keyword_key, keywords in self.CONTENT_SECTION_MAP.items():
                if any(kw in line_upper for kw in keywords):
                    # Save previous content
                    if current_content:
                        key = f"chapter{current_chapter}_{current_section}"
                        self.content_structure[key] = '\n'.join(current_content)
                        current_content = []
                    
                    current_section = keyword_key
                    break
            else:
                # Add to current section content
                if current_section:
                    current_content.append(line)
        
        # Save final content
        if current_content:
            key = f"chapter{current_chapter}_{current_section}"
            self.content_structure[key] = '\n'.join(current_content)
        
        print(f"[SIMPLE_BUILDER] Parsed {len(self.content_structure)} content sections")
    
    def _copy_template_styles(self, target_doc: Document, source_doc: Document) -> None:
        """Copy styles from template document."""
        try:
            for style in source_doc.styles:
                if style.name not in [s.name for s in target_doc.styles]:
                    try:
                        target_doc.styles.add_style(style.name, style.type)
                    except:
                        pass
        except Exception as e:
            print(f"[SIMPLE_BUILDER] Could not copy styles: {e}")
    
    def _build_front_matter(self, doc: Document, user_data: Dict[str, Any]) -> None:
        """Build front matter (cover, approval, abstract)."""
        print("[SIMPLE_BUILDER] Building front matter...")
        
        # Title Page
        self._add_title_page(doc, user_data)
        
        # Approval Page
        self._add_approval_page(doc, user_data)
        
        # Abstract
        abstract_id = user_data.get('abstract_id', '')
        if abstract_id:
            self._add_abstract(doc, 'ABSTRAK', abstract_id)
        
        abstract_en = user_data.get('abstract_en', '')
        if abstract_en:
            self._add_abstract(doc, 'ABSTRACT', abstract_en)
    
    def _add_title_page(self, doc: Document, user_data: Dict[str, Any]) -> None:
        """Add title page."""
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Title
        title = user_data.get('title', 'JUDUL SKRIPSI')
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.bold = True
            run.font.size = Pt(14)
        
        # Spacing
        for _ in range(8):
            doc.add_paragraph()
        
        # Author info
        author = user_data.get('author', 'Nama Penulis')
        nim = user_data.get('nim', 'NIM')
        author_para = doc.add_paragraph(f"Oleh\n\n{author}\n{nim}")
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in author_para.runs:
            run.font.size = Pt(12)
        
        # Spacing
        for _ in range(6):
            doc.add_paragraph()
        
        # Institution
        institution = user_data.get('institution', 'Universitas')
        year = user_data.get('date', datetime.now().year)
        inst_para = doc.add_paragraph(f"{institution}\n{year}")
        inst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in inst_para.runs:
            run.font.size = Pt(12)
        
        doc.add_page_break()
    
    def _add_approval_page(self, doc: Document, user_data: Dict[str, Any]) -> None:
        """Add approval page."""
        # Title
        title_para = doc.add_paragraph("HALAMAN PENGESAHAN DOSEN PEMBIMBING")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.bold = True
            run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Content
        title = user_data.get('title', 'JUDUL')
        author = user_data.get('author', 'Penulis')
        nim = user_data.get('nim', 'NIM')
        advisor = user_data.get('advisor', 'Dosen Pembimbing')
        
        content = f"""Judul: {title}

Penulis: {author}

NIM: {nim}



Dosen Pembimbing:



_____________________________
{advisor}"""
        
        content_para = doc.add_paragraph(content)
        for run in content_para.runs:
            run.font.size = Pt(11)
        
        doc.add_page_break()
    
    def _add_abstract(self, doc: Document, lang: str, content: str) -> None:
        """Add abstract."""
        title_para = doc.add_paragraph(lang)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.bold = True
            run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        content_para = doc.add_paragraph(content)
        for run in content_para.runs:
            run.font.size = Pt(11)
        
        doc.add_page_break()
    
    def _build_main_content(self, doc: Document, user_data: Dict[str, Any]) -> None:
        """Build main content chapters."""
        print("[SIMPLE_BUILDER] Building main content...")
        
        chapter_titles = {
            1: "PENDAHULUAN",
            2: "TINJAUAN PUSTAKA",
            3: "METODOLOGI PENELITIAN",
            4: "ANALISIS DAN PERANCANGAN",
            5: "HASIL DAN PEMBAHASAN",
            6: "PENUTUP"
        }
        
        subsection_order = {
            1: ['latar_belakang', 'rumusan_masalah', 'tujuan', 'manfaat', 'batasan'],
            2: ['tinjauan_pustaka', 'penelitian_terkait'],
            3: ['metodologi', 'desain'],
            4: ['analisis_kebutuhan', 'desain'],
            5: ['implementasi', 'hasil', 'pembahasan'],
            6: ['kesimpulan', 'saran']
        }
        
        for chapter_num in range(1, 7):
            # Add chapter heading
            chapter_title = chapter_titles.get(chapter_num, f"BAB {chapter_num}")
            chapter_para = doc.add_paragraph(f"BAB {self._to_roman(chapter_num)}\n{chapter_title}")
            chapter_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in chapter_para.runs:
                run.bold = True
                run.font.size = Pt(14)
            
            doc.add_paragraph()
            
            # Add subsections
            subsections = subsection_order.get(chapter_num, [])
            for sub_idx, subsection_key in enumerate(subsections, 1):
                content_key = f"chapter{chapter_num}_{subsection_key}"
                content = self.content_structure.get(content_key, '')
                
                if content and len(content.strip()) > 50:
                    # Add subsection heading
                    subsection_title = subsection_key.replace('_', ' ').title()
                    sub_para = doc.add_paragraph(f"{chapter_num}.{sub_idx} {subsection_title}")
                    for run in sub_para.runs:
                        run.bold = True
                        run.font.size = Pt(12)
                    
                    # Add content paragraphs
                    for line in content.split('\n'):
                        if line.strip():
                            p = doc.add_paragraph(line)
                            p.paragraph_format.first_line_indent = Inches(0.5)
                            p.paragraph_format.line_spacing = 1.5
                            for run in p.runs:
                                run.font.size = Pt(11)
                else:
                    print(f"[SIMPLE_BUILDER] No content for chapter{chapter_num}_{subsection_key}")
            
            doc.add_page_break()
    
    def _build_back_matter(self, doc: Document, user_data: Dict[str, Any]) -> None:
        """Build back matter (references, appendices)."""
        print("[SIMPLE_BUILDER] Building back matter...")
        
        # References
        references = user_data.get('references', [])
        if references:
            ref_title = doc.add_paragraph("DAFTAR PUSTAKA")
            ref_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in ref_title.runs:
                run.bold = True
                run.font.size = Pt(12)
            
            doc.add_paragraph()
            
            for ref in references:
                if isinstance(ref, dict):
                    ref_text = ref.get('text', '')
                elif isinstance(ref, str):
                    ref_text = ref
                else:
                    continue
                
                if ref_text:
                    ref_para = doc.add_paragraph(ref_text)
                    ref_para.paragraph_format.hanging_indent = Inches(0.5)
                    for run in ref_para.runs:
                        run.font.size = Pt(11)
            
            doc.add_page_break()
        
        # Appendices
        appendices = user_data.get('appendices', [])
        if appendices:
            app_title = doc.add_paragraph("LAMPIRAN")
            app_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in app_title.runs:
                run.bold = True
                run.font.size = Pt(12)
            
            doc.add_paragraph()
            
            for idx, appendix in enumerate(appendices, 1):
                if isinstance(appendix, dict):
                    app_content = appendix.get('content', '')
                elif isinstance(appendix, str):
                    app_content = appendix
                else:
                    continue
                
                if app_content:
                    app_para = doc.add_paragraph(f"Lampiran {chr(64+idx)}")
                    for run in app_para.runs:
                        run.bold = True
                        run.font.size = Pt(12)
                    
                    content_para = doc.add_paragraph(app_content)
                    for run in content_para.runs:
                        run.font.size = Pt(11)
                    
                    if idx < len(appendices):
                        doc.add_page_break()
    
    def _to_roman(self, num: int) -> str:
        """Convert number to Roman numeral."""
        val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
        syms = ['M', 'CM', 'D', 'CD', 'C', 'XC', 'L', 'XL', 'X', 'IX', 'V', 'IV', 'I']
        roman_num = ''
        i = 0
        while num > 0:
            for _ in range(num // val[i]):
                roman_num += syms[i]
                num -= val[i]
            i += 1
        return roman_num
