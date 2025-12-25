"""
Front Matter Generator
Creates all required preliminary pages for thesis documents.
Supports Indonesian thesis formats (skripsi/tesis).
"""

from typing import Dict, List, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class FrontMatterGenerator:
    """Generates front matter pages (preliminary pages) for thesis."""
    
    def __init__(self, user_data: Dict[str, Any] = None):
        """Initialize with user data for personalization."""
        self.user_data = user_data or {}
        self.title = user_data.get("title", "JUDUL SKRIPSI")
        self.author = user_data.get("author", "Nama Penulis")
        self.nim = user_data.get("nim", "000000")
        self.advisor = user_data.get("advisor", "Nama Dosen Pembimbing")
        self.institution = user_data.get("institution", "Universitas")
        self.date = user_data.get("date", datetime.now().strftime("%Y"))
        self.abstract_id = user_data.get("abstract_id", "")
        self.abstract_en = user_data.get("abstract_en", "")
        self.keywords = user_data.get("keywords", [])
        self.preface = user_data.get("preface", "")
    
    def create_title_page(self, doc: Document) -> None:
        """Create Halaman Judul (Title Page) - pyramid inverted style."""
        # Clear existing content or add to new document
        doc.add_page_break()
        
        # Add spacing
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Title in pyramid inverted format (longest at top)
        title_para = doc.add_paragraph(self.title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        
        # Add spacing
        for _ in range(4):
            doc.add_paragraph()
        
        # Author info
        author_para = doc.add_paragraph(f"Oleh:\n{self.author}\nNIM: {self.nim}")
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in author_para.runs:
            run.font.size = Pt(12)
        
        # Add spacing
        for _ in range(6):
            doc.add_paragraph()
        
        # Institution info
        institution_para = doc.add_paragraph(
            f"{self.institution}\n{datetime.now().year}"
        )
        institution_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in institution_para.runs:
            run.font.size = Pt(12)
    
    def create_approval_page(self, doc: Document, approval_type: str = "supervisor") -> None:
        """Create Halaman Pengesahan (Approval Page).
        
        Args:
            doc: Document object
            approval_type: "supervisor" or "examiner"
        """
        doc.add_page_break()
        
        # Title
        title = "HALAMAN PENGESAHAN DOSEN PEMBIMBING" if approval_type == "supervisor" else "HALAMAN PENGESAHAN DOSEN PENGUJI"
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.bold = True
        title_run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Content
        if approval_type == "supervisor":
            content = f"""Judul: {self.title}
Penulis: {self.author}
NIM: {self.nim}

Dosen Pembimbing:

_________________________
{self.advisor}"""
        else:
            content = f"""Judul: {self.title}
Penulis: {self.author}
NIM: {self.nim}

Telah dipertahankan di depan Tim Penguji Skripsi pada hari _____ tanggal _____ tahun _____

Tim Penguji:
1. _________________________ (Ketua)
2. _________________________ (Anggota)
3. _________________________ (Anggota)"""
        
        content_para = doc.add_paragraph(content)
        content_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in content_para.runs:
            run.font.size = Pt(11)
    
    def create_originality_statement(self, doc: Document) -> None:
        """Create Halaman Pernyataan Keaslian (Originality Statement)."""
        doc.add_page_break()
        
        # Title
        title_para = doc.add_paragraph("PERNYATAAN KEASLIAN SKRIPSI")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.bold = True
        title_run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Statement text
        statement = f"""Saya yang bertanda tangan di bawah ini:

Nama: {self.author}
NIM: {self.nim}
Program Studi: [Tuliskan Program Studi]
Universitas: {self.institution}

Dengan ini menyatakan bahwa skripsi yang berjudul:

"{self.title}"

adalah benar-benar karya saya sendiri dan bukan merupakan duplikasi atau plagiasi dari karya orang lain. Apabila di kemudian hari terbukti bahwa skripsi ini adalah hasil plagiasi, maka saya siap menerima konsekuensi hukum yang berlaku.

Demikian pernyataan ini saya buat dengan sebenar-benarnya.

Jakarta, {datetime.now().strftime("%d %B %Y")}

_________________________
{self.author}
NIM: {self.nim}"""
        
        statement_para = doc.add_paragraph(statement)
        statement_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in statement_para.runs:
            run.font.size = Pt(11)
    
    def create_dedication_page(self, doc: Document, dedication_text: str = None) -> None:
        """Create Halaman Persembahan (Dedication Page)."""
        doc.add_page_break()
        
        # Title (optional)
        title_para = doc.add_paragraph("PERSEMBAHAN")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.bold = True
        title_run.font.size = Pt(12)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Dedication content
        dedication = dedication_text or f"Skripsi ini saya persembahkan untuk:\n\n[Tuliskan persembahan Anda di sini]\n\nDari {self.author}"
        
        ded_para = doc.add_paragraph(dedication)
        ded_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in ded_para.runs:
            run.font.size = Pt(11)
    
    def create_motto_page(self, doc: Document, motto_text: str = None) -> None:
        """Create Halaman Moto (Motto Page)."""
        doc.add_page_break()
        
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Motto content
        motto = motto_text or '"[Tuliskan moto atau kutipan inspiratif Anda di sini]"'
        
        motto_para = doc.add_paragraph(motto)
        motto_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in motto_para.runs:
            run.font.size = Pt(11)
            run.font.italic = True
    
    def create_preface(self, doc: Document) -> None:
        """Create Kata Pengantar (Preface)."""
        doc.add_page_break()
        
        # Title
        title_para = doc.add_paragraph("KATA PENGANTAR")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.bold = True
        title_run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Preface content
        preface_content = self.preface or f"""Puji syukur kami panjatkan kepada Tuhan Yang Maha Esa atas segala rahmat dan hidayahnya sehingga penulis dapat menyelesaikan skripsi ini dengan baik.

Skripsi dengan judul "{self.title}" ini dibuat sebagai salah satu persyaratan untuk memperoleh gelar sarjana pada Program Studi [Tuliskan Program Studi] di {self.institution}.

Penulis mengucapkan terima kasih kepada semua pihak yang telah membantu dalam penyelesaian skripsi ini, khususnya kepada:

1. {self.advisor} selaku dosen pembimbing
2. Orang tua dan keluarga yang memberikan dukungan moral dan material
3. Teman-teman yang telah memberikan inspirasi dan semangat

Penulis menyadari bahwa skripsi ini masih memiliki banyak kekurangan. Oleh karena itu, penulis dengan hati yang tulus menerima segala kritik dan saran untuk penyempurnaan skripsi ini.

Semoga skripsi ini dapat memberikan manfaat bagi pengembangan ilmu pengetahuan.

Jakarta, {datetime.now().strftime("%B %Y")}

_________________________
{self.author}"""
        
        preface_para = doc.add_paragraph(preface_content)
        preface_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._set_paragraph_indent(preface_para, left=Inches(1.0))
        for run in preface_para.runs:
            run.font.size = Pt(11)
    
    def create_abstract(self, doc: Document, language: str = "id") -> None:
        """Create Sari/Abstract page.
        
        Args:
            doc: Document object
            language: "id" for Indonesian or "en" for English
        """
        doc.add_page_break()
        
        # Title
        title = "SARI" if language == "id" else "ABSTRACT"
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.bold = True
        title_run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Abstract content
        abstract_text = self.abstract_id if language == "id" else self.abstract_en
        abstract_text = abstract_text or f"[Tuliskan abstrak dalam bahasa {'Indonesia' if language == 'id' else 'Inggris'} di sini. Maksimal 250 kata yang merangkum latar belakang, metodologi, hasil, dan kesimpulan penelitian.]"
        
        abstract_para = doc.add_paragraph(abstract_text)
        abstract_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._set_paragraph_indent(abstract_para, left=Inches(1.0))
        for run in abstract_para.runs:
            run.font.size = Pt(11)
        
        # Keywords
        doc.add_paragraph()
        keywords_text = ", ".join(self.keywords) if self.keywords else "[tuliskan kata kunci dipisahkan dengan koma]"
        keywords_para = doc.add_paragraph(f"Kata Kunci: {keywords_text}")
        for run in keywords_para.runs:
            run.font.size = Pt(11)
    
    def create_table_of_contents(self, doc: Document) -> None:
        """Create Daftar Isi (Table of Contents).
        
        Note: In a real implementation, this would be generated from headings.
        For now, we create a placeholder that can be auto-updated in Word.
        """
        doc.add_page_break()
        
        # Title
        title_para = doc.add_paragraph("DAFTAR ISI")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.bold = True
        title_run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Note about auto-generation
        note_para = doc.add_paragraph("[Daftar isi akan diperbarui secara otomatis. Tekan Ctrl+A kemudian F9 di Microsoft Word untuk memperbarui.]")
        for run in note_para.runs:
            run.font.italic = True
            run.font.size = Pt(10)
    
    def create_glossary(self, doc: Document, glossary_items: Dict[str, str] = None) -> None:
        """Create Glosarium (Glossary)."""
        if not glossary_items:
            glossary_items = {}
        
        doc.add_page_break()
        
        # Title
        title_para = doc.add_paragraph("GLOSARIUM")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.bold = True
        title_run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Glossary items
        if glossary_items:
            for term, definition in sorted(glossary_items.items()):
                term_para = doc.add_paragraph()
                term_run = term_para.add_run(f"{term}: ")
                term_run.font.bold = True
                term_run.font.size = Pt(11)
                
                def_run = term_para.add_run(definition)
                def_run.font.size = Pt(11)
        else:
            placeholder = doc.add_paragraph("[Tuliskan istilah-istilah khusus dan definisinya secara alfabetis]")
            for run in placeholder.runs:
                run.font.italic = True
                run.font.size = Pt(11)
    
    def _set_paragraph_indent(self, paragraph, left: Inches = None, first_line: Inches = None):
        """Set paragraph indentation."""
        pPr = paragraph._element.get_or_add_pPr()
        if left:
            pInd = OxmlElement('w:ind')
            pInd.set(qn('w:left'), str(int(left.twips)))
            pPr.append(pInd)
        if first_line:
            pInd = pPr.find(qn('w:ind'))
            if pInd is None:
                pInd = OxmlElement('w:ind')
                pPr.append(pInd)
            pInd.set(qn('w:firstLine'), str(int(first_line.twips)))


class BackMatterGenerator:
    """Generates back matter pages (appendices, references, etc.)."""
    
    def __init__(self, user_data: Dict[str, Any] = None):
        """Initialize with user data."""
        self.user_data = user_data or {}
    
    def create_list_of_tables(self, doc: Document) -> None:
        """Create Daftar Tabel (List of Tables)."""
        doc.add_page_break()
        
        # Title
        title_para = doc.add_paragraph("DAFTAR TABEL")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.bold = True
        title_run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        note_para = doc.add_paragraph("[Daftar tabel akan diperbarui otomatis berdasarkan caption tabel dalam dokumen]")
        for run in note_para.runs:
            run.font.italic = True
            run.font.size = Pt(10)
    
    def create_list_of_figures(self, doc: Document) -> None:
        """Create Daftar Gambar (List of Figures)."""
        doc.add_page_break()
        
        # Title
        title_para = doc.add_paragraph("DAFTAR GAMBAR")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.bold = True
        title_run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        note_para = doc.add_paragraph("[Daftar gambar akan diperbarui otomatis berdasarkan caption gambar dalam dokumen]")
        for run in note_para.runs:
            run.font.italic = True
            run.font.size = Pt(10)
    
    def create_references(self, doc: Document, references: List[Dict[str, str]] = None) -> None:
        """Create Daftar Pustaka (References/Bibliography) in APA format."""
        doc.add_page_break()
        
        # Title
        title_para = doc.add_paragraph("DAFTAR PUSTAKA")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.bold = True
        title_run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        if references:
            for ref in references:
                # Skip empty references
                if not ref:
                    continue
                    
                ref_text = self._format_apa_reference(ref)
                if ref_text:  # Only add if we got valid text
                    ref_para = doc.add_paragraph(ref_text)
                    # Hanging indent
                    self._set_hanging_indent(ref_para, Inches(0.5))
                    for run in ref_para.runs:
                        run.font.size = Pt(11)
        else:
            placeholder = doc.add_paragraph("[Tuliskan referensi dalam format APA 6th edition]")
            for run in placeholder.runs:
                run.font.italic = True
                run.font.size = Pt(11)
    
    def create_appendices(self, doc: Document, appendix_items: List[Dict[str, str]] = None) -> None:
        """Create Lampiran (Appendices)."""
        if not appendix_items:
            appendix_items = []
        
        doc.add_page_break()
        
        # Title
        title_para = doc.add_paragraph("LAMPIRAN")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.bold = True
        title_run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        if appendix_items and isinstance(appendix_items, list):
            for idx, item in enumerate(appendix_items, 1):
                # Skip empty items
                if not item:
                    continue
                
                try:
                    # Handle both string and dict items
                    if isinstance(item, str):
                        # If item is a string, use it as content
                        heading = f"Lampiran {chr(64 + idx)}"
                        head_para = doc.add_paragraph(heading)
                        head_run = head_para.runs[0]
                        head_run.font.bold = True
                        head_run.font.size = Pt(12)
                        
                        content_para = doc.add_paragraph(item)
                        for run in content_para.runs:
                            run.font.size = Pt(11)
                    elif isinstance(item, dict):
                        # If item is dict, extract title and content
                        heading = f"Lampiran {chr(64 + idx)}: {item.get('title', f'Appendix {idx}')}"
                        head_para = doc.add_paragraph(heading)
                        head_run = head_para.runs[0]
                        head_run.font.bold = True
                        head_run.font.size = Pt(12)
                        
                        # Content
                        content = item.get("content", "")
                        if content:
                            content_para = doc.add_paragraph(content)
                            for run in content_para.runs:
                                run.font.size = Pt(11)
                    else:
                        # For other types, convert to string
                        heading = f"Lampiran {chr(64 + idx)}"
                        head_para = doc.add_paragraph(heading)
                        head_run = head_para.runs[0]
                        head_run.font.bold = True
                        head_run.font.size = Pt(12)
                        
                        content_para = doc.add_paragraph(str(item))
                        for run in content_para.runs:
                            run.font.size = Pt(11)
                    
                    # Page break between appendices
                    if idx < len(appendix_items):
                        doc.add_page_break()
                except Exception as e:
                    print(f"[WARNING] Error processing appendix {idx}: {e}")
                    continue
        else:
            placeholder = doc.add_paragraph("[Lampiran disusun dengan huruf besar (A, B, C, dst.) dan diletakkan setelah Daftar Pustaka]")
            for run in placeholder.runs:
                run.font.italic = True
                run.font.size = Pt(11)
    
    def _format_apa_reference(self, ref: Any) -> str:
        """Format a reference in APA 6th edition style.
        
        Args:
            ref: Reference as dict or string
            
        Returns:
            Formatted reference string
        """
        # Handle string references (already formatted)
        if isinstance(ref, str):
            return ref.strip() if ref else ""
        
        # Handle dictionary references
        if not isinstance(ref, dict):
            return ""
        
        ref_type = ref.get("type", "book").lower() if ref else "book"
        
        if ref_type == "book":
            return f"{ref.get('author', 'Unknown')}. ({ref.get('year', 'n.d.')}). {ref.get('title', 'Unknown title')}. {ref.get('publisher', '')}."
        elif ref_type == "journal":
            return f"{ref.get('author', 'Unknown')}. ({ref.get('year', 'n.d.')}). {ref.get('title', 'Unknown title')}. {ref.get('journal', '')}, {ref.get('volume', '')}({ref.get('issue', '')}), {ref.get('pages', '')}."
        elif ref_type == "website":
            return f"{ref.get('author', 'Unknown')}. ({ref.get('year', 'n.d.')}). {ref.get('title', 'Unknown title')}. Retrieved from {ref.get('url', '')}"
        else:
            return f"{ref.get('author', 'Unknown')}. ({ref.get('year', 'n.d.')}). {ref.get('title', 'Unknown title')}."
    
    def _set_hanging_indent(self, paragraph, indent: Inches):
        """Set hanging indent for bibliography entries."""
        pPr = paragraph._element.get_or_add_pPr()
        pInd = OxmlElement('w:ind')
        pInd.set(qn('w:left'), str(int(indent.twips)))
        pInd.set(qn('w:hanging'), str(int(indent.twips)))
        pPr.append(pInd)
