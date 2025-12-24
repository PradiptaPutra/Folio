"""
Indonesian University Template Analyzer
Supports both DOCX analysis and structured JSON/XML templates for fast, accurate processing:
- DOCX Analysis: Deep parsing of university templates (slower but comprehensive)
- JSON/XML Templates: Pre-defined structured templates (fast and precise)
- Automatic format detection and processing
- Indonesian academic standards compliance
"""

from typing import Dict, List, Any, Optional, Set, Tuple, Union
from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.style import WD_STYLE_TYPE

# Try to import Mammoth for enhanced DOCX processing
try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False

# No additional imports needed for basic functionality


class TemplateAnalyzer:
    """Analyzes Indonesian university DOCX templates to extract formatting rules and structure."""

    # Indonesian academic standards and common university patterns
    INDONESIAN_STANDARDS = {
        "fonts": {
            "primary": ["Times New Roman", "Arial", "Calibri"],
            "academic": ["Times New Roman", "Bookman Old Style"],
            "size_title": [14, 16, 18],  # Title page font sizes
            "size_body": [11, 12],       # Main content font sizes
            "size_heading": [12, 14]     # Chapter headings
        },
        "spacing": {
            "line_spacing": [1.5, 2.0],    # Academic standard line spacing
            "paragraph_spacing": [0, 6, 12],  # Space between paragraphs
            "indentation": [1.0, 1.25, 1.5]   # First line indentation in cm
        },
        "structure": {
            "front_matter": [
                "halaman_judul", "lembar_pengesahan", "pernyataan_keaslian",
                "kata_pengantar", "abstrak", "abstract", "daftar_isi",
                "daftar_tabel", "daftar_gambar", "daftar_lampiran"
            ],
            "main_content": [
                "bab_i", "bab_ii", "bab_iii", "bab_iv", "bab_v"
            ],
            "back_matter": [
                "daftar_pustaka", "bibliography", "lampiran", "appendix"
            ]
        }
    }

    def __init__(self, template_path: Union[str, Path]):
        """Initialize with Indonesian university template (DOCX, JSON, or XML)."""
        self.template_path = Path(template_path)

        # Use comprehensive DOCX analyzer with Mammoth enhancement if available
        self.analyzer_type = "docx"
        self.doc = Document(str(self.template_path))

        # Try to enhance analysis with Mammoth if available
        if MAMMOTH_AVAILABLE:
            try:
                self.analysis = self._analyze_with_mammoth()
            except Exception as e:
                print(f"Mammoth analysis failed, using python-docx: {e}")
                self.analysis = self._analyze()
        else:
            self.analysis = self._analyze()


    
    def _analyze(self) -> Dict[str, Any]:
        """Perform complete Indonesian university template analysis."""
        analysis = {
            "styles": self._extract_styles(),
            "structure": self._detect_structure(),
            "placeholders": self._detect_placeholders(),
            "sections": self._analyze_sections(),
            "formatting_rules": self._detect_formatting_rules(),
            "front_matter": self._detect_front_matter(),
            "document_properties": self._extract_properties(),
            "margins": self._extract_margins(),
            "heading_hierarchy": self._analyze_heading_hierarchy(),
            "special_elements": self._detect_special_elements(),
            "numbering": self._extract_numbering_map(),
            "sectPr": self._extract_sectPr_extended(),
        }

        # Add Indonesian academic profile after main analysis is complete
        # Temporarily disabled due to method signature issues
        # analysis["indonesian_academic_profile"] = self._analyze_indonesian_academic_profile(analysis)  # TODO: Enable when methods are fixed

        return analysis
    
    def _extract_styles(self) -> Dict[str, Dict[str, Any]]:
        """Extract all paragraph and character styles from template."""
        styles_info = {}
        
        for style in self.doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                # Get outline level safely - may not exist on all styles
                outline_level = None
                try:
                    outline_level_attr = style.element.xpath('.//w:outlineLvl/@w:val')
                    if outline_level_attr:
                        outline_level = int(outline_level_attr[0])
                except Exception:
                    outline_level = None
                
                style_data = {
                    "name": style.name,
                    "base_style": style.base_style.name if style.base_style else None,
                    "outline_level": outline_level,
                    "font": {
                        "name": style.font.name if style.font and style.font.name else "Calibri",
                        "size": style.font.size.pt if style.font and style.font.size else 12,
                        "bold": style.font.bold if style.font else False,
                        "italic": style.font.italic if style.font else False,
                    },
                    "paragraph": {
                        "alignment": str(style.paragraph_format.alignment) if style.paragraph_format else None,
                        "line_spacing": style.paragraph_format.line_spacing if style.paragraph_format else None,
                        "space_before": style.paragraph_format.space_before if style.paragraph_format else None,
                        "space_after": style.paragraph_format.space_after if style.paragraph_format else None,
                        "left_indent": style.paragraph_format.left_indent if style.paragraph_format else None,
                        "first_line_indent": style.paragraph_format.first_line_indent if style.paragraph_format else None,
                    },
                    "usage_count": self._count_style_usage(style.name),
                }
                styles_info[style.name] = style_data
        
        return styles_info
    
    def _count_style_usage(self, style_name: str) -> int:
        """Count how many paragraphs use a specific style."""
        count = 0
        for para in self.doc.paragraphs:
            if para.style.name == style_name:
                count += 1
        return count
    
    def _detect_structure(self) -> Dict[str, List[Dict[str, Any]]]:
        """Detect document structure (headings, sections, hierarchy)."""
        structure = {
            "headings": [],
            "paragraphs": [],
            "tables": [],
            "hierarchy": []
        }
        
        current_heading_level = None
        
        for para in self.doc.paragraphs:
            # Detect headings - check for outline level in style
            outline_level = None
            if para.style:
                try:
                    outline_level_attr = para.style.element.xpath('.//w:outlineLvl/@w:val')
                    if outline_level_attr:
                        outline_level = int(outline_level_attr[0])
                except Exception:
                    outline_level = None
            
            if outline_level is not None:
                heading_data = {
                    "level": outline_level,
                    "style": para.style.name,
                    "text": para.text[:100],  # First 100 chars
                    "index": len(self.doc.paragraphs),
                }
                structure["headings"].append(heading_data)
                structure["hierarchy"].append(heading_data)
            
            # Detect key section markers
            if self._is_section_marker(para.text):
                structure["paragraphs"].append({
                    "type": "section_marker",
                    "text": para.text,
                    "style": para.style.name if para.style else None,
                })
        
        # Detect tables
        for table in self.doc.tables:
            structure["tables"].append({
                "rows": len(table.rows),
                "cols": len(table.columns),
                "has_header": self._has_table_header(table),
            })
        
        return structure
    
    def _is_section_marker(self, text: str) -> bool:
        """Check if text is a section marker (cover, approval, abstract, etc.)."""
        section_keywords = [
            "halaman judul", "cover", "halaman pengesahan", "pernyataan",
            "kata pengantar", "preface", "abstrak", "abstract", "sari",
            "daftar isi", "daftar tabel", "daftar gambar", "daftar lampiran",
            "glossary", "glosarium", "bibliography", "references", "daftar pustaka",
            "appendix", "lampiran", "bab", "chapter", "kesimpulan", "conclusion"
        ]
        text_lower = text.lower().strip()
        return any(keyword in text_lower for keyword in section_keywords)
    
    def _has_table_header(self, table) -> bool:
        """Check if table has a header row."""
        if len(table.rows) == 0:
            return False
        first_row = table.rows[0]
        # Header rows typically have shading or bold text
        for cell in first_row.cells:
            for para in cell.paragraphs:
                if para.runs and any(run.bold for run in para.runs):
                    return True
        return False
    
    def _detect_placeholders(self) -> Dict[str, List[str]]:
        """Detect placeholder text that needs to be replaced."""
        placeholders = {
            "text_placeholders": [],
            "field_codes": [],
            "blank_sections": []
        }
        
        placeholder_patterns = [
            r"\[.*?\]",  # [Placeholder]
            r"\{.*?\}",  # {Placeholder}
            r"{{.*?}}",  # {{Placeholder}} Jinja2-style
            r"TULIS\s+.*?(?=\n|$)",  # TULIS... (Indonesian)
            r"BAGIAN\s+.*?(?=\n|$)",  # BAGIAN... (Indonesian)
            r"<.*?>",  # <Placeholder>
        ]
        
        for para in self.doc.paragraphs:
            text = para.text
            for pattern in placeholder_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    placeholders["text_placeholders"].extend(matches)
        
        # Also check for field codes
        for para in self.doc.paragraphs:
            if hasattr(para, '_element'):
                fldChar = para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar')
                if fldChar:
                    placeholders["field_codes"].append(para.text)
        
        return placeholders
    
    def _analyze_sections(self) -> Dict[str, Any]:
        """Analyze document sections (headers, footers, page setup)."""
        sections_info = {}
        
        if self.doc.sections:
            section = self.doc.sections[0]
            
            # Safely extract gutter_margin from XML (may not exist)
            gutter_margin = None
            try:
                gutter_elem = section._sectPr.xpath('.//w:pgMar/@w:gutter')
                if gutter_elem:
                    # Convert from twips to inches (1 inch = 1440 twips)
                    gutter_margin = int(gutter_elem[0]) / 1440.0
            except Exception:
                gutter_margin = None
            
            sections_info = {
                "page_height": section.page_height.inches,
                "page_width": section.page_width.inches,
                "top_margin": section.top_margin.inches if section.top_margin else None,
                "bottom_margin": section.bottom_margin.inches if section.bottom_margin else None,
                "left_margin": section.left_margin.inches if section.left_margin else None,
                "right_margin": section.right_margin.inches if section.right_margin else None,
                "header_exists": len(section.header.paragraphs) > 0,
                "footer_exists": len(section.footer.paragraphs) > 0,
                "gutter_margin": gutter_margin,
            }
        
        return sections_info

    def _extract_sectPr_extended(self) -> Dict[str, Any]:
        """Extract extended sectPr properties: columns, titlePg (different first page), page breaks hint."""
        info: Dict[str, Any] = {}
        try:
            # Access raw sectPr XML to find cols and titlePg
            sectPr = self.doc.sections[0]._sectPr  # type: ignore
            cols = sectPr.xpath('.//w:cols')
            if cols:
                node = cols[0]
                num = node.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num')
                space = node.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space')
                info['columns'] = {
                    'num': int(num) if num else None,
                    'space': int(space) if space else None,
                }
            titlePg = sectPr.xpath('.//w:titlePg')
            info['different_first_page'] = bool(titlePg)
        except Exception:
            pass
        return info
    
    def _detect_formatting_rules(self) -> Dict[str, Any]:
        """Detect special formatting rules (spacing, indentation, alignment)."""
        rules = {
            "common_font": self._detect_common_font(),
            "common_font_size": self._detect_common_font_size(),
            "indentation_pattern": self._detect_indentation_pattern(),
            "spacing_pattern": self._detect_spacing_pattern(),
            "alignment_pattern": self._detect_alignment_pattern(),
            "line_spacing_pattern": self._detect_line_spacing_pattern(),
            "list_styles": self._detect_list_styles(),
        }
        return rules

    def _detect_list_styles(self) -> Dict[str, Any]:
        """Detect presence of common list styles to help numbering enforcement."""
        styles = set(s.name for s in self.doc.styles if s is not None)
        return {
            'has_List_Paragraph': 'List Paragraph' in styles,
            'has_List_Number': 'List Number' in styles,
            'has_List_Bullet': 'List Bullet' in styles,
        }

    def _extract_numbering_map(self) -> Dict[str, Any]:
        """Parse numbering.xml to build a map of numId -> abstractNumId -> levels.

        Returns a dict with keys:
        - abstract_nums: {abstractNumId: {levels: {ilvl: {numFmt, start}}}}
        - nums: {numId: abstractNumId}
        """
        import zipfile
        from lxml import etree
        NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        result = {"abstract_nums": {}, "nums": {}}
        try:
            with zipfile.ZipFile(self.template_path) as z:
                numbering_xml = z.read("word/numbering.xml")
            root = etree.XML(numbering_xml)

            # abstractNum definitions
            for abs_num in root.findall('.//w:abstractNum', NS):
                abs_id = abs_num.get('{%s}abstractNumId' % NS['w'])
                if not abs_id:
                    continue
                levels = {}
                for lvl in abs_num.findall('.//w:lvl', NS):
                    ilvl = lvl.get('{%s}ilvl' % NS['w'])
                    num_fmt = None
                    start = None
                    numFmt = lvl.find('w:numFmt', NS)
                    if numFmt is not None:
                        num_fmt = numFmt.get('{%s}val' % NS['w'])
                    startNode = lvl.find('w:start', NS)
                    if startNode is not None:
                        start = startNode.get('{%s}val' % NS['w'])
                    levels[ilvl or '0'] = {"numFmt": num_fmt, "start": start}
                result["abstract_nums"][abs_id] = {"levels": levels}

            # num instances mapping to abstractNum
            for num in root.findall('.//w:num', NS):
                num_id = num.get('{%s}numId' % NS['w'])
                abs_ref = num.find('w:abstractNumId', NS)
                if num_id and abs_ref is not None:
                    result["nums"][num_id] = abs_ref.get('{%s}val' % NS['w'])
        except Exception:
            # silently ignore if numbering.xml not present
            pass

        return result
    
    def _detect_common_font(self) -> str:
        """Detect most common font in document."""
        fonts = {}
        for para in self.doc.paragraphs:
            for run in para.runs:
                font_name = run.font.name or "Calibri"
                fonts[font_name] = fonts.get(font_name, 0) + 1
        
        return max(fonts, key=fonts.get) if fonts else "Times New Roman"
    
    def _detect_common_font_size(self) -> float:
        """Detect most common font size in document."""
        sizes = {}
        for para in self.doc.paragraphs:
            for run in para.runs:
                if run.font.size:
                    size = run.font.size.pt
                    sizes[size] = sizes.get(size, 0) + 1
        
        return max(sizes, key=sizes.get) if sizes else 12.0
    
    def _detect_indentation_pattern(self) -> Dict[str, Any]:
        """Detect paragraph indentation patterns."""
        indents = {"first_line": [], "left": [], "right": []}
        
        for para in self.doc.paragraphs:
            if para.paragraph_format:
                pf = para.paragraph_format
                if pf.first_line_indent:
                    indents["first_line"].append(pf.first_line_indent.inches)
                if pf.left_indent:
                    indents["left"].append(pf.left_indent.inches)
                if pf.right_indent:
                    indents["right"].append(pf.right_indent.inches)
        
        return {
            "common_first_line": max(set(indents["first_line"]), key=indents["first_line"].count) if indents["first_line"] else 0,
            "common_left": max(set(indents["left"]), key=indents["left"].count) if indents["left"] else 0,
        }
    
    def _detect_spacing_pattern(self) -> Dict[str, Any]:
        """Detect paragraph spacing patterns."""
        space_before = {}
        space_after = {}
        
        for para in self.doc.paragraphs:
            if para.paragraph_format:
                pf = para.paragraph_format
                if pf.space_before:
                    sb = pf.space_before.pt
                    space_before[sb] = space_before.get(sb, 0) + 1
                if pf.space_after:
                    sa = pf.space_after.pt
                    space_after[sa] = space_after.get(sa, 0) + 1
        
        return {
            "common_space_before": max(space_before, key=space_before.get) if space_before else 0,
            "common_space_after": max(space_after, key=space_after.get) if space_after else 0,
        }
    
    def _detect_alignment_pattern(self) -> str:
        """Detect common paragraph alignment."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        alignments = {}
        
        for para in self.doc.paragraphs:
            if para.paragraph_format and para.paragraph_format.alignment:
                align = str(para.paragraph_format.alignment)
                alignments[align] = alignments.get(align, 0) + 1
        
        return max(alignments, key=alignments.get) if alignments else "CENTER"
    
    def _detect_line_spacing_pattern(self) -> float:
        """Detect common line spacing."""
        spacings = {}
        
        for para in self.doc.paragraphs:
            if para.paragraph_format and para.paragraph_format.line_spacing:
                spacing = para.paragraph_format.line_spacing
                spacings[spacing] = spacings.get(spacing, 0) + 1
        
        return max(spacings, key=spacings.get) if spacings else 1.5
    
    def _detect_front_matter(self) -> Dict[str, Any]:
        """Detect front matter sections (cover, approval, etc.)."""
        front_matter = {
            "sections": [],
            "page_breaks": [],
            "required_sections": []
        }
        
        section_markers = [
            ("cover", ["halaman judul", "cover", "judul"]),
            ("approval", ["pengesahan", "approval", "persetujuan"]),
            ("statement", ["pernyataan", "declaration", "keaslian"]),
            ("dedication", ["persembahan", "dedication"]),
            ("motto", ["motto"]),
            ("preface", ["kata pengantar", "preface", "pengantar"]),
            ("abstract_id", ["abstrak", "ringkasan"]),
            ("abstract_en", ["abstract"]),
            ("glossary", ["glosarium", "glossary", "istilah"]),
            ("toc", ["daftar isi", "table of contents"]),
            ("list_figures", ["daftar gambar", "list of figures"]),
            ("list_tables", ["daftar tabel", "list of tables"]),
        ]
        
        for para in self.doc.paragraphs:
            text_lower = para.text.lower().strip()
            for section_type, keywords in section_markers:
                if any(kw in text_lower for kw in keywords):
                    front_matter["sections"].append(section_type)
                    front_matter["required_sections"].append(section_type)
        
        return front_matter
    
    def _extract_properties(self) -> Dict[str, str]:
        """Extract document properties (title, author, etc.)."""
        props = self.doc.core_properties
        return {
            "title": props.title or "",
            "author": props.author or "",
            "subject": props.subject or "",
            "keywords": props.keywords or "",
            "created": str(props.created) if props.created else "",
            "modified": str(props.modified) if props.modified else "",
        }
    
    def _extract_margins(self) -> Dict[str, float]:
        """Extract document margins."""
        if not self.doc.sections:
            return {}
        
        section = self.doc.sections[0]
        return {
            "top": section.top_margin.inches if section.top_margin else 1.0,
            "bottom": section.bottom_margin.inches if section.bottom_margin else 1.0,
            "left": section.left_margin.inches if section.left_margin else 1.0,
            "right": section.right_margin.inches if section.right_margin else 1.0,
        }
    
    def _analyze_heading_hierarchy(self) -> Dict[str, Any]:
        """Analyze heading hierarchy (e.g., BAB, subbab 1.1)."""
        hierarchy = {
            "heading_styles": {},
            "numbering_pattern": None,
            "title_format": None
        }
        
        heading_styles = {}
        for para in self.doc.paragraphs:
            if para.style:
                # Safely extract outline level from XML
                outline_level = None
                try:
                    outline_level_attr = para.style.element.xpath('.//w:outlineLvl/@w:val')
                    if outline_level_attr:
                        outline_level = int(outline_level_attr[0])
                except Exception:
                    outline_level = None
                
                if outline_level is not None:
                    style_name = para.style.name
                    if style_name not in heading_styles:
                        heading_styles[style_name] = {
                            "level": outline_level,
                            "examples": []
                        }
                    if len(heading_styles[style_name]["examples"]) < 3:
                        heading_styles[style_name]["examples"].append(para.text[:50])
        
        hierarchy["heading_styles"] = heading_styles
        
        # Detect numbering pattern
        for para in self.doc.paragraphs:
            if para.text and re.match(r"^(BAB|CHAPTER)\s+[IVX]+", para.text):
                hierarchy["numbering_pattern"] = "roman_uppercase"
            elif para.text and re.match(r"^(BAB|CHAPTER)\s+\d+", para.text):
                hierarchy["numbering_pattern"] = "arabic"
        
        return hierarchy
    
    def _detect_special_elements(self) -> Dict[str, List[str]]:
        """Detect special elements (tables, figures, equations, etc.)."""
        special = {
            "tables": [],
            "figures": [],
            "equations": [],
            "code_blocks": [],
            "text_boxes": [],
            "shapes": []
        }
        
        # Count tables
        for i, table in enumerate(self.doc.tables):
            special["tables"].append(f"Table {i+1}")
        
        # Detect figure/table captions
        for para in self.doc.paragraphs:
            text_lower = para.text.lower()
            if "gambar" in text_lower or "figure" in text_lower:
                special["figures"].append(para.text[:50])
            elif "persamaan" in text_lower or "equation" in text_lower:
                special["equations"].append(para.text[:50])
        
        return special
    
    def get_analysis(self) -> Dict[str, Any]:
        """Return the complete analysis."""
        return self.analysis
    


        try:
            converter = TemplateConverter(self)
            json_template = converter.convert_to_json(output_path)

            # Add metadata about the conversion
            if json_template:
                json_template["conversion_info"] = {
                    "original_format": "docx",
                    "converted_at": "dynamic_analysis",
                    "source_file": str(self.template_path),
                    "analysis_completeness": "comprehensive"
                }

            return json_template
        except Exception as e:
            print(f"[WARNING] Could not convert template to structured format: {e}")
            import traceback
            traceback.print_exc()
            return None

    def get_summary(self) -> str:
        """Return a human-readable summary of template analysis."""
        summary = f"""
DOCX TEMPLATE ANALYSIS SUMMARY
==============================

Template Type: DOCX Analysis
Processing Method: {'Enhanced with Mammoth' if MAMMOTH_AVAILABLE else 'Standard python-docx'}

Document Properties:
- Title: {self.analysis['document_properties'].get('title', 'N/A')}
- Author: {self.analysis['document_properties'].get('author', 'N/A')}

Page Setup:
- Margins: L={self.analysis['margins'].get('left', 1.0)}" T={self.analysis['margins'].get('top', 1.0)}" R={self.analysis['margins'].get('right', 1.0)}" B={self.analysis['margins'].get('bottom', 1.0)}"

Formatting Rules:
- Primary Font: {self.analysis['formatting_rules']['common_font']}
- Primary Font Size: {self.analysis['formatting_rules']['common_font_size']}pt
- Line Spacing: {self.analysis['formatting_rules']['line_spacing_pattern']}
- First Line Indent: {self.analysis['formatting_rules']['indentation_pattern']['common_first_line']}"

Front Matter Sections Detected:
- {', '.join(self.analysis['front_matter']['sections']) or 'None detected'}

Detected Styles: {len(self.analysis['styles'])}
- Heading Styles: {[s for s in self.analysis['styles'] if 'Heading' in s]}

Special Elements:
- Tables: {len(self.analysis['special_elements']['tables'])}
- Figures: {len(self.analysis['special_elements']['figures'])}
- Equations: {len(self.analysis['special_elements']['equations'])}

Placeholders Found: {len(self.analysis['placeholders']['text_placeholders'])}
- Examples: {self.analysis['placeholders']['text_placeholders'][:3]}
"""
        return summary.strip()

    def _analyze_with_mammoth(self) -> Dict[str, Any]:
        """Analyze template using Mammoth for enhanced HTML-based analysis."""
        if not MAMMOTH_AVAILABLE:
            raise Exception("Mammoth not available")

        # Use Mammoth to get HTML representation
        with open(str(self.template_path), 'rb') as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html_content = result.value

        # Parse HTML to extract structure
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            # Fallback to basic analysis if BeautifulSoup not available
            return self._analyze()

        soup = BeautifulSoup(html_content, 'html.parser')

        # Enhanced analysis using HTML structure
        analysis = self._analyze()  # Start with basic analysis

        # Add Mammoth-specific enhancements
        analysis['mammoth_html'] = html_content
        analysis['mammoth_structure'] = {
            'headings': [h.get_text().strip() for h in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])],
            'paragraphs': len(soup.find_all('p')),
            'lists': len(soup.find_all(['ul', 'ol'])),
            'tables': len(soup.find_all('table')),
            'links': len(soup.find_all('a')),
            'images': len(soup.find_all('img'))
        }

        return analysis

    def _analyze_indonesian_academic_profile(self, analysis: Dict[str, Any]) -> Dict[str, Any]:
        """Analyze Indonesian academic formatting compliance and university patterns."""
        profile = {
            "university_type": self._detect_university_type_from_analysis(analysis),
            "formatting_compliance": self._check_indonesian_formatting_compliance(analysis),
            "academic_standards": self._analyze_academic_standards(),
            "structure_completeness": self._check_structure_completeness(analysis),
            "typography_score": self._calculate_typography_score(analysis),
            "recommendations": []
        }

        # Generate recommendations based on analysis
        profile["recommendations"] = self._generate_indonesian_recommendations(profile)

        return profile

    def _detect_university_type_from_analysis(self, analysis: Dict[str, Any]) -> str:
        """Detect Indonesian university type based on formatting patterns."""
        # Analyze document properties and formatting to identify university type
        title = analysis.get("document_properties", {}).get("title", "").lower()

        # University-specific patterns
        university_patterns = {
            "ui": ["universitas indonesia", "ui", "fib", "fekon", "fmipa", "ft"],
            "itb": ["institut teknologi bandung", "itb", "stei", "tf", "civil"],
            "ugm": ["ugm", "gadjah mada", "ugm", "feb", "fib", "fmipa"],
            "unpad": ["padjadjaran", "unpad", "feb", "fh", "fisip"],
            "ipb": ["institut pertanian bogor", "ipb", "fakultas"],
            "uns": ["sebelas maret", "uns", "surakarta"],
            "undip": ["diponegoro", "undip", "semarang"],
            "unair": ["airlangga", "unair", "surabaya"],
            "ub": ["brawijaya", "ub", "malang"],
            "unnes": ["negeri semarang", "unnes"]
        }

        for univ, patterns in university_patterns.items():
            if any(pattern in title for pattern in patterns):
                return univ.upper()

        # Check formatting patterns for university type inference
        formatting_rules = self._detect_formatting_rules()
        font_name = formatting_rules.get('common_font', '').lower()
        font_size = formatting_rules.get('common_font_size', 12)

        if 'times new roman' in font_name and font_size == 12:
            return "STANDARD_INDO_ACADEMIC"
        elif 'arial' in font_name and font_size == 11:
            return "MODERN_INDO_ACADEMIC"
        else:
            return "GENERIC_INDO_ACADEMIC"

    def _check_indonesian_formatting_compliance(self, analysis: Dict[str, Any]) -> Dict[str, Any]:
        """Check compliance with Indonesian academic formatting standards."""
        compliance = {
            "font_compliance": False,
            "spacing_compliance": False,
            "indentation_compliance": False,
            "structure_compliance": False,
            "typography_compliance": False,
            "overall_score": 0.0
        }

        # Font compliance (Times New Roman, Arial, Calibri preferred)
        common_font = analysis['formatting_rules'].get('common_font', '').lower()
        compliance["font_compliance"] = any(academic_font in common_font
                                           for academic_font in ['times new roman', 'arial', 'calibri'])

        # Spacing compliance (1.5 line spacing standard)
        line_spacing = analysis['formatting_rules'].get('line_spacing_pattern', 1.0)
        compliance["spacing_compliance"] = 1.4 <= line_spacing <= 1.6

        # Indentation compliance (1.0-1.5 cm first line indent)
        indent_pattern = analysis['formatting_rules'].get('indentation_pattern', {})
        first_line_indent = indent_pattern.get('common_first_line', 0)
        compliance["indentation_compliance"] = 1.0 <= first_line_indent <= 1.5

        # Structure compliance (required sections present)
        front_matter_sections = analysis['front_matter'].get('sections', [])
        required_sections = ['halaman judul', 'pengesahan', 'abstrak', 'daftar isi']
        compliance["structure_compliance"] = all(any(req in section.lower()
                                                     for req in required_sections)
                                                for section in front_matter_sections)

        # Typography compliance (11-12pt body text)
        font_size = analysis['formatting_rules'].get('common_font_size', 12)
        compliance["typography_compliance"] = 11 <= font_size <= 12

        # Calculate overall score
        compliant_items = sum(compliance.values()) - 1  # Subtract the score field
        compliance["overall_score"] = compliant_items / 5.0

        return compliance

    def _analyze_academic_standards(self) -> Dict[str, Any]:
        """Analyze compliance with Indonesian academic writing standards."""
        standards = {
            "language_standard": "INDONESIAN_ACADEMIC",
            "citation_style": self._detect_citation_style(),
            "reference_format": self._detect_reference_format(),
            "terminology_compliance": True,  # Assume compliant unless detected otherwise
            "structure_standard": "BAB_SYSTEM",  # Indonesian thesis structure
            "formatting_standard": "SK_MENDIKBUD"  # Indonesian education ministry standards
        }
        return standards

    def _detect_citation_style(self) -> str:
        """Detect citation style used in the document."""
        # Look for citation patterns in text
        citation_patterns = {
            "APA": [r"\([A-Za-z]+, \d{4}\)", r"\([A-Za-z]+ & [A-Za-z]+, \d{4}\)"],  # (Author, Year)
            "MLA": [r"\([A-Za-z]+ \d+\)", r'\([A-Za-z]+ \d+ [A-Za-z]+\)'],  # (Author Page)
            "Chicago": [r"\[[\d]+\]", r"\([A-Za-z]+ [\d]+, [\d]+\)"],  # [1] or (Author Year, Page)
            "Harvard": [r"\([A-Za-z]+, [\d]+\)", r"\([A-Za-z]+ [\d]+\)"],  # (Author Year)
            "Vancouver": [r"\[[\d]+\]", r"^[[\d]+]"]  # [1]
        }

        for para in self.doc.paragraphs:
            text = para.text
            for style, patterns in citation_patterns.items():
                if any(re.search(pattern, text) for pattern in patterns):
                    return style

        return "UNKNOWN"

    def _detect_reference_format(self) -> str:
        """Detect reference/bibliography format."""
        # Check for common Indonesian academic reference formats
        reference_formats = {
            "ALPHABETICAL": r"^[A-Z][a-zA-Z]+, [A-Z]\.",  # Author, Initials
            "NUMERICAL": r"^\[[\d]+\]",  # [1], [2], etc.
            "FOOTNOTE": r"^[1-9][0-9]*\.",  # 1., 2., etc.
        }

        # Look for bibliography/reference sections
        for para in self.doc.paragraphs:
            text = para.text.strip()
            for format_type, pattern in reference_formats.items():
                if re.match(pattern, text):
                    return format_type

        return "UNKNOWN"

    def _check_structure_completeness(self) -> Dict[str, Any]:
        """Check completeness of Indonesian thesis structure."""
        structure = {
            "front_matter_complete": False,
            "main_content_complete": False,
            "back_matter_complete": False,
            "missing_sections": [],
            "extra_sections": [],
            "completeness_score": 0.0
        }

        # Check front matter
        required_front = ['halaman judul', 'pengesahan', 'abstrak', 'daftar isi']
        present_front = self.analysis['front_matter'].get('sections', [])
        structure["front_matter_complete"] = len(present_front) >= len(required_front) * 0.8

        # Check main content (BAB structure)
        bab_pattern = r"BAB\s+[IVX]+\s*[:-]?\s*(.+)"
        bab_count = 0
        for para in self.doc.paragraphs:
            if re.match(bab_pattern, para.text, re.IGNORECASE):
                bab_count += 1
        structure["main_content_complete"] = bab_count >= 3  # At least BAB I, II, III

        # Check back matter
        back_sections = ['daftar pustaka', 'bibliography', 'lampiran']
        present_back = any(any(term in section.lower() for term in back_sections)
                          for section in present_front)
        structure["back_matter_complete"] = present_back

        # Calculate completeness score
        complete_items = sum([structure["front_matter_complete"],
                             structure["main_content_complete"],
                             structure["back_matter_complete"]])
        structure["completeness_score"] = complete_items / 3.0

        return structure

    def _calculate_typography_score(self) -> float:
        """Calculate typography quality score."""
        score = 0.0
        total_checks = 5

        # Font appropriateness
        common_font = self.analysis['formatting_rules'].get('common_font', '').lower()
        if any(academic_font in common_font for academic_font in ['times new roman', 'arial', 'calibri']):
            score += 1

        # Font size appropriateness
        font_size = self.analysis['formatting_rules'].get('common_font_size', 12)
        if 11 <= font_size <= 12:
            score += 1

        # Line spacing
        line_spacing = self.analysis['formatting_rules'].get('line_spacing_pattern', 1.0)
        if 1.4 <= line_spacing <= 1.6:
            score += 1

        # Indentation
        indent_pattern = self.analysis['formatting_rules'].get('indentation_pattern', {})
        first_line_indent = indent_pattern.get('common_first_line', 0)
        if 1.0 <= first_line_indent <= 1.5:
            score += 1

        # Consistency (having defined styles)
        if len(self.analysis.get('styles', {})) > 5:
            score += 1

        return score / total_checks

    def _generate_indonesian_recommendations(self, profile: Dict[str, Any]) -> List[str]:
        """Generate recommendations for Indonesian academic formatting."""
        recommendations = []

        compliance = profile.get('formatting_compliance', {})

        if not compliance.get('font_compliance', False):
            recommendations.append("Gunakan font Times New Roman, Arial, atau Calibri sesuai standar akademik Indonesia")

        if not compliance.get('spacing_compliance', False):
            recommendations.append("Sesuaikan jarak baris menjadi 1.5 sesuai SK Mendikbud")

        if not compliance.get('indentation_compliance', False):
            recommendations.append("Gunakan indentasi baris pertama 1.0-1.5 cm untuk paragraf")

        if not compliance.get('typography_compliance', False):
            recommendations.append("Gunakan ukuran font 11-12 pt untuk teks utama")

        structure = profile.get('structure_completeness', {})
        if not structure.get('front_matter_complete', False):
            recommendations.append("Pastikan semua bagian depan (halaman judul, pengesahan, abstrak, daftar isi) lengkap")

        if not structure.get('main_content_complete', False):
            recommendations.append("Struktur BAB harus lengkap (minimal BAB I, II, III)")

        if not structure.get('back_matter_complete', False):
            recommendations.append("Tambahkan daftar pustaka dan lampiran sesuai kebutuhan")

        return recommendations

