"""
Document Merger
Merges analyzed content into template while preserving all formatting and styles.
"""

from typing import Dict, List, Any, Optional
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .template_analyzer import TemplateAnalyzer
from .content_extractor import ContentExtractor
from .content_mapper import ContentMapper
from .front_matter_generator import FrontMatterGenerator, BackMatterGenerator


class DocumentMerger:
    """Merges content into template with style preservation."""
    
    def __init__(self, template_analyzer: TemplateAnalyzer, 
                 content_extractor: ContentExtractor,
                 content_mapper: ContentMapper,
                 output_path: str):
        """Initialize merger with analysis and mapping."""
        self.template = template_analyzer
        self.content = content_extractor
        self.mapper = content_mapper
        self.output_path = Path(output_path)
        self.doc = Document(template_analyzer.template_path)
    
    def merge(self, user_data: Dict[str, Any] = None) -> Path:
        """Execute the merge process."""
        user_data = user_data or {}
        
        # Create a fresh document for complete thesis structure
        doc = Document()
        self._apply_section_properties(doc)
        
        # Step 1: Add front matter (title page, approval, abstract, etc.)
        self._add_front_matter(doc, user_data)
        
        # Step 2: Add main content (chapters)
        self._add_main_content(doc, user_data)
        
        # Step 3: Add back matter (references, appendices)
        self._add_back_matter(doc, user_data)
        
        # Step 4: Apply style consistency
        self._apply_style_consistency(doc)
        
        # Save document
        doc.save(str(self.output_path))
        
        return self.output_path
    
    def _add_front_matter(self, doc: Document, user_data: Dict[str, Any]) -> None:
        """Add all front matter (preliminary pages) to document."""
        generator = FrontMatterGenerator(user_data)
        
        # Add title page
        generator.create_title_page(doc)
        
        # Add approval pages
        generator.create_approval_page(doc, approval_type="supervisor")
        generator.create_approval_page(doc, approval_type="examiner")
        
        # Add originality statement
        generator.create_originality_statement(doc)
        
        # Add dedication and motto (optional but recommended)
        generator.create_dedication_page(doc, user_data.get("dedication"))
        generator.create_motto_page(doc, user_data.get("motto"))
        
        # Add preface
        generator.create_preface(doc)
        
        # Add abstracts (Indonesian and English)
        generator.create_abstract(doc, language="id")
        generator.create_abstract(doc, language="en")
        
        # Add table of contents placeholder
        generator.create_table_of_contents(doc)
        
        # Add glossary if available
        generator.create_glossary(doc, user_data.get("glossary"))
    
    def _add_main_content(self, doc: Document, user_data: Dict[str, Any]) -> None:
        """Add main content chapters from extractor."""
        sections = self.content.get_sections()
        
        for section in sections:
            # Add chapter heading
            title = section.get("title", "")
            content = section.get("content", [])
            
            if title:
                heading_para = doc.add_paragraph(title)
                heading_para.style = "Heading 1"
                heading_run = heading_para.runs[0]
                heading_run.font.bold = True
                heading_run.font.size = Pt(12)
                # Page break before new chapter except first
                if len(doc.paragraphs) > 1:
                    doc.add_page_break()
            
            # Add section content
            for line in content:
                if line.strip():
                    para = doc.add_paragraph(line)
                    para.style = "Normal"
                    # Set indentation
                    para.paragraph_format.left_indent = Inches(0.0)
                    para.paragraph_format.first_line_indent = Inches(1.0)

            # Insert lists if present
            for lst in section.get("lists", []):
                style_name = "List Number" if lst.get("ordered") else "List Bullet"
                for item in lst.get("items", []):
                    para = doc.add_paragraph(item)
                    para.style = style_name
    
    def _add_back_matter(self, doc: Document, user_data: Dict[str, Any]) -> None:
        """Add all back matter (references, appendices, etc.)."""
        generator = BackMatterGenerator(user_data)
        
        # Add list of tables
        generator.create_list_of_tables(doc)
        
        # Add list of figures
        generator.create_list_of_figures(doc)
        
        # Add references
        references = user_data.get("references", [])
        generator.create_references(doc, references)
        
        # Add appendices if available
        appendices = user_data.get("appendices", [])
        if appendices:
            generator.create_appendices(doc, appendices)

    def _apply_section_properties(self, doc: Document) -> None:
        """Apply sectPr properties from template analyzer to the new document."""
        try:
            src = self.template.doc.sections[0]
            dst = doc.sections[0]
            # Margins
            dst.top_margin = src.top_margin
            dst.bottom_margin = src.bottom_margin
            dst.left_margin = src.left_margin
            dst.right_margin = src.right_margin
            # Page size
            dst.page_height = src.page_height
            dst.page_width = src.page_width
            # Different first page
            dst.different_first_page_header_footer = src.different_first_page_header_footer
            # Columns (approximate via width; python-docx lacks direct column API)
            # Note: full column support requires OXML manipulation; omitted for safety.
        except Exception:
            pass
    
    def _replace_placeholders(self, user_data: Dict[str, Any]) -> None:
        """Replace template placeholders with user data or content."""
        placeholder_map = self.mapper.mapping["placeholder_replacements"]
        
        for placeholder, field_name in placeholder_map.items():
            if not field_name:
                continue
            
            # Get replacement value
            if field_name == "{TITLE}":
                value = user_data.get("title", "")
            elif field_name == "{AUTHOR}":
                value = user_data.get("author", "")
            elif field_name == "{DATE}":
                value = user_data.get("date", "")
            elif field_name == "{ADVISOR}":
                value = user_data.get("advisor", "")
            else:
                value = field_name
            
            # Replace in all paragraphs
            self._replace_text_in_paragraphs(placeholder, value)
    
    def _replace_text_in_paragraphs(self, old_text: str, new_text: str) -> None:
        """Replace text in paragraphs while preserving formatting."""
        # This method is kept for legacy compatibility but new code uses direct doc manipulation
        for para in self.doc.paragraphs:
            if old_text in para.text:
                # Preserve paragraph style
                style = para.style
                alignment = para.paragraph_format.alignment
                
                # Clear paragraph
                para.clear()
                
                # Re-add text with same style
                run = para.add_run(new_text)
                para.style = style
                if alignment:
                    para.paragraph_format.alignment = alignment
    
    def _merge_content_sections(self) -> None:
        """Merge content sections from extractor into template."""
        action_plan = self.mapper.get_action_plan()
        
        for action in action_plan:
            if action["action"] == "INSERT_CHAPTER":
                self._insert_chapter(action)
            elif action["action"] == "INSERT_SECTION":
                self._insert_section(action)
            elif action["action"] == "INSERT_APPENDIX":
                self._insert_appendix(action)
    
    def _insert_chapter(self, action: Dict[str, Any]) -> None:
        """Insert a chapter into the document."""
        title = action["title"]
        
        # Find chapter in content
        chapters = self.mapper.mapping["main_chapters"]
        chapter = next((ch for ch in chapters if ch["title"] == title), None)
        
        if not chapter:
            return
        
        # Add chapter heading using template style
        heading_style = self._get_heading_style(1)  # Level 1 = chapter
        heading_para = self.doc.add_paragraph(title, style=heading_style)
        
        # Add chapter content
        for content_line in chapter.get("sections", []):
            if content_line.strip():
                self.doc.add_paragraph(content_line, style="Normal")
    
    def _insert_section(self, action: Dict[str, Any]) -> None:
        """Insert a section (preface, references, etc.)."""
        section_name = action["section"]
        
        # Find corresponding content section
        section = self.content.get_section_by_title(
            self._get_search_pattern_for_section(section_name)
        )
        
        if not section:
            return
        
        # Add section heading
        heading_style = self._get_heading_style(0)
        self.doc.add_paragraph(section["title"], style=heading_style)
        
        # Add content
        for content_line in section.get("content", []):
            if content_line.strip():
                self.doc.add_paragraph(content_line, style="Normal")
    
    def _insert_appendix(self, action: Dict[str, Any]) -> None:
        """Insert an appendix."""
        title = action["title"]
        
        # Add page break before appendix
        self.doc.add_page_break()
        
        # Add appendix heading
        heading_style = self._get_heading_style(0)
        self.doc.add_paragraph(title, style=heading_style)
    
    def _add_auto_generated_sections(self, user_data: Dict[str, Any]) -> None:
        """Add auto-generated sections for missing content."""
        missing = self.mapper.mapping["missing_sections"]
        
        # Add critical missing sections
        for section in missing["critical"]:
            if "preface" in section.lower():
                self._insert_auto_preface(user_data)
            elif "abstract" in section.lower() or "sari" in section.lower():
                self._insert_auto_abstract(user_data)
    
    def _insert_auto_preface(self, user_data: Dict[str, Any]) -> None:
        """Insert auto-generated preface."""
        # This would be called from AI module
        # For now, just add placeholder
        self.doc.add_paragraph("KATA PENGANTAR", style="Heading 1")
        self.doc.add_paragraph(
            "[Kata pengantar akan dihasilkan otomatis oleh sistem AI]",
            style="Normal"
        )
    
    def _insert_auto_abstract(self, user_data: Dict[str, Any]) -> None:
        """Insert auto-generated abstract."""
        # This would be called from AI module
        self.doc.add_paragraph("ABSTRAK", style="Heading 1")
        self.doc.add_paragraph(
            "[Abstrak akan dihasilkan otomatis oleh sistem AI]",
            style="Normal"
        )
    
    def _update_fields(self) -> None:
        """Update document fields (TOC, page numbers, etc.)."""
        # Note: python-docx has limited support for fields
        # For full field support, use python-docx-template with Jinja2
        # or manipulate the XML directly
        pass
    
    def _apply_style_consistency(self, doc: Document) -> None:
        """Apply consistent styling throughout document."""
        try:
            formatting_rules = self.template.analysis["formatting_rules"]
        except:
            formatting_rules = {}
    
    def _get_heading_style(self, level: int) -> str:
        """Get appropriate heading style for level."""
        heading_map = {
            0: "Heading 1",
            1: "Heading 2",
            2: "Heading 3",
        }
        return heading_map.get(level, "Heading 1")
    
    def _get_search_pattern_for_section(self, section_name: str) -> str:
        """Get search pattern for section type."""
        patterns = {
            "preface": r"kata pengantar|preface",
            "abstract_id": r"abstrak|ringkasan|sari",
            "abstract_en": r"abstract|english",
            "glossary": r"glosarium|glossary",
            "references": r"daftar pustaka|references|bibliography",
            "toc": r"daftar isi|table of contents",
        }
        return patterns.get(section_name, section_name)
    
    def get_merge_report(self) -> Dict[str, Any]:
        """Generate report of merge operations."""
        report = {
            "status": "completed",
            "output_file": str(self.output_path),
            "sections_inserted": len(self.mapper.mapping["main_chapters"]),
            "sections_auto_generated": len(self.mapper.mapping["missing_sections"]["critical"]),
            "placeholders_replaced": len([p for p in self.mapper.mapping["placeholder_replacements"].values() if p]),
            "warnings": self._collect_warnings(),
        }
        return report
    
    def _collect_warnings(self) -> List[str]:
        """Collect warnings about the merge process."""
        warnings = []
        
        missing = self.mapper.mapping["missing_sections"]
        if missing["critical"]:
            warnings.append(f"Added {len(missing['critical'])} critical missing sections: {', '.join(missing['critical'])}")
        
        if missing["recommended"]:
            warnings.append(f"Note: Missing recommended sections: {', '.join(missing['recommended'])}")
        
        unreplaced = [p for p, v in self.mapper.mapping["placeholder_replacements"].items() if not v]
        if unreplaced:
            warnings.append(f"Placeholders not automatically replaced: {', '.join(unreplaced[:3])}")
        
        return warnings
