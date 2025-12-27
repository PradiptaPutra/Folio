"""
Analyze different templates to understand their structures and create adaptive system.
"""

from docx import Document
import re
from collections import defaultdict

template1_path = r"C:\Folio\form-memory\storage\references\Template-TA-UI.docx"
template2_path = r"C:\Folio\form-memory\storage\references\Template-skripsi-final-versi2020.docx"
output1_path = r"C:\Folio\form-memory\storage\outputs\Skripsi_Bary.docx"
output2_path = r"C:\Folio\form-memory\storage\outputs\Skripsi_Santi.docx"

def analyze_template(template_path, name):
    """Analyze a template structure."""
    doc = Document(template_path)
    
    print("=" * 80)
    print(f"ANALYZING TEMPLATE: {name}")
    print("=" * 80)
    
    # 1. Basic structure
    print(f"\n1. BASIC STRUCTURE:")
    print(f"   Total paragraphs: {len(doc.paragraphs)}")
    print(f"   Total styles: {len(doc.styles)}")
    print(f"   Total tables: {len(doc.tables)}")
    
    # 2. Chapter detection patterns
    print(f"\n2. CHAPTER DETECTION:")
    chapters = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # Check various chapter patterns
        if re.match(r'^BAB\s+[IVX\d]+', text, re.I) or re.match(r'^CHAPTER\s+\d+', text, re.I):
            chapters.append((i, text[:60], para.style.name if para.style else 'None'))
    
    print(f"   Found {len(chapters)} chapter headings:")
    for idx, (line, text, style) in enumerate(chapters[:10], 1):
        print(f"      {idx}. Line {line}: '{text}' (Style: {style})")
    
    # 3. Subsection patterns
    print(f"\n3. SUBSECTION PATTERNS:")
    subsections = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # Check various subsection patterns
        if (re.search(r'\b(Subbab|Anak Subbab|Sub-section|Subsection)\b', text, re.I) or
            re.match(r'^\d+\.\d+', text) or
            re.match(r'^[A-Z]\.\d+', text)):
            subsections.append((i, text[:60], para.style.name if para.style else 'None'))
    
    print(f"   Found {len(subsections)} subsection patterns:")
    for idx, (line, text, style) in enumerate(subsections[:15], 1):
        print(f"      {idx}. Line {line}: '{text}' (Style: {style})")
    
    # 4. Style analysis
    print(f"\n4. KEY STYLES:")
    style_usage = defaultdict(int)
    for para in doc.paragraphs:
        if para.style:
            style_usage[para.style.name] += 1
    
    print("   Most used styles:")
    for style, count in sorted(style_usage.items(), key=lambda x: x[1], reverse=True)[:10]:
        print(f"      {style}: {count} paragraphs")
    
    # 5. Content placeholder patterns
    print(f"\n5. PLACEHOLDER PATTERNS:")
    placeholders = []
    placeholder_patterns = [
        r'TULISKAN',
        r'KETIK',
        r'ISI',
        r'\[.*\]',
        r'\.\.\.',
        r'___+',
    ]
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for pattern in placeholder_patterns:
            if re.search(pattern, text, re.I):
                placeholders.append((i, text[:60]))
                break
    
    print(f"   Found {len(placeholders)} placeholder patterns:")
    for idx, (line, text) in enumerate(placeholders[:10], 1):
        print(f"      {idx}. Line {line}: '{text}'")
    
    # 6. Font analysis
    print(f"\n6. FONT ANALYSIS:")
    fonts = defaultdict(int)
    for para in doc.paragraphs[:100]:
        for run in para.runs:
            if run.font.name:
                fonts[run.font.name] += 1
    
    print("   Fonts used:")
    for font, count in sorted(fonts.items(), key=lambda x: x[1], reverse=True):
        print(f"      {font}: {count} runs")
    
    return {
        'paragraphs': len(doc.paragraphs),
        'styles': len(doc.styles),
        'chapters': len(chapters),
        'subsections': len(subsections),
        'placeholders': len(placeholders),
        'chapter_patterns': [text for _, text, _ in chapters],
        'subsection_patterns': [text for _, text, _ in subsections[:10]],
        'key_styles': dict(sorted(style_usage.items(), key=lambda x: x[1], reverse=True)[:10])
    }

# Analyze both templates
template1_info = analyze_template(template1_path, "Template-TA-UI.docx")
print("\n\n")
template2_info = analyze_template(template2_path, "Template-skripsi-final-versi2020.docx")

# Compare templates
print("\n" + "=" * 80)
print("TEMPLATE COMPARISON")
print("=" * 80)

print(f"\nDifferences:")
print(f"   Paragraphs: Template1={template1_info['paragraphs']}, Template2={template2_info['paragraphs']}")
print(f"   Chapters: Template1={template1_info['chapters']}, Template2={template2_info['chapters']}")
print(f"   Subsections: Template1={template1_info['subsections']}, Template2={template2_info['subsections']}")
print(f"   Placeholders: Template1={template1_info['placeholders']}, Template2={template2_info['placeholders']}")

print(f"\nTemplate1 chapter patterns:")
for pattern in template1_info['chapter_patterns'][:5]:
    print(f"   - {pattern}")

print(f"\nTemplate2 chapter patterns:")
for pattern in template2_info['chapter_patterns'][:5]:
    print(f"   - {pattern}")

print(f"\nTemplate1 subsection patterns:")
for pattern in template1_info['subsection_patterns'][:5]:
    print(f"   - {pattern}")

print(f"\nTemplate2 subsection patterns:")
for pattern in template2_info['subsection_patterns'][:5]:
    print(f"   - {pattern}")

