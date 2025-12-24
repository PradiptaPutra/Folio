#!/usr/bin/env python
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from engine.analyzer.complete_thesis_builder import create_complete_thesis

BASE = Path(__file__).parent
UPLOAD = BASE / "storage" / "uploads"
OUTPUTS = BASE / "storage" / "outputs"
UPLOAD.mkdir(parents=True, exist_ok=True)
OUTPUTS.mkdir(parents=True, exist_ok=True)

def make_template(path: Path):
    doc = Document()
    # Add a heading style usage
    h = doc.add_heading('Template Heading Sample', level=1)
    # Normal paragraph sample
    p = doc.add_paragraph('Template body paragraph sample')
    p.style = 'Normal'
    doc.save(str(path))


def make_content_docx(path: Path):
    doc = Document()
    # Heading 1 for chapter
    h1 = doc.add_heading('BAB I Pendahuluan', level=1)
    # Subsection heading emulation (will be detected by normalized extractor)
    h2 = doc.add_heading('1.1 Latar Belakang', level=2)
    # Body paragraphs
    para = doc.add_paragraph('Ini adalah paragraf isi skripsi yang cukup panjang untuk pengujian.')
    para.paragraph_format.first_line_indent = Inches(1.0)
    para.paragraph_format.left_indent = Inches(0.0)
    # Numbered list using built-in style
    for item in ['Tujuan Penelitian', 'Metode Penelitian', 'Kontribusi']:
        li = doc.add_paragraph(item)
        li.style = 'List Number'
    # Bullet list
    for item in ['Kata Kunci 1', 'Kata Kunci 2']:
        li = doc.add_paragraph(item)
        li.style = 'List Bullet'
    doc.save(str(path))


def main():
    template = UPLOAD / 'test_template.docx'
    content = UPLOAD / 'test_content.docx'
    output = OUTPUTS / 'test_output.docx'
    make_template(template)
    make_content_docx(content)
    result = create_complete_thesis(str(template), str(content), str(output), use_ai=False)
    print('Result:', result)
    print('Output exists:', output.exists(), 'size:', output.stat().st_size if output.exists() else 0)

if __name__ == '__main__':
    main()
