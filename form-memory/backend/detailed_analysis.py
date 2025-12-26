"""
Detailed analysis of template vs output differences.
"""

from docx import Document
from pathlib import Path
import re

template_path = r"C:\Folio\form-memory\storage\references\Template-skripsi-final-versi2020.docx"
output_path = r"C:\Folio\form-memory\storage\outputs\Skripsi_Tarik.docx"

template_doc = Document(template_path)
output_doc = Document(output_path)

print("=" * 80)
print("DETAILED TEMPLATE vs OUTPUT ANALYSIS")
print("=" * 80)

# 1. Missing Subsections Analysis
print("\n" + "=" * 80)
print("1. MISSING SUBSECTIONS DETAILED ANALYSIS")
print("=" * 80)

template_subbabs = []
for i, para in enumerate(template_doc.paragraphs):
    text = para.text.strip()
    if re.search(r'\b(Subbab|Anak Subbab)\b', text, re.I) or re.match(r'^\d+\.\d+', text):
        template_subbabs.append((i, text, para.style.name if para.style else 'None'))

output_subbabs = []
for i, para in enumerate(output_doc.paragraphs):
    text = para.text.strip()
    if re.search(r'\b(Subbab|Anak Subbab)\b', text, re.I) or re.match(r'^\d+\.\d+', text):
        output_subbabs.append((i, text, para.style.name if para.style else 'None'))

print(f"\nTemplate has {len(template_subbabs)} subsections:")
for idx, (para_idx, text, style) in enumerate(template_subbabs[:30], 1):
    print(f"   {idx}. Para {para_idx}: '{text[:60]}' (Style: {style})")

print(f"\nOutput has {len(output_subbabs)} subsections:")
for idx, (para_idx, text, style) in enumerate(output_subbabs[:30], 1):
    print(f"   {idx}. Para {para_idx}: '{text[:60]}' (Style: {style})")

# 2. Style Usage Comparison
print("\n" + "=" * 80)
print("2. STYLE USAGE COMPARISON")
print("=" * 80)

template_style_count = {}
for para in template_doc.paragraphs:
    if para.style:
        style_name = para.style.name
        template_style_count[style_name] = template_style_count.get(style_name, 0) + 1

output_style_count = {}
for para in output_doc.paragraphs:
    if para.style:
        style_name = para.style.name
        output_style_count[style_name] = output_style_count.get(style_name, 0) + 1

print("\nKey style differences:")
key_styles = ['Isi Paragraf', 'Normal', 'Heading 2', 'Heading 3', 'JUDUL']
for style in key_styles:
    template_count = template_style_count.get(style, 0)
    output_count = output_style_count.get(style, 0)
    diff = output_count - template_count
    status = "[OK]" if diff >= 0 else "[X]"
    print(f"   {style}: Template={template_count}, Output={output_count}, Diff={diff:+.0f} {status}")

# 3. Content Paragraph Style Analysis
print("\n" + "=" * 80)
print("3. CONTENT PARAGRAPH STYLE ANALYSIS")
print("=" * 80)

template_content = []
for para in template_doc.paragraphs:
    text = para.text.strip()
    if len(text) > 50 and not re.match(r'^BAB\s+[IVX\d]+', text, re.I) and not re.search(r'\b(Subbab|Anak Subbab)\b', text, re.I):
        template_content.append((text[:50], para.style.name if para.style else 'None'))

output_content = []
for para in output_doc.paragraphs:
    text = para.text.strip()
    if len(text) > 50 and not re.match(r'^BAB\s+[IVX\d]+', text, re.I) and not re.search(r'\b(Subbab|Anak Subbab)\b', text, re.I):
        output_content.append((text[:50], para.style.name if para.style else 'None'))

template_content_styles = {}
for text, style in template_content[:20]:
    template_content_styles[style] = template_content_styles.get(style, 0) + 1

output_content_styles = {}
for text, style in output_content[:20]:
    output_content_styles[style] = output_content_styles.get(style, 0) + 1

print("\nTemplate content paragraph styles (sample):")
for style, count in sorted(template_content_styles.items(), key=lambda x: x[1], reverse=True):
    print(f"   {style}: {count}")

print("\nOutput content paragraph styles (sample):")
for style, count in sorted(output_content_styles.items(), key=lambda x: x[1], reverse=True):
    print(f"   {style}: {count}")

# 4. Formatting Details
print("\n" + "=" * 80)
print("4. FORMATTING DETAILS - 'Isi Paragraf' STYLE")
print("=" * 80)

template_isi_paras = [p for p in template_doc.paragraphs if p.style and p.style.name == 'Isi Paragraf']
output_isi_paras = [p for p in output_doc.paragraphs if p.style and p.style.name == 'Isi Paragraf']

if template_isi_paras:
    sample_template = template_isi_paras[0]
    pf_template = sample_template.paragraph_format
    print("\nTemplate 'Isi Paragraf' formatting:")
    print(f"   Line spacing: {pf_template.line_spacing}")
    print(f"   First line indent: {pf_template.first_line_indent}")
    print(f"   Left indent: {pf_template.left_indent}")
    print(f"   Alignment: {pf_template.alignment}")
    print(f"   Space before: {pf_template.space_before}")
    print(f"   Space after: {pf_template.space_after}")

if output_isi_paras:
    sample_output = output_isi_paras[0]
    pf_output = sample_output.paragraph_format
    print("\nOutput 'Isi Paragraf' formatting:")
    print(f"   Line spacing: {pf_output.line_spacing}")
    print(f"   First line indent: {pf_output.first_line_indent}")
    print(f"   Left indent: {pf_output.left_indent}")
    print(f"   Alignment: {pf_output.alignment}")
    print(f"   Space before: {pf_output.space_before}")
    print(f"   Space after: {pf_output.space_after}")

# 5. Font Analysis
print("\n" + "=" * 80)
print("5. FONT ANALYSIS")
print("=" * 80)

template_fonts = {}
for para in template_doc.paragraphs[:100]:
    for run in para.runs:
        if run.font.name:
            template_fonts[run.font.name] = template_fonts.get(run.font.name, 0) + 1

output_fonts = {}
for para in output_doc.paragraphs[:100]:
    for run in para.runs:
        if run.font.name:
            output_fonts[run.font.name] = output_fonts.get(run.font.name, 0) + 1

print("\nTemplate fonts (first 100 paragraphs):")
for font, count in sorted(template_fonts.items(), key=lambda x: x[1], reverse=True):
    print(f"   {font}: {count} runs")

print("\nOutput fonts (first 100 paragraphs):")
for font, count in sorted(output_fonts.items(), key=lambda x: x[1], reverse=True):
    print(f"   {font}: {count} runs")

# 6. Summary of Issues
print("\n" + "=" * 80)
print("6. SUMMARY OF ISSUES")
print("=" * 80)

issues = []

# Missing subsections
if len(output_subbabs) < len(template_subbabs):
    issues.append(f"MISSING SUBSECTIONS: {len(template_subbabs) - len(output_subbabs)} subsections are missing")
    issues.append(f"   Template: {len(template_subbabs)} subsections")
    issues.append(f"   Output: {len(output_subbabs)} subsections")

# Isi Paragraf style usage
isi_diff = template_style_count.get('Isi Paragraf', 0) - output_style_count.get('Isi Paragraf', 0)
if isi_diff > 0:
    issues.append(f"ISI PARAGRAF STYLE: {isi_diff} fewer paragraphs using 'Isi Paragraf' style")
    issues.append(f"   Template: {template_style_count.get('Isi Paragraf', 0)} paragraphs")
    issues.append(f"   Output: {output_style_count.get('Isi Paragraf', 0)} paragraphs")

# Font mismatch
if 'Arial' in template_fonts and 'Times New Roman' in output_fonts:
    issues.append("FONT MISMATCH: Template uses Arial, but output uses Times New Roman")

# Content style issues
missing_content_styles = set(template_content_styles.keys()) - set(output_content_styles.keys())
if missing_content_styles:
    issues.append(f"CONTENT STYLES: Missing styles in output: {missing_content_styles}")

for issue in issues:
    print(f"   [X] {issue}")

print("\n" + "=" * 80)
print("ANALYSIS COMPLETE")
print("=" * 80)

