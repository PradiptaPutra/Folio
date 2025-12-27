"""
Debug template structure to understand where content is being inserted incorrectly.
"""

from docx import Document
import re
from pathlib import Path

def analyze_template(template_path, name):
    """Analyze template structure."""
    if not Path(template_path).exists():
        print(f"[SKIP] File not found: {template_path}")
        return
    
    print("=" * 80)
    print(f"ANALYZING TEMPLATE: {name}")
    print("=" * 80)
    
    doc = Document(template_path)
    
    # Find key sections
    daftar_isi_start = None
    daftar_isi_end = None
    bab_i_start = None
    
    front_matter_keywords = ['DAFTAR ISI', 'DAFTAR GAMBAR', 'DAFTAR TABEL', 'DAFTAR LAMPIRAN']
    
    print(f"\n[STRUCTURE] Document has {len(doc.paragraphs)} paragraphs\n")
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        text_upper = text.upper()
        
        # Find DAFTAR ISI
        if 'DAFTAR ISI' in text_upper and daftar_isi_start is None:
            daftar_isi_start = i
            print(f"[FOUND] DAFTAR ISI at paragraph {i}: {text[:60]}")
            
            # Find where DAFTAR ISI section ends
            for j in range(i + 1, min(i + 200, len(doc.paragraphs))):
                next_text = doc.paragraphs[j].text.strip().upper()
                # Check if we hit another front matter section or main content
                if any(kw in next_text for kw in ['DAFTAR GAMBAR', 'DAFTAR TABEL', 'BAB I', 'CHAPTER 1']):
                    daftar_isi_end = j
                    print(f"[FOUND] DAFTAR ISI likely ends around paragraph {j}")
                    break
        
        # Find BAB I
        if bab_i_start is None and ('BAB I' in text_upper or 'BAB 1' in text_upper):
            # Check if this is a TOC entry
            is_toc = '\t' in text or re.search(r'\s+\d+\s*$', text)
            if not is_toc:
                bab_i_start = i
                print(f"[FOUND] BAB I (main content) at paragraph {i}: {text[:60]}")
    
    # Show paragraphs around DAFTAR ISI
    if daftar_isi_start is not None:
        print(f"\n[DAFTAR ISI AREA] Showing paragraphs {daftar_isi_start} to {daftar_isi_end or daftar_isi_start + 30}:")
        end_show = daftar_isi_end if daftar_isi_end else daftar_isi_start + 30
        for i in range(daftar_isi_start, min(end_show, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            text = para.text.strip()
            style = para.style.name if para.style else "None"
            has_tab = '\t' in text
            has_page_num = bool(re.search(r'\s+\d+\s*$', text))
            
            print(f"  Para {i:3d} [Style: {style:15s}] [Tab: {has_tab}] [PageNum: {has_page_num}] {text[:70]}")
    
    # Show paragraphs around BAB I
    if bab_i_start is not None:
        print(f"\n[BAB I AREA] Showing paragraphs {bab_i_start} to {bab_i_start + 20}:")
        for i in range(bab_i_start, min(bab_i_start + 20, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            text = para.text.strip()
            style = para.style.name if para.style else "None"
            print(f"  Para {i:3d} [Style: {style:15s}] {text[:70]}")
    
    print()

def analyze_output(output_path, name):
    """Analyze output to see where content was incorrectly placed."""
    if not Path(output_path).exists():
        print(f"[SKIP] File not found: {output_path}")
        return
    
    print("=" * 80)
    print(f"ANALYZING OUTPUT: {name}")
    print("=" * 80)
    
    doc = Document(output_path)
    
    daftar_isi_start = None
    bab_i_start = None
    
    print(f"\n[STRUCTURE] Document has {len(doc.paragraphs)} paragraphs\n")
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        text_upper = text.upper()
        
        if 'DAFTAR ISI' in text_upper and daftar_isi_start is None:
            daftar_isi_start = i
            print(f"[FOUND] DAFTAR ISI at paragraph {i}")
        
        if bab_i_start is None and ('BAB I' in text_upper or 'BAB 1' in text_upper):
            is_toc = '\t' in text or re.search(r'\s+\d+\s*$', text)
            if not is_toc:
                bab_i_start = i
                print(f"[FOUND] BAB I at paragraph {i}")
    
    # Check for content in DAFTAR ISI area
    if daftar_isi_start is not None:
        print(f"\n[CHECKING] Content in DAFTAR ISI area (paragraphs {daftar_isi_start} to {daftar_isi_start + 50}):")
        content_found = []
        
        for i in range(daftar_isi_start, min(daftar_isi_start + 50, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            if len(text) < 20:
                continue
            
            # Check if this looks like actual content (not TOC)
            is_toc_entry = '\t' in text or re.search(r'\s+\d+\s*$', text) or len(text) < 50
            
            # Check if this looks like chapter content
            is_content = (
                re.match(r'^\d+\.\d+', text) or  # Subsection like "1.1"
                (len(text) > 200) or  # Long paragraph
                any(kw in text.upper() for kw in ['LATAR BELAKANG', 'RUMUSAN MASALAH', 'TUJUAN PENELITIAN'])
            )
            
            if is_content and not is_toc_entry:
                content_found.append({
                    'para_idx': i,
                    'text': text[:150],
                    'length': len(text)
                })
        
        if content_found:
            print(f"  [ERROR] Found {len(content_found)} paragraphs with content in DAFTAR ISI area:")
            for item in content_found[:10]:
                print(f"    Para {item['para_idx']} ({item['length']} chars): {item['text']}")
        else:
            print(f"  [OK] No content found in DAFTAR ISI area")
    
    print()

# Analyze templates
templates = [
    (r"C:\Folio\form-memory\storage\uploads\Template-skripsi-final-versi2020.docx", "Template-skripsi-final-versi2020"),
    (r"C:\Folio\form-memory\storage\uploads\Template-TA-UI.docx", "Template-TA-UI")
]

for template_path, name in templates:
    analyze_template(template_path, name)

# Analyze output
analyze_output(r"C:\Folio\form-memory\storage\outputs\Skripsi_Awe.docx", "Skripsi_Awe")

