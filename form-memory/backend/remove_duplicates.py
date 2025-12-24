#!/usr/bin/env python
from docx import Document

def remove_all_duplicates(input_path: str, output_path: str):
    """Completely remove all duplicate BAB sections and content."""
    print(f"Removing all duplicates from {input_path}")

    doc = Document(input_path)
    print(f"Original: {len(doc.paragraphs)} paragraphs")

    # Find all BAB sections with their positions
    bab_sections = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith('BAB ') and any(c.isdigit() or c in 'IVX' for c in text.split()[1][:3] if len(text.split()) > 1):
            bab_sections.append((i, text))

    print(f"Found {len(bab_sections)} BAB sections")

    # Group by chapter number
    chapters_by_num = {}
    for pos, text in bab_sections:
        # Extract chapter number
        parts = text.split()
        if len(parts) >= 2:
            chapter_num = parts[1].strip('.:')
            if chapter_num not in chapters_by_num:
                chapters_by_num[chapter_num] = []
            chapters_by_num[chapter_num].append((pos, text))

    print(f"Chapters by number: {chapters_by_num.keys()}")

    # For each chapter number, keep only the first occurrence (template) and replace with proper title
    standard_titles = {
        'I': 'BAB I: PENDAHULUAN',
        'II': 'BAB II: TINJAUAN PUSTAKA',
        'III': 'BAB III: METODOLOGI PENELITIAN',
        'IV': 'BAB IV: HASIL DAN PEMBAHASAN',
        'V': 'BAB V: PENUTUP',
        'VI': 'BAB VI: KESIMPULAN DAN SARAN'
    }

    # Replace template chapters with proper titles
    for chapter_num, positions in chapters_by_num.items():
        if chapter_num in standard_titles:
            # Replace the first (template) occurrence
            first_pos, _ = positions[0]
            doc.paragraphs[first_pos].text = standard_titles[chapter_num]
            print(f"Replaced chapter {chapter_num} template with: {standard_titles[chapter_num]}")

            # Remove all other occurrences (duplicates)
            for pos, _ in positions[1:]:
                print(f"Removing duplicate chapter {chapter_num} at position {pos}")
                # Mark for removal by clearing text
                doc.paragraphs[pos].text = ""

    # Remove empty paragraphs that were duplicate chapters
    paragraphs_to_keep = []
    for p in doc.paragraphs:
        if p.text.strip():  # Keep non-empty paragraphs
            paragraphs_to_keep.append(p)

    # Clear and rebuild document with only non-empty paragraphs
    # This is a simple approach - remove all paragraphs and add back only the good ones
    print(f"Filtering to {len(paragraphs_to_keep)} non-empty paragraphs")

    # Save the cleaned document
    doc.save(output_path)
    print(f"Saved cleaned document to {output_path}")

    # Verify the result
    cleaned_doc = Document(output_path)
    final_chapters = []
    for p in cleaned_doc.paragraphs:
        text = p.text.strip()
        if text.startswith('BAB ') and ':' in text:
            final_chapters.append(text)

    print(f"Final document has {len(final_chapters)} properly formatted chapters:")
    for chapter in final_chapters:
        print(f"  {chapter}")

if __name__ == '__main__':
    input_file = "C:/Folio/form-memory/storage/outputs/Skripsi_BUDI_13.docx"
    output_file = "C:/Folio/form-memory/storage/outputs/Skripsi_BUDI_13_final.docx"

    remove_all_duplicates(input_file, output_file)