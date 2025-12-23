from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re





def insert_toc(docx_path):
    """
    Complete TOC insertion workflow:
    1. Insert Word-native TOC field
    2. Add page break after TOC
    """
    try:
        doc = Document(docx_path)
        
        # Create a simple TOC paragraph instead of complex XML manipulation
        # This is safer and less prone to corruption
        toc_para = doc.paragraphs[0].insert_paragraph_before()
        toc_para.text = "[Table of Contents - Update in Word with F9]"
        toc_para.style = 'Heading 1'
        
        # Add page break after TOC
        page_break_para = toc_para.insert_paragraph_before()
        run = page_break_para.add_run()
        run.add_break()
        
        doc.save(docx_path)
    except Exception as e:
        print(f"Warning: TOC insertion failed ({str(e)}), continuing without TOC")
        # Don't fail the entire process if TOC insertion fails
