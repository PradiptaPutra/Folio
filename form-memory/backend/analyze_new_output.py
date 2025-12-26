"""
Analyze the new output to see what went wrong.
"""

from docx import Document
import re
from collections import defaultdict

output_path = r"C:\Folio\form-memory\storage\outputs\Skripsi_Toast.docx"
template_path = r"C:\Folio\form-memory\storage\references\Template-skripsi-final-versi2020.docx"

output_doc = Document(output_path)
template_doc = Document(template_path)

print("=" * 80)
print("ANALYZING NEW OUTPUT - IDENTIFYING ISSUES")
print("=" * 80)

# 1. Check paragraph count
print(f"\n1. BASIC STATS:")
print(f"   Template paragraphs: {len(template_doc.paragraphs)}")
print(f"   Output paragraphs: {len(output_doc.paragraphs)}")
print(f"   Difference: {len(output_doc.paragraphs) - len(template_doc.paragraphs)}")

# 2. Check for problematic content
print("\n2. PROBLEMATIC CONTENT DETECTION:")

issues = []
empty_subsections = []
malformed_subsections = []
unwanted_text = []

for i, para in enumerate(output_doc.paragraphs):
    text = para.text.strip()
    
    # Check for empty subsections
    if re.match(r'^\d+\.\d+', text):
        if len(text) < 30:  # Very short subsection
            next_idx = i + 1
            if next_idx < len(output_doc.paragraphs):
                next_text = output_doc.paragraphs[next_idx].text.strip()
                # If next is also a subsection or BAB, this one is empty
                if re.match(r'^\d+\.\d+', next_text) or re.match(r'^BAB', next_text, re.I):
                    empty_subsections.append((i, text))
    
    # Check for malformed subsections (missing chapter number)
    if re.match(r'^\.\d+', text):
        malformed_subsections.append((i, text[:60]))
    
    # Check for unwanted text
    if any(pattern in text for pattern in ['Disusun Oleh:', 'DISUSUN OLEH', 'Tujuan dari penelitian ini adalah:']):
        unwanted_text.append((i, text[:60]))
    
    # Check for placeholder content
    if '[Content untuk' in text or 'akan ditambahkan di sini' in text:
        issues.append((i, "Placeholder content found", text[:60]))

print(f"\n   Empty subsections: {len(empty_subsections)}")
for line, text in empty_subsections[:5]:
    print(f"      Line {line}: '{text}'")

print(f"\n   Malformed subsections (missing chapter): {len(malformed_subsections)}")
for line, text in malformed_subsections[:5]:
    print(f"      Line {line}: '{text}'")

print(f"\n   Unwanted text: {len(unwanted_text)}")
for line, text in unwanted_text[:5]:
    print(f"      Line {line}: '{text}'")

print(f"\n   Placeholder content: {len(issues)}")
for line, issue_type, text in issues[:5]:
    print(f"      Line {line}: {issue_type} - '{text}'")

# 3. Check subsection structure
print("\n3. SUBSECTION STRUCTURE:")
subsections = []
for i, para in enumerate(output_doc.paragraphs):
    text = para.text.strip()
    if re.match(r'^\d+\.\d+', text) or re.match(r'^\.\d+', text):
        subsections.append((i, text[:60], para.style.name if para.style else 'None'))

print(f"\n   Total subsections found: {len(subsections)}")
for idx, (line, text, style) in enumerate(subsections[:25], 1):
    print(f"      {idx}. Line {line}: {text} (Style: {style})")

# 4. Check content paragraphs
print("\n4. CONTENT PARAGRAPH ANALYSIS:")
content_paras = []
for i, para in enumerate(output_doc.paragraphs):
    text = para.text.strip()
    if len(text) > 50 and not re.match(r'^BAB\s+[IVX\d]+', text, re.I) and not re.match(r'^\d+\.\d+', text):
        style_name = para.style.name if para.style else 'None'
        content_paras.append((i, text[:50], style_name))

print(f"\n   Content paragraphs: {len(content_paras)}")
style_dist = defaultdict(int)
for _, _, style in content_paras:
    style_dist[style] += 1

print("\n   Content paragraph styles:")
for style, count in sorted(style_dist.items(), key=lambda x: x[1], reverse=True)[:10]:
    print(f"      {style}: {count}")

# 5. Check for duplicate or repeated content
print("\n5. DUPLICATE CONTENT CHECK:")
seen_texts = {}
duplicates = []
for i, para in enumerate(output_doc.paragraphs):
    text = para.text.strip()[:100]  # First 100 chars
    if text and len(text) > 20:
        if text in seen_texts:
            duplicates.append((i, text[:50], seen_texts[text]))
        else:
            seen_texts[text] = i

if duplicates:
    print(f"\n   Found {len(duplicates)} potential duplicates:")
    for line, text, first_line in duplicates[:5]:
        print(f"      Line {line} duplicates line {first_line}: '{text}'")
else:
    print("\n   No duplicates found")

# 6. Sample problematic paragraphs
print("\n6. SAMPLE PROBLEMATIC PARAGRAPHS:")
problem_count = 0
for i, para in enumerate(output_doc.paragraphs):
    text = para.text.strip()
    
    # Check various issues
    is_problem = False
    problem_type = ""
    
    if '[Content untuk' in text:
        is_problem = True
        problem_type = "Placeholder"
    elif re.match(r'^\.\d+', text):
        is_problem = True
        problem_type = "Missing chapter number"
    elif 'Disusun Oleh:' in text and i > 50:
        is_problem = True
        problem_type = "Unwanted text"
    elif len(text) > 10 and text.count('\n') > 5:
        is_problem = True
        problem_type = "Too many line breaks"
    
    if is_problem and problem_count < 10:
        print(f"\n   Line {i} ({problem_type}):")
        print(f"      '{text[:100]}...'")
        problem_count += 1

print("\n" + "=" * 80)
print("ANALYSIS COMPLETE")
print("=" * 80)

