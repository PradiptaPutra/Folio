#!/usr/bin/env python
from pathlib import Path
from docx import Document

def test_minimal():
    """Test just the document loading and basic operations."""
    print("Testing minimal document operations...")

    template_path = "C:\\Folio\\form-memory\\storage\\references\\Template-skripsi-final-versi2020.docx"
    output_path = "C:\\Folio\\form-memory\\storage\\outputs\\minimal_test.docx"

    try:
        # Load template
        print("Loading template...")
        doc = Document(template_path)
        print(f"Template loaded: {len(doc.paragraphs)} paragraphs")

        # Simple chapter replacement
        print("Doing simple chapter replacement...")
        replacements = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if text.startswith('BAB I'):
                para.text = 'BAB I: PENDAHULUAN'
                replacements += 1
                print("Replaced BAB I")
            elif text.startswith('BAB II'):
                para.text = 'BAB II: TINJAUAN PUSTAKA'
                replacements += 1
                print("Replaced BAB II")

        print(f"Made {replacements} replacements")

        # Save
        print("Saving document...")
        doc.save(output_path)
        print(f"Saved to {output_path}")

        # Verify
        saved_doc = Document(output_path)
        print(f"Saved document has {len(saved_doc.paragraphs)} paragraphs")

        print("SUCCESS: Minimal test completed")

    except Exception as e:
        print(f"FAILED: {e}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    test_minimal()