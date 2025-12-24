#!/usr/bin/env python
from pathlib import Path
from docx import Document

def clean_duplicate_chapters(input_path: str, output_path: str):
    """Clean up duplicate chapters in the generated document."""
    print(f"Cleaning duplicate chapters in {input_path}")

    doc = Document(input_path)
    print(f"Original document: {len(doc.paragraphs)} paragraphs")

    # Find all BAB sections
    bab_positions = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith('BAB ') and any(c.isdigit() or c in 'IVX' for c in text.split()[1][:3] if len(text.split()) > 1):
            bab_positions.append((i, text))

    print(f"Found {len(bab_positions)} BAB sections")

    # Separate template chapters (with TULISKAN) from content chapters (with proper titles)
    template_chapters = []
    content_chapters = []

    for pos, text in bab_positions:
        if 'TULISKAN' in text.upper():
            template_chapters.append((pos, text))
        else:
            content_chapters.append((pos, text))

    print(f"Template chapters: {len(template_chapters)}")
    print(f"Content chapters: {len(content_chapters)}")

    # If we have both template and content chapters, replace template with content
    if template_chapters and content_chapters and len(template_chapters) == len(content_chapters):
        print("Replacing template chapters with content chapters...")

        for (template_pos, _), (content_pos, content_text) in zip(template_chapters, content_chapters):
            # Replace template chapter title with content chapter title
            doc.paragraphs[template_pos].text = content_text
            print(f"Replaced template at {template_pos} with content: {content_text}")

        # Remove the duplicate content chapters
        content_positions = [pos for pos, _ in content_chapters]
        content_positions.sort(reverse=True)  # Remove from end first

        for pos in content_positions:
            if pos < len(doc.paragraphs):
                del doc.paragraphs[pos]
                print(f"Removed duplicate content at {pos}")

    # Save cleaned document
    doc.save(output_path)
    print(f"Saved cleaned document to {output_path}")

    # Verify
    cleaned_doc = Document(output_path)
    print(f"Cleaned document: {len(cleaned_doc.paragraphs)} paragraphs")

if __name__ == '__main__':
    input_file = "C:/Folio/form-memory/storage/outputs/Skripsi_BUDI_13.docx"
    output_file = "C:/Folio/form-memory/storage/outputs/Skripsi_BUDI_13_clean.docx"

    if Path(input_file).exists():
        clean_duplicate_chapters(input_file, output_file)
    else:
        print(f"Input file not found: {input_file}")