
from pathlib import Path
from pandoc_runner import markdown_to_docx
import os
import shutil
from docx import Document

# Setup
TEST_DIR = Path("test_workflow_fm")
TEST_DIR.mkdir(exist_ok=True)

# 1. Create Mock Template
TEMPLATE_PATH = TEST_DIR / "template.docx"
doc = Document()
doc.add_heading("Heading 1", level=1)
doc.add_heading("Heading 2", level=2)
doc.add_paragraph("Normal text")
doc.save(TEMPLATE_PATH)

# 2. Create Markdown Content
MD_PATH = TEST_DIR / "content.md"
MD_PATH.write_text("""
# BAB I PENDAHULUAN

## Latar Belakang
Ini latar belakang.

## Rumusan Masalah
Ini rumusan masalah.
""", encoding="utf-8")

# 3. Define Front Matter Data
frontmatter_data = {
    'judul': "SISTEM PENJUALAN AI",
    'penulis': "RUDI TABUTI",
    'nim': "123456789",
    'universitas': "UNIVERSITAS TEKNOLOGI",
    'tahun': 2024,
    'abstrak_teks': "Abstrak indo...",
    'abstrak_en_teks': "English abstract...",
    'kata_kunci': "AI, Penjualan"
}

# 4. Run Pipeline
OUTPUT_DOCX = TEST_DIR / "output_fm.docx"

print("Running markdown_to_docx with frontmatter...")
try:
    markdown_to_docx(
        MD_PATH,
        TEMPLATE_PATH,
        OUTPUT_DOCX,
        style_config=None,
        frontmatter_data=frontmatter_data
    )
    print("Generation successful.")
except Exception as e:
    print(f"Generation failed: {e}")
    exit(1)

# 5. Inspect Output
print("\nInspecting Output:")
doc = Document(OUTPUT_DOCX)
for p in doc.paragraphs[:20]: # Check first 20 paragraphs (Front matter should be here)
    if 'w:pageBreakBefore' in p._element.xml:
        print("[PAGE BREAK]")
    if p.text.strip():
        print(f"[{p.style.name}] {p.text[:50]}")

# Cleanup (Optional)
# shutil.rmtree(TEST_DIR)
