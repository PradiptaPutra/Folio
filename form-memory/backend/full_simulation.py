
from docx import Document
from skripsi_enforcer import SkripsiEnforcer
import os

# 1. Setup Dummy Input
DRAFT_PATH = "dummy_draft.docx"
doc = Document()
# Raw content simulation
doc.add_paragraph("BAB I PENDAHULUAN")
doc.add_paragraph("Latar Belakang") # Should become 1.1 Heading 2
doc.add_paragraph("Perkembangan teknologi informasi saat ini sangat pesat. Hal ini menuntut adanya sistem yang efektif.")
doc.add_paragraph("Rumusan Masalah") # Should become 1.2 Heading 2
doc.add_paragraph("Berdasarkan latar belakang di atas, maka rumusan masalah adalah bagaimana membangun sistem yang baik?")
doc.add_paragraph("BAB II TINJAUAN PUSTAKA")
doc.add_paragraph("Landasan Teori") # Should become 2.1 Heading 2
doc.add_paragraph("Sistem Informasi") # Should become 2.1.1 Heading 3 (if detected) or just text
doc.add_paragraph("Sistem informasi adalah kombinasi dari teknologi informasi dan aktivitas orang.")
doc.add_paragraph("BAB III METODOLOGI PENELITIAN")
doc.add_paragraph("Analisis Kebutuhan")
doc.add_paragraph("Tahap ini melakukan analisis terhadap kebutuhan sistem.")
doc.save(DRAFT_PATH)

# 2. Run Enforcer
enforcer = SkripsiEnforcer(DRAFT_PATH)

# Phase 3.1 - 3.5
enforcer.phase_3_1_merge_bab_judul() # Won't do much on this clean input, but good practice
enforcer.phase_3_2_enforce_isiparagraf()
enforcer.phase_3_4_enforce_heading_styles()
enforcer.phase_3_5_page_breaks_numbering()
enforcer.phase_3_3_generate_toc()

# Phase 4: Front Matter
enforcer.execute_phase_4(
    judul="RANCANG BANGUN SISTEM INFORMASI AKADEMIK BERBASIS WEB PADA UNIVERSITAS CONTOH",
    penulis="AHMAD DAHLAN",
    nim="1900018291",
    universitas="UNIVERSITAS AHMAD DAHLAN",
    tahun=2025,
    abstrak_teks="Penelitian ini bertujuan untuk membangun sistem informasi akademik. Metode yang digunakan adalah Waterfall. Hasil pengujian menunjukkan sistem berjalan dengan baik.",
    abstrak_en_teks="This research aims to build an academic information system. The method used is Waterfall. Testing results show the system runs well.",
    kata_kunci="Sistem Informasi, Web, Akademik"
)

enforcer.save()

# 3. Output Representation
final_doc = Document(DRAFT_PATH)

print('<DOCUMENT filename="improved_formatted.doc">')

def get_style_repr(p):
    s = p.style.name
    if s == 'Heading 1': return '# '
    if s == 'Heading 2': return '## '
    if s == 'Heading 3': return '### '
    if 'List' in s: return '- '
    return ''

for p in final_doc.paragraphs:
    # visual page break
    if 'w:pageBreakBefore' in p._element.xml:
        print("\n[PAGE BREAK]\n")
        
    text = p.text.strip()
    if not text: continue
    
    prefix = get_style_repr(p)
    # Simulate content
    if p.style.name == 'IsiParagraf':
        print(f"{text}")
    elif 'Heading' in p.style.name:
        print(f"{prefix}{text.upper() if '1' in p.style.name else text.title()}")
    elif 'TOC' in p.style.name or 'instr' in p._element.xml:
        print(f"[TOC FIELD: {text}]")
    else:
        print(f"[{p.style.name}] {text}")

print('</DOCUMENT>')

# Cleanup
try:
    os.remove(DRAFT_PATH)
except:
    pass
