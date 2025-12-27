"""
Check output files for content placement errors.
"""

from docx import Document
import re
from pathlib import Path

def analyze_file(doc_path):
    """Analyze a document to find where content is incorrectly placed."""
    if not Path(doc_path).exists():
        print(f"[SKIP] File not found: {doc_path}")
        return
    
    print("=" * 80)
    print(f"ANALYZING: {Path(doc_path).name}")
    print("=" * 80)
    
    doc = Document(doc_path)
    
    # Find DAFTAR ISI section
    daftar_isi_start = None
    daftar_isi_end = None
    main_content_start = None
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().upper()
        
        if 'DAFTAR ISI' in text or 'TABLE OF CONTENTS' in text:
            daftar_isi_start = i
            print(f"[FOUND] DAFTAR ISI starts at paragraph {i}")
            # Find where DAFTAR ISI ends (look for next major section)
            for j in range(i + 1, min(i + 100, len(doc.paragraphs))):
                next_text = doc.paragraphs[j].text.strip().upper()
                if any(kw in next_text for kw in ['DAFTAR GAMBAR', 'DAFTAR TABEL', 'BAB I', 'CHAPTER 1', 'PENDAHULUAN']):
                    daftar_isi_end = j
                    print(f"[FOUND] DAFTAR ISI likely ends around paragraph {j}")
                    break
        
        if not main_content_start and ('BAB I' in text or 'BAB 1' in text or 'CHAPTER 1' in text):
            # Check if this is not a TOC entry
            if not ('\t' in para.text or re.search(r'\s+\d+\s*$', para.text)):
                main_content_start = i
                print(f"[FOUND] Main content (BAB I) starts at paragraph {i}")
    
    # Check for content in DAFTAR ISI area
    if daftar_isi_start is not None:
        print(f"\n[CHECKING] Content in DAFTAR ISI area (paragraphs {daftar_isi_start} to {daftar_isi_end or daftar_isi_start + 50}):")
        content_in_daftar_isi = []
        
        end_check = daftar_isi_end if daftar_isi_end else daftar_isi_start + 50
        
        for i in range(daftar_isi_start, min(end_check, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            # Skip empty or very short paragraphs
            if len(text) < 20:
                continue
            
            # Check if this looks like actual content (not TOC entry)
            is_toc_entry = '\t' in text or re.search(r'\s+\d+\s*$', text) or len(text) < 50
            
            # Check if this looks like chapter content
            is_content = (
                re.match(r'^\d+\.\d+', text) or  # Subsection numbering like "1.1"
                len(text) > 200 or  # Long paragraphs
                any(kw in text.upper() for kw in ['LATAR BELAKANG', 'RUMUSAN MASALAH', 'TUJUAN PENELITIAN', 'MANFAAT PENELITIAN'])
            )
            
            if is_content and not is_toc_entry:
                content_in_daftar_isi.append({
                    'para_idx': i,
                    'text': text[:150],
                    'length': len(text)
                })
        
        if content_in_daftar_isi:
            print(f"  [ERROR] Found {len(content_in_daftar_isi)} paragraphs with content in DAFTAR ISI area:")
            for item in content_in_daftar_isi[:5]:
                print(f"    Para {item['para_idx']} ({item['length']} chars): {item['text']}")
        else:
            print(f"  [OK] No content found in DAFTAR ISI area")
    
    print()

# Analyze the files
files = [
    r"C:\Folio\form-memory\storage\outputs\Skripsi_Dany.docx",
    r"C:\Folio\form-memory\storage\outputs\Skripsi_Pican.docx"
]

for file_path in files:
    analyze_file(file_path)

